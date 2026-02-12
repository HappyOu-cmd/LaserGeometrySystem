#!/usr/bin/env python3
"""
 –û–°–ù–û–í–ù–ê–Ø –°–ò–°–¢–ï–ú–ê –õ–ê–ó–ï–†–ù–û–ô –ì–ï–û–ú–ï–¢–†–ò–ò
–ò–Ω—Ç–µ–≥—Ä–∞—Ü–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ –†–§602, Modbus —Å–µ—Ä–≤–µ—Ä–∞ –∏ –∞–≤—Ç–æ–º–∞—Ç–∞ —Å–æ—Å—Ç–æ—è–Ω–∏–π
"""

import time
import threading
import struct
import os
import ctypes
import math
import contextlib
import fcntl
import sys
import shutil
from datetime import datetime
from queue import Queue, Empty
from enum import Enum
from typing import Optional, Tuple, List
from collections import deque

# –ü—Ä–æ–≤–µ—Ä—è–µ–º –¥–æ—Å—Ç—É–ø–Ω–æ—Å—Ç—å psutil
try:
    import psutil
    HAS_PSUTIL = True
except ImportError:
    HAS_PSUTIL = False

# –ò–º–ø–æ—Ä—Ç—ã –∏–∑ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏—Ö –º–æ–¥—É–ª–µ–π
import sys

try:
    from main import HighSpeedRiftekSensor, apply_system_optimizations
except ImportError:
    script_dir = os.path.dirname(os.path.abspath(__file__))
    if script_dir not in sys.path:
        sys.path.insert(0, script_dir)
    from main import HighSpeedRiftekSensor, apply_system_optimizations
from modbus_slave_server import ModbusSlaveServer
from modbus_database_integration import ModbusDatabaseIntegration

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def apply_laser_system_optimizations():
    """
    –ü—Ä–∏–º–µ–Ω–µ–Ω–∏–µ —Å–∏—Å—Ç–µ–º–Ω—ã—Ö –æ–ø—Ç–∏–º–∏–∑–∞—Ü–∏–π –¥–ª—è –ª–∞–∑–µ—Ä–Ω–æ–π —Å–∏—Å—Ç–µ–º—ã (–±–µ–∑ pyftdi)
    """
    print("[SYSTEM] –ü–†–ò–ú–ï–ù–ï–ù–ò–ï –°–ò–°–¢–ï–ú–ù–´–• –û–ü–¢–ò–ú–ò–ó–ê–¶–ò–ô...")
    
    # 1. –í—ã—Å–æ–∫–∏–π –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç –ø—Ä–æ—Ü–µ—Å—Å–∞ (REALTIME_PRIORITY_CLASS –¥–ª—è –º–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–π –ø—Ä–æ–∏–∑–≤–æ–¥–∏—Ç–µ–ª—å–Ω–æ—Å—Ç–∏)
    if HAS_PSUTIL:
        try:
            p = psutil.Process(os.getpid())
            # –ò—Å–ø–æ–ª—å–∑—É–µ–º REALTIME_PRIORITY_CLASS –¥–ª—è –º–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–π –ø—Ä–æ–∏–∑–≤–æ–¥–∏—Ç–µ–ª—å–Ω–æ—Å—Ç–∏
            # –≠—Ç–æ –∫—Ä–∏—Ç–∏—á–Ω–æ –¥–ª—è —Å—Ç–∞–±–∏–ª—å–Ω–æ–π —Ä–∞–±–æ—Ç—ã –Ω–∞ —Å–ª–∞–±—ã—Ö –ø—Ä–æ—Ü–µ—Å—Å–æ—Ä–∞—Ö
            try:
                p.nice(psutil.REALTIME_PRIORITY_CLASS)
                print("[OK] –£—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç REALTIME (–º–∞–∫—Å–∏–º–∞–ª—å–Ω—ã–π)")
            except:
                # –ï—Å–ª–∏ REALTIME –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω, –∏—Å–ø–æ–ª—å–∑—É–µ–º HIGH_PRIORITY
                p.nice(psutil.HIGH_PRIORITY_CLASS)
                print("[OK] –£—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –≤—ã—Å–æ–∫–∏–π –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç –ø—Ä–æ—Ü–µ—Å—Å–∞")
        except Exception as e:
            print(f"[WARNING] –ù–µ —É–¥–∞–ª–æ—Å—å —É—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç: {e}")
    else:
        print("[WARNING] psutil –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω - –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç –Ω–µ –∏–∑–º–µ–Ω–µ–Ω")
    
    # 2. –í—ã—Å–æ–∫–æ–µ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ —Ç–∞–π–º–µ—Ä–∞ Windows
    print("[TIMER] –ù–∞—Å—Ç—Ä–æ–π–∫–∞ —Ç–æ—á–Ω–æ—Å—Ç–∏ —Ç–∞–π–º–µ—Ä–∞ Windows...")
    try:
        # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–∏–Ω–∏–º–∞–ª—å–Ω–æ–µ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ —Ç–∞–π–º–µ—Ä–∞ (1 –º—Å)
        result = ctypes.windll.winmm.timeBeginPeriod(1)
        if result == 0:  # TIMERR_NOERROR
            print("[OK] –£—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–æ –≤—ã—Å–æ–∫–æ–µ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ —Ç–∞–π–º–µ—Ä–∞ (1 –º—Å)")
        else:
            print(f"[WARNING] –û—à–∏–±–∫–∞ —É—Å—Ç–∞–Ω–æ–≤–∫–∏ —Ç–∞–π–º–µ—Ä–∞: {result}")
    except Exception as e:
        print(f"[WARNING] –ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞—Å—Ç—Ä–æ–∏—Ç—å —Ç–∞–π–º–µ—Ä: {e}")
    
    # 3. –û—Ç–∫–ª—é—á–µ–Ω–∏–µ —Å–ø—è—â–µ–≥–æ —Ä–µ–∂–∏–º–∞
    try:
        # –û—Ç–∫–ª—é—á–∞–µ–º —Å–ø—è—â–∏–π —Ä–µ–∂–∏–º –¥–ª—è —Ç–µ–∫—É—â–µ–≥–æ –ø—Ä–æ—Ü–µ—Å—Å–∞
        ctypes.windll.kernel32.SetThreadExecutionState(0x80000001)  # ES_CONTINUOUS | ES_SYSTEM_REQUIRED
        print("[OK] –û—Ç–∫–ª—é—á–µ–Ω —Å–ø—è—â–∏–π —Ä–µ–∂–∏–º —Å–∏—Å—Ç–µ–º—ã")
    except Exception as e:
        print(f"[WARNING] –ù–µ —É–¥–∞–ª–æ—Å—å –æ—Ç–∫–ª—é—á–∏—Ç—å —Å–ø—è—â–∏–π —Ä–µ–∂–∏–º: {e}")
    
    print("[SYSTEM] –°–∏—Å—Ç–µ–º–Ω—ã–µ –æ–ø—Ç–∏–º–∏–∑–∞—Ü–∏–∏ –ø—Ä–∏–º–µ–Ω–µ–Ω—ã\n")


def cleanup_laser_system_optimizations():
    """
    –û—á–∏—Å—Ç–∫–∞ —Å–∏—Å—Ç–µ–º–Ω—ã—Ö –æ–ø—Ç–∏–º–∏–∑–∞—Ü–∏–π –ø—Ä–∏ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–∏
    """
    try:
        # –í–æ—Å—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –æ–±—ã—á–Ω–æ–µ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ —Ç–∞–π–º–µ—Ä–∞
        ctypes.windll.winmm.timeEndPeriod(1)
        print("üîß –í–æ—Å—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–æ —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω–æ–µ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ —Ç–∞–π–º–µ—Ä–∞")
    except:
        pass
    
    try:
        # –í–æ—Å—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –æ–±—ã—á–Ω–æ–µ —É–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø–∏—Ç–∞–Ω–∏–µ–º
        ctypes.windll.kernel32.SetThreadExecutionState(0x80000000)  # ES_CONTINUOUS
        print("üîß –í–æ—Å—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–æ —É–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø–∏—Ç–∞–Ω–∏–µ–º")
    except:
        pass


class SystemState(Enum):
    """–°–æ—Å—Ç–æ—è–Ω–∏—è —Å–∏—Å—Ç–µ–º—ã"""
    IDLE = "IDLE"
    
    # –ö–∞–ª–∏–±—Ä–æ–≤–∫–∏
    CALIBRATE_WALL = "CALIBRATE_WALL"
    CALIBRATE_BOTTOM = "CALIBRATE_BOTTOM"
    CALIBRATE_FLANGE = "CALIBRATE_FLANGE"
    CALIBRATE_HEIGHT = "CALIBRATE_HEIGHT"
    CALIBRATE_FLANGE_DIAMETER = "CALIBRATE_FLANGE_DIAMETER"
    CALIBRATE_BODY_DIAMETER_SEPARATE = "CALIBRATE_BODY_DIAMETER_SEPARATE"  # CMD=107
    CALIBRATE_BODY2_DIAMETER = "CALIBRATE_BODY2_DIAMETER"  # CMD=108
    DEBUG_REGISTERS = "DEBUG_REGISTERS"
    CONFIGURE_SENSOR3_RANGE = "CONFIGURE_SENSOR3_RANGE"  # CMD=106: –Ω–∞—Å—Ç—Ä–æ–π–∫–∞ –¥–∏–∞–ø–∞–∑–æ–Ω–æ–≤ –¥–∞—Ç—á–∏–∫–∞ 3
    
    # –ò–∑–º–µ—Ä–µ–Ω–∏–µ –≤—ã—Å–æ—Ç—ã
    MEASURE_HEIGHT_PROCESS = "MEASURE_HEIGHT_PROCESS"      # CMD=9: –ø–æ–∏—Å–∫ –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏—è –∏ —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
    
    # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - –≤–µ—Ä—Ö–Ω—è—è —Å—Ç–µ–Ω–∫–∞
    MEASURE_WALL_PROCESS = "MEASURE_WALL_PROCESS"      # CMD=10: —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
    MEASURE_WALL_CALCULATE = "MEASURE_WALL_CALCULATE"  # CMD=11: –ø–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤
    
    # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - —Ñ–ª–∞–Ω–µ—Ü
    MEASURE_FLANGE_PROCESS = "MEASURE_FLANGE_PROCESS"      # CMD=12: —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
    MEASURE_FLANGE_CALCULATE = "MEASURE_FLANGE_CALCULATE"  # CMD=13: –ø–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤
    MEASURE_FLANGE_ONLY_PROCESS = "MEASURE_FLANGE_ONLY_PROCESS"      # CMD=20
    MEASURE_FLANGE_ONLY_CALCULATE = "MEASURE_FLANGE_ONLY_CALCULATE"  # CMD=21
    MEASURE_BODY_ONLY_PROCESS = "MEASURE_BODY_ONLY_PROCESS"      # CMD=30
    MEASURE_BODY_ONLY_CALCULATE = "MEASURE_BODY_ONLY_CALCULATE"  # CMD=31
    MEASURE_BODY2_PROCESS = "MEASURE_BODY2_PROCESS"      # CMD=40
    MEASURE_BODY2_CALCULATE = "MEASURE_BODY2_CALCULATE"  # CMD=41
    
    # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - –Ω–∏–∂–Ω—è—è —Å—Ç–µ–Ω–∫–∞
    MEASURE_BOTTOM_PROCESS = "MEASURE_BOTTOM_PROCESS"      # CMD=14: —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
    MEASURE_BOTTOM_CALCULATE = "MEASURE_BOTTOM_CALCULATE"  # CMD=15: –ø–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤
    
    # –û—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞ –∏–∑–¥–µ–ª–∏—è
    QUALITY_EVALUATION = "QUALITY_EVALUATION"  # CMD=16: –æ—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞
    
    # –ü–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º (QUAD - –≤—Å–µ 4 –¥–∞—Ç—á–∏–∫–∞)
    STREAM_QUAD = "STREAM_QUAD"  # CMD=200: QUAD —Ä–µ–∂–∏–º –≤—Å–µ—Ö –¥–∞—Ç—á–∏–∫–æ–≤
    
    ERROR = "ERROR"


class LaserGeometrySystem:
    """–û—Å–Ω–æ–≤–Ω–∞—è —Å–∏—Å—Ç–µ–º–∞ –ª–∞–∑–µ—Ä–Ω–æ–π –≥–µ–æ–º–µ—Ç—Ä–∏–∏"""
    
    def __init__(self, port: str = '/dev/ttyUSB0', baudrate: int = 921600, modbus_port: int = 502, 
                 test_mode: bool = False):
        """
        –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è —Å–∏—Å—Ç–µ–º—ã
        
        Args:
            port: COM –ø–æ—Ä—Ç –¥–ª—è –¥–∞—Ç—á–∏–∫–æ–≤
            baudrate: –°–∫–æ—Ä–æ—Å—Ç—å –ø–µ—Ä–µ–¥–∞—á–∏ –¥–∞–Ω–Ω—ã—Ö
            modbus_port: –ü–æ—Ä—Ç Modbus —Å–µ—Ä–≤–µ—Ä–∞
            test_mode: –†–µ–∂–∏–º —Ç–µ—Å—Ç–∏—Ä–æ–≤–∞–Ω–∏—è –±–µ–∑ —Ä–µ–∞–ª—å–Ω—ã—Ö –¥–∞—Ç—á–∏–∫–æ–≤
        """
        # –ù–∞—Å—Ç—Ä–æ–π–∫–∏ –¥–∞—Ç—á–∏–∫–æ–≤
        self.port = port
        self.baudrate = baudrate
        self.test_mode = test_mode
        
        # –ö–æ–º–ø–æ–Ω–µ–Ω—Ç—ã —Å–∏—Å—Ç–µ–º—ã
        self.sensors = None
        self.modbus_server = None
        self.db_integration = None
        self.data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
        try:
            os.makedirs(self.data_dir, exist_ok=True)
        except Exception as e:
            print(f" [REPORT] –ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ–∑–¥–∞—Ç—å –ø–∞–ø–∫—É –¥–ª—è –æ—Ç—á—ë—Ç–æ–≤ {self.data_dir}: {e}")
        
        # –°–æ—Å—Ç–æ—è–Ω–∏–µ —Å–∏—Å—Ç–µ–º—ã
        self.current_state = SystemState.IDLE
        self.previous_cmd = 0
        
        # –î–∞–Ω–Ω—ã–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏
        self.calibration_data = {
            'wall_distance_1_2': 0.0,  # –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,2
            'wall_distance_1_3': 0.0,  # –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,3
            'bottom_distance_4': 0.0,  # –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫ 4 –¥–æ –ø–æ–≤–µ—Ä—Ö–Ω–æ—Å—Ç–∏
            'flange_distance_1_center': 0.0,  # –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫ 1 –¥–æ —Ü–µ–Ω—Ç—Ä–∞
        }
        
        # –ë—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è
        self.measurement_buffer = {
            'sensor1': deque(maxlen=1000),
            'sensor2': deque(maxlen=1000),
            'sensor3': deque(maxlen=1000),
            'sensor4': deque(maxlen=1000),
        }
        
        
        # –ü–∞—Ä–∞–º–µ—Ç—Ä—ã –∏–∑–º–µ—Ä–µ–Ω–∏–π
        self.measurement_duration = 4.0  # —Å–µ–∫—É–Ω–¥ –¥–ª—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏
        self.sensor_range_mm = 25.0      # –¥–∏–∞–ø–∞–∑–æ–Ω –¥–∞—Ç—á–∏–∫–æ–≤
        self.base_distance_mm = 20.0     # –±–∞–∑–æ–≤–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ
        
        # –§–ª–∞–≥–∏
        self.is_running = False
        self.calibration_in_progress = False
        self.stream_active_quad = False  # –§–ª–∞–≥ –∞–∫—Ç–∏–≤–Ω–æ–≥–æ QUAD —Ä–µ–∂–∏–º–∞
        
        # –ü–æ—Ç–æ–∫–∏ –¥–ª—è —Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ –∏ –æ–±—Ä–∞–±–æ—Ç–∫–∏ –¥–∞–Ω–Ω—ã—Ö
        self.sensor_reading_thread = None
        self.sensor_data_queue = Queue(maxsize=1000)  # –û—á–µ—Ä–µ–¥—å –¥–ª—è –ø–µ—Ä–µ–¥–∞—á–∏ –¥–∞–Ω–Ω—ã—Ö –æ—Ç –¥–∞—Ç—á–∏–∫–æ–≤
        self.sensor_reading_active = False  # –§–ª–∞–≥ –∞–∫—Ç–∏–≤–Ω–æ—Å—Ç–∏ –ø–æ—Ç–æ–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤
        self.sensor_reading_lock = threading.Lock()  # –ë–ª–æ–∫–∏—Ä–æ–≤–∫–∞ –¥–ª—è —Å–∏–Ω—Ö—Ä–æ–Ω–∏–∑–∞—Ü–∏–∏ –¥–æ—Å—Ç—É–ø–∞ –∫ –¥–∞—Ç—á–∏–∫–∞–º
        self.height_calibration_nonzero_count = 0  # –°—á–µ—Ç—á–∏–∫ –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏–π –¥–ª—è CMD=103
        self.distance_to_plane_calculated = False  # –§–ª–∞–≥ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è —Ä–∞—Å—á—ë—Ç–∞ –¥–∏—Å—Ç–∞–Ω—Ü–∏–∏ (CMD=103)
        self.recent_measurements = []  # –ë—É—Ñ–µ—Ä –ø–æ—Å–ª–µ–¥–Ω–∏—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–ª—è CMD=103
        
        # –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–µ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–æ–≤
        self.last_reconnect_attempt = 0  # –í—Ä–µ–º—è –ø–æ—Å–ª–µ–¥–Ω–µ–π –ø–æ–ø—ã—Ç–∫–∏ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è
        self.reconnect_interval = 5.0  # –ò–Ω—Ç–µ—Ä–≤–∞–ª –ø–æ–ø—ã—Ç–æ–∫ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è (—Å–µ–∫—É–Ω–¥—ã)
        
        # –°—á–µ—Ç—á–∏–∫–∏ –¥–ª—è –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ —Ä–µ–∂–∏–º–∞ QUAD
        self.stream_measurement_count = 0
        self.stream_start_time = None
        # –ë—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è –ø–æ 10 –∏–∑–º–µ—Ä–µ–Ω–∏—è–º –¥–ª—è –∫–∞–∂–¥–æ–≥–æ –¥–∞—Ç—á–∏–∫–∞
        self.stream_temp_sensor1_buffer = []
        self.stream_temp_sensor2_buffer = []
        self.stream_temp_sensor3_buffer = []
        self.stream_temp_sensor4_buffer = []
        
        # –ë—É—Ñ–µ—Ä—ã –¥–ª—è –æ—Å–Ω–æ–≤–Ω–æ–≥–æ —Ü–∏–∫–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è
        self.sensor1_measurements = []  # –ë—É—Ñ–µ—Ä —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1
        self.sensor2_measurements = []  # –ë—É—Ñ–µ—Ä —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 2
        self.wall_thickness_buffer = []  # –ë—É—Ñ–µ—Ä —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–Ω–æ–π —Ç–æ–ª—â–∏–Ω—ã —Å—Ç–µ–Ω–∫–∏
        self.measurement_cycle_active = False  # –§–ª–∞–≥ –∞–∫—Ç–∏–≤–Ω–æ–≥–æ —Ü–∏–∫–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è
        
        # –§–ª–∞–≥–∏ –≤—ã–ø–æ–ª–Ω–µ–Ω–∏—è —Ä–∞—Å—á—ë—Ç–æ–≤ (—á—Ç–æ–±—ã –Ω–µ –≤—ã–ø–æ–ª–Ω—è—Ç—å –º–Ω–æ–≥–æ–∫—Ä–∞—Ç–Ω–æ)
        self.wall_calculated = False
        self.flange_calculated = False
        self.bottom_calculated = False
        self.flange_only_calculated = False
        self.body_only_calculated = False
        self.body2_calculated = False
        self.quality_evaluated = False
        self.body2_quality_required = False
        
        # –ú–æ–Ω–∏—Ç–æ—Ä–∏–Ω–≥ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–∏ –æ–∂–∏–¥–∞–Ω–∏—è (IDLE)
        self.idle_monitor_last_time = 0.0

        # –ù–æ–º–µ—Ä —Å–º–µ–Ω—ã (–¥–ª—è —Å–±—Ä–æ—Å–∞ —Å—á—ë—Ç—á–∏–∫–æ–≤ –ø—Ä–∏ —Å–º–µ–Ω–µ)
        self.last_shift_number = None
        
        # –ö–µ—à –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω—ã—Ö —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–π (–¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è —Ü–∏–∫–ª–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è)
        self.cached_distance_1_2 = None
        self.cached_distance_to_center = None
        self.cached_distance_1_3 = None
        self.cached_distance_sensor4 = None
        self.cached_distance_sensor3_to_center = None
        self.cached_distance_sensor3_to_center_body = None
        self.cached_distance_sensor3_to_center_body2 = None
        
        # –û—Ç—Å–ª–µ–∂–∏–≤–∞–Ω–∏–µ —Å–º–µ–Ω—ã –¥–ª—è —Å–±—Ä–æ—Å–∞ —Å—á—ë—Ç—á–∏–∫–æ–≤
        self.current_shift_number = 1  # –¢–µ–∫—É—â–∞—è —Å–º–µ–Ω–∞
        self.shift_initialized = False  # –§–ª–∞–≥, —á—Ç–æ–±—ã –Ω–µ —Å–±—Ä–∞—Å—ã–≤–∞—Ç—å –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∑–∞–ø—É—Å–∫–µ
        
        # –ú–æ–Ω–∏—Ç–æ—Ä–∏–Ω–≥ —á–∞—Å—Ç–æ—Ç—ã –æ–ø—Ä–æ—Å–∞
        self.frequency_counter = 0
        self.frequency_start_time = None
        self.last_frequency_display = 0
        
        # –ë—É—Ñ–µ—Ä—ã –¥–ª—è –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤—ã—Å–æ—Ç—ã (–∫–æ–º–∞–Ω–¥–∞ 9)
        self.height_measurements = []  # –ë—É—Ñ–µ—Ä –∏–∑–º–µ—Ä–µ–Ω–∏–π –≤—ã—Å–æ—Ç—ã
        self.obstacle_detected = False  # –§–ª–∞–≥ –æ–±–Ω–∞—Ä—É–∂–µ–Ω–∏—è –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏—è
        self.obstacle_filter_count = 0  # –°—á–µ—Ç—á–∏–∫ –¥–ª—è —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏–∏ –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏–π
        self.height_calculated = False  # –§–ª–∞–≥ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è —Ä–∞—Å—á–µ—Ç–∞ –≤—ã—Å–æ—Ç—ã
        
        # –í—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è (–∫–æ–º–∞–Ω–¥–∞ 10)
        self.temp_sensor1_buffer = []  # –í—Ä–µ–º–µ–Ω–Ω—ã–π –±—É—Ñ–µ—Ä –¥–ª—è 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1
        self.temp_sensor2_buffer = []  # –í—Ä–µ–º–µ–Ω–Ω—ã–π –±—É—Ñ–µ—Ä –¥–ª—è 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 2
        
        # –ë—É—Ñ–µ—Ä—ã –¥–ª—è –∫–æ–º–∞–Ω–¥—ã 11 (–∏–∑–º–µ—Ä–µ–Ω–∏–µ —Ñ–ª–∞–Ω—Ü–∞)
        self.sensor1_flange_measurements = []  # –ë—É—Ñ–µ—Ä —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1 –¥–ª—è –∫–æ–º–∞–Ω–¥—ã 11
        self.sensor3_measurements = []  # –ë—É—Ñ–µ—Ä —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 3
        self.sensor4_measurements = []  # –ë—É—Ñ–µ—Ä —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 4
        
        # –í—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è (–∫–æ–º–∞–Ω–¥–∞ 11)
        self.temp_sensor1_flange_buffer = []  # –í—Ä–µ–º–µ–Ω–Ω—ã–π –±—É—Ñ–µ—Ä –¥–ª—è 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1 (–∫–æ–º–∞–Ω–¥–∞ 11)
        self.temp_sensor3_buffer = []  # –í—Ä–µ–º–µ–Ω–Ω—ã–π –±—É—Ñ–µ—Ä –¥–ª—è 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 3
        self.temp_sensor4_buffer = []  # –í—Ä–µ–º–µ–Ω–Ω—ã–π –±—É—Ñ–µ—Ä –¥–ª—è 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 4
        
        # –†–∞—Å—á–µ—Ç–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è –∫–æ–º–∞–Ω–¥—ã 11
        self.body_diameter_buffer = []    # –ë—É—Ñ–µ—Ä –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ (–¥–∞—Ç—á–∏–∫ 1)
        self.flange_diameter_buffer = []  # –ë—É—Ñ–µ—Ä –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞ (–¥–∞—Ç—á–∏–∫ 3)
        self.bottom_thickness_buffer = [] # –ë—É—Ñ–µ—Ä —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞ (–¥–∞—Ç—á–∏–∫ 4)
        
        # –ë—É—Ñ–µ—Ä—ã –¥–ª—è –∫–æ–º–∞–Ω–¥—ã 12 (–∏–∑–º–µ—Ä–µ–Ω–∏–µ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏)
        self.sensor1_bottom_measurements = []  # –ë—É—Ñ–µ—Ä —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1 –¥–ª—è –∫–æ–º–∞–Ω–¥—ã 12
        self.sensor2_bottom_measurements = []  # –ë—É—Ñ–µ—Ä —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 2 –¥–ª—è –∫–æ–º–∞–Ω–¥—ã 12
        
        # –í—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è (–∫–æ–º–∞–Ω–¥–∞ 12)
        self.temp_sensor1_bottom_buffer = []  # –í—Ä–µ–º–µ–Ω–Ω—ã–π –±—É—Ñ–µ—Ä –¥–ª—è 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1 (–∫–æ–º–∞–Ω–¥–∞ 12)
        self.temp_sensor2_bottom_buffer = []  # –í—Ä–µ–º–µ–Ω–Ω—ã–π –±—É—Ñ–µ—Ä –¥–ª—è 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 2 (–∫–æ–º–∞–Ω–¥–∞ 12)
        
        # –†–∞—Å—á–µ—Ç–Ω—ã–π –±—É—Ñ–µ—Ä –¥–ª—è –∫–æ–º–∞–Ω–¥—ã 12
        self.bottom_wall_thickness_buffer = []  # –ë—É—Ñ–µ—Ä —Ç–æ–ª—â–∏–Ω—ã –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏

        # –ë—É—Ñ–µ—Ä—ã –¥–ª—è —Ä–∞–∑–¥–µ–ª—å–Ω—ã—Ö –∫–æ–º–∞–Ω–¥ –¥–∏–∞–º–µ—Ç—Ä–æ–≤ (20/30/40)
        self.sensor3_flange_only_measurements = []
        self.sensor3_body_only_measurements = []
        self.sensor3_body2_measurements = []
        self.temp_sensor3_flange_only_buffer = []
        self.temp_sensor3_body_only_buffer = []
        self.temp_sensor3_body2_buffer = []
        self.body_only_diameter_buffer = []
        self.body2_diameter_buffer = []
        
    def start_system(self):
        """–ó–∞–ø—É—Å–∫ —Å–∏—Å—Ç–µ–º—ã"""
        print("–ó–ê–ü–£–°–ö –°–ò–°–¢–ï–ú–´ –õ–ê–ó–ï–†–ù–û–ô –ì–ï–û–ú–ï–¢–†–ò–ò")
        print("=" * 50)
        
        # –ü—Ä–∏–º–µ–Ω—è–µ–º —Å–∏—Å—Ç–µ–º–Ω—ã–µ –æ–ø—Ç–∏–º–∏–∑–∞—Ü–∏–∏ –¥–ª—è –ª–∞–∑–µ—Ä–Ω–æ–π —Å–∏—Å—Ç–µ–º—ã
        apply_laser_system_optimizations()
        
        # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –¥–∞—Ç—á–∏–∫–æ–≤
        sensors_connected = False
        if self.test_mode:
            print("–¢–ï–°–¢–û–í–´–ô –†–ï–ñ–ò–ú - –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã")
            self.sensors = None
        else:
            print(f"–ü–æ–¥–∫–ª—é—á–µ–Ω–∏–µ –∫ –¥–∞—Ç—á–∏–∫–∞–º –Ω–∞ –ø–æ—Ä—Ç—É {self.port}...")
            self.sensors = HighSpeedRiftekSensor(self.port, self.baudrate, timeout=0.002)
            
            if not self.sensors.connect():
                print(" –í–ù–ò–ú–ê–ù–ò–ï: –û—à–∏–±–∫–∞ –ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è –∫ –¥–∞—Ç—á–∏–∫–∞–º!")
                print(" –ü—Ä–æ–≥—Ä–∞–º–º–∞ –ø—Ä–æ–¥–æ–ª–∂–∏—Ç —Ä–∞–±–æ—Ç—É, –Ω–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –±—É–¥—É—Ç –Ω–µ–¥–æ—Å—Ç—É–ø–Ω—ã.")
                self.sensors = None  # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —É–∫–∞–∑–∞—Ç–µ–ª—å –Ω–∞ –¥–∞—Ç—á–∏–∫–∏
                sensors_connected = False
            else:
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–æ—Å–ª–µ –ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è
                self.clear_serial_buffers()
        print("OK –î–∞—Ç—á–∏–∫–∏ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã")
        sensors_connected = True
        
        # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è Modbus —Å–µ—Ä–≤–µ—Ä–∞ (–±–µ–∑ GUI, —Ç–∞–∫ –∫–∞–∫ —É –Ω–∞—Å –µ—Å—Ç—å Debug GUI)
        self.modbus_server = ModbusSlaveServer(enable_gui=False)
        
        # –ó–∞–ø—É—Å–∫–∞–µ–º Modbus —Å–µ—Ä–≤–µ—Ä –±–µ–∑ GUI
        try:
            self.modbus_server.start_modbus_server()
            print("OK Modbus —Å–µ—Ä–≤–µ—Ä –∑–∞–ø—É—â–µ–Ω")
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ –∑–∞–ø—É—Å–∫–∞ Modbus —Å–µ—Ä–≤–µ—Ä–∞: {e}")
            return False
        
        # –¢–µ–ø–µ—Ä—å –º–æ–∂–µ–º —É—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –±–∏—Ç –æ—à–∏–±–∫–∏ –ø–æ—Å–ª–µ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏ Modbus —Å–µ—Ä–≤–µ—Ä–∞
        if not sensors_connected:
            self.set_error_bit(0, True)  # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –±–∏—Ç 0 - –æ—à–∏–±–∫–∞ –ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤
        else:
            self.set_error_bit(0, False)  # –°–±—Ä–∞—Å—ã–≤–∞–µ–º –±–∏—Ç –æ—à–∏–±–∫–∏ –µ—Å–ª–∏ –ø–æ–¥–∫–ª—é—á–µ–Ω–∏–µ —É—Å–ø–µ—à–Ω–æ

        # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º –Ω–æ–º–µ—Ä —Å–º–µ–Ω—ã –¥–ª—è –¥–µ—Ç–µ–∫—Ç–æ—Ä–∞ —Å–º–µ–Ω—ã
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # 40100 -> –∏–Ω–¥–µ–∫—Å 99 –≤ Holding —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ö
                current_shift = self.modbus_server.slave_context.getValues(3, 99, 1)[0]
                self.last_shift_number = int(current_shift)
                print(f" [SHIFT] –¢–µ–∫—É—â–∞—è —Å–º–µ–Ω–∞: {self.last_shift_number}")
        except Exception as e:
            print(f" [SHIFT] –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ç–µ–∫—É—â–µ–π —Å–º–µ–Ω—ã –ø—Ä–∏ —Å—Ç–∞—Ä—Ç–µ: {e}")
        
        # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –∏–Ω—Ç–µ–≥—Ä–∞—Ü–∏–∏ —Å –±–∞–∑–æ–π –¥–∞–Ω–Ω—ã—Ö
        self.db_integration = ModbusDatabaseIntegration(self.modbus_server)
        
        # –ó–∞–≥—Ä—É–∂–∞–µ–º —Å–æ—Ö—Ä–∞–Ω–µ–Ω–Ω—ã–µ —Ä–µ–≥–∏—Å—Ç—Ä—ã –∏–∑ –±–∞–∑—ã –¥–∞–Ω–Ω—ã—Ö
        self.db_integration.load_all_registers_from_db()
        
        # –ü–æ—Å–ª–µ –∑–∞–≥—Ä—É–∑–∫–∏ –ë–î –ø–æ–≤—Ç–æ—Ä–Ω–æ —Å–∏–Ω—Ö—Ä–æ–Ω–∏–∑–∏—Ä—É–µ–º –Ω–æ–º–µ—Ä —Å–º–µ–Ω—ã,
        # —á—Ç–æ–±—ã –∏–∑–±–µ–∂–∞—Ç—å –ª–æ–∂–Ω–æ–≥–æ —Å—Ä–∞–±–∞—Ç—ã–≤–∞–Ω–∏—è –¥–µ—Ç–µ–∫—Ç–æ—Ä–∞ —Å–º–µ–Ω—ã
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                current_shift = self.modbus_server.slave_context.getValues(3, 99, 1)[0]
                self.last_shift_number = int(current_shift)
                print(f" [SHIFT] –°–º–µ–Ω—ã —Å–∏–Ω—Ö—Ä–æ–Ω–∏–∑–∏—Ä–æ–≤–∞–Ω—ã –ø–æ—Å–ª–µ –∑–∞–≥—Ä—É–∑–∫–∏ –ë–î: {self.last_shift_number}")
        except Exception as e:
            print(f" [SHIFT] –û—à–∏–±–∫–∞ –ø–æ–≤—Ç–æ—Ä–Ω–æ–π —Å–∏–Ω—Ö—Ä–æ–Ω–∏–∑–∞—Ü–∏–∏ —Å–º–µ–Ω—ã: {e}")
        
        # –ó–∞–ø—É—Å–∫–∞–µ–º –º–æ–Ω–∏—Ç–æ—Ä–∏–Ω–≥ –∏–∑–º–µ–Ω–µ–Ω–∏–π —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤
        self.db_integration.start_monitoring(interval=1.0)
        print("OK –ò–Ω—Ç–µ–≥—Ä–∞—Ü–∏—è —Å –±–∞–∑–æ–π –¥–∞–Ω–Ω—ã—Ö –∑–∞–ø—É—â–µ–Ω–∞")
        
        # –§–ª–∞–≥ —Ä–∞–±–æ—Ç—ã —Å–∏—Å—Ç–µ–º—ã –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –ø–æ–¥–Ω—è—Ç –î–û –∑–∞–ø—É—Å–∫–∞ –≤—Å–ø–æ–º–æ–≥–∞—Ç–µ–ª—å–Ω—ã—Ö –ø–æ—Ç–æ–∫–æ–≤,
        # –∏–Ω–∞—á–µ –ø–æ—Ç–æ–∫ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ –∑–∞–≤–µ—Ä—à–∏—Ç—Å—è, –Ω–µ –≤–æ–π–¥—è –≤ —Ü–∏–∫–ª
        self.is_running = True

        # –ó–∞–ø—É—Å–∫ –ø–æ—Ç–æ–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ –¥–ª—è –¥–æ—Å—Ç–∞–≤–∫–∏ –¥–∞–Ω–Ω—ã—Ö —á–µ—Ä–µ–∑ –æ—á–µ—Ä–µ–¥—å
        if self.sensors and not self.test_mode and not self.sensor_reading_active:
            self.sensor_reading_active = True
            self.sensor_reading_thread = threading.Thread(target=self.sensor_reading_loop, daemon=True)
            self.sensor_reading_thread.start()
            print("OK –ü–æ—Ç–æ–∫ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ –∑–∞–ø—É—â–µ–Ω")
        
        # –ó–∞–ø—É—Å–∫ –æ—Å–Ω–æ–≤–Ω–æ–≥–æ —Ü–∏–∫–ª–∞ –≤ –æ—Ç–¥–µ–ª—å–Ω–æ–º –ø–æ—Ç–æ–∫–µ
        main_thread = threading.Thread(target=self.main_loop, daemon=True)
        main_thread.start()
        print("OK –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∑–∞–ø—É—â–µ–Ω")
        
        # –ñ–¥–µ–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –æ—Å–Ω–æ–≤–Ω–æ–≥–æ –ø–æ—Ç–æ–∫–∞
        main_thread.join()
        
        return True
    
    def stop_system(self):
        """–û—Å—Ç–∞–Ω–æ–≤–∫–∞ —Å–∏—Å—Ç–µ–º—ã"""
        print("\n –û–°–¢–ê–ù–û–í–ö–ê –°–ò–°–¢–ï–ú–´")
        self.is_running = False
        
        # –û—Å—Ç–∞–Ω–æ–≤–∫–∞ –ø–æ—Ç–æ–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤
        if self.sensor_reading_thread:
            try:
                self.sensor_reading_active = False
                # –ñ–¥–µ–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –ø–æ—Ç–æ–∫–∞ (–º–∞–∫—Å–∏–º—É–º 1 —Å–µ–∫—É–Ω–¥–∞)
                self.sensor_reading_thread.join(timeout=1.0)
                print(" –û—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –ø–æ—Ç–æ–∫ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤")
            except Exception as e:
                print(f" –û—à–∏–±–∫–∞ –æ—Å—Ç–∞–Ω–æ–≤–∫–∏ –ø–æ—Ç–æ–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤: {e}")
        
        # –û—Å—Ç–∞–Ω–æ–≤–∫–∞ QUAD –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ —Ä–µ–∂–∏–º–∞
        if self.sensors and self.stream_active_quad:
            try:
                self.stream_active_quad = False
                print(" –û—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω QUAD –ø–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º")
            except Exception as e:
                print(f" –û—à–∏–±–∫–∞ –æ—Å—Ç–∞–Ω–æ–≤–∫–∏ QUAD —Ä–µ–∂–∏–º–∞: {e}")
        
        if self.sensors:
            try:
                self.sensors.disconnect()
                print(" –î–∞—Ç—á–∏–∫–∏ –æ—Ç–∫–ª—é—á–µ–Ω—ã")
            except Exception as e:
                print(f" –û—à–∏–±–∫–∞ –æ—Ç–∫–ª—é—á–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤: {e}")
            finally:
                self.sensors = None
            
        if self.db_integration:
            try:
                self.db_integration.stop_monitoring()
            except Exception as e:
                print(f"–û—à–∏–±–∫–∞ –æ—Å—Ç–∞–Ω–æ–≤–∫–∏ –º–æ–Ω–∏—Ç–æ—Ä–∏–Ω–≥–∞ –ë–î: {e}")
            
        if self.modbus_server:
            try:
                self.modbus_server.stop_modbus_server()
            except Exception as e:
                print(f"–û—à–∏–±–∫–∞ –æ—Å—Ç–∞–Ω–æ–≤–∫–∏ Modbus —Å–µ—Ä–≤–µ—Ä–∞: {e}")
            
        # –û—á–∏—â–∞–µ–º —Å–∏—Å—Ç–µ–º–Ω—ã–µ –æ–ø—Ç–∏–º–∏–∑–∞—Ü–∏–∏
        cleanup_laser_system_optimizations()
            
        print(" –°–∏—Å—Ç–µ–º–∞ –æ—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–∞")
    
    def sensor_reading_loop(self):
        """
        –û—Ç–¥–µ–ª—å–Ω—ã–π –ø–æ—Ç–æ–∫ –¥–ª—è —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ —Å –≤—ã—Å–æ–∫–∏–º –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç–æ–º
        –¢–æ–ª—å–∫–æ —á—Ç–µ–Ω–∏–µ –¥–∞–Ω–Ω—ã—Ö, –±–µ–∑ –æ–±—Ä–∞–±–æ—Ç–∫–∏ –∏ –∑–∞–ø–∏—Å–∏ –≤ Modbus
        """
        print(" –ó–∞–ø—É—Å–∫ –ø–æ—Ç–æ–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤...")
        
        # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≤—ã—Å–æ–∫–∏–π –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç –¥–ª—è –ø–æ—Ç–æ–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤
        if HAS_PSUTIL:
            try:
                thread_id = threading.current_thread().ident
                if thread_id:
                    # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≤—ã—Å–æ–∫–∏–π –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç –¥–ª—è —Ç–µ–∫—É—â–µ–≥–æ –ø–æ—Ç–æ–∫–∞
                    import ctypes
                    kernel32 = ctypes.windll.kernel32
                    THREAD_PRIORITY_HIGHEST = 2
                    kernel32.SetThreadPriority(kernel32.OpenThread(0x1F03FF, False, thread_id), THREAD_PRIORITY_HIGHEST)
                    print(" [SENSOR THREAD] –£—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –≤—ã—Å–æ–∫–∏–π –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç –¥–ª—è –ø–æ—Ç–æ–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤")
            except Exception as e:
                print(f" [SENSOR THREAD] –ù–µ —É–¥–∞–ª–æ—Å—å —É—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç: {e}")
        
        try:
            while self.sensor_reading_active and self.is_running:
                if not self.sensors:
                    time.sleep(0.01)  # –ù–µ–±–æ–ª—å—à–∞—è –ø–∞—É–∑–∞ –µ—Å–ª–∏ –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã
                    continue
                
                # –ß–∏—Ç–∞–µ–º –¥–∞–Ω–Ω—ã–µ —Å –¥–∞—Ç—á–∏–∫–æ–≤ (—Ç–æ–ª—å–∫–æ —á—Ç–µ–Ω–∏–µ, –±–µ–∑ –æ–±—Ä–∞–±–æ—Ç–∫–∏)
                try:
                    with self.sensor_reading_lock:
                        sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.sensors.perform_quad_sensor_measurement(
                            self.sensor_range_mm, self.sensor_range_mm, 
                            self.sensor_range_mm, self.sensor_range_mm
                        )
                    
                    # –ü–æ–º–µ—â–∞–µ–º –¥–∞–Ω–Ω—ã–µ –≤ –æ—á–µ—Ä–µ–¥—å (–Ω–µ–±–ª–æ–∫–∏—Ä—É—é—â–∞—è –∑–∞–ø–∏—Å—å)
                    try:
                        self.sensor_data_queue.put_nowait({
                            'sensor1': sensor1_mm,
                            'sensor2': sensor2_mm,
                            'sensor3': sensor3_mm,
                            'sensor4': sensor4_mm,
                            'timestamp': time.time()
                        })
                    except:
                        # –û—á–µ—Ä–µ–¥—å –ø–µ—Ä–µ–ø–æ–ª–Ω–µ–Ω–∞ - —É–¥–∞–ª—è–µ–º —Å—Ç–∞—Ä—ã–µ –¥–∞–Ω–Ω—ã–µ –∏ –¥–æ–±–∞–≤–ª—è–µ–º –Ω–æ–≤—ã–µ
                        try:
                            self.sensor_data_queue.get_nowait()
                            self.sensor_data_queue.put_nowait({
                                'sensor1': sensor1_mm,
                                'sensor2': sensor2_mm,
                                'sensor3': sensor3_mm,
                                'sensor4': sensor4_mm,
                                'timestamp': time.time()
                            })
                        except:
                            pass  # –ò–≥–Ω–æ—Ä–∏—Ä—É–µ–º –æ—à–∏–±–∫–∏ –ø—Ä–∏ –ø–µ—Ä–µ–ø–æ–ª–Ω–µ–Ω–∏–∏ –æ—á–µ—Ä–µ–¥–∏
                            
                except Exception as e:
                    # –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ - –Ω–µ–±–æ–ª—å—à–∞—è –ø–∞—É–∑–∞ –∏ –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º
                    if self.sensor_reading_active:
                        time.sleep(0.001)  # –ú–∏–Ω–∏–º–∞–ª—å–Ω–∞—è –ø–∞—É–∑–∞ –ø—Ä–∏ –æ—à–∏–±–∫–µ
                    continue
                    
        except Exception as e:
            print(f" [SENSOR THREAD] –ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ –≤ –ø–æ—Ç–æ–∫–µ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤: {e}")
            import traceback
            traceback.print_exc()
        finally:
            print(" [SENSOR THREAD] –ü–æ—Ç–æ–∫ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ –∑–∞–≤–µ—Ä—à–µ–Ω")
    
    def get_sensor_data(self, timeout=0.001):
        """
        –ü–æ–ª—É—á–µ–Ω–∏–µ –¥–∞–Ω–Ω—ã—Ö —Å –¥–∞—Ç—á–∏–∫–æ–≤ –∏–∑ –æ—á–µ—Ä–µ–¥–∏ (–Ω–µ–±–ª–æ–∫–∏—Ä—É—é—â–µ–µ)
        
        Args:
            timeout: –¢–∞–π–º–∞—É—Ç –æ–∂–∏–¥–∞–Ω–∏—è –¥–∞–Ω–Ω—ã—Ö (—Å–µ–∫—É–Ω–¥—ã)
            
        Returns:
            dict —Å –¥–∞–Ω–Ω—ã–º–∏ –¥–∞—Ç—á–∏–∫–æ–≤ –∏–ª–∏ None –µ—Å–ª–∏ –¥–∞–Ω–Ω—ã—Ö –Ω–µ—Ç
        """
        try:
            return self.sensor_data_queue.get(timeout=timeout)
        except Empty:
            return None
    
    def read_sensors_safe(self):
        """
        –ë–µ–∑–æ–ø–∞—Å–Ω–æ–µ —á—Ç–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–æ–≤ —Å –±–ª–æ–∫–∏—Ä–æ–≤–∫–æ–π
        –ü—Ä–µ–¥–æ—Ç–≤—Ä–∞—â–∞–µ—Ç –∫–æ–Ω—Ñ–ª–∏–∫—Ç—ã –ø—Ä–∏ –æ–¥–Ω–æ–≤—Ä–µ–º–µ–Ω–Ω–æ–º —á—Ç–µ–Ω–∏–∏ –∏–∑ —Ä–∞–∑–Ω—ã—Ö –ø–æ—Ç–æ–∫–æ–≤
        
        Returns:
            Tuple[sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm]
        """
        if not self.sensors:
            return None, None, None, None

        # –ï—Å–ª–∏ —Ä–∞–±–æ—Ç–∞–µ—Ç –ø–æ—Ç–æ–∫ —á—Ç–µ–Ω–∏—è, –ø—Ä–æ–±—É–µ–º –≤–∑—è—Ç—å —Å–≤–µ–∂–∏–µ –¥–∞–Ω–Ω—ã–µ –∏–∑ –æ—á–µ—Ä–µ–¥–∏
        if self.sensor_reading_active and self.sensor_reading_thread and self.sensor_reading_thread.is_alive():
            attempts = 0
            while attempts < 5:
                try:
                    data = self.sensor_data_queue.get(timeout=0.01)
                    return (
                        data.get('sensor1'),
                        data.get('sensor2'),
                        data.get('sensor3'),
                        data.get('sensor4'),
                    )
                except Empty:
                    attempts += 1

        # –î–∞–Ω–Ω—ã—Ö –≤ –æ—á–µ—Ä–µ–¥–∏ –Ω–µ—Ç ‚Äî –≤–æ–∑–≤—Ä–∞—â–∞–µ–º None, —á—Ç–æ–±—ã –≤—ã–∑—ã–≤–∞—é—â–∏–π –∫–æ–¥ –ø–æ–≤—Ç–æ—Ä–∏–ª –ø–æ–ø—ã—Ç–∫—É
        return None, None, None, None
    
    def main_loop(self):
        """–û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª —Å–∏—Å—Ç–µ–º—ã"""
        print(" –ó–∞–ø—É—Å–∫ –æ—Å–Ω–æ–≤–Ω–æ–≥–æ —Ü–∏–∫–ª–∞...")
        
        try:
            while self.is_running:
                # –ü–†–û–í–ï–†–ö–ê –°–ë–†–û–°–ê –û–®–ò–ë–û–ö (—Ä–µ–≥–∏—Å—Ç—Ä 40024)
                try:
                    reset_flag = self.modbus_server.slave_context.getValues(3, 23, 1)[0]  # 40024 -> index 23
                    if reset_flag == 1:
                        print(" [RESET] –ü–æ–ª—É—á–µ–Ω –∑–∞–ø—Ä–æ—Å –Ω–∞ —Å–±—Ä–æ—Å –æ—à–∏–±–æ–∫ (40024=1)")
                        # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å –≤ 0
                        self.modbus_server.slave_context.setValues(4, 8, [0])  # 30009 -> index 8
                        # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ñ–ª–∞–≥ —Å–±—Ä–æ—Å–∞ –æ–±—Ä–∞—Ç–Ω–æ –≤ 0
                        self.modbus_server.slave_context.setValues(3, 23, [0])  # 40024 -> index 23
                        print(" [RESET] –°—Ç–∞—Ç—É—Å 30009 —Å–±—Ä–æ—à–µ–Ω –≤ 0, —Ñ–ª–∞–≥ 40024 —Å–±—Ä–æ—à–µ–Ω")
                except Exception as e:
                    print(f" –û—à–∏–±–∫–∞ –ø—Ä–æ–≤–µ—Ä–∫–∏ —Ä–µ–≥–∏—Å—Ç—Ä–∞ —Å–±—Ä–æ—Å–∞ 40024: {e}")
                
                # –ê–í–¢–û–ú–ê–¢–ò–ß–ï–°–ö–û–ï –ü–ï–†–ï–ü–û–î–ö–õ–Æ–ß–ï–ù–ò–ï –î–ê–¢–ß–ò–ö–û–í (–µ—Å–ª–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã)
                if not self.test_mode:
                    self.check_and_reconnect_sensors()
                
                # –ü—Ä–æ–≤–µ—Ä—è–µ–º –∫–æ–º–∞–Ω–¥—É –æ—Ç Modbus
                current_cmd = self.get_current_command()
               
                
                if current_cmd != self.previous_cmd:
                    print(f"üì® –ü–æ–ª—É—á–µ–Ω–∞ –∫–æ–º–∞–Ω–¥–∞: {current_cmd}")
                    self.handle_command(current_cmd)
                    self.previous_cmd = current_cmd

                # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å–º–µ–Ω—É —Å–º–µ–Ω—ã –≤ –æ—Å–Ω–æ–≤–Ω–æ–º —Ü–∏–∫–ª–µ
                self.check_shift_change()
                
                # –í—ã–ø–æ–ª–Ω—è–µ–º –¥–µ–π—Å—Ç–≤–∏—è –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç —Å–æ—Å—Ç–æ—è–Ω–∏—è
                self.execute_state_actions()
                
                # –ü–∞—É–∑–∞ —Ç–æ–ª—å–∫–æ –µ—Å–ª–∏ –ù–ï –ø–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º (–∏–Ω–∞—á–µ —Ç–æ—Ä–º–æ–∑–∏—Ç –ø–æ—Ç–æ–∫!)
                if self.current_state not in [SystemState.STREAM_QUAD, 
                                             SystemState.MEASURE_HEIGHT_PROCESS, SystemState.MEASURE_WALL_PROCESS, 
                                             SystemState.MEASURE_FLANGE_PROCESS, SystemState.MEASURE_FLANGE_ONLY_PROCESS,
                                             SystemState.MEASURE_BODY_ONLY_PROCESS, SystemState.MEASURE_BODY2_PROCESS,
                                             SystemState.MEASURE_BOTTOM_PROCESS,
                                             SystemState.CALIBRATE_HEIGHT,SystemState.CALIBRATE_WALL,SystemState.CALIBRATE_FLANGE,
                                             SystemState.CALIBRATE_FLANGE_DIAMETER, SystemState.CALIBRATE_BODY_DIAMETER_SEPARATE,
                                             SystemState.CALIBRATE_BODY2_DIAMETER, SystemState.CALIBRATE_BOTTOM]:
                    time.sleep(0.1)
                elif self.current_state == SystemState.STREAM_QUAD:
                    # –ú–∏–∫—Ä–æ-–ø–∞—É–∑–∞ –≤ QUAD —Ä–µ–∂–∏–º–µ –¥–ª—è —Å–Ω–∏–∂–µ–Ω–∏—è –Ω–∞–≥—Ä—É–∑–∫–∏ –Ω–∞ CPU (5 –º—Å)
                    # –≠—Ç–æ –ø–æ–º–æ–∂–µ—Ç –∏–∑–±–µ–∂–∞—Ç—å –ø–µ—Ä–µ–ø–æ–ª–Ω–µ–Ω–∏—è –±—É—Ñ–µ—Ä–æ–≤ –Ω–∞ —Å–ª–∞–±—ã—Ö –ø—Ä–æ—Ü–µ—Å—Å–æ—Ä–∞—Ö
                    time.sleep(0.005)
                
        except KeyboardInterrupt:
            print("\n –û—Å—Ç–∞–Ω–æ–≤–∫–∞ –ø–æ –∑–∞–ø—Ä–æ—Å—É –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –≤ –æ—Å–Ω–æ–≤–Ω–æ–º —Ü–∏–∫–ª–µ: {e}")
            self.current_state = SystemState.ERROR
        finally:
            self.stop_system()
    
    def get_current_command(self) -> int:
        """–ü–æ–ª—É—á–µ–Ω–∏–µ —Ç–µ–∫—É—â–µ–π –∫–æ–º–∞–Ω–¥—ã –∏–∑ Modbus —Ä–µ–≥–∏—Å—Ç—Ä–∞ 40001"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 1, 1)  # Holding Register 40001
                if values:
                    return int(values[0])
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è –∫–æ–º–∞–Ω–¥—ã: {e}")
        return 0
    
    def handle_command(self, cmd: int):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ –∫–æ–º–∞–Ω–¥—ã –∏ –ø–µ—Ä–µ—Ö–æ–¥ –≤ —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–µ–µ —Å–æ—Å—Ç–æ—è–Ω–∏–µ"""
        # –û—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≤—Å–µ –∞–∫—Ç–∏–≤–Ω—ã–µ –ø–æ—Ç–æ–∫–∏ –ø—Ä–∏ —Å–º–µ–Ω–µ –∫–æ–º–∞–Ω–¥—ã
        self.stop_all_streams()
        
        # –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ —Ñ–ª–∞–≥–æ–º —Ü–∏–∫–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è (—Ä–µ–≥–∏—Å—Ç—Ä 30009)
        self.manage_measurement_cycle_flag(cmd)
        
        if cmd == 0:
            self.current_state = SystemState.IDLE
            # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ –≤ IDLE
            self.clear_serial_buffers()
            
        # –ö–∞–ª–∏–±—Ä–æ–≤–∫–∏
        elif cmd == 100:
            self.current_state = SystemState.CALIBRATE_WALL
        elif cmd == 101:
            self.current_state = SystemState.CALIBRATE_BOTTOM
        elif cmd == 102:
            self.current_state = SystemState.CALIBRATE_FLANGE
        elif cmd == 103:
            self.current_state = SystemState.CALIBRATE_HEIGHT
        elif cmd == 104:
            self.current_state = SystemState.DEBUG_REGISTERS
        elif cmd == 105:
            self.current_state = SystemState.CALIBRATE_FLANGE_DIAMETER
        elif cmd == 106:
            # –ö–æ–º–∞–Ω–¥–∞ 106: –ù–∞—Å—Ç—Ä–æ–π–∫–∞ –¥–∏–∞–ø–∞–∑–æ–Ω–æ–≤ –¥–ª—è –¥–∏—Å–∫—Ä–µ—Ç–Ω–æ–≥–æ —Å–∏–≥–Ω–∞–ª–∞ –¥–∞—Ç—á–∏–∫–∞ 3
            self.current_state = SystemState.CONFIGURE_SENSOR3_RANGE
        elif cmd == 107:
            self.current_state = SystemState.CALIBRATE_BODY_DIAMETER_SEPARATE
        elif cmd == 108:
            self.current_state = SystemState.CALIBRATE_BODY2_DIAMETER
            
        # –ò–∑–º–µ—Ä–µ–Ω–∏–µ –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        elif cmd == 10:
            self.current_state = SystemState.MEASURE_WALL_PROCESS
            
        # –ü–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º (QUAD - –≤—Å–µ 4 –¥–∞—Ç—á–∏–∫–∞)
        elif cmd == 200:
            self.current_state = SystemState.STREAM_QUAD
            # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ –≤ QUAD —Ä–µ–∂–∏–º
            self.clear_serial_buffers()
            
        # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - –ø–æ–¥—Å—á—ë—Ç –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        elif cmd == 11:
            self.current_state = SystemState.MEASURE_WALL_CALCULATE
            
        # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - —Ñ–ª–∞–Ω–µ—Ü
        elif cmd == 12:
            self.current_state = SystemState.MEASURE_FLANGE_PROCESS
        elif cmd == 13:
            self.current_state = SystemState.MEASURE_FLANGE_CALCULATE
        elif cmd == 20:
            self.current_state = SystemState.MEASURE_FLANGE_ONLY_PROCESS
        elif cmd == 21:
            self.current_state = SystemState.MEASURE_FLANGE_ONLY_CALCULATE
        elif cmd == 30:
            self.current_state = SystemState.MEASURE_BODY_ONLY_PROCESS
        elif cmd == 31:
            self.current_state = SystemState.MEASURE_BODY_ONLY_CALCULATE
        elif cmd == 40:
            self.current_state = SystemState.MEASURE_BODY2_PROCESS
            self.body2_quality_required = True
        elif cmd == 41:
            self.current_state = SystemState.MEASURE_BODY2_CALCULATE
            
        # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - –Ω–∏–∂–Ω—è—è —Å—Ç–µ–Ω–∫–∞
        elif cmd == 14:
            self.current_state = SystemState.MEASURE_BOTTOM_PROCESS
        elif cmd == 15:
            self.current_state = SystemState.MEASURE_BOTTOM_CALCULATE
            
        # –û—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞ –∏–∑–¥–µ–ª–∏—è
        elif cmd == 16:
            self.current_state = SystemState.QUALITY_EVALUATION
            
        else:
            print(f" –ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –∫–æ–º–∞–Ω–¥–∞: {cmd}")
            self.current_state = SystemState.ERROR
        
        print(f" –ü–µ—Ä–µ—Ö–æ–¥ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–µ: {self.current_state.value}")
    
    def manage_measurement_cycle_flag(self, new_cmd: int):
        """
        –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ —Ñ–ª–∞–≥–æ–º —Ü–∏–∫–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤ —Ä–µ–≥–∏—Å—Ç—Ä–µ 30009
        
        –ù–æ–≤–∞—è –ª–æ–≥–∏–∫–∞ —Å—Ç–∞—Ç—É—Å–æ–≤:
        0   - –≥–æ—Ç–æ–≤ –∫ —Å–ª–µ–¥—É—é—â–µ–π –∫–æ–º–∞–Ω–¥–µ
        10  - –∏–∑–º–µ—Ä–µ–Ω–∏–µ –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        11  - –ø–æ–¥—Å—á—ë—Ç –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        110 - –ø–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –≥–æ—Ç–æ–≤ –∫ CMD=12
        12  - –∏–∑–º–µ—Ä–µ–Ω–∏–µ —Ñ–ª–∞–Ω—Ü–∞
        13  - –ø–æ–¥—Å—á—ë—Ç —Ñ–ª–∞–Ω—Ü–∞
        112 - –ø–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –≥–æ—Ç–æ–≤ –∫ CMD=14
        20  - —Ä–∞–∑–¥–µ–ª—å–Ω–æ–µ –∏–∑–º–µ—Ä–µ–Ω–∏–µ —Ñ–ª–∞–Ω—Ü–∞
        21  - —Ä–∞–∑–¥–µ–ª—å–Ω—ã–π –ø–æ–¥—Å—á—ë—Ç —Ñ–ª–∞–Ω—Ü–∞
        212 - —Ä–∞–∑–¥–µ–ª—å–Ω—ã–π –ø–æ–¥—Å—á—ë—Ç —Ñ–ª–∞–Ω—Ü–∞ –∑–∞–≤–µ—Ä—à—ë–Ω
        30  - —Ä–∞–∑–¥–µ–ª—å–Ω–æ–µ –∏–∑–º–µ—Ä–µ–Ω–∏–µ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞
        31  - —Ä–∞–∑–¥–µ–ª—å–Ω—ã–π –ø–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞
        312 - —Ä–∞–∑–¥–µ–ª—å–Ω—ã–π –ø–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ –∑–∞–≤–µ—Ä—à—ë–Ω
        40  - –∏–∑–º–µ—Ä–µ–Ω–∏–µ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2
        41  - –ø–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2
        412 - –ø–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 –∑–∞–≤–µ—Ä—à—ë–Ω
        14  - –∏–∑–º–µ—Ä–µ–Ω–∏–µ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        15  - –ø–æ–¥—Å—á—ë—Ç –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        114 - –ø–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –≥–æ—Ç–æ–≤ –∫ CMD=16
        16  - –æ—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞
        116 - –æ—Ü–µ–Ω–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∞, –≥–æ—Ç–æ–≤ –∫ CMD=0
        -1  - –æ—à–∏–±–∫–∞
        """
        try:
            if not self.modbus_server or not self.modbus_server.slave_context:
                return
                
            current_state_value = self.current_state.value if hasattr(self.current_state, 'value') else str(self.current_state)
            
            # === –ö–ê–õ–ò–ë–†–û–í–ö–ê –í–´–°–û–¢–´ (CMD=103) ===
            if current_state_value == "IDLE" and new_cmd == 103:
                # 0 ‚Üí 103: –Ω–∞—á–∞–ª–æ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –≤—ã—Å–æ—Ç—ã (–ø–æ–∏—Å–∫ 3 –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1)
                self.write_cycle_flag(103)
                self.clear_measurement_buffers()
                self.height_calibration_nonzero_count = 0
                self.distance_to_plane_calculated = False
                self.recent_measurements = []
                print(" [0‚Üí103] –ù–∞—á–∞–ª–æ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –≤—ã—Å–æ—Ç—ã: –ø–æ–∏—Å–∫ 3 –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1")
            
            # === –ö–ê–õ–ò–ë–†–û–í–ö–ò –°–¢–ï–ù–ö–ò/–î–ù–ê/–§–õ–ê–ù–¶–ê (100/101/102) ===
            elif current_state_value == "IDLE" and new_cmd in [100, 101, 102]:
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å —Ä–∞–≤–Ω—ã–º –Ω–æ–º–µ—Ä—É –∫–æ–º–∞–Ω–¥—ã
                self.write_cycle_flag(new_cmd)
                self.clear_measurement_buffers()
                print(f" [0‚Üí{new_cmd}] –ù–∞—á–∞–ª–æ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏")
            
            # === –ù–ê–ß–ê–õ–û –¶–ò–ö–õ–ê –ò–ó–ú–ï–†–ï–ù–ò–ô ===
            elif current_state_value == "IDLE" and new_cmd == 10:
                # 0 ‚Üí 10: –Ω–∞—á–∞–ª–æ —Ü–∏–∫–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
                self.write_cycle_flag(10)
                self.measurement_cycle_active = True
                self.clear_measurement_buffers()
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ñ–ª–∞–≥–∏ –≤—ã–ø–æ–ª–Ω–µ–Ω–∏—è —Ä–∞—Å—á—ë—Ç–æ–≤
                self.wall_calculated = False
                self.flange_calculated = False
                self.bottom_calculated = False
                self.flange_only_calculated = False
                self.body_only_calculated = False
                self.body2_calculated = False
                self.quality_evaluated = False
                self.body2_quality_required = False
                # –û—á–∏—â–∞–µ–º –∫–µ—à –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω—ã—Ö —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–π
                self.cached_distance_1_2 = None
                self.cached_distance_to_center = None
                self.cached_distance_1_3 = None
                self.cached_distance_sensor4 = None
                self.cached_distance_sensor3_to_center = None
                self.cached_distance_sensor3_to_center_body = None
                self.cached_distance_sensor3_to_center_body2 = None
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫–∏ —á–∞—Å—Ç–æ—Ç—ã
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                print(" [0‚Üí10] –ù–∞—á–∞–ª–æ —Ü–∏–∫–ª–∞: –∏–∑–º–µ—Ä–µ–Ω–∏–µ –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏")
            
            # === –í–ï–†–•–ù–Ø–Ø –°–¢–ï–ù–ö–ê ===
            elif current_state_value == "MEASURE_WALL_PROCESS" and new_cmd == 11:
                # 10 ‚Üí 11: –∫–æ–º–∞–Ω–¥–∞ –Ω–∞ –ø–æ–¥—Å—á—ë—Ç –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
                self.write_cycle_flag(11)
                # –û—á–∏—â–∞–µ–º —Ñ–ª–∞–≥ –Ω–∞—á–∞–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–ª—è –≤–æ–∑–º–æ–∂–Ω–æ—Å—Ç–∏ –ø–æ–≤—Ç–æ—Ä–Ω–æ–≥–æ –∑–∞–ø—É—Å–∫–∞
                if hasattr(self, '_wall_measurement_started'):
                    delattr(self, '_wall_measurement_started')
                print(" [10‚Üí11] –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏...")
                
            elif current_state_value == "MEASURE_WALL_CALCULATE" and new_cmd == 12:
                # –ü–æ—Å–ª–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –ø–æ–¥—Å—á—ë—Ç–∞, HMI –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç CMD=12
                self.write_cycle_flag(12)
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ñ–ª–∞–≥ —Ä–∞—Å—á—ë—Ç–∞ —Ñ–ª–∞–Ω—Ü–∞ –¥–ª—è —Å–ª–µ–¥—É—é—â–µ–≥–æ —ç—Ç–∞–ø–∞
                self.flange_calculated = False
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫–∏ —á–∞—Å—Ç–æ—Ç—ã –¥–ª—è –Ω–æ–≤–æ–≥–æ —ç—Ç–∞–ø–∞
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                # –û—á–∏—â–∞–µ–º –∫–µ—à –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω—ã—Ö —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–π –¥–ª—è –ø–µ—Ä–µ–∑–∞–≥—Ä—É–∑–∫–∏ —Å–≤–µ–∂–∏—Ö –∑–Ω–∞—á–µ–Ω–∏–π
                self.cached_distance_to_center = None
                self.cached_distance_1_3 = None
                self.cached_distance_sensor4 = None
                # –í–ê–ñ–ù–û: –û—á–∏—â–∞–µ–º –í–°–ï –±—É—Ñ–µ—Ä—ã –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –Ω–æ–≤–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è
                print(" [11‚Üí12] –ü–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –Ω–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞, –∫–µ—à –æ—á–∏—â–µ–Ω")
                print(" [11‚Üí12] –û–ß–ò–°–¢–ö–ê –ë–£–§–ï–†–û–í –ò–ó–ú–ï–†–ï–ù–ò–Ø –§–õ–ê–ù–¶–ê –ü–ï–†–ï–î –ù–û–í–´–ú –ò–ó–ú–ï–†–ï–ù–ò–ï–ú")
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–æ–≤ –¥–ª—è —Ñ–ª–∞–Ω—Ü–∞
                self.sensor1_flange_measurements = []
                self.sensor3_measurements = []
                self.sensor4_measurements = []
                self.temp_sensor1_flange_buffer = []
                self.temp_sensor3_buffer = []
                self.temp_sensor4_buffer = []
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π –¥–ª—è —Ñ–ª–∞–Ω—Ü–∞
                self.body_diameter_buffer = []
                self.flange_diameter_buffer = []
                self.bottom_thickness_buffer = []
                print(" [11‚Üí12] –í—Å–µ –±—É—Ñ–µ—Ä—ã –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞ –æ—á–∏—â–µ–Ω—ã")
            elif current_state_value == "MEASURE_WALL_CALCULATE" and new_cmd == 20:
                self.write_cycle_flag(20)
                self.flange_only_calculated = False
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                self.cached_distance_sensor3_to_center = None
                self.sensor3_flange_only_measurements = []
                self.temp_sensor3_flange_only_buffer = []
                self.flange_diameter_buffer = []
                print(" [11‚Üí20] –ù–∞—á–∞–ª–æ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞")
            elif current_state_value == "MEASURE_WALL_CALCULATE" and new_cmd == 30:
                self.write_cycle_flag(30)
                self.body_only_calculated = False
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                self.cached_distance_sensor3_to_center_body = None
                self.sensor3_body_only_measurements = []
                self.temp_sensor3_body_only_buffer = []
                self.body_only_diameter_buffer = []
                print(" [11‚Üí30] –ù–∞—á–∞–ª–æ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞")
            elif current_state_value == "MEASURE_WALL_CALCULATE" and new_cmd == 40:
                self.write_cycle_flag(40)
                self.body2_calculated = False
                self.body2_quality_required = True
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                self.cached_distance_sensor3_to_center_body2 = None
                self.sensor3_body2_measurements = []
                self.temp_sensor3_body2_buffer = []
                self.body2_diameter_buffer = []
                print(" [11‚Üí40] –ù–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2")
            
            # === –§–õ–ê–ù–ï–¶ ===
            elif current_state_value == "MEASURE_FLANGE_PROCESS" and new_cmd == 13:
                # 12 ‚Üí 13: –∫–æ–º–∞–Ω–¥–∞ –Ω–∞ –ø–æ–¥—Å—á—ë—Ç —Ñ–ª–∞–Ω—Ü–∞
                self.write_cycle_flag(13)
                # –û—á–∏—â–∞–µ–º —Ñ–ª–∞–≥ –Ω–∞—á–∞–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–ª—è –≤–æ–∑–º–æ–∂–Ω–æ—Å—Ç–∏ –ø–æ–≤—Ç–æ—Ä–Ω–æ–≥–æ –∑–∞–ø—É—Å–∫–∞
                if hasattr(self, '_flange_measurement_started'):
                    delattr(self, '_flange_measurement_started')
                print(" [12‚Üí13] –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ñ–ª–∞–Ω—Ü–∞...")
                
            elif current_state_value == "MEASURE_FLANGE_CALCULATE" and new_cmd == 14:
                # –ü–æ—Å–ª–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –ø–æ–¥—Å—á—ë—Ç–∞, HMI –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç CMD=14
                self.write_cycle_flag(14)
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ñ–ª–∞–≥ —Ä–∞—Å—á—ë—Ç–∞ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ –¥–ª—è —Å–ª–µ–¥—É—é—â–µ–≥–æ —ç—Ç–∞–ø–∞
                self.bottom_calculated = False
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫–∏ —á–∞—Å—Ç–æ—Ç—ã –¥–ª—è –Ω–æ–≤–æ–≥–æ —ç—Ç–∞–ø–∞
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ñ–ª–∞–≥ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏ —á–∞—Å—Ç–æ—Ç—ã –¥–ª—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
                if hasattr(self, '_bottom_frequency_initialized'):
                    delattr(self, '_bottom_frequency_initialized')
                print(" [13‚Üí14] –ü–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –Ω–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏")
            elif current_state_value == "MEASURE_FLANGE_CALCULATE" and new_cmd == 30:
                self.write_cycle_flag(30)
                self.body_only_calculated = False
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                self.cached_distance_sensor3_to_center_body = None
                self.sensor3_body_only_measurements = []
                self.temp_sensor3_body_only_buffer = []
                self.body_only_diameter_buffer = []
                print(" [13‚Üí30] –ù–∞—á–∞–ª–æ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞")
            elif current_state_value == "MEASURE_FLANGE_CALCULATE" and new_cmd == 40:
                self.write_cycle_flag(40)
                self.body2_calculated = False
                self.body2_quality_required = True
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                self.cached_distance_sensor3_to_center_body2 = None
                self.sensor3_body2_measurements = []
                self.temp_sensor3_body2_buffer = []
                self.body2_diameter_buffer = []
                print(" [13‚Üí40] –ù–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2")

            # === –†–ê–ó–î–ï–õ–¨–ù–´–ô –§–õ–ê–ù–ï–¶ ===
            elif current_state_value == "MEASURE_FLANGE_ONLY_PROCESS" and new_cmd == 21:
                self.write_cycle_flag(21)
                if hasattr(self, '_flange_only_measurement_started'):
                    delattr(self, '_flange_only_measurement_started')
                print(" [20‚Üí21] –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ —Ñ–ª–∞–Ω—Ü–∞...")
            elif current_state_value == "MEASURE_FLANGE_ONLY_CALCULATE" and new_cmd == 14:
                self.write_cycle_flag(14)
                self.bottom_calculated = False
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                if hasattr(self, '_bottom_frequency_initialized'):
                    delattr(self, '_bottom_frequency_initialized')
                print(" [21‚Üí14] –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ —Ñ–ª–∞–Ω—Ü–∞ –∑–∞–≤–µ—Ä—à—ë–Ω, –Ω–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏")
            elif current_state_value == "MEASURE_FLANGE_ONLY_CALCULATE" and new_cmd == 30:
                self.write_cycle_flag(30)
                self.body_only_calculated = False
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                self.cached_distance_sensor3_to_center_body = None
                self.sensor3_body_only_measurements = []
                self.temp_sensor3_body_only_buffer = []
                self.body_only_diameter_buffer = []
                print(" [21‚Üí30] –ù–∞—á–∞–ª–æ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞")
            elif current_state_value == "MEASURE_FLANGE_ONLY_CALCULATE" and new_cmd == 40:
                self.write_cycle_flag(40)
                self.body2_calculated = False
                self.body2_quality_required = True
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                self.cached_distance_sensor3_to_center_body2 = None
                self.sensor3_body2_measurements = []
                self.temp_sensor3_body2_buffer = []
                self.body2_diameter_buffer = []
                print(" [21‚Üí40] –ù–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2")

            # === –†–ê–ó–î–ï–õ–¨–ù–´–ô –î–ò–ê–ú–ï–¢–† –ö–û–†–ü–£–°–ê ===
            elif current_state_value == "MEASURE_BODY_ONLY_PROCESS" and new_cmd == 31:
                self.write_cycle_flag(31)
                if hasattr(self, '_body_only_measurement_started'):
                    delattr(self, '_body_only_measurement_started')
                print(" [30‚Üí31] –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞...")
            elif current_state_value == "MEASURE_BODY_ONLY_CALCULATE" and new_cmd == 14:
                self.write_cycle_flag(14)
                self.bottom_calculated = False
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                if hasattr(self, '_bottom_frequency_initialized'):
                    delattr(self, '_bottom_frequency_initialized')
                print(" [31‚Üí14] –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ –∑–∞–≤–µ—Ä—à—ë–Ω, –Ω–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏")
            elif current_state_value == "MEASURE_BODY_ONLY_CALCULATE" and new_cmd == 40:
                self.write_cycle_flag(40)
                self.body2_calculated = False
                self.body2_quality_required = True
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                self.cached_distance_sensor3_to_center_body2 = None
                self.sensor3_body2_measurements = []
                self.temp_sensor3_body2_buffer = []
                self.body2_diameter_buffer = []
                print(" [31‚Üí40] –ù–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2")

            # === –î–ò–ê–ú–ï–¢–† –ö–û–†–ü–£–°–ê 2 ===
            elif current_state_value == "MEASURE_BODY2_PROCESS" and new_cmd == 41:
                self.write_cycle_flag(41)
                if hasattr(self, '_body2_measurement_started'):
                    delattr(self, '_body2_measurement_started')
                print(" [40‚Üí41] –ü–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2...")
            elif current_state_value == "MEASURE_BODY2_CALCULATE" and new_cmd == 14:
                self.write_cycle_flag(14)
                self.bottom_calculated = False
                self.frequency_counter = 0
                self.frequency_start_time = None
                self.last_frequency_display = 0
                if hasattr(self, '_bottom_frequency_initialized'):
                    delattr(self, '_bottom_frequency_initialized')
                print(" [41‚Üí14] –ü–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 –∑–∞–≤–µ—Ä—à—ë–Ω, –Ω–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏")
            
            # === –ù–ò–ñ–ù–Ø–Ø –°–¢–ï–ù–ö–ê ===
            elif current_state_value == "MEASURE_BOTTOM_PROCESS" and new_cmd == 15:
                # 14 ‚Üí 15: –∫–æ–º–∞–Ω–¥–∞ –Ω–∞ –ø–æ–¥—Å—á—ë—Ç –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
                self.write_cycle_flag(15)
                # –û—á–∏—â–∞–µ–º —Ñ–ª–∞–≥ –Ω–∞—á–∞–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–ª—è –≤–æ–∑–º–æ–∂–Ω–æ—Å—Ç–∏ –ø–æ–≤—Ç–æ—Ä–Ω–æ–≥–æ –∑–∞–ø—É—Å–∫–∞
                if hasattr(self, '_bottom_measurement_started'):
                    delattr(self, '_bottom_measurement_started')
                print(" [14‚Üí15] –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏...")
                
            elif current_state_value == "MEASURE_BOTTOM_CALCULATE" and new_cmd == 16:
                # –ü–æ—Å–ª–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –ø–æ–¥—Å—á—ë—Ç–∞, HMI –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç CMD=16
                self.write_cycle_flag(16)
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ñ–ª–∞–≥ –æ—Ü–µ–Ω–∫–∏ –∫–∞—á–µ—Å—Ç–≤–∞ –¥–ª—è —Å–ª–µ–¥—É—é—â–µ–≥–æ —ç—Ç–∞–ø–∞
                self.quality_evaluated = False
                print(" [15‚Üí16] –ü–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –Ω–∞—á–∞–ª–æ –æ—Ü–µ–Ω–∫–∏ –∫–∞—á–µ—Å—Ç–≤–∞")
            
            # === –û–¶–ï–ù–ö–ê –ö–ê–ß–ï–°–¢–í–ê ===
            elif current_state_value == "QUALITY_EVALUATION" and new_cmd == 0:
                # 16 ‚Üí 0: –∑–∞–≤–µ—Ä—à–µ–Ω–∏–µ —Ü–∏–∫–ª–∞ –∏ –≤–æ–∑–≤—Ä–∞—Ç –≤ IDLE
                self.write_cycle_flag(0)
                self.measurement_cycle_active = False
                self.clear_measurement_buffers()
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ –≤ IDLE
                self.clear_serial_buffers()
                print(" [16‚Üí0] –û—Ü–µ–Ω–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∞, —Ü–∏–∫–ª –∑–∞–≤–µ—Ä—à—ë–Ω, –≤–æ–∑–≤—Ä–∞—Ç –≤ IDLE")
            
            # === –ó–ê–í–ï–†–®–ï–ù–ò–ï –ö–ê–õ–ò–ë–†–û–í–ö–ò –í–´–°–û–¢–´ ===
            elif current_state_value == "CALIBRATE_HEIGHT" and new_cmd == 0:
                # 103 ‚Üí 0: –∑–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –≤—ã—Å–æ—Ç—ã
                self.write_cycle_flag(0)
                self.clear_measurement_buffers()
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ –≤ IDLE
                self.clear_serial_buffers()
                print(" [103‚Üí0] –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –≤—ã—Å–æ—Ç—ã –∑–∞–≤–µ—Ä—à–µ–Ω–∞, –≤–æ–∑–≤—Ä–∞—Ç –≤ IDLE")
            
            # === –ó–ê–í–ï–†–®–ï–ù–ò–ï –û–¢–õ–ê–î–ö–ò –†–ï–ì–ò–°–¢–†–û–í ===
            elif current_state_value == "DEBUG_REGISTERS" and new_cmd == 0:
                # 104 ‚Üí 0: –∑–∞–≤–µ—Ä—à–µ–Ω–∏–µ –æ—Ç–ª–∞–¥–∫–∏
                self.write_cycle_flag(0)
                self.clear_measurement_buffers()
                # –û—á–∏—â–∞–µ–º –∞—Ç—Ä–∏–±—É—Ç—ã –æ—Ç–ª–∞–¥–∫–∏
                if hasattr(self, 'debug_start_time'):
                    delattr(self, 'debug_start_time')
                if hasattr(self, 'debug_last_display'):
                    delattr(self, 'debug_last_display')
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ –≤ IDLE
                self.clear_serial_buffers()
                print(" [104‚Üí0] –û—Ç–ª–∞–¥–∫–∞ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ –∑–∞–≤–µ—Ä—à–µ–Ω–∞, –≤–æ–∑–≤—Ä–∞—Ç –≤ IDLE")
            
            # === –ó–ê–í–ï–†–®–ï–ù–ò–ï –ö–ê–õ–ò–ë–†–û–í–û–ö 100/101/102/105/107/108 –ü–û –ö–û–ú–ê–ù–î–ï 0 ===
            # –ü—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ CMD -> 0 –∑–∞–≤–µ—Ä—à–∞–µ–º –∫–∞–ª–∏–±—Ä–æ–≤–∫—É: —Ä–∞—Å—Å—á–∏—Ç—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã –∏ –∑–∞–ø–∏—Å—ã–≤–∞–µ–º –∏—Ö
            elif current_state_value == "CALIBRATE_WALL" and new_cmd == 0:
                print(" [CALIBRATE_WALL‚Üí0] –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Å—Ç–µ–Ω–∫–∏, —Ä–∞—Å—á–µ—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤...")
                self._finish_calibration_wall()
                return
            elif current_state_value == "CALIBRATE_BOTTOM" and new_cmd == 0:
                print(" [CALIBRATE_BOTTOM‚Üí0] –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–Ω–∞, —Ä–∞—Å—á–µ—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤...")
                self._finish_calibration_bottom()
                return
            elif current_state_value == "CALIBRATE_FLANGE" and new_cmd == 0:
                print(" [CALIBRATE_FLANGE‚Üí0] –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Ñ–ª–∞–Ω—Ü–∞, —Ä–∞—Å—á–µ—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤...")
                self._finish_calibration_flange()
                return
            elif current_state_value == "CALIBRATE_FLANGE_DIAMETER" and new_cmd == 0:
                print(" [CALIBRATE_FLANGE_DIAMETER‚Üí0] –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞, —Ä–∞—Å—á–µ—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤...")
                self._finish_calibration_flange_diameter()
                return
            elif current_state_value == "CALIBRATE_BODY_DIAMETER_SEPARATE" and new_cmd == 0:
                print(" [CALIBRATE_BODY_DIAMETER_SEPARATE‚Üí0] –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ (—Ä–∞–∑–¥–µ–ª—å–Ω–æ), —Ä–∞—Å—á–µ—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤...")
                self._finish_calibration_body_diameter_separate()
                return
            elif current_state_value == "CALIBRATE_BODY2_DIAMETER" and new_cmd == 0:
                print(" [CALIBRATE_BODY2_DIAMETER‚Üí0] –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2, —Ä–∞—Å—á–µ—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤...")
                self._finish_calibration_body2_diameter()
                return
            
            # === –ü–†–ï–†–´–í–ê–ù–ò–ï –¶–ò–ö–õ–ê (–û–®–ò–ë–ö–ò) ===
            elif new_cmd == 0 and self.measurement_cycle_active:
                # –õ—é–±–æ–π –ø–µ—Ä–µ—Ö–æ–¥ –≤ 0 –≤–æ –≤—Ä–µ–º—è –∞–∫—Ç–∏–≤–Ω–æ–≥–æ —Ü–∏–∫–ª–∞ (–∫—Ä–æ–º–µ —É—Å–ø–µ—à–Ω–æ–≥–æ 16‚Üí0) = –ø—Ä–µ—Ä—ã–≤–∞–Ω–∏–µ
                self.write_cycle_flag(-1)
                self.measurement_cycle_active = False
                self.clear_measurement_buffers()
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ –≤ IDLE
                self.clear_serial_buffers()
                print(f" [{current_state_value}‚Üí0] –¶–∏–∫–ª –ø—Ä–µ—Ä–≤–∞–Ω! –û—à–∏–±–∫–∞.")

            # === –ü–û–¢–û–ö–û–í–´–ô –†–ï–ñ–ò–ú (CMD=200 - QUAD –≤—Å–µ—Ö –¥–∞—Ç—á–∏–∫–æ–≤) ===
            elif current_state_value == "IDLE" and new_cmd == 200:
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å —Ä–∞–≤–Ω—ã–º –Ω–æ–º–µ—Ä—É –∫–æ–º–∞–Ω–¥—ã
                self.write_cycle_flag(200)
                self.clear_measurement_buffers()
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º QUAD —Ä–µ–∂–∏–º–∞
                self.clear_serial_buffers()
                print(f" [0‚Üí200] –ù–∞—á–∞–ª–æ QUAD –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ —Ä–µ–∂–∏–º–∞ (–≤—Å–µ 4 –¥–∞—Ç—á–∏–∫–∞)")
            elif current_state_value == "STREAM_QUAD" and new_cmd == 0:
                # –í—ã—Ö–æ–¥ –∏–∑ –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ —Ä–µ–∂–∏–º–∞ ‚Üí 0
                self.write_cycle_flag(0)
                self.clear_measurement_buffers()
                # –û—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º QUAD —Ä–µ–∂–∏–º
                if self.sensors and self.stream_active_quad:
                    try:
                        # QUAD —Ä–µ–∂–∏–º –Ω–µ –∏—Å–ø–æ–ª—å–∑—É–µ—Ç –ø–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º –¥–∞—Ç—á–∏–∫–æ–≤, –ø—Ä–æ—Å—Ç–æ –æ—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –∏–∑–º–µ—Ä–µ–Ω–∏—è
                        self.stream_active_quad = False
                        print(" QUAD –ø–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º –æ—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω")
                    except Exception as e:
                        print(f" –û—à–∏–±–∫–∞ –æ—Å—Ç–∞–Ω–æ–≤–∫–∏ QUAD —Ä–µ–∂–∏–º–∞: {e}")
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ –≤ IDLE
                self.clear_serial_buffers()
                print(f" [STREAM_QUAD‚Üí0] –í—ã—Ö–æ–¥ –∏–∑ –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ —Ä–µ–∂–∏–º–∞")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è —Ñ–ª–∞–≥–æ–º —Ü–∏–∫–ª–∞: {e}")
    
    def write_cycle_flag(self, flag_value: int):
        """–ó–∞–ø–∏—Å—å —Ñ–ª–∞–≥–∞ —Ü–∏–∫–ª–∞ –≤ —Ä–µ–≥–∏—Å—Ç—Ä 30009"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # 30009 -> –∏–Ω–¥–µ–∫—Å 8 –≤ pymodbus (—Å–º. modbus_slave_server.py init_registers)
                # –î–ª—è –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π –∏—Å–ø–æ–ª—å–∑—É–µ–º –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–π –∫–æ–¥ (two's complement)
                if flag_value < 0:
                    register_value = 65536 + flag_value  # –î–ª—è -1 –ø–æ–ª—É—á–∏—Ç—Å—è 65535
                else:
                    register_value = flag_value
                
                # –ò–Ω–¥–µ–∫—Å 8 (–∫–∞–∫ –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–æ –≤ modbus_slave_server.py)
                self.modbus_server.slave_context.setValues(4, 8, [int(register_value)])
                
                # –ü—Ä–æ–≤–µ—Ä–∫–∞ –∑–∞–ø–∏—Å–∏
                verify = self.modbus_server.slave_context.getValues(4, 8, 1)
                print(f" [–§–õ–ê–ì 30009] –ó–∞–ø–∏—Å–∞–Ω–æ: {flag_value} | –ü—Ä–æ–≤–µ—Ä–∫–∞: {verify[0] if verify else 'ERROR'}")
        except Exception as e:
            print(f" [–§–õ–ê–ì 30009] –û–®–ò–ë–ö–ê –∑–∞–ø–∏—Å–∏: {e}")
            import traceback
            traceback.print_exc()
    
    def set_error_bit(self, bit_number: int, value: bool):
        """
        –£—Å—Ç–∞–Ω–æ–≤–∫–∞ —Å–æ—Å—Ç–æ—è–Ω–∏—è –æ—à–∏–±–∫–∏ –≤ —Ä–µ–≥–∏—Å—Ç—Ä–µ 30058 –∫–∞–∫ —Ü–µ–ª–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è:
        1 = –µ—Å—Ç—å –æ—à–∏–±–∫–∞, 0 = –Ω–µ—Ç –æ—à–∏–±–∫–∏. bit_number –∏–≥–Ω–æ—Ä–∏—Ä—É–µ—Ç—Å—è.
        """
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ò–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º: 1 = –ü–û–î–ö–õ–Æ–ß–ï–ù–û (–Ω–µ—Ç –æ—à–∏–±–∫–∏), 0 = –û–¢–ö–õ–Æ–ß–ï–ù–û (–æ—à–∏–±–∫–∞)
                new_value = 0 if value else 1
                # 30058 -> –∞–¥—Ä–µ—Å 58 –≤ –±–ª–æ–∫–µ Input (—Å—Ç–∞—Ä—Ç —Å 1)
                self.modbus_server.slave_context.setValues(4, 58, [int(new_value)])
                print(f" [–°–û–°–¢–û–Ø–ù–ò–ï –î–ê–¢–ß–ò–ö–û–í] 30058 = {new_value} ({'OK' if new_value == 1 else '–ù–ï–¢'})")
        except Exception as e:
            print(f" [–û–®–ò–ë–ö–ê] –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–≥–∏—Å—Ç—Ä–∞ 30058: {e}")
    
    def get_error_bit(self, bit_number: int) -> bool:
        """
        –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç True, –µ—Å–ª–∏ —Ä–µ–≥–∏—Å—Ç—Ä 30058 —Ä–∞–≤–µ–Ω 1 (–µ—Å—Ç—å –æ—à–∏–±–∫–∞), –∏–Ω–∞—á–µ False.
        bit_number –∏–≥–Ω–æ—Ä–∏—Ä—É–µ—Ç—Å—è.
        """
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # 30058 -> –∞–¥—Ä–µ—Å 58 –≤ –±–ª–æ–∫–µ Input (—Å—Ç–∞—Ä—Ç —Å 1)
                current_value = self.modbus_server.slave_context.getValues(4, 58, 1)[0]
                # –í–æ–∑–≤—Ä–∞—â–∞–µ–º True, –µ—Å–ª–∏ –ï–°–¢–¨ –û–®–ò–ë–ö–ê (–¥–ª—è —Å–æ–≤–º–µ—Å—Ç–∏–º–æ—Å—Ç–∏ —Å –Ω–∞–∑–≤–∞–Ω–∏–µ–º —Ñ—É–Ω–∫—Ü–∏–∏)
                # –û—à–∏–±–∫–∞ —Ç–µ–ø–µ—Ä—å –∫–æ–≥–¥–∞ 30058 == 0
                return int(current_value) == 0
        except Exception as e:
            print(f" [–û–®–ò–ë–ö–ê] –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–µ–≥–∏—Å—Ç—Ä–∞ 30058: {e}")
        return False
    
    def _is_port_available(self, port_name: str) -> bool:
        """
        –ü—Ä–æ–≤–µ—Ä—è–µ—Ç –Ω–∞–ª–∏—á–∏–µ COM-–ø–æ—Ä—Ç–∞ –≤ —Å–∏—Å—Ç–µ–º–µ —á–µ—Ä–µ–∑ pyserial.tools.list_ports.
        –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç True, –µ—Å–ª–∏ –ø–æ—Ä—Ç –µ—Å—Ç—å –≤ —Å–ø–∏—Å–∫–µ –¥–æ—Å—Ç—É–ø–Ω—ã—Ö.
        """
        try:
            import serial.tools.list_ports as list_ports
            ports = [p.device.upper() for p in list_ports.comports()]
            return port_name.upper() in ports
        except Exception:
            return True  # –µ—Å–ª–∏ –Ω–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ–≤–µ—Ä–∏—Ç—å ‚Äî –Ω–µ –±–ª–æ–∫–∏—Ä—É–µ–º –ª–æ–≥–∏–∫—É
    
    def _is_sensor_connection_alive(self) -> bool:
        """
        –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç True, –µ—Å–ª–∏ –æ–±—ä–µ–∫—Ç –¥–∞—Ç—á–∏–∫–æ–≤ –µ—Å—Ç—å, –ø–æ—Ä—Ç —Å—É—â–µ—Å—Ç–≤—É–µ—Ç –≤ —Å–∏—Å—Ç–µ–º–µ,
        –∏ (–µ—Å–ª–∏ –¥–æ—Å—Ç—É–ø–Ω–æ) serial.is_open == True.
        """
        if self.sensors is None:
            return False
        if not self._is_port_available(self.port):
            return False
        ser = getattr(self.sensors, 'ser', None)
        if ser is None:
            return False
        if hasattr(ser, 'is_open'):
            try:
                return bool(ser.is_open)
            except Exception:
                return False
        return True
    
    def check_and_reconnect_sensors(self):
        """
        –ü—Ä–æ–≤–µ—Ä–∫–∞ –∏ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–µ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–æ–≤, –µ—Å–ª–∏ –æ–Ω–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã
        –í—ã–∑—ã–≤–∞–µ—Ç—Å—è –ø–µ—Ä–∏–æ–¥–∏—á–µ—Å–∫–∏ –≤ –æ—Å–Ω–æ–≤–Ω–æ–º —Ü–∏–∫–ª–µ
        """
        try:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ Modbus —Å–µ—Ä–≤–µ—Ä –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω
            if not self.modbus_server or not self.modbus_server.slave_context:
                return  # Modbus —Å–µ—Ä–≤–µ—Ä –µ—â–µ –Ω–µ –≥–æ—Ç–æ–≤
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω—É–∂–Ω–æ –ª–∏ –ø—ã—Ç–∞—Ç—å—Å—è –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–∏—Ç—å—Å—è
            if self.sensors is not None:
                # –ö–æ–º–ø–ª–µ–∫—Å–Ω–∞—è –ø—Ä–æ–≤–µ—Ä–∫–∞ —Å–æ–µ–¥–∏–Ω–µ–Ω–∏—è
                if not self._is_sensor_connection_alive():
                    print(" [–ü–û–î–ö–õ–Æ–ß–ï–ù–ò–ï] –°–æ–µ–¥–∏–Ω–µ–Ω–∏–µ —Å –¥–∞—Ç—á–∏–∫–∞–º–∏ –ø–æ—Ç–µ—Ä—è–Ω–æ. –ü–æ–º–µ—á–∞–µ–º –∫–∞–∫ –æ—Ç–∫–ª—é—á–µ–Ω–Ω—ã–µ.")
                    self.sensors = None
                    self.set_error_bit(0, True)
                    self.last_reconnect_attempt = 0  # —Ä–∞–∑—Ä–µ—à–∏—Ç—å –Ω–µ–º–µ–¥–ª–µ–Ω–Ω—É—é –ø–æ–ø—ã—Ç–∫—É
                else:
                    if self.get_error_bit(0):
                        self.set_error_bit(0, False)
                    return
            
            # –î–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã - –ø—Ä–æ–≤–µ—Ä—è–µ–º –∏–Ω—Ç–µ—Ä–≤–∞–ª –ø–µ—Ä–µ–¥ —Å–ª–µ–¥—É—é—â–µ–π –ø–æ–ø—ã—Ç–∫–æ–π
            current_time = time.time()
            time_since_last_attempt = current_time - self.last_reconnect_attempt
            
            # –ü–æ–ø—ã—Ç–∫–∏ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è –¥–µ–ª–∞–µ–º –ø–æ —Ç–∞–π–º–µ—Ä—É, –¥–∞–∂–µ –µ—Å–ª–∏ –±–∏—Ç –æ—à–∏–±–∫–∏ –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω
            should_attempt = (self.last_reconnect_attempt == 0) or (time_since_last_attempt >= self.reconnect_interval)
            
            # –û—Ç–ª–∞–¥–æ—á–Ω–æ–µ —Å–æ–æ–±—â–µ–Ω–∏–µ –¥–ª—è –¥–∏–∞–≥–Ω–æ—Å—Ç–∏–∫–∏
            if should_attempt:
                current_err = None
                try:
                    current_err = self.get_error_bit(0)
                except Exception:
                    current_err = None
                print(f" [–ü–ï–†–ï–ü–û–î–ö–õ–Æ–ß–ï–ù–ò–ï] –ü—Ä–æ–≤–µ—Ä–∫–∞: –¥–∞—Ç—á–∏–∫–∏=None, –±–∏—Ç_–æ—à–∏–±–∫–∏={current_err}, –≤—Ä–µ–º—è_—Å_–ø–æ—Å–ª–µ–¥–Ω–µ–π_–ø–æ–ø—ã—Ç–∫–∏={time_since_last_attempt:.1f}—Å")
            
            if should_attempt:
                # –ü—ã—Ç–∞–µ–º—Å—è –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–∏—Ç—å—Å—è
                self.last_reconnect_attempt = current_time
                print(f" [–ü–ï–†–ï–ü–û–î–ö–õ–Æ–ß–ï–ù–ò–ï] –ü–æ–ø—ã—Ç–∫–∞ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ –Ω–∞ –ø–æ—Ä—Ç—É {self.port}...")
                
                self.sensors = HighSpeedRiftekSensor(self.port, self.baudrate, timeout=0.002)
                
                if self.sensors.connect():
                    # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–æ—Å–ª–µ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è
                    self.clear_serial_buffers()
                    print(" [–ü–ï–†–ï–ü–û–î–ö–õ–Æ–ß–ï–ù–ò–ï] ‚úÖ –î–∞—Ç—á–∏–∫–∏ —É—Å–ø–µ—à–Ω–æ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
                    self.set_error_bit(0, False)  # –°–±—Ä–∞—Å—ã–≤–∞–µ–º –±–∏—Ç –æ—à–∏–±–∫–∏
                else:
                    print(f" [–ü–ï–†–ï–ü–û–î–ö–õ–Æ–ß–ï–ù–ò–ï] ‚ùå –û—à–∏–±–∫–∞ –ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è, –ø–æ–≤—Ç–æ—Ä–∏–º —á–µ—Ä–µ–∑ {self.reconnect_interval:.0f} —Å–µ–∫")
                    self.sensors = None  # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —É–∫–∞–∑–∞—Ç–µ–ª—å
                    self.set_error_bit(0, True)  # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –±–∏—Ç –æ—à–∏–±–∫–∏
            # else: –Ω–µ –ø—Ä–æ—à–ª–æ –¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –≤—Ä–µ–º–µ–Ω–∏ —Å –ø–æ—Å–ª–µ–¥–Ω–µ–π –ø–æ–ø—ã—Ç–∫–∏ - –Ω–∏—á–µ–≥–æ –Ω–µ –¥–µ–ª–∞–µ–º
                
        except Exception as e:
            print(f" [–ü–ï–†–ï–ü–û–î–ö–õ–Æ–ß–ï–ù–ò–ï] –û—à–∏–±–∫–∞ –ø—Ä–∏ –ø—Ä–æ–≤–µ—Ä–∫–µ/–ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏–∏ –¥–∞—Ç—á–∏–∫–æ–≤: {e}")
            import traceback
            traceback.print_exc()
            self.sensors = None
            try:
                self.set_error_bit(0, True)
            except:
                pass  # –ï—Å–ª–∏ Modbus —Å–µ—Ä–≤–µ—Ä –µ—â–µ –Ω–µ –≥–æ—Ç–æ–≤, –∏–≥–Ω–æ—Ä–∏—Ä—É–µ–º –æ—à–∏–±–∫—É
    
    def increment_product_number(self):
        """–£–≤–µ–ª–∏—á–µ–Ω–∏–µ –Ω–æ–º–µ—Ä–∞ –∏–∑–¥–µ–ª–∏—è –ø—Ä–∏ —É—Å–ø–µ—à–Ω–æ–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏–∏ —Ü–∏–∫–ª–∞"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # 40101 -> –∏–Ω–¥–µ–∫—Å 100 –≤ Holding —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ö (40001=–∏–Ω–¥–µ–∫—Å 0, 40101=–∏–Ω–¥–µ–∫—Å 100)
                current_number = self.modbus_server.slave_context.getValues(3, 100, 1)[0]
                new_number = current_number + 1
                self.modbus_server.slave_context.setValues(3, 100, [new_number])
                print(f" –ù–æ–º–µ—Ä –∏–∑–¥–µ–ª–∏—è —É–≤–µ–ª–∏—á–µ–Ω: {current_number} ‚Üí {new_number}")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —É–≤–µ–ª–∏—á–µ–Ω–∏—è –Ω–æ–º–µ—Ä–∞ –∏–∑–¥–µ–ª–∏—è: {e}")
    
    def evaluate_product_quality(self) -> dict:
        """
        –û—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞ –∏–∑–¥–µ–ª–∏—è –ø–æ—Å–ª–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è —Ü–∏–∫–ª–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è
        
        Returns:
            –°–ª–æ–≤–∞—Ä—å —Å –∫–ª—é—á–∞–º–∏:
            - 'result': "GOOD", "CONDITIONALLY_GOOD" –∏–ª–∏ "BAD"
            - 'param_details': –°–ø–∏—Å–æ–∫ —Å–ª–æ–≤–∞—Ä–µ–π —Å –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–µ–π –æ –∫–∞–∂–¥–æ–º –ø–∞—Ä–∞–º–µ—Ç—Ä–µ
                (name, status, measured, base, check_type)
        """
        try:
            if not self.modbus_server or not self.modbus_server.slave_context:
                return "BAD"
            
            # –ß–∏—Ç–∞–µ–º –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ –ø—Ä–æ–≤–µ—Ä–∫–∏ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40049-40051
            check_mode = self.modbus_server.slave_context.getValues(3, 48, 1)[0]  # 40049
            allowed_conditionally_bad = self.modbus_server.slave_context.getValues(3, 49, 1)[0]  # 40050
            allowed_bad = self.modbus_server.slave_context.getValues(3, 50, 1)[0]  # 40051
            
            # –ß–∏—Ç–∞–µ–º –Ω–æ–º–µ—Ä —Å–º–µ–Ω—ã –∏ –∏–∑–¥–µ–ª–∏—è
            shift_number = self.modbus_server.slave_context.getValues(3, 99, 1)[0]  # 40100
            product_number = self.modbus_server.slave_context.getValues(3, 100, 1)[0]  # 40101
            
            # –°—á—ë—Ç—á–∏–∫–∏ –æ—à–∏–±–æ–∫
            conditionally_bad_count = 0
            bad_count = 0
            
            # –°—Ç—Ä—É–∫—Ç—É—Ä–∞ –¥–ª—è —Ö—Ä–∞–Ω–µ–Ω–∏—è —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –ø—Ä–æ–≤–µ—Ä–∫–∏
            measurement_data = {
                'shift_number': shift_number,
                'product_number': product_number,
                'check_mode': check_mode,
                'allowed_conditionally_bad': allowed_conditionally_bad,
                'allowed_bad': allowed_bad
            }
            
            # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –∫–∞–∫–∏–µ –∑–Ω–∞—á–µ–Ω–∏—è –ø—Ä–æ–≤–µ—Ä—è—Ç—å
            check_indices = []
            if check_mode == 0:  # –≤—Å–µ –∑–Ω–∞—á–µ–Ω–∏—è (–º–∏–Ω, —Å—Ä–µ–¥, –º–∞–∫—Å)
                check_indices = [0, 1, 2]  # –º–∞–∫—Å, —Å—Ä–µ–¥, –º–∏–Ω
            elif check_mode == 1:  # —Ç–æ–ª—å–∫–æ —Å—Ä–µ–¥–Ω–µ–µ
                check_indices = [1]
            elif check_mode == 2:  # –º–∞–∫—Å + —Å—Ä–µ–¥
                check_indices = [0, 1]
            elif check_mode == 3:  # –º–∏–Ω + —Å—Ä–µ–¥
                check_indices = [1, 2]
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º 7 –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ (–¥–æ–±–∞–≤–ª–µ–Ω—ã: –≤—ã—Å–æ—Ç–∞, –Ω–∏–∂–Ω—è—è —Å—Ç–µ–Ω–∫–∞, –¥–Ω–æ)
            parameters = [
                {
                    'name': 'height',
                    'measured_regs': [(40057, 40058)],  # –∏–∑–º–µ—Ä–µ–Ω–Ω–∞—è –≤—ã—Å–æ—Ç–∞ —Å –ü–õ–ö (–æ–¥–Ω–æ –∑–Ω–∞—á–µ–Ω–∏–µ)
                    'base_regs': (40376, 40377),  # –±–∞–∑–æ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ
                    'cond_bad_regs': (40378, 40379),  # —É—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'bad_regs': (40380, 40381),  # –Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'check_type': 'one_sided',  # –æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω—è—è –ø—Ä–æ–≤–µ—Ä–∫–∞
                    'single_value': True  # –∑–Ω–∞—á–µ–Ω–∏–µ –ø—Ä–∏—Ö–æ–¥–∏—Ç –æ—Ç –ü–õ–ö –≤ 40057-40058
                },
                {
                    'name': 'upper_wall',
                    'measured_regs': [(30016, 30017), (30018, 30019), (30020, 30021)],  # –º–∞–∫—Å, —Å—Ä–µ–¥, –º–∏–Ω
                    'base_regs': (40352, 40353),  # –±–∞–∑–æ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ
                    'cond_bad_regs': (40354, 40355),  # —É—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'bad_regs': (40356, 40357),  # –Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'check_type': 'one_sided'  # –æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω—è—è –ø—Ä–æ–≤–µ—Ä–∫–∞
                },
                {
                    'name': 'flange_thickness',
                    'measured_regs': [(40059, 40060)],  # –ò–∑–º–µ—Ä–µ–Ω–Ω–∞—è —Ç–æ–ª—â–∏–Ω–∞ —Ñ–ª–∞–Ω—Ü–∞ —Å –ü–ö (–æ–¥–∏–Ω —Ä–µ–≥–∏—Å—Ç—Ä)
                    'base_regs': (40370, 40371),
                    'cond_bad_regs': (40372, 40373),
                    'bad_regs': (40374, 40375),
                    'check_type': 'one_sided',  # –æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω—è—è –ø—Ä–æ–≤–µ—Ä–∫–∞
                    'single_value': True  # –§–ª–∞–≥: –æ–¥–Ω–æ –∑–Ω–∞—á–µ–Ω–∏–µ –≤–º–µ—Å—Ç–æ —Ç—Ä–µ—Ö (–º–∞–∫—Å/—Å—Ä–µ–¥/–º–∏–Ω)
                },
                {
                    'name': 'body_diameter',
                    'measured_regs': [(30046, 30047), (30048, 30049), (30050, 30051)],
                    'base_regs': (40382, 40383),
                    'cond_bad_regs': (40384, 40385),
                    'bad_regs': (40386, 40387),
                    'check_type': 'one_sided'  # –æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω—è—è –ø—Ä–æ–≤–µ—Ä–∫–∞
                },
                {
                    'name': 'flange_diameter',
                    'measured_regs': [(30054, 30055), (30052, 30053), (30056, 30057)],  # –º–∞–∫—Å, —Å—Ä–µ–¥, –º–∏–Ω
                    'base_regs': (40388, 40389),
                    'cond_bad_regs': (40390, 40391),
                    'bad_regs': (40392, 40393),
                    'check_type': 'one_sided'  # –æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω—è—è –ø—Ä–æ–≤–µ—Ä–∫–∞
                },
                {
                    'name': 'bottom_wall',
                    'measured_regs': [(30022, 30023), (30024, 30025), (30026, 30027)],  # –º–∞–∫—Å, —Å—Ä–µ–¥, –º–∏–Ω
                    'base_regs': (40358, 40359),  # –±–∞–∑–æ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ
                    'cond_bad_regs': (40360, 40361),  # —É—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'bad_regs': (40362, 40363),  # –Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'positive_bad_regs': (40402, 40403),  # –ø–æ–ª–æ–∂–∏—Ç–µ–ª—å–Ω–∞—è –Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'check_type': 'two_sided'  # –¥–≤—É—Å—Ç–æ—Ä–æ–Ω–Ω—è—è –ø—Ä–æ–≤–µ—Ä–∫–∞
                },
                {
                    'name': 'bottom',
                    'measured_regs': [(30028, 30029), (30030, 30031), (30032, 30033)],  # –º–∞–∫—Å, —Å—Ä–µ–¥, –º–∏–Ω
                    'base_regs': (40364, 40365),  # –±–∞–∑–æ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ
                    'cond_bad_regs': (40366, 40367),  # —É—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'bad_regs': (40368, 40369),  # –Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'positive_bad_regs': (40400, 40401),  # –ø–æ–ª–æ–∂–∏—Ç–µ–ª—å–Ω–∞—è –Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                    'check_type': 'two_sided'  # –¥–≤—É—Å—Ç–æ—Ä–æ–Ω–Ω—è—è –ø—Ä–æ–≤–µ—Ä–∫–∞
                }
            ]

            # –î–∏–∞–º–µ—Ç—Ä –∫–æ—Ä–ø—É—Å–∞ 2 –æ—Ü–µ–Ω–∏–≤–∞–µ–º —Ç–æ–ª—å–∫–æ –µ—Å–ª–∏ –≤ —Ü–∏–∫–ª–µ –±—ã–ª–∞ –∫–æ–º–∞–Ω–¥–∞ 40
            if self.body2_quality_required:
                parameters.append(
                    {
                        'name': 'body_diameter_2',
                        'measured_regs': [(30059, 30060), (30061, 30062), (30063, 30064)],
                        'base_regs': (40346, 40347),
                        'cond_bad_regs': (40348, 40349),
                        'bad_regs': (40350, 40351),
                        'check_type': 'one_sided'
                    }
                )
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –∫–∞–∂–¥—ã–π –ø–∞—Ä–∞–º–µ—Ç—Ä
            for param in parameters:
                # –ß–∏—Ç–∞–µ–º –±–∞–∑–æ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ –∏ –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç–∏
                base_value = self.read_float_from_registers(param['base_regs'], 'holding')
                cond_bad_error = self.read_float_from_registers(param['cond_bad_regs'], 'holding')  # –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω–æ–µ
                bad_error = self.read_float_from_registers(param['bad_regs'], 'holding')  # –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω–æ–µ
                
                # –î–ª—è –¥–≤—É—Å—Ç–æ—Ä–æ–Ω–Ω–µ–π –ø—Ä–æ–≤–µ—Ä–∫–∏ —á–∏—Ç–∞–µ–º –ø–æ–ª–æ–∂–∏—Ç–µ–ª—å–Ω—É—é –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å
                positive_bad_error = None
                if param['check_type'] == 'two_sided':
                    positive_bad_error = self.read_float_from_registers(param['positive_bad_regs'], 'holding')
                
                # –ß–∏—Ç–∞–µ–º –∏–∑–º–µ—Ä–µ–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è
                measured_values = []
                
                # –°–ø–µ—Ü–∏–∞–ª—å–Ω–∞—è –æ–±—Ä–∞–±–æ—Ç–∫–∞ –¥–ª—è flange_thickness (—á–∏—Ç–∞–µ–º –∏–∑ holding —Ä–µ–≥–∏—Å—Ç—Ä–∞ 40059-40060)
                if param.get('single_value', False):
                    # –î–ª—è –æ–¥–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è —á–∏—Ç–∞–µ–º –∏–∑ —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–µ–≥–æ —Ä–µ–≥–∏—Å—Ç—Ä–∞
                    if param['name'] == 'flange_thickness':
                        value = self.read_measured_flange_thickness()
                        if value is None:
                            print(f" [CMD=16] –û–®–ò–ë–ö–ê: –ù–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å –∏–∑–º–µ—Ä–µ–Ω–Ω—É—é —Ç–æ–ª—â–∏–Ω—É —Ñ–ª–∞–Ω—Ü–∞ –∏–∑ 40059-40060!")
                            value = 0.0  # –ó–Ω–∞—á–µ–Ω–∏–µ –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é
                        measured_values = [value]
                    elif param['name'] == 'height':
                        value = self.read_measured_height()
                        if value is None:
                            print(f" [CMD=16] –û–®–ò–ë–ö–ê: –ù–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å –∏–∑–º–µ—Ä–µ–Ω–Ω—É—é –≤—ã—Å–æ—Ç—É –∏–∑ 40057-40058!")
                            value = 0.0  # –ó–Ω–∞—á–µ–Ω–∏–µ –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é
                        measured_values = [value]
                    else:
                        # –î–ª—è –¥—Ä—É–≥–∏—Ö –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ —Å –æ–¥–Ω–∏–º –∑–Ω–∞—á–µ–Ω–∏–µ–º —á–∏—Ç–∞–µ–º –∫–∞–∫ —Ä–∞–Ω—å—à–µ (–∏–∑ input —Ä–µ–≥–∏—Å—Ç—Ä–∞)
                        value = self.read_float_from_registers(param['measured_regs'][0], 'input')
                        measured_values = [value]
                else:
                    # –î–ª—è –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ —Å —Ç—Ä–µ–º—è –∑–Ω–∞—á–µ–Ω–∏—è–º–∏ (–º–∞–∫—Å, —Å—Ä–µ–¥, –º–∏–Ω)
                    for reg_pair in param['measured_regs']:
                        value = self.read_float_from_registers(reg_pair, 'input')
                        measured_values.append(value)
                
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∑–Ω–∞—á–µ–Ω–∏—è –≤ measurement_data
                if param.get('single_value', False):
                    # –î–ª—è –æ–¥–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è —Å–æ—Ö—Ä–∞–Ω—è–µ–º –∫–∞–∫ —Å—Ä–µ–¥–Ω–µ–µ
                    measurement_data[f"{param['name']}_avg"] = measured_values[0]
                    measurement_data[f"{param['name']}_max"] = measured_values[0]
                    measurement_data[f"{param['name']}_min"] = measured_values[0]
                else:
                    # –î–ª—è —Ç—Ä–µ—Ö –∑–Ω–∞—á–µ–Ω–∏–π
                    measurement_data[f"{param['name']}_max"] = measured_values[0]
                    measurement_data[f"{param['name']}_avg"] = measured_values[1]
                    measurement_data[f"{param['name']}_min"] = measured_values[2]
                
                # –í—ã–≤–æ–¥–∏–º –ø—Ä–æ–≤–µ—Ä—è–µ–º—ã–π –ø–∞—Ä–∞–º–µ—Ç—Ä –∏ –¥–∏–∞–ø–∞–∑–æ–Ω—ã –¥–æ–ø—É—Å–∫–∞
                print(f"\n === {param['name'].upper()} ({param['check_type']}) ===")
                if param['check_type'] == 'two_sided':
                    print(f"   –ë–ê–ó–ê={base_value:.3f}–º–º | "
                          f"–£—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–∞—è (–Ω–∏–∂–µ)={base_value + cond_bad_error:.3f}–º–º | "
                          f"–ù–µ–≥–æ–¥–Ω–∞—è (–Ω–∏–∂–µ)={base_value + bad_error:.3f}–º–º | "
                          f"–ù–µ–≥–æ–¥–Ω–∞—è (–≤—ã—à–µ)={base_value + (positive_bad_error if positive_bad_error is not None else 0):.3f}–º–º")
                else:
                    print(f"   –ë–ê–ó–ê={base_value:.3f}–º–º | "
                          f"–£—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–∞—è –≥—Ä–∞–Ω–∏—Ü–∞={base_value + cond_bad_error:.3f}–º–º | "
                          f"–ù–µ–≥–æ–¥–Ω–∞—è –≥—Ä–∞–Ω–∏—Ü–∞={base_value + bad_error:.3f}–º–º")
                
                # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Ç–æ–ª—å–∫–æ –≤—ã–±—Ä–∞–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è
                param_errors = []
                if param.get('single_value', False):
                    # –î–ª—è –æ–¥–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è –ø—Ä–æ–≤–µ—Ä—è–µ–º —Ç–æ–ª—å–∫–æ –µ–≥–æ
                    measured = measured_values[0]
                    if param['check_type'] == 'two_sided':
                        status = self.check_single_value_with_upper_limit(
                            measured, base_value, cond_bad_error, bad_error, positive_bad_error
                        )
                    else:  # one_sided
                        status = self.check_single_value(measured, base_value, cond_bad_error, bad_error)
                    param_errors.append(status)
                    print(f" [–ò–ó–ú–ï–†–ï–ù–ò–ï] {measured:.3f} ‚Üí {status}")
                else:
                    # –î–ª—è —Ç—Ä–µ—Ö –∑–Ω–∞—á–µ–Ω–∏–π –ø—Ä–æ–≤–µ—Ä—è–µ–º –≤—ã–±—Ä–∞–Ω–Ω—ã–µ –∏–Ω–¥–µ–∫—Å—ã
                    value_names = ['–ú–ê–ö–°', '–°–†–ï–î', '–ú–ò–ù']
                    for idx in check_indices:
                        measured = measured_values[idx]
                        
                        # –í—ã–±–∏—Ä–∞–µ–º –º–µ—Ç–æ–¥ –ø—Ä–æ–≤–µ—Ä–∫–∏ –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç —Ç–∏–ø–∞
                        if param['check_type'] == 'two_sided':
                            status = self.check_single_value_with_upper_limit(
                                measured, base_value, cond_bad_error, bad_error, positive_bad_error
                            )
                        else:  # one_sided
                            status = self.check_single_value(measured, base_value, cond_bad_error, bad_error)
                        
                        param_errors.append(status)
                        print(f" [{value_names[idx]}] {measured:.3f} ‚Üí {status}")
                
                # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Å—Ç–∞—Ç—É—Å –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ (—Ö—É–¥—à–∏–π –∏–∑ –ø—Ä–æ–≤–µ—Ä–µ–Ω–Ω—ã—Ö)
                if "BAD" in param_errors:
                    param_status = "BAD"
                    bad_count += 1
                elif "CONDITIONALLY_GOOD" in param_errors:
                    param_status = "CONDITIONALLY_GOOD"
                    conditionally_bad_count += 1
                else:
                    param_status = "GOOD"
                
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ –ø–∞—Ä–∞–º–µ—Ç—Ä–µ –¥–ª—è —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –≤—Å–µ –ø—Ä–æ–≤–µ—Ä–µ–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è —Å –∏—Ö —Å—Ç–∞—Ç—É—Å–∞–º–∏ –¥–ª—è –ø—Ä–∞–≤–∏–ª—å–Ω–æ–≥–æ –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∏—è –Ω–∞–ø—Ä–∞–≤–ª–µ–Ω–∏—è
                checked_values = []
                if param.get('single_value', False):
                    # –î–ª—è –æ–¥–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è
                    checked_values.append({
                        'value': measured_values[0],
                        'status': param_errors[0] if param_errors else 'GOOD',
                        'index': 0
                    })
                else:
                    # –î–ª—è –Ω–µ—Å–∫–æ–ª—å–∫–∏—Ö –∑–Ω–∞—á–µ–Ω–∏–π —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –ø—Ä–æ–≤–µ—Ä–µ–Ω–Ω—ã–µ (–ø–æ check_indices)
                    value_names = ['–ú–ê–ö–°', '–°–†–ï–î', '–ú–ò–ù']
                    for i, idx in enumerate(check_indices):
                        checked_values.append({
                            'value': measured_values[idx],
                            'status': param_errors[i] if i < len(param_errors) else 'GOOD',
                            'index': idx,
                            'name': value_names[idx]
                        })
                
                param_info = {
                    'name': param['name'],
                    'status': param_status,
                    'base': base_value,
                    'check_type': param['check_type'],
                    'cond_bad_error': cond_bad_error,
                    'bad_error': bad_error,
                    'positive_bad_error': positive_bad_error if param['check_type'] == 'two_sided' else None,
                    'checked_values': checked_values
                }
                if 'param_details' not in measurement_data:
                    measurement_data['param_details'] = []
                measurement_data['param_details'].append(param_info)
                
                print(f" –ò–¢–û–ì: {param_status}")
                measurement_data[f"{param['name']}_status"] = param_status
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Å—á—ë—Ç—á–∏–∫–∏ –æ—à–∏–±–æ–∫
            measurement_data['conditionally_bad_count'] = conditionally_bad_count
            measurement_data['bad_count'] = bad_count
            
            # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –∏—Ç–æ–≥–æ–≤—ã–π —Ä–µ–∑—É–ª—å—Ç–∞—Ç
            if bad_count > allowed_bad:
                result = "BAD"
            elif conditionally_bad_count > allowed_conditionally_bad:
                result = "CONDITIONALLY_GOOD"
            else:
                result = "GOOD"
            
            print(f"\n –ò–¢–û–ì: {result} | –°–º–µ–Ω–∞: {shift_number} | –ò–∑–¥–µ–ª–∏–µ: {product_number}")
            
            measurement_data['result'] = result
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Å—Ä–µ–¥–Ω–∏–µ –∑–Ω–∞—á–µ–Ω–∏—è –≤ –ë–î –¥–ª—è –æ—Ç—á—ë—Ç–∞
            if self.db_integration:
                try:
                    self.db_integration.db.save_quality_measurement(shift_number, {
                        'height_avg': measurement_data.get('height_avg'),
                        'upper_wall_avg': measurement_data.get('upper_wall_avg'),
                        'body_diameter_avg': measurement_data.get('body_diameter_avg'),
                        'flange_diameter_avg': measurement_data.get('flange_diameter_avg'),
                        'bottom_wall_avg': measurement_data.get('bottom_wall_avg'),
                        'flange_thickness_avg': measurement_data.get('flange_thickness_avg'),
                        'bottom_avg': measurement_data.get('bottom_avg')
                    })
                except Exception as e:
                    print(f" [CMD=16] –û—à–∏–±–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –∏–∑–º–µ—Ä–µ–Ω–∏–π –≤ –ë–î: {e}")
            
            # –í–æ–∑–≤—Ä–∞—â–∞–µ–º —Å–ª–æ–≤–∞—Ä—å —Å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–º –∏ –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–µ–π –æ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞—Ö
            return {
                'result': result,
                'param_details': measurement_data.get('param_details', [])
            }
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –æ—Ü–µ–Ω–∫–∏ –∫–∞—á–µ—Å—Ç–≤–∞: {e}")
            return {'result': 'BAD', 'param_details': []}
    
    def check_single_value(self, measured: float, base: float, cond_bad_error: float, bad_error: float) -> str:
        """
        –ü—Ä–æ–≤–µ—Ä–∫–∞ –æ–¥–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è (–æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω—è—è - —Ç–æ–ª—å–∫–æ –Ω–∏–∂–Ω–∏–µ –≥—Ä–∞–Ω–∏—Ü—ã)
        
        Args:
            measured: –ò–∑–º–µ—Ä–µ–Ω–Ω–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ
            base: –ë–∞–∑–æ–≤–æ–µ (—ç—Ç–∞–ª–æ–Ω–Ω–æ–µ) –∑–Ω–∞—á–µ–Ω–∏–µ
            cond_bad_error: –£—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å (–æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω–∞—è)
            bad_error: –ù–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å (–æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω–∞—è, –º–µ–Ω—å—à–µ cond_bad_error)
        
        Returns:
            "GOOD", "CONDITIONALLY_GOOD" –∏–ª–∏ "BAD"
        """
        # –î–∏–∞–ø–∞–∑–æ–Ω—ã:
        # –ì–û–î–ù–ê–Ø: [base + cond_bad_error, base]
        # –£–°–õ–û–í–ù–û-–ì–û–î–ù–ê–Ø: [base + bad_error, base + cond_bad_error)
        # –ù–ï–ì–û–î–ù–ê–Ø: < base + bad_error –∏–ª–∏ > base
        
        if measured > base:
            return "BAD"  # –ë–æ–ª—å—à–µ –±–∞–∑–æ–≤–æ–≥–æ - –Ω–µ–≥–æ–¥–Ω–∞—è
        elif measured >= base + cond_bad_error:
            return "GOOD"  # –í –¥–∏–∞–ø–∞–∑–æ–Ω–µ –≥–æ–¥–Ω—ã—Ö
        elif measured >= base + bad_error:
            return "CONDITIONALLY_GOOD"  # –í –¥–∏–∞–ø–∞–∑–æ–Ω–µ —É—Å–ª–æ–≤–Ω–æ-–≥–æ–¥–Ω—ã—Ö
        else:
            return "BAD"  # –ú–µ–Ω—å—à–µ –º–∏–Ω–∏–º—É–º–∞ - –Ω–µ–≥–æ–¥–Ω–∞—è
    
    def check_single_value_with_upper_limit(self, measured: float, base: float, 
                                           cond_bad_error: float, bad_error: float, 
                                           positive_bad_error: float) -> str:
        """
        –ü—Ä–æ–≤–µ—Ä–∫–∞ –æ–¥–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è (–¥–≤—É—Å—Ç–æ—Ä–æ–Ω–Ω—è—è - –Ω–∏–∂–Ω–∏–µ –∏ –≤–µ—Ä—Ö–Ω—è—è –≥—Ä–∞–Ω–∏—Ü—ã)
        –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –¥–ª—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ –∏ –¥–Ω–∞
        
        Args:
            measured: –ò–∑–º–µ—Ä–µ–Ω–Ω–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ
            base: –ë–∞–∑–æ–≤–æ–µ (—ç—Ç–∞–ª–æ–Ω–Ω–æ–µ) –∑–Ω–∞—á–µ–Ω–∏–µ
            cond_bad_error: –£—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å (–æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω–∞—è)
            bad_error: –ù–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å (–æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω–∞—è, –º–µ–Ω—å—à–µ cond_bad_error)
            positive_bad_error: –ü–æ–ª–æ–∂–∏—Ç–µ–ª—å–Ω–∞—è –Ω–µ–≥–æ–¥–Ω–∞—è –ø–æ–≥—Ä–µ—à–Ω–æ—Å—Ç—å (–≤–µ—Ä—Ö–Ω—è—è –≥—Ä–∞–Ω–∏—Ü–∞)
        
        Returns:
            "GOOD", "CONDITIONALLY_GOOD" –∏–ª–∏ "BAD"
        """
        # –î–∏–∞–ø–∞–∑–æ–Ω—ã:
        # –ù–ï–ì–û–î–ù–ê–Ø: measured < (base + bad_error)
        # –£–°–õ–û–í–ù–û-–ì–û–î–ù–ê–Ø: (base + bad_error) <= measured < (base + cond_bad_error)
        # –ì–û–î–ù–ê–Ø: (base + cond_bad_error) <= measured <= (base + positive_bad_error)
        # –ù–ï–ì–û–î–ù–ê–Ø: measured > (base + positive_bad_error)
        
        if measured > base + positive_bad_error:
            return "BAD"  # –ë–æ–ª—å—à–µ –≤–µ—Ä—Ö–Ω–µ–π –≥—Ä–∞–Ω–∏—Ü—ã - –Ω–µ–≥–æ–¥–Ω–∞—è
        elif measured >= base + cond_bad_error:
            return "GOOD"  # –í –¥–∏–∞–ø–∞–∑–æ–Ω–µ –≥–æ–¥–Ω—ã—Ö
        elif measured >= base + bad_error:
            return "CONDITIONALLY_GOOD"  # –í –¥–∏–∞–ø–∞–∑–æ–Ω–µ —É—Å–ª–æ–≤–Ω–æ-–≥–æ–¥–Ω—ã—Ö
        else:
            return "BAD"  # –ú–µ–Ω—å—à–µ –Ω–∏–∂–Ω–µ–π –≥—Ä–∞–Ω–∏—Ü—ã - –Ω–µ–≥–æ–¥–Ω–∞—è
    
    def read_float_from_registers(self, reg_pair: Tuple[int, int], reg_type: str = 'holding') -> float:
        """
        –ß—Ç–µ–Ω–∏–µ float –∑–Ω–∞—á–µ–Ω–∏—è –∏–∑ –ø–∞—Ä—ã —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤
        
        Args:
            reg_pair: –ö–æ—Ä—Ç–µ–∂ (low_register, high_register)
            reg_type: 'holding' –∏–ª–∏ 'input'
        
        Returns:
            float –∑–Ω–∞—á–µ–Ω–∏–µ
        """
        try:
            # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Ç–∏–ø —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤
            if reg_type == 'holding':
                function_code = 3
                base_offset = 40000
            else:  # input
                function_code = 4
                base_offset = 30000
            
            # –í—ã—á–∏—Å–ª—è–µ–º –∏–Ω–¥–µ–∫—Å—ã (–í–ê–ñ–ù–û: 40001 = –∏–Ω–¥–µ–∫—Å 1, 30001 = –∏–Ω–¥–µ–∫—Å 1 –≤ pymodbus)
            # –ü—Ä–∏ –∑–∞–ø–∏—Å–∏: base_address - 1 = –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ, base_address = —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ
            # –ü—Ä–∏ —á—Ç–µ–Ω–∏–∏: reg_pair[0] = base_address = —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ, reg_pair[0] - 1 = –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ
            first_idx = reg_pair[0] - base_offset      # –°—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ (base_address)
            second_idx = reg_pair[0] - base_offset + 1  # –ú–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ (base_address)
            
            # –ß–∏—Ç–∞–µ–º –∑–Ω–∞—á–µ–Ω–∏—è (–í HMI: base_address = —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ, base_address - 1 = –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ)
            high_word = self.modbus_server.slave_context.getValues(function_code, first_idx, 1)[0]
            low_word = self.modbus_server.slave_context.getValues(function_code, second_idx, 1)[0]
            
            # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º –≤ float
            combined = (high_word << 16) | low_word
            float_value = struct.unpack('!f', struct.pack('!I', combined))[0]
            
            return float_value
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è float –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ {reg_pair}: {e}")
            return 0.0
    
    def read_input_register(self, register_number: int) -> int:
        """
        –ß—Ç–µ–Ω–∏–µ –∑–Ω–∞—á–µ–Ω–∏—è Input —Ä–µ–≥–∏—Å—Ç—Ä–∞ –ø–æ –µ–≥–æ –∞–¥—Ä–µ—Å—É (30001+)
        """
        try:
            if not self.modbus_server or not self.modbus_server.slave_context:
                return 0
            
            index = register_number - 30001
            if index < 0:
                return 0
            
            value = self.modbus_server.slave_context.getValues(4, index, 1)[0]
            return int(value)
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è Input —Ä–µ–≥–∏—Å—Ç—Ä–∞ {register_number}: {e}")
            return 0
    
    def generate_shift_report(self, shift_number: int):
        """
        –§–æ—Ä–º–∏—Ä—É–µ—Ç DOCX –æ—Ç—á—ë—Ç –ø–æ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞–º —Å–º–µ–Ω—ã
        """
        if not DOCX_AVAILABLE:
            print(" [REPORT] –ù–µ–≤–æ–∑–º–æ–∂–Ω–æ —Å—Ñ–æ—Ä–º–∏—Ä–æ–≤–∞—Ç—å –æ—Ç—á—ë—Ç: –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–∞ –±–∏–±–ª–∏–æ—Ç–µ–∫–∞ python-docx")
            return
        
        if not self.modbus_server or not self.modbus_server.slave_context:
            print(" [REPORT] Modbus —Å–µ—Ä–≤–µ—Ä –Ω–µ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω, –æ—Ç—á—ë—Ç –Ω–µ —Å–æ–∑–¥–∞–Ω")
            return
        
        try:
            timestamp = datetime.now()
            formatted_time = timestamp.strftime("%d.%m.%Y %H:%M:%S")
            filename = f"shift_{shift_number}_{timestamp.strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(self.data_dir, filename)
            
            # –°—á—ë—Ç—á–∏–∫–∏ –∏–∑–¥–µ–ª–∏–π
            total_measured = self.read_input_register(30101)
            total_good = self.read_input_register(30102)
            total_conditionally_good = self.read_input_register(30103)
            total_bad = self.read_input_register(30104)
            
            # –ö–∞—Ä—Ç–∞ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ –¥–ª—è —Ç–∞–±–ª–∏—Ü
            parameter_config = [
                {
                    'label': '–í—ã—Å–æ—Ç–∞',
                    'cond_less': 30201,
                    'cond_greater': None,
                    'bad_less': 30210,
                    'bad_greater': 30219
                },
                {
                    'label': '–¢–æ–ª—â–∏–Ω–∞ —Å—Ç–µ–Ω–∫–∏ –≤–≤–µ—Ä—Ö',
                    'cond_less': 30202,
                    'cond_greater': None,
                    'bad_less': 30211,
                    'bad_greater': 30220
                },
                {
                    'label': '–¢–æ–ª—â–∏–Ω–∞ —Å—Ç–µ–Ω–∫–∏ –≤–Ω–∏–∑',
                    'cond_less': 30205,
                    'cond_greater': 30204,
                    'bad_less': 30213,
                    'bad_greater': 30214
                },
                {
                    'label': '–î–∏–∞–º–µ—Ç—Ä –∫–æ—Ä–ø—É—Å–∞',
                    'cond_less': 30209,
                    'cond_greater': None,
                    'bad_less': 30218,
                    'bad_greater': 30223
                },
                {
                    'label': '–¢–æ–ª—â–∏–Ω–∞ –¥–Ω–∞',
                    'cond_less': 30206,
                    'cond_greater': 30207,
                    'bad_less': 30215,
                    'bad_greater': 30216
                },
                {
                    'label': '–¢–æ–ª—â–∏–Ω–∞ —Ñ–ª–∞–Ω—Ü–∞',
                    'cond_less': 30203,
                    'cond_greater': None,
                    'bad_less': 30212,
                    'bad_greater': 30221
                },
                {
                    'label': '–î–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞',
                    'cond_less': 30208,
                    'cond_greater': None,
                    'bad_less': 30217,
                    'bad_greater': 30222
                },
            ]
            
            def get_reg_value(reg_number: Optional[int]) -> int:
                if reg_number is None:
                    return 0
                return self.read_input_register(reg_number)
            
            conditional_rows = []
            bad_rows = []
            for param in parameter_config:
                conditional_rows.append({
                    'label': param['label'],
                    'less': get_reg_value(param['cond_less']),
                    'greater': get_reg_value(param['cond_greater'])
                })
                bad_rows.append({
                    'label': param['label'],
                    'less': get_reg_value(param['bad_less']),
                    'greater': get_reg_value(param['bad_greater'])
                })
            
            # –°–æ–∑–¥–∞—ë–º –¥–æ–∫—É–º–µ–Ω—Ç
            document = Document()
            style = document.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            title = document.add_paragraph("–ü—Ä–æ—Ç–æ–∫–æ–ª –∏–∑–º–µ—Ä–µ–Ω–∏–π")
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(16)
            
            document.add_paragraph("")  # –û—Ç—Å—Ç—É–ø
            
            header_table = document.add_table(rows=2, cols=4)
            header_table.style = 'Table Grid'
            header_cells = header_table.rows[0].cells
            header_cells[0].text = "–¢–∏–ø"
            header_cells[1].text = ""
            header_cells[2].text = "–ò–∑–¥–µ–ª–∏–µ"
            header_cells[3].text = ""
            
            row2 = header_table.rows[1].cells
            row2[0].text = "–°–º–µ–Ω–∞"
            row2[1].text = str(shift_number)
            row2[2].text = "–î–∞—Ç–∞"
            row2[3].text = formatted_time
            
            document.add_paragraph("")
            
            summary_table = document.add_table(rows=2, cols=4)
            summary_table.style = 'Table Grid'
            summary_table.rows[0].cells[0].text = "–ò–∑–º–µ—Ä–µ–Ω–æ, —à—Ç."
            summary_table.rows[0].cells[1].text = "–ì–æ–¥–Ω—ã–µ, —à—Ç."
            summary_table.rows[0].cells[2].text = "–£—Å–ª–æ–≤–Ω–æ –≥–æ–¥–Ω—ã–µ, —à—Ç."
            summary_table.rows[0].cells[3].text = "–ë—Ä–∞–∫, —à—Ç."
            
            summary_table.rows[1].cells[0].text = str(total_measured)
            summary_table.rows[1].cells[1].text = str(total_good)
            summary_table.rows[1].cells[2].text = str(total_conditionally_good)
            summary_table.rows[1].cells[3].text = str(total_bad)
            
            document.add_paragraph("")
            
            heading = document.add_paragraph("–°–≤–æ–¥–Ω—ã–µ —Ç–∞–±–ª–∏—Ü—ã —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏–π")
            heading.runs[0].bold = True
            
            def add_parameter_table(title_text: str, rows: List[dict]):
                document.add_paragraph("")
                table = document.add_table(rows=len(rows) + 2, cols=3)
                table.style = 'Table Grid'
                
                top_row = table.rows[0].cells
                top_row[0].text = "–ü–∞—Ä–∞–º–µ—Ç—Ä"
                top_row[1].text = title_text
                top_row[1].merge(top_row[2])
                
                second_row = table.rows[1].cells
                second_row[0].text = ""
                second_row[1].text = "< –Ω–æ—Ä–º—ã"
                second_row[2].text = "> –Ω–æ—Ä–º—ã"
                
                for idx, row in enumerate(rows, start=2):
                    table.rows[idx].cells[0].text = row['label']
                    table.rows[idx].cells[1].text = str(row['less'])
                    table.rows[idx].cells[2].text = str(row['greater'])
            
            add_parameter_table("–£—Å–ª–æ–≤–Ω–æ –≥–æ–¥–Ω—ã–µ, —à—Ç.", conditional_rows)
            add_parameter_table("–ó–∞–±—Ä–∞–∫–æ–≤–∞–Ω–æ, —à—Ç.", bad_rows)
            
            # –î–æ–±–∞–≤–ª—è–µ–º –≤—Ç–æ—Ä–æ–π –ª–∏—Å—Ç —Å —Ç–∞–±–ª–∏—Ü–µ–π –∏–∑–º–µ—Ä–µ–Ω–∏–π
            document.add_page_break()
            
            # –ó–∞–≥–æ–ª–æ–≤–æ–∫ –≤—Ç–æ—Ä–æ–≥–æ –ª–∏—Å—Ç–∞
            title2 = document.add_paragraph("–ü—Ä–æ—Ç–æ–∫–æ–ª –∏–∑–º–µ—Ä–µ–Ω–∏–π")
            title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title2.runs[0].bold = True
            title2.runs[0].font.size = Pt(16)
            
            document.add_paragraph("")  # –û—Ç—Å—Ç—É–ø
            
            # –®–∞–ø–∫–∞ –≤—Ç–æ—Ä–æ–≥–æ –ª–∏—Å—Ç–∞ (–°–º–µ–Ω–∞ –∏ –î–∞—Ç–∞)
            header_table2 = document.add_table(rows=2, cols=4)
            header_table2.style = 'Table Grid'
            header_cells2 = header_table2.rows[0].cells
            header_cells2[0].text = "–¢–∏–ø"
            header_cells2[1].text = ""
            header_cells2[2].text = "–ò–∑–¥–µ–ª–∏–µ"
            header_cells2[3].text = ""
            
            row2_2 = header_table2.rows[1].cells
            row2_2[0].text = "–°–º–µ–Ω–∞"
            row2_2[1].text = str(shift_number)
            row2_2[2].text = "–î–∞—Ç–∞"
            row2_2[3].text = formatted_time
            
            document.add_paragraph("")
            
            # –ü–æ–ª—É—á–∞–µ–º –∏–∑–º–µ—Ä–µ–Ω–∏—è –∏–∑ –ë–î
            measurements = []
            if self.db_integration:
                try:
                    measurements = self.db_integration.db.get_shift_measurements(shift_number)
                except Exception as e:
                    print(f" [REPORT] –û—à–∏–±–∫–∞ –ø–æ–ª—É—á–µ–Ω–∏—è –∏–∑–º–µ—Ä–µ–Ω–∏–π –∏–∑ –ë–î: {e}")
            
            # –°–æ–∑–¥–∞—ë–º —Ç–∞–±–ª–∏—Ü—É –∏–∑–º–µ—Ä–µ–Ω–∏–π
            if measurements:
                # –ó–∞–≥–æ–ª–æ–≤–æ–∫ —Ç–∞–±–ª–∏—Ü—ã
                heading2 = document.add_paragraph("–°–≤–æ–¥–Ω—ã–µ —Ç–∞–±–ª–∏—Ü—ã —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏–π")
                heading2.runs[0].bold = True
                heading2.runs[0].italic = True
                heading2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                document.add_paragraph("")
                
                # –¢–∞–±–ª–∏—Ü–∞ –∏–∑–º–µ—Ä–µ–Ω–∏–π
                measurements_table = document.add_table(rows=len(measurements) + 1, cols=8)
                measurements_table.style = 'Table Grid'
                
                # –ó–∞–≥–æ–ª–æ–≤–∫–∏ —Å—Ç–æ–ª–±—Ü–æ–≤
                header_row = measurements_table.rows[0].cells
                header_row[0].text = "‚Ññ –ø/–ø"
                header_row[1].text = "–í—ã—Å–æ—Ç–∞, –º–º"
                header_row[2].text = "–¢–æ–ª—â–∏–Ω–∞ —Å—Ç–µ–Ω–∫–∏ –≤–≤–µ—Ä—Ö—É, –º–º —Å—Ä–µ–¥–Ω–µ–µ"
                header_row[3].text = "–î–∏–∞–º–µ—Ç—Ä –∫–æ—Ä–ø—É—Å–∞, –º–º —Å—Ä–µ–¥–Ω–µ–µ"
                header_row[4].text = "–î–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞, –º–º —Å—Ä–µ–¥–Ω–µ–µ"
                header_row[5].text = "–¢–æ–ª—â–∏–Ω–∞ —Å—Ç–µ–Ω–∫–∏ –≤–Ω–∏–∑—É, –º–º —Å—Ä–µ–¥–Ω–µ–µ"
                header_row[6].text = "–¢–æ–ª—â–∏–Ω–∞ —Ñ–ª–∞–Ω—Ü–∞, –º–º"
                header_row[7].text = "–¢–æ–ª—â–∏–Ω–∞ –¥–Ω–∞, –º–º —Å—Ä–µ–¥–Ω–µ–µ"
                
                # –í—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏–µ –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤
                for cell in header_row:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                
                # –ó–∞–ø–æ–ª–Ω—è–µ–º –¥–∞–Ω–Ω—ã–µ
                for idx, meas in enumerate(measurements, start=1):
                    row = measurements_table.rows[idx].cells
                    row[0].text = str(idx)
                    row[1].text = f"{meas.get('height_avg', 0):.3f}" if meas.get('height_avg') is not None else "0.000"
                    row[2].text = f"{meas.get('upper_wall_avg', 0):.3f}" if meas.get('upper_wall_avg') is not None else "0.000"
                    row[3].text = f"{meas.get('body_diameter_avg', 0):.3f}" if meas.get('body_diameter_avg') is not None else "0.000"
                    row[4].text = f"{meas.get('flange_diameter_avg', 0):.3f}" if meas.get('flange_diameter_avg') is not None else "0.000"
                    row[5].text = f"{meas.get('bottom_wall_avg', 0):.3f}" if meas.get('bottom_wall_avg') is not None else "0.000"
                    row[6].text = f"{meas.get('flange_thickness_avg', 0):.3f}" if meas.get('flange_thickness_avg') is not None else "0.000"
                    row[7].text = f"{meas.get('bottom_avg', 0):.3f}" if meas.get('bottom_avg') is not None else "0.000"
                    
                    # –í—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏–µ: –ø–µ—Ä–≤—ã–π —Å—Ç–æ–ª–±–µ—Ü –ø–æ —Ü–µ–Ω—Ç—Ä—É, –æ—Å—Ç–∞–ª—å–Ω—ã–µ –ø–æ —Ü–µ–Ω—Ç—Ä—É
                    row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for i in range(1, 8):
                        row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            document.save(filepath)
            print(f" [REPORT] –°—Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω –æ—Ç—á—ë—Ç –ø–æ —Å–º–µ–Ω–µ {shift_number}: {filepath}")
            
            # –ö–æ–ø–∏—Ä—É–µ–º –æ—Ç—á—ë—Ç –Ω–∞ —Ñ–ª–µ—à–∫—É, –µ—Å–ª–∏ –æ–Ω–∞ –¥–æ—Å—Ç—É–ø–Ω–∞
            self.copy_report_to_flash(filepath, filename)
        
        except Exception as e:
            print(f" [REPORT] –û—à–∏–±–∫–∞ —Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—è –æ—Ç—á—ë—Ç–∞: {e}")
    
    def copy_report_to_flash(self, source_filepath: str, filename: str):
        """
        –ö–æ–ø–∏—Ä—É–µ—Ç –æ—Ç—á—ë—Ç –Ω–∞ —Ñ–ª–µ—à–∫—É, –µ—Å–ª–∏ –æ–Ω–∞ –¥–æ—Å—Ç—É–ø–Ω–∞
        
        Args:
            source_filepath: –ü–æ–ª–Ω—ã–π –ø—É—Ç—å –∫ –∏—Å—Ö–æ–¥–Ω–æ–º—É —Ñ–∞–π–ª—É –æ—Ç—á—ë—Ç–∞
            filename: –ò–º—è —Ñ–∞–π–ª–∞ –æ—Ç—á—ë—Ç–∞
        """
        # –ü—É—Ç—å –∫ —Ñ–ª–µ—à–∫–µ
        flash_drive_path = "/media/stend_1/ARS"
        
        try:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —Å—É—â–µ—Å—Ç–≤—É–µ—Ç –ª–∏ —Ñ–ª–µ—à–∫–∞
            if not os.path.exists(flash_drive_path):
                print(f" [REPORT] –§–ª–µ—à–∫–∞ –Ω–µ –Ω–∞–π–¥–µ–Ω–∞: {flash_drive_path}")
                return
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ —ç—Ç–æ –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏—è –∏ –¥–æ—Å—Ç—É–ø–Ω–∞ –¥–ª—è –∑–∞–ø–∏—Å–∏
            if not os.path.isdir(flash_drive_path):
                print(f" [REPORT] –ü—É—Ç—å –Ω–µ —è–≤–ª—è–µ—Ç—Å—è –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏–µ–π: {flash_drive_path}")
                return
            
            if not os.access(flash_drive_path, os.W_OK):
                print(f" [REPORT] –ù–µ—Ç –ø—Ä–∞–≤ –Ω–∞ –∑–∞–ø–∏—Å—å –≤: {flash_drive_path}")
                return
            
            # –°–æ–∑–¥–∞—ë–º –ø—É—Ç—å –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –Ω–∞ —Ñ–ª–µ—à–∫–µ
            flash_filepath = os.path.join(flash_drive_path, filename)
            
            # –ö–æ–ø–∏—Ä—É–µ–º —Ñ–∞–π–ª
            shutil.copy2(source_filepath, flash_filepath)
            
            print(f" [REPORT] –û—Ç—á—ë—Ç —Å–∫–æ–ø–∏—Ä–æ–≤–∞–Ω –Ω–∞ —Ñ–ª–µ—à–∫—É: {flash_filepath}")
            
        except PermissionError:
            print(f" [REPORT] –û—à–∏–±–∫–∞ –¥–æ—Å—Ç—É–ø–∞ –∫ —Ñ–ª–µ—à–∫–µ {flash_drive_path}: –Ω–µ—Ç –ø—Ä–∞–≤ –Ω–∞ –∑–∞–ø–∏—Å—å")
        except Exception as e:
            print(f" [REPORT] –û—à–∏–±–∫–∞ –∫–æ–ø–∏—Ä–æ–≤–∞–Ω–∏—è –æ—Ç—á—ë—Ç–∞ –Ω–∞ —Ñ–ª–µ—à–∫—É: {e}")
    
    def update_product_counters(self, result: str):
        """
        –û–±–Ω–æ–≤–ª–µ–Ω–∏–µ —Å—á—ë—Ç—á–∏–∫–æ–≤ –∏–∑–¥–µ–ª–∏–π
        
        Args:
            result: "GOOD", "CONDITIONALLY_GOOD" –∏–ª–∏ "BAD"
        """
        try:
            if not self.modbus_server or not self.modbus_server.slave_context:
                return
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å–º–µ–Ω—É —Å–º–µ–Ω—ã
            self.check_shift_change()
            
            # 30101 (–∏–Ω–¥–µ–∫—Å 100) - –æ–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ
            total = self.modbus_server.slave_context.getValues(4, 100, 1)[0]
            self.modbus_server.slave_context.setValues(4, 100, [total + 1])
            
            # –û–±–Ω–æ–≤–ª—è–µ–º —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–∏–π —Å—á—ë—Ç—á–∏–∫
            if result == "GOOD":
                # 30102 (–∏–Ω–¥–µ–∫—Å 101) - –≥–æ–¥–Ω—ã–µ
                good = self.modbus_server.slave_context.getValues(4, 101, 1)[0]
                self.modbus_server.slave_context.setValues(4, 101, [good + 1])
            elif result == "CONDITIONALLY_GOOD":
                # 30103 (–∏–Ω–¥–µ–∫—Å 102) - —É—Å–ª–æ–≤–Ω–æ-–≥–æ–¥–Ω—ã–µ
                cond_good = self.modbus_server.slave_context.getValues(4, 102, 1)[0]
                self.modbus_server.slave_context.setValues(4, 102, [cond_good + 1])
            else:  # BAD
                # 30104 (–∏–Ω–¥–µ–∫—Å 103) - –Ω–µ–≥–æ–¥–Ω—ã–µ
                bad = self.modbus_server.slave_context.getValues(4, 103, 1)[0]
                self.modbus_server.slave_context.setValues(4, 103, [bad + 1])
            
            print(f" –°—á—ë—Ç—á–∏–∫–∏ –æ–±–Ω–æ–≤–ª–µ–Ω—ã: –í—Å–µ–≥–æ={total + 1}, –†–µ–∑—É–ª—å—Ç–∞—Ç={result}")
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è —Å—á—ë—Ç—á–∏–∫–æ–≤: {e}")
    
    def increment_parameter_statistics(self, quality_result: dict):
        """
        –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∞—Ü–∏—è —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤
        
        Args:
            quality_result: –°–ª–æ–≤–∞—Ä—å —Å –∫–ª—é—á–∞–º–∏ 'result' –∏ 'param_details'
                - result: "GOOD", "CONDITIONALLY_GOOD" –∏–ª–∏ "BAD"
                - param_details: –°–ø–∏—Å–æ–∫ —Å–ª–æ–≤–∞—Ä–µ–π —Å –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–µ–π –æ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞—Ö
        """
        try:
            if not self.modbus_server or not self.modbus_server.slave_context:
                return
            
            result = quality_result.get('result', 'GOOD')
            param_details = quality_result.get('param_details', [])
            
            # –ï—Å–ª–∏ –¥–µ—Ç–∞–ª—å –≥–æ–¥–Ω–∞—è - –Ω–∏—á–µ–≥–æ –Ω–µ –∏–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä—É–µ–º
            if result == 'GOOD':
                return
            
            # –ú–∞–ø–ø–∏–Ω–≥ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ –Ω–∞ —Ä–µ–≥–∏—Å—Ç—Ä—ã —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏
            # –§–æ—Ä–º–∞—Ç: (param_name, check_type) -> (—É—Å–ª–æ–≤–Ω–æ_–Ω–µ–≥–æ–¥–Ω—ã–π_—Ä–µ–≥–∏—Å—Ç—Ä, –Ω–µ–≥–æ–¥–Ω—ã–π_–º–µ–Ω—å—à–µ_—Ä–µ–≥–∏—Å—Ç—Ä, –Ω–µ–≥–æ–¥–Ω—ã–π_–±–æ–ª—å—à–µ_—Ä–µ–≥–∏—Å—Ç—Ä)
            param_registers = {
                ('height', 'one_sided'): (200, 209, 218),  # 30201, 30210, 30219
                ('upper_wall', 'one_sided'): (201, 210, 219),  # 30202, 30211, 30220
                ('flange_thickness', 'one_sided'): (202, 211, 220),  # 30203, 30212, 30221
                ('body_diameter', 'one_sided'): (208, 217, 222),  # 30209, 30218, 30223
                ('flange_diameter', 'one_sided'): (207, 216, 221),  # 30208, 30217, 30222
                ('bottom_wall', 'two_sided'): {
                    'cond_bad_greater': 203,  # 30204
                    'cond_bad_less': 204,  # 30205
                    'bad_less': 212,  # 30213
                    'bad_greater': 213  # 30214
                },
                ('bottom', 'two_sided'): {
                    'cond_bad_less': 205,  # 30206
                    'cond_bad_greater': 206,  # 30207
                    'bad_less': 214,  # 30215
                    'bad_greater': 215  # 30216
                }
            }
            
            # –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä—É–µ–º —Ä–µ–≥–∏—Å—Ç—Ä—ã –¥–ª—è –ø—Ä–æ–±–ª–µ–º–Ω—ã—Ö –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤
            for param_info in param_details:
                param_name = param_info['name']
                param_status = param_info['status']
                check_type = param_info['check_type']
                base = param_info['base']
                cond_bad_error = param_info.get('cond_bad_error', 0)
                bad_error = param_info.get('bad_error', 0)
                positive_bad_error = param_info.get('positive_bad_error')
                checked_values = param_info.get('checked_values', [])
                
                # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –≥–æ–¥–Ω—ã–µ –ø–∞—Ä–∞–º–µ—Ç—Ä—ã
                if param_status == 'GOOD':
                    continue
                
                # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –∫–ª—é—á –¥–ª—è –ø–æ–∏—Å–∫–∞ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤
                key = (param_name, check_type)
                
                if check_type == 'one_sided':
                    # –î–ª—è –æ–¥–Ω–æ—Å—Ç–æ—Ä–æ–Ω–Ω–∏—Ö –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤
                    if key in param_registers:
                        cond_reg, bad_less_reg, bad_greater_reg = param_registers[key]
                        
                        # –û—Ç—Å–ª–µ–∂–∏–≤–∞–µ–º –Ω–∞–ø—Ä–∞–≤–ª–µ–Ω–∏—è –æ—Ç–∫–ª–æ–Ω–µ–Ω–∏—è –¥–ª—è –≤—Å–µ—Ö –ø—Ä–æ–≤–µ—Ä–µ–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π
                        has_cond_bad = False
                        has_bad_greater = False
                        has_bad_less = False
                        
                        for checked in checked_values:
                            value = checked['value']
                            value_status = checked['status']
                            
                            if result == 'CONDITIONALLY_GOOD' and value_status == 'CONDITIONALLY_GOOD':
                                has_cond_bad = True
                            
                            elif result == 'BAD' and value_status == 'BAD':
                                # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –Ω–∞–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –æ—Ç–∫–ª–æ–Ω–µ–Ω–∏—è –¥–ª—è –Ω–µ–≥–æ–¥–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è
                                if value > base:
                                    # –ë–æ–ª—å—à–µ –Ω–æ—Ä–º—ã
                                    has_bad_greater = True
                                elif value < base + bad_error:
                                    # –ú–µ–Ω—å—à–µ –Ω–æ—Ä–º—ã (–º–µ–Ω—å—à–µ –Ω–µ–≥–æ–¥–Ω–æ–π –≥—Ä–∞–Ω–∏—Ü—ã)
                                    has_bad_less = True
                        
                        # –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä—É–µ–º —Ä–µ–≥–∏—Å—Ç—Ä—ã
                        if result == 'CONDITIONALLY_GOOD' and param_status == 'CONDITIONALLY_GOOD' and has_cond_bad:
                            current = self.modbus_server.slave_context.getValues(4, cond_reg, 1)[0]
                            self.modbus_server.slave_context.setValues(4, cond_reg, [current + 1])
                            print(f" [–°–¢–ê–¢–ò–°–¢–ò–ö–ê] –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω —Ä–µ–≥–∏—Å—Ç—Ä {30001 + cond_reg} ({param_name}, —É—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω—ã–π)")
                        
                        if result == 'BAD' and param_status == 'BAD':
                            if has_bad_greater:
                                current = self.modbus_server.slave_context.getValues(4, bad_greater_reg, 1)[0]
                                self.modbus_server.slave_context.setValues(4, bad_greater_reg, [current + 1])
                                print(f" [–°–¢–ê–¢–ò–°–¢–ò–ö–ê] –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω —Ä–µ–≥–∏—Å—Ç—Ä {30001 + bad_greater_reg} ({param_name}, –Ω–µ–≥–æ–¥–Ω—ã–π, –±–æ–ª—å—à–µ –Ω–æ—Ä–º—ã)")
                            if has_bad_less:
                                current = self.modbus_server.slave_context.getValues(4, bad_less_reg, 1)[0]
                                self.modbus_server.slave_context.setValues(4, bad_less_reg, [current + 1])
                                print(f" [–°–¢–ê–¢–ò–°–¢–ò–ö–ê] –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω —Ä–µ–≥–∏—Å—Ç—Ä {30001 + bad_less_reg} ({param_name}, –Ω–µ–≥–æ–¥–Ω—ã–π, –º–µ–Ω—å—à–µ –Ω–æ—Ä–º—ã)")
                
                elif check_type == 'two_sided':
                    # –î–ª—è –¥–≤—É—Å—Ç–æ—Ä–æ–Ω–Ω–∏—Ö –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤
                    if key in param_registers:
                        regs = param_registers[key]
                        
                        # –û—Ç—Å–ª–µ–∂–∏–≤–∞–µ–º –Ω–∞–ø—Ä–∞–≤–ª–µ–Ω–∏—è –æ—Ç–∫–ª–æ–Ω–µ–Ω–∏—è –¥–ª—è –≤—Å–µ—Ö –ø—Ä–æ–≤–µ—Ä–µ–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π
                        has_cond_bad_greater = False
                        has_cond_bad_less = False
                        has_bad_greater = False
                        has_bad_less = False
                        
                        for checked in checked_values:
                            value = checked['value']
                            value_status = checked['status']
                            
                            if result == 'CONDITIONALLY_GOOD' and value_status == 'CONDITIONALLY_GOOD':
                                # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –Ω–∞–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –¥–ª—è —É—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è
                                if value > base:
                                    has_cond_bad_greater = True
                                elif value < base:
                                    has_cond_bad_less = True
                            
                            elif result == 'BAD' and value_status == 'BAD':
                                # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –Ω–∞–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –¥–ª—è –Ω–µ–≥–æ–¥–Ω–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è
                                if value > base + (positive_bad_error if positive_bad_error else 0):
                                    # –ë–æ–ª—å—à–µ –≤–µ—Ä—Ö–Ω–µ–π –≥—Ä–∞–Ω–∏—Ü—ã
                                    has_bad_greater = True
                                elif value < base + bad_error:
                                    # –ú–µ–Ω—å—à–µ –Ω–∏–∂–Ω–µ–π –≥—Ä–∞–Ω–∏—Ü—ã
                                    has_bad_less = True
                        
                        # –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä—É–µ–º —Ä–µ–≥–∏—Å—Ç—Ä—ã
                        if result == 'CONDITIONALLY_GOOD' and param_status == 'CONDITIONALLY_GOOD':
                            if has_cond_bad_greater:
                                reg_idx = regs['cond_bad_greater']
                                current = self.modbus_server.slave_context.getValues(4, reg_idx, 1)[0]
                                self.modbus_server.slave_context.setValues(4, reg_idx, [current + 1])
                                print(f" [–°–¢–ê–¢–ò–°–¢–ò–ö–ê] –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω —Ä–µ–≥–∏—Å—Ç—Ä {30001 + reg_idx} ({param_name}, —É—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω—ã–π, –±–æ–ª—å—à–µ –Ω–æ—Ä–º—ã)")
                            if has_cond_bad_less:
                                reg_idx = regs['cond_bad_less']
                                current = self.modbus_server.slave_context.getValues(4, reg_idx, 1)[0]
                                self.modbus_server.slave_context.setValues(4, reg_idx, [current + 1])
                                print(f" [–°–¢–ê–¢–ò–°–¢–ò–ö–ê] –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω —Ä–µ–≥–∏—Å—Ç—Ä {30001 + reg_idx} ({param_name}, —É—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω—ã–π, –º–µ–Ω—å—à–µ –Ω–æ—Ä–º—ã)")
                        
                        if result == 'BAD' and param_status == 'BAD':
                            if has_bad_greater:
                                reg_idx = regs['bad_greater']
                                current = self.modbus_server.slave_context.getValues(4, reg_idx, 1)[0]
                                self.modbus_server.slave_context.setValues(4, reg_idx, [current + 1])
                                print(f" [–°–¢–ê–¢–ò–°–¢–ò–ö–ê] –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω —Ä–µ–≥–∏—Å—Ç—Ä {30001 + reg_idx} ({param_name}, –Ω–µ–≥–æ–¥–Ω—ã–π, –±–æ–ª—å—à–µ –Ω–æ—Ä–º—ã)")
                            if has_bad_less:
                                reg_idx = regs['bad_less']
                                current = self.modbus_server.slave_context.getValues(4, reg_idx, 1)[0]
                                self.modbus_server.slave_context.setValues(4, reg_idx, [current + 1])
                                print(f" [–°–¢–ê–¢–ò–°–¢–ò–ö–ê] –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä–æ–≤–∞–Ω —Ä–µ–≥–∏—Å—Ç—Ä {30001 + reg_idx} ({param_name}, –Ω–µ–≥–æ–¥–Ω—ã–π, –º–µ–Ω—å—à–µ –Ω–æ—Ä–º—ã)")
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∏–Ω–∫—Ä–µ–º–µ–Ω—Ç–∞—Ü–∏–∏ —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤: {e}")
            import traceback
            traceback.print_exc()
    
    def check_shift_change(self):
        """
        –ü—Ä–æ–≤–µ—Ä–∫–∞ —Å–º–µ–Ω—ã —Å–º–µ–Ω—ã –∏ —Å–±—Ä–æ—Å —Å—á—ë—Ç—á–∏–∫–æ–≤ –ø—Ä–∏ –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ—Å—Ç–∏
        """
        try:
            if not self.modbus_server or not self.modbus_server.slave_context:
                return
            
            # –ß–∏—Ç–∞–µ–º —Ç–µ–∫—É—â–µ–µ –∑–Ω–∞—á–µ–Ω–∏–µ —Å–º–µ–Ω—ã –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–∞ 40100
            new_shift_number = self.modbus_server.slave_context.getValues(3, 99, 1)[0]
            
            # –ü—Ä–∏ –ø–µ—Ä–≤–æ–º –∑–∞–ø—É—Å–∫–µ –ø—Ä–æ—Å—Ç–æ —Å–∏–Ω—Ö—Ä–æ–Ω–∏–∑–∏—Ä—É–µ–º—Å—è –±–µ–∑ —Å–±—Ä–æ—Å–∞
            if not self.shift_initialized:
                self.current_shift_number = new_shift_number
                self.shift_initialized = True
                print(f" [SHIFT] –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è —Å–º–µ–Ω—ã: —Ç–µ–∫—É—â–∞—è —Å–º–µ–Ω–∞ = {new_shift_number} (–±–µ–∑ —Å–±—Ä–æ—Å–∞)")
                return

            # –ï—Å–ª–∏ —Å–º–µ–Ω–∞ –∏–∑–º–µ–Ω–∏–ª–∞—Å—å –ø–æ—Å–ª–µ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏
            if new_shift_number != self.current_shift_number:
                print(f" –û–±–Ω–∞—Ä—É–∂–µ–Ω–∞ —Å–º–µ–Ω–∞ —Å–º–µ–Ω—ã: {self.current_shift_number} -> {new_shift_number}")
                
                previous_shift = self.current_shift_number
                if previous_shift is not None:
                    self.generate_shift_report(previous_shift)
                    # –û—á–∏—â–∞–µ–º –∏–∑–º–µ—Ä–µ–Ω–∏—è –ø—Ä–µ–¥—ã–¥—É—â–µ–π —Å–º–µ–Ω—ã –ø–æ—Å–ª–µ —Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏—è –æ—Ç—á—ë—Ç–∞
                    if self.db_integration:
                        try:
                            self.db_integration.db.clear_shift_measurements(previous_shift)
                        except Exception as e:
                            print(f" [SHIFT] –û—à–∏–±–∫–∞ –æ—á–∏—Å—Ç–∫–∏ –∏–∑–º–µ—Ä–µ–Ω–∏–π —Å–º–µ–Ω—ã {previous_shift}: {e}")
                
                # –ü—Ä–∏–Ω—É–¥–∏—Ç–µ–ª—å–Ω–æ —Å–æ—Ö—Ä–∞–Ω—è–µ–º –Ω–æ–≤—ã–π –Ω–æ–º–µ—Ä —Å–º–µ–Ω—ã –≤ –ë–î
                if self.db_integration:
                    try:
                        self.db_integration.db.save_single_register(
                            40100, 'holding', int(new_shift_number), "–ù–æ–º–µ—Ä —Å–º–µ–Ω—ã"
                        )
                        print(f" [SHIFT] –°–æ—Ö—Ä–∞–Ω–µ–Ω –Ω–æ–º–µ—Ä —Å–º–µ–Ω—ã {new_shift_number} –≤ –ë–î")
                    except Exception as e:
                        print(f" [SHIFT] –û—à–∏–±–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –Ω–æ–º–µ—Ä–∞ —Å–º–µ–Ω—ã –≤ –ë–î: {e}")
                
                # –°–±—Ä–∞—Å—ã–≤–∞–µ–º –≤—Å–µ —Å—á—ë—Ç—á–∏–∫–∏ –∏–∑–¥–µ–ª–∏–π
                self.reset_product_counters()
                
                # –û–±–Ω–æ–≤–ª—è–µ–º —Ç–µ–∫—É—â—É—é —Å–º–µ–Ω—É
                self.current_shift_number = new_shift_number
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –ø—Ä–æ–≤–µ—Ä–∫–∏ —Å–º–µ–Ω—ã: {e}")
    
    def reset_product_counters(self):
        """
        –°–±—Ä–æ—Å –≤—Å–µ—Ö —Å—á—ë—Ç—á–∏–∫–æ–≤ –∏–∑–¥–µ–ª–∏–π –∏ —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ –ø—Ä–∏ —Å–º–µ–Ω–µ —Å–º–µ–Ω—ã
        """
        try:
            if not self.modbus_server or not self.modbus_server.slave_context:
                return
            
            # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—á—ë—Ç—á–∏–∫–∏ –≤ Input —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ö 30101-30104
            self.modbus_server.slave_context.setValues(4, 100, [0])  # 30101 - –≤—Å–µ–≥–æ
            self.modbus_server.slave_context.setValues(4, 101, [0])  # 30102 - –≥–æ–¥–Ω—ã—Ö
            self.modbus_server.slave_context.setValues(4, 102, [0])  # 30103 - —É—Å–ª–æ–≤–Ω–æ-–≥–æ–¥–Ω—ã—Ö
            self.modbus_server.slave_context.setValues(4, 103, [0])  # 30104 - –Ω–µ–≥–æ–¥–Ω—ã—Ö
            
            # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ä–µ–≥–∏—Å—Ç—Ä—ã —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ 30201-30223
            # –£—Å–ª–æ–≤–Ω–æ-–Ω–µ–≥–æ–¥–Ω—ã–µ (30201-30209): –∏–Ω–¥–µ–∫—Å—ã 200-208
            for idx in range(200, 209):
                self.modbus_server.slave_context.setValues(4, idx, [0])
            
            # –ù–µ–≥–æ–¥–Ω—ã–µ (30210-30223): –∏–Ω–¥–µ–∫—Å—ã 209-222
            for idx in range(209, 223):
                self.modbus_server.slave_context.setValues(4, idx, [0])
            
            print(" –°—á—ë—Ç—á–∏–∫–∏ –∏–∑–¥–µ–ª–∏–π –∏ —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ —Å–±—Ä–æ—à–µ–Ω—ã –¥–ª—è –Ω–æ–≤–æ–π —Å–º–µ–Ω—ã")
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —Å–±—Ä–æ—Å–∞ —Å—á—ë—Ç—á–∏–∫–æ–≤: {e}")
    
    def clear_serial_buffers(self):
        """–û—á–∏—Å—Ç–∫–∞ –±—É—Ñ–µ—Ä–æ–≤ —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –û–° Windows"""
        try:
            if self.sensors and self.sensors.ser:
                lock = getattr(self, 'sensor_reading_lock', None)
                if lock:
                    lock_context = lock
                else:
                    lock_context = contextlib.nullcontext()

                with lock_context:
                    if hasattr(self.sensors.ser, 'reset_input_buffer'):
                        self.sensors.ser.reset_input_buffer()
                    if hasattr(self.sensors.ser, 'reset_output_buffer'):
                        self.sensors.ser.reset_output_buffer()
        except Exception as e:
            # –ò–≥–Ω–æ—Ä–∏—Ä—É–µ–º –æ—à–∏–±–∫–∏ –æ—á–∏—Å—Ç–∫–∏ –±—É—Ñ–µ—Ä–æ–≤
            pass
    
    def flush_sensor_queue(self):
        """–ü–æ–ª–Ω–æ—Å—Ç—å—é –æ—á–∏—â–∞–µ—Ç –æ—á–µ—Ä–µ–¥—å –∏–∑–º–µ—Ä–µ–Ω–∏–π —Å–µ–Ω—Å–æ—Ä–æ–≤"""
        if not hasattr(self, 'sensor_data_queue') or self.sensor_data_queue is None:
            return
        try:
            while True:
                self.sensor_data_queue.get_nowait()
        except Empty:
            pass

    def finalize_calibration_failure(
        self,
        cmd_code: int,
        message: str,
        cleanup_flags: List[str] = None,
        register_doublewords: List[int] = None,
        cache_attrs: List[str] = None,
    ):
        """–û–±—â–∏–π –æ–±—Ä–∞–±–æ—Ç—á–∏–∫ –æ—à–∏–±–æ–∫ –∫–∞–ª–∏–±—Ä–æ–≤–æ–∫: –ª–æ–≥–∏—Ä—É–µ—Ç, —Å–±—Ä–∞—Å—ã–≤–∞–µ—Ç –∫–æ–º–∞–Ω–¥—É –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Å–∏—Å—Ç–µ–º—É –≤ IDLE"""
        if message:
            print(message)
        try:
            self.write_cycle_flag(-1)
        except Exception:
            pass

        # –û–±–Ω—É–ª—è–µ–º –∫–∞–ª–∏–±—Ä–æ–≤–æ—á–Ω—ã–µ —Ä–µ–≥–∏—Å—Ç—Ä—ã, –µ—Å–ª–∏ —É–∫–∞–∑–∞–Ω–æ
        if register_doublewords and self.modbus_server and self.modbus_server.slave_context:
            for base_addr in register_doublewords:
                try:
                    # base_addr - —Ä–µ–∞–ª—å–Ω—ã–π –∞–¥—Ä–µ—Å (–Ω–∞–ø—Ä–∏–º–µ—Ä 40010). –í setValues –∏—Å–ø–æ–ª—å–∑—É–µ–º –∏–Ω–¥–µ–∫—Å (addr-40000)
                    idx = base_addr - 40000
                    self.modbus_server.slave_context.setValues(3, idx, [0])      # —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ
                    self.modbus_server.slave_context.setValues(3, idx + 1, [0])  # –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ
                except Exception as e:
                    print(f" [CALIBRATION] –û—à–∏–±–∫–∞ –æ–±–Ω—É–ª–µ–Ω–∏—è —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ {base_addr}-{base_addr+1}: {e}")

        self.reset_command()
        self.clear_measurement_buffers()
        self.clear_serial_buffers()
        if cleanup_flags:
            for flag in cleanup_flags:
                if hasattr(self, flag):
                    delattr(self, flag)
        if cache_attrs:
            for attr in cache_attrs:
                if hasattr(self, attr):
                    setattr(self, attr, None)
        if hasattr(self, 'calibration_in_progress') and self.calibration_in_progress:
            self.calibration_in_progress = False
        self.current_state = SystemState.IDLE

    def is_valid_measurement(self, value: float, max_range: float = None, min_range: float = None) -> bool:
        """
        –ü—Ä–æ–≤–µ—Ä–∫–∞ –≤–∞–ª–∏–¥–Ω–æ—Å—Ç–∏ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞
        
        Args:
            value: –ó–Ω–∞—á–µ–Ω–∏–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤ –º–º
            max_range: –ú–∞–∫—Å–∏–º–∞–ª—å–Ω—ã–π –¥–∏–∞–ø–∞–∑–æ–Ω –¥–∞—Ç—á–∏–∫–∞ (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é sensor_range_mm * 2)
            min_range: –ú–∏–Ω–∏–º–∞–ª—å–Ω—ã–π –¥–∏–∞–ø–∞–∑–æ–Ω –¥–∞—Ç—á–∏–∫–∞ (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é 20 –º–º - –±–∞–∑–æ–≤–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ)
        
        Returns:
            True –µ—Å–ª–∏ –∑–Ω–∞—á–µ–Ω–∏–µ –≤–∞–ª–∏–¥–Ω–æ, False –µ—Å–ª–∏ –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ (None, 0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω–æ–µ, –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞)
        """
        if value is None:
            return False
        if value <= 0.0:
            return False  # –ù—É–ª–µ–≤—ã–µ –∏ –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã
        
        # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–∏–Ω–∏–º–∞–ª—å–Ω—ã–π –¥–∏–∞–ø–∞–∑–æ–Ω (–±–∞–∑–æ–≤–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ - 20 –º–º)
        if min_range is None:
            min_range = 20.0  # –ë–∞–∑–æ–≤–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞
        
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º –º–∏–Ω–∏–º–∞–ª—å–Ω—ã–π –¥–∏–∞–ø–∞–∑–æ–Ω
        if value < min_range:
            return False  # –ó–Ω–∞—á–µ–Ω–∏–µ –º–µ–Ω—å—à–µ –º–∏–Ω–∏–º–∞–ª—å–Ω–æ–≥–æ –¥–∏–∞–ø–∞–∑–æ–Ω–∞
        
        # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–∞–∫—Å–∏–º–∞–ª—å–Ω—ã–π –¥–∏–∞–ø–∞–∑–æ–Ω
        if max_range is None:
            max_range = self.sensor_range_mm * 2.0  # –ú–∞–∫—Å–∏–º–∞–ª—å–Ω—ã–π –¥–∏–∞–ø–∞–∑–æ–Ω (25 * 2 = 50 –º–º)
        
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º –º–∞–∫—Å–∏–º–∞–ª—å–Ω—ã–π –¥–∏–∞–ø–∞–∑–æ–Ω
        if value > max_range:
            return False  # –ó–Ω–∞—á–µ–Ω–∏–µ –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞ –¥–∞—Ç—á–∏–∫–∞
        
        return True
    
    def clear_measurement_buffers(self):
        """–û—á–∏—Å—Ç–∫–∞ –±—É—Ñ–µ—Ä–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏–π"""
        # –û—á–∏—â–∞–µ–º –æ—Å–Ω–æ–≤–Ω–æ–π –±—É—Ñ–µ—Ä –∫–∞–ª–∏–±—Ä–æ–≤–æ–∫ (measurement_buffer)
        self.measurement_buffer['sensor1'].clear()
        self.measurement_buffer['sensor2'].clear()
        self.measurement_buffer['sensor3'].clear()
        self.measurement_buffer['sensor4'].clear()
        
        # –ë—É—Ñ–µ—Ä—ã –∫–æ–º–∞–Ω–¥—ã 9 (–≤—ã—Å–æ—Ç–∞)
        self.height_measurements = []
        self.obstacle_detected = False
        self.obstacle_filter_count = 0
        self.height_calculated = False
        # –ë—É—Ñ–µ—Ä—ã –∏ —Ñ–ª–∞–≥–∏ –∫–æ–º–∞–Ω–¥—ã 103 (–∫–∞–ª–∏–±—Ä–æ–≤–∫–∞ –≤—ã—Å–æ—Ç—ã)
        self.height_calibration_nonzero_count = 0
        self.distance_to_plane_calculated = False
        self.recent_measurements = []
        
        # –ë—É—Ñ–µ—Ä—ã –∫–æ–º–∞–Ω–¥—ã 10
        self.sensor1_measurements = []
        self.sensor2_measurements = []
        self.wall_thickness_buffer = []
        self.temp_sensor1_buffer = []
        self.temp_sensor2_buffer = []
        
        # –ë—É—Ñ–µ—Ä—ã –∫–æ–º–∞–Ω–¥—ã 11 (—Ñ–ª–∞–Ω–µ—Ü)
        self.sensor1_flange_measurements = []
        self.sensor3_measurements = []
        self.sensor4_measurements = []
        self.temp_sensor1_flange_buffer = []
        self.temp_sensor3_buffer = []
        self.temp_sensor4_buffer = []
        # –ë—É—Ñ–µ—Ä—ã –¥–∏–∞–º–µ—Ç—Ä–æ–≤ –∏ —Ç–æ–ª—â–∏–Ω - –æ—á–∏—â–∞–µ–º –ø—Ä–∏ –∫–∞–∂–¥–æ–º –Ω–æ–≤–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–∏
        self.body_diameter_buffer = []
        self.flange_diameter_buffer = []
        self.bottom_thickness_buffer = []
        
        # –ë—É—Ñ–µ—Ä—ã –∫–æ–º–∞–Ω–¥—ã 12
        self.sensor1_bottom_measurements = []
        self.sensor2_bottom_measurements = []
        self.temp_sensor1_bottom_buffer = []
        self.temp_sensor2_bottom_buffer = []
        self.bottom_wall_thickness_buffer = []

        # –ë—É—Ñ–µ—Ä—ã —Ä–∞–∑–¥–µ–ª—å–Ω—ã—Ö –∫–æ–º–∞–Ω–¥ –¥–∏–∞–º–µ—Ç—Ä–æ–≤ (20/30/40)
        self.sensor3_flange_only_measurements = []
        self.sensor3_body_only_measurements = []
        self.sensor3_body2_measurements = []
        self.temp_sensor3_flange_only_buffer = []
        self.temp_sensor3_body_only_buffer = []
        self.temp_sensor3_body2_buffer = []
        self.body_only_diameter_buffer = []
        self.body2_diameter_buffer = []
        
        # –ë—É—Ñ–µ—Ä—ã QUAD –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ —Ä–µ–∂–∏–º–∞ (CMD=200)
        self.stream_temp_sensor1_buffer = []
        self.stream_temp_sensor2_buffer = []
        self.stream_temp_sensor3_buffer = []
        self.stream_temp_sensor4_buffer = []
        
        print(" –ë—É—Ñ–µ—Ä—ã –∏–∑–º–µ—Ä–µ–Ω–∏–π –æ—á–∏—â–µ–Ω—ã")
        self.flush_sensor_queue()
    
    def stop_all_streams(self):
        """–û—Å—Ç–∞–Ω–æ–≤–∫–∞ –≤—Å–µ—Ö –∞–∫—Ç–∏–≤–Ω—ã—Ö –ø–æ—Ç–æ–∫–æ–≤—ã—Ö —Ä–µ–∂–∏–º–æ–≤"""
        if self.sensors and self.stream_active_quad:
            try:
                self.stream_active_quad = False
                print(" –û—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω QUAD –ø–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º")
            except Exception as e:
                print(f" –û—à–∏–±–∫–∞ –æ—Å—Ç–∞–Ω–æ–≤–∫–∏ QUAD —Ä–µ–∂–∏–º–∞: {e}")
    
    def execute_state_actions(self):
        """–í—ã–ø–æ–ª–Ω–µ–Ω–∏–µ –¥–µ–π—Å—Ç–≤–∏–π –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç —Ç–µ–∫—É—â–µ–≥–æ —Å–æ—Å—Ç–æ—è–Ω–∏—è"""
        if self.current_state == SystemState.IDLE:
            self.handle_idle_state()
            
        # –ö–∞–ª–∏–±—Ä–æ–≤–∫–∏
        elif self.current_state == SystemState.CALIBRATE_WALL:
            self.handle_calibrate_wall_state()
        elif self.current_state == SystemState.CALIBRATE_BOTTOM:
            self.handle_calibrate_bottom_state()
        elif self.current_state == SystemState.CALIBRATE_FLANGE:
            self.handle_calibrate_flange_state()
        elif self.current_state == SystemState.CALIBRATE_HEIGHT:
            self.handle_calibrate_height_state()
        elif self.current_state == SystemState.CALIBRATE_FLANGE_DIAMETER:
            self.handle_calibrate_flange_diameter_state()
        elif self.current_state == SystemState.CALIBRATE_BODY_DIAMETER_SEPARATE:
            self.handle_calibrate_body_diameter_separate_state()
        elif self.current_state == SystemState.CALIBRATE_BODY2_DIAMETER:
            self.handle_calibrate_body2_diameter_state()
        elif self.current_state == SystemState.CONFIGURE_SENSOR3_RANGE:
            self.handle_configure_sensor3_range_state()
        elif self.current_state == SystemState.DEBUG_REGISTERS:
            self.handle_debug_registers_state()
            
        # –ò–∑–º–µ—Ä–µ–Ω–∏–µ –≤—ã—Å–æ—Ç—ã - –û–¢–ö–õ–Æ–ß–ï–ù–û: —Ä–∞—Å—á–µ—Ç –≤—ã—Å–æ—Ç—ã –ø—Ä–æ–∏–∑–≤–æ–¥–∏—Ç—Å—è –Ω–∞ –ü–õ–ö
        # –†–µ–≥–∏—Å—Ç—Ä 40057-40058 –∑–∞–ø–∏—Å—ã–≤–∞–µ—Ç—Å—è —Ç–æ–ª—å–∫–æ –ü–õ–ö/HMI
        # elif self.current_state == SystemState.MEASURE_HEIGHT_PROCESS:
        #     self.handle_measure_height_process_state()
            
        # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - –≤–µ—Ä—Ö–Ω—è—è —Å—Ç–µ–Ω–∫–∞
        elif self.current_state == SystemState.MEASURE_WALL_PROCESS:
            self.handle_measure_wall_process_state()
        elif self.current_state == SystemState.MEASURE_WALL_CALCULATE:
            self.handle_calculate_wall_state()
            
        # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - —Ñ–ª–∞–Ω–µ—Ü
        elif self.current_state == SystemState.MEASURE_FLANGE_PROCESS:
            self.handle_measure_flange_process_state()
        elif self.current_state == SystemState.MEASURE_FLANGE_CALCULATE:
            self.handle_calculate_flange_state()
        elif self.current_state == SystemState.MEASURE_FLANGE_ONLY_PROCESS:
            self.handle_measure_flange_only_process_state()
        elif self.current_state == SystemState.MEASURE_FLANGE_ONLY_CALCULATE:
            self.handle_calculate_flange_only_state()
        elif self.current_state == SystemState.MEASURE_BODY_ONLY_PROCESS:
            self.handle_measure_body_only_process_state()
        elif self.current_state == SystemState.MEASURE_BODY_ONLY_CALCULATE:
            self.handle_calculate_body_only_state()
        elif self.current_state == SystemState.MEASURE_BODY2_PROCESS:
            self.handle_measure_body2_process_state()
        elif self.current_state == SystemState.MEASURE_BODY2_CALCULATE:
            self.handle_calculate_body2_state()
            
        # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª –∏–∑–º–µ—Ä–µ–Ω–∏—è - –Ω–∏–∂–Ω—è—è —Å—Ç–µ–Ω–∫–∞
        elif self.current_state == SystemState.MEASURE_BOTTOM_PROCESS:
            self.handle_measure_bottom_process_state()
        elif self.current_state == SystemState.MEASURE_BOTTOM_CALCULATE:
            self.handle_calculate_bottom_state()
            
        # –û—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞ –∏–∑–¥–µ–ª–∏—è
        elif self.current_state == SystemState.QUALITY_EVALUATION:
            self.handle_quality_evaluation_state()
            
        # –ü–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º (QUAD - –≤—Å–µ 4 –¥–∞—Ç—á–∏–∫–∞)
        elif self.current_state == SystemState.STREAM_QUAD:
            self.handle_stream_quad_state()
            
        elif self.current_state == SystemState.ERROR:
            self.handle_error_state()
    
    def handle_calibrate_height_state(self):
        """
        CMD=103: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –≤—ã—Å–æ—Ç—ã
        - –û–∂–∏–¥–∞–µ–º 3 –ø–æ—Å–ª–µ–¥–æ–≤–∞—Ç–µ–ª—å–Ω—ã—Ö –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏—è —Å –¥–∞—Ç—á–∏–∫–∞ 1
        - –ß–∏—Ç–∞–µ–º —à–∞–≥–∏ (40052-40053) –∏ –∏–º–ø—É–ª—å—Å—ã –Ω–∞ –º–º (40054)
        - –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—É—é –≤—ã—Å–æ—Ç—É (40008-40009)
        - –í—ã—á–∏—Å–ª—è–µ–º –¥–∏—Å—Ç–∞–Ω—Ü–∏—é –¥–æ –Ω–∞—á–∞–ª–∞ –ø–ª–æ—Å–∫–æ—Å—Ç–∏ –∏ –∑–∞–ø–∏—Å—ã–≤–∞–µ–º –≤ 40055-40056
        - –û–±–Ω–æ–≤–ª—è–µ–º —Å—Ç–∞—Ç—É—Å: 103 (–ø–æ–∏—Å–∫) ‚Üí 931 (—Ä–∞—Å—Å—á–∏—Ç–∞–Ω–æ), –æ–∂–∏–¥–∞–µ–º CMD=0
        """
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            self.current_state = SystemState.ERROR
            return
        
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è —Ç–∞–π–º–µ—Ä–∞ —á–∞—Å—Ç–æ—Ç—ã –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–∏
            if self.frequency_start_time is None:
                self.frequency_start_time = time.time()
                self.last_frequency_display = self.frequency_start_time
            
            # –ß–∏—Ç–∞–µ–º —Ç–æ–ª—å–∫–æ –¥–∞—Ç—á–∏–∫ 1
            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor1_mm is None:
                time.sleep(0.001)
                return
            
            # –û—Ç–ª–∞–¥–æ—á–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ —Å—ã—Ä—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏—è—Ö (—Ç–æ–ª—å–∫–æ –ø–µ—Ä–≤—ã–µ –Ω–µ—Å–∫–æ–ª—å–∫–æ –∏–∑–º–µ—Ä–µ–Ω–∏–π)
            if self.frequency_counter <= 5:
                print(f" [CMD=103] –û—Ç–ª–∞–¥–∫–∞ #{self.frequency_counter}: sensor1_mm={sensor1_mm}")
            
            
            # –£–≤–µ–ª–∏—á–∏–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫ –∏–∑–º–µ—Ä–µ–Ω–∏–π –∏ –≤—ã–≤–æ–¥–∏–º —á–∞—Å—Ç–æ—Ç—É —Ä–∞–∑ –≤ —Å–µ–∫—É–Ω–¥—É
            self.frequency_counter += 1
            current_time = time.time()
            if current_time - self.last_frequency_display >= 1.0:
                elapsed = current_time - self.frequency_start_time
                if elapsed > 0:
                    instant_freq = self.frequency_counter / elapsed
                    status = "–ü–æ–∏—Å–∫ 3 –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏–π" if not self.distance_to_plane_calculated else "–†–∞—Å—Å—á–∏—Ç–∞–Ω–æ"
                    print(f" [CMD=103] {status}: {instant_freq:.1f} –ì—Ü | –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.frequency_counter}")
                self.last_frequency_display = current_time
            
            # –õ–æ–≥–∏–∫–∞ –ø–æ–∏—Å–∫–∞ 3 –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏–π –∏–∑ –ø–æ—Å–ª–µ–¥–Ω–∏—Ö 5 –∏–∑–º–µ—Ä–µ–Ω–∏–π
            # –î–æ–±–∞–≤–ª—è–µ–º —Ç–µ–∫—É—â–µ–µ –∏–∑–º–µ—Ä–µ–Ω–∏–µ –≤ –±—É—Ñ–µ—Ä (—Ö—Ä–∞–Ω–∏–º –ø–æ—Å–ª–µ–¥–Ω–∏–µ 5)
            # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è –ø–µ—Ä–µ–¥ –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ–º
            if self.is_valid_measurement(sensor1_mm):
                self.recent_measurements.append(sensor1_mm)
            if len(self.recent_measurements) > 5:
                self.recent_measurements.pop(0)  # –£–¥–∞–ª—è–µ–º —Å–∞–º–æ–µ —Å—Ç–∞—Ä–æ–µ
            
            # –ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ–º –Ω–µ–Ω—É–ª–µ–≤—ã–µ –ø–æ–∫–∞–∑–∞–Ω–∏—è –≤ –±—É—Ñ–µ—Ä–µ
            nonzero_count = sum(1 for m in self.recent_measurements if m is not None and m > 0)
            
            # –í—ã–≤–æ–¥–∏–º —Ç–æ–ª—å–∫–æ –ø—Ä–∏ –∏–∑–º–µ–Ω–µ–Ω–∏–∏ –∫–æ–ª–∏—á–µ—Å—Ç–≤–∞ –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –≤ –±—É—Ñ–µ—Ä–µ
            if hasattr(self, '_last_nonzero_count') and self._last_nonzero_count != nonzero_count:
                if sensor1_mm is not None and sensor1_mm > 0:
                    print(f" [CMD=103] –ù–µ–Ω—É–ª–µ–≤–æ–µ –ø–æ–∫–∞–∑–∞–Ω–∏–µ: {sensor1_mm:.3f}–º–º | –ù–µ–Ω—É–ª–µ–≤—ã—Ö –≤ –±—É—Ñ–µ—Ä–µ: {nonzero_count}/5")
                else:
                    print(f" [CMD=103] –ù—É–ª–µ–≤–æ–µ –ø–æ–∫–∞–∑–∞–Ω–∏–µ: {sensor1_mm} | –ù–µ–Ω—É–ª–µ–≤—ã—Ö –≤ –±—É—Ñ–µ—Ä–µ: {nonzero_count}/5")
            self._last_nonzero_count = nonzero_count
            
            # –ï—Å–ª–∏ –Ω–∞–π–¥–µ–Ω–æ 3+ –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏—è –≤ –±—É—Ñ–µ—Ä–µ –∏ –µ—â–µ –Ω–µ —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–æ
            if nonzero_count >= 3 and not self.distance_to_plane_calculated:
                print(f" [CMD=103] –ù–∞–π–¥–µ–Ω–æ {nonzero_count} –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏–π –≤ –±—É—Ñ–µ—Ä–µ! –ù–∞—á–∏–Ω–∞–µ–º —Ä–∞—Å—á–µ—Ç...")
                # –ß–∏—Ç–∞–µ–º —Ä–µ–≥–∏—Å—Ç—Ä—ã
                steps = self.read_register_40020()          # 40052-40053
                pulses_per_mm = self.read_register_40021()  # 40054
                reference_height = self.read_reference_height()  # 40008-40009
                
                print(f" [CMD=103] –î–∞–Ω–Ω—ã–µ –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞: —à–∞–≥–∏={steps}, –∏–º–ø—É–ª—å—Å—ã/–º–º={pulses_per_mm}, —ç—Ç–∞–ª–æ–Ω–Ω–∞—è_–≤—ã—Å–æ—Ç–∞={reference_height}")
                
                if steps is not None and pulses_per_mm is not None and pulses_per_mm > 0 and reference_height is not None:
                    # –î–∏—Å—Ç–∞–Ω—Ü–∏—è –¥–æ –Ω–∞—á–∞–ª–∞ –ø–ª–æ—Å–∫–æ—Å—Ç–∏ = —à–∞–≥–∏/–∏–º–ø—É–ª—å—Å—ã_–Ω–∞_–º–º + —ç—Ç–∞–ª–æ–Ω–Ω–∞—è –≤—ã—Å–æ—Ç–∞
                    distance_to_plane = (steps / float(pulses_per_mm)) + float(reference_height)
                    self.write_distance_to_plane(distance_to_plane)
                    self.write_cycle_flag(931)  # –°—Ç–∞—Ç—É—Å: —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–æ
                    # –í—ã–≤–æ–¥–∏–º —Ç–æ–ª—å–∫–æ –∫–æ–≥–¥–∞ —Ñ–ª–∞–≥ –ø–µ—Ä–µ—Ö–æ–¥–∏—Ç –∏–∑ False –≤ True
                    if not self.distance_to_plane_calculated:
                        print(f" [CMD=103] –§–õ–ê–ì –ò–ó–ú–ï–ù–ò–õ–°–Ø: distance_to_plane_calculated False ‚Üí True")
                        print(f" [CMD=103] –î–∏—Å—Ç–∞–Ω—Ü–∏—è –¥–æ –ø–ª–æ—Å–∫–æ—Å—Ç–∏ —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–∞: {distance_to_plane:.3f}–º–º")
                        print(f" [CMD=103] –û–∂–∏–¥–∞–µ–º CMD=0 –¥–ª—è –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏")
                    self.distance_to_plane_calculated = True
                else:
                    print(f" [CMD=103] –û–®–ò–ë–ö–ê: –ù–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å –¥–∞–Ω–Ω—ã–µ –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!")
            elif nonzero_count >= 3:
                pass  # –£–∂–µ —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–æ, –Ω–∏—á–µ–≥–æ –Ω–µ –≤—ã–≤–æ–¥–∏–º
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –≤—ã—Å–æ—Ç—ã (CMD=103): {e}")
            self.current_state = SystemState.ERROR

    def handle_debug_registers_state(self):
        """CMD=104: –û—Ç–ª–∞–¥–∫–∞ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ - –ø–æ–∫–∞–∑—ã–≤–∞–µ—Ç –¥–∞–Ω–Ω—ã–µ —Ä–∞–∑ –≤ —Å–µ–∫—É–Ω–¥—É"""
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è —Ç–∞–π–º–µ—Ä–∞ –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –≤—ã–∑–æ–≤–µ
            if not hasattr(self, 'debug_start_time'):
                self.debug_start_time = time.time()
                self.debug_last_display = self.debug_start_time
                print(" [CMD=104] –ù–∞—á–∞–ª–æ –æ—Ç–ª–∞–¥–∫–∏ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤...")
            
            current_time = time.time()
            # –í—ã–≤–æ–¥–∏–º –¥–∞–Ω–Ω—ã–µ —Ä–∞–∑ –≤ —Å–µ–∫—É–Ω–¥—É
            if current_time - self.debug_last_display >= 1.0:
                # –ß–∏—Ç–∞–µ–º —Ä–µ–≥–∏—Å—Ç—Ä —Å—Ç–∞—Ç—É—Å–∞ (30009)
                try:
                    status_values = self.modbus_server.slave_context.getValues(4, 8, 1)  # 30009 -> index 8
                    status = status_values[0] if status_values else "None"
                except:
                    status = "–û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è"
                
                # –ß–∏—Ç–∞–µ–º –≤—Å–µ —Ä–µ–≥–∏—Å—Ç—Ä—ã
                steps_raw = self.read_register_40020_raw()
                pulses_per_mm_raw = self.read_register_40021_raw()
                reference_height_raw = self.read_reference_height_raw()
                
                print(f" [CMD=104] === –û–¢–õ–ê–î–ö–ê –†–ï–ì–ò–°–¢–†–û–í ===")
                print(f" [CMD=104] –°—Ç–∞—Ç—É—Å (30009): {status}")
                print(f" [CMD=104] 40052-40053 (—à–∞–≥–∏): {steps_raw}")
                print(f" [CMD=104] 40054 (–∏–º–ø—É–ª—å—Å—ã/–º–º): {pulses_per_mm_raw}")
                print(f" [CMD=104] 40008-40009 (—ç—Ç–∞–ª–æ–Ω–Ω–∞—è –≤—ã—Å–æ—Ç–∞): {reference_height_raw}")
                print(f" [CMD=104] =========================")
                
                self.debug_last_display = current_time
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –æ—Ç–ª–∞–¥–∫–∏ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ (CMD=104): {e}")
            self.current_state = SystemState.ERROR
    
    def handle_calibrate_flange_diameter_state(self):
        """
        CMD=105: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞
        - –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40030-40031
        - –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω–æ —Å–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ –ø–æ–ª—É—á–µ–Ω–∏—è CMD=0
        - –†–∞—Å—á–µ—Ç –∏ –∑–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –≤—ã–ø–æ–ª–Ω—è–µ—Ç—Å—è –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ CMD -> 0
        """
        if not self.sensors:
            self.finalize_calibration_failure(
                105,
                " [CMD=105] –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!",
                cleanup_flags=['calibrate_flange_diameter_started'],
                register_doublewords=[40032],
                cache_attrs=['cached_distance_sensor3_to_center'],
            )
            return
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∑–∞–ø—É—Å–∫–µ
            if not hasattr(self, 'calibrate_flange_diameter_started'):
                self.calibration_in_progress = True
                self.calibrate_flange_diameter_started = True
                self.calibrate_flange_diameter_sensor3_buffer = []
                self.calibrate_flange_diameter_measurement_count = 0
                self.calibrate_flange_diameter_last_log_time = time.time()
                
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π
                self.clear_serial_buffers()
                self.flush_sensor_queue()
                
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä –¥–∞—Ç—á–∏–∫–∞ 3
                self.measurement_buffer['sensor3'].clear()
                
                # –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40030-40031
                reference_flange_diameter = self.read_reference_flange_diameter()
                print(f"üîß –ù–ê–ß–ê–õ–û –ö–ê–õ–ò–ë–†–û–í–ö–ò –î–ò–ê–ú–ï–¢–†–ê –§–õ–ê–ù–¶–ê (CMD=105)")
                print(f" [CMD=105] –≠—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞: {reference_flange_diameter:.3f} –º–º")
                
                if reference_flange_diameter <= 0:
                    self.finalize_calibration_failure(
                        105,
                        " [CMD=105] –û–®–ò–ë–ö–ê: –≠—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞ –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –±–æ–ª—å—à–µ 0!",
                        cleanup_flags=['calibrate_flange_diameter_started'],
                        register_doublewords=[40032],
                        cache_attrs=['cached_distance_sensor3_to_center'],
                    )
                    return
                
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏
                self.write_cycle_flag(105)
                print(f" [CMD=105] –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω—ã–π —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 3...")
            
            # –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω—ã–π —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö (–±–µ–∑ —Ç–∞–π–º–∞—É—Ç–∞)
            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor3_mm is None:
                time.sleep(0.002)
                return
                
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –≤–∞–ª–∏–¥–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 3
            if self.is_valid_measurement(sensor3_mm):
                self.measurement_buffer['sensor3'].append(sensor3_mm)
                self.calibrate_flange_diameter_sensor3_buffer.append(sensor3_mm)
                self.calibrate_flange_diameter_measurement_count += 1

            # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø—Ä–æ–≥—Ä–µ—Å—Å —Ä–∞–∑ –≤ —Å–µ–∫—É–Ω–¥—É
            current_time = time.time()
            if current_time - getattr(self, 'calibrate_flange_diameter_last_log_time', 0) >= 1.0:
                print(f" [CMD=105] –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö... –ò–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 3: {self.calibrate_flange_diameter_measurement_count}")
                self.calibrate_flange_diameter_last_log_time = current_time
                
        except Exception as e:
            print(f" [CMD=105] –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞: {e}")
            import traceback
            traceback.print_exc()
            self.finalize_calibration_failure(
                105,
                f" [CMD=105] –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_flange_diameter_started'],
                register_doublewords=[40032],
                cache_attrs=['cached_distance_sensor3_to_center'],
            )
    
    def read_reference_flange_diameter(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40030, 40031"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # HMI: —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ –≤ 40030, –º–ª–∞–¥—à–µ–µ –≤ 40031
                values = self.modbus_server.slave_context.getValues(3, 30, 2)  # 40030-40031 -> –∏–Ω–¥–µ–∫—Å—ã 29-30
                if values and len(values) >= 2:
                    high_word = int(values[0])  # 40030 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])  # 40031 - –º–ª–∞–¥—à–∏–π
                    diameter = self.doubleword_to_float(low_word, high_word)
                    return diameter
        except Exception as e:
            print(f" [CMD=105] –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞: {e}")
        return 0.0

    def handle_calibrate_body_diameter_separate_state(self):
        """
        CMD=107: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞
        - –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40034-40035
        - –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω–æ —Å–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ –ø–æ–ª—É—á–µ–Ω–∏—è CMD=0
        - –†–∞—Å—á–µ—Ç –∏ –∑–∞–ø–∏—Å—å –≤—ã–ø–æ–ª–Ω—è—é—Ç—Å—è –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ CMD -> 0
        """
        if not self.sensors:
            self.finalize_calibration_failure(
                107,
                " [CMD=107] –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!",
                cleanup_flags=['calibrate_body_diameter_separate_started'],
                register_doublewords=[40038],
                cache_attrs=['cached_distance_sensor3_to_center_body'],
            )
            return

        try:
            if not hasattr(self, 'calibrate_body_diameter_separate_started'):
                self.calibration_in_progress = True
                self.calibrate_body_diameter_separate_started = True
                self.calibrate_body_diameter_separate_sensor3_buffer = []
                self.calibrate_body_diameter_separate_measurement_count = 0
                self.calibrate_body_diameter_separate_last_log_time = time.time()

                self.clear_serial_buffers()
                self.flush_sensor_queue()
                self.measurement_buffer['sensor3'].clear()

                reference_body_diameter = self.read_reference_body_diameter_separate()
                print("üîß –ù–ê–ß–ê–õ–û –ö–ê–õ–ò–ë–†–û–í–ö–ò –†–ê–ó–î–ï–õ–¨–ù–û–ì–û –î–ò–ê–ú–ï–¢–†–ê –ö–û–†–ü–£–°–ê (CMD=107)")
                print(f" [CMD=107] –≠—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä: {reference_body_diameter:.3f} –º–º")

                if reference_body_diameter <= 0:
                    self.finalize_calibration_failure(
                        107,
                        " [CMD=107] –û–®–ò–ë–ö–ê: –≠—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –±–æ–ª—å—à–µ 0!",
                        cleanup_flags=['calibrate_body_diameter_separate_started'],
                        register_doublewords=[40038],
                        cache_attrs=['cached_distance_sensor3_to_center_body'],
                    )
                    return

                self.write_cycle_flag(107)
                print(" [CMD=107] –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω—ã–π —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 3...")

            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor3_mm is None:
                time.sleep(0.002)
                return

            if self.is_valid_measurement(sensor3_mm):
                self.measurement_buffer['sensor3'].append(sensor3_mm)
                self.calibrate_body_diameter_separate_sensor3_buffer.append(sensor3_mm)
                self.calibrate_body_diameter_separate_measurement_count += 1

            current_time = time.time()
            if current_time - getattr(self, 'calibrate_body_diameter_separate_last_log_time', 0) >= 1.0:
                print(f" [CMD=107] –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö... –ò–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 3: {self.calibrate_body_diameter_separate_measurement_count}")
                self.calibrate_body_diameter_separate_last_log_time = current_time

        except Exception as e:
            print(f" [CMD=107] –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞: {e}")
            import traceback
            traceback.print_exc()
            self.finalize_calibration_failure(
                107,
                f" [CMD=107] –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_body_diameter_separate_started'],
                register_doublewords=[40038],
                cache_attrs=['cached_distance_sensor3_to_center_body'],
            )

    def handle_calibrate_body2_diameter_state(self):
        """
        CMD=108: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2
        - –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40036-40037
        - –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω–æ —Å–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ –ø–æ–ª—É—á–µ–Ω–∏—è CMD=0
        - –†–∞—Å—á–µ—Ç –∏ –∑–∞–ø–∏—Å—å –≤—ã–ø–æ–ª–Ω—è—é—Ç—Å—è –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ CMD -> 0
        """
        if not self.sensors:
            self.finalize_calibration_failure(
                108,
                " [CMD=108] –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!",
                cleanup_flags=['calibrate_body2_diameter_started'],
                register_doublewords=[40040],
                cache_attrs=['cached_distance_sensor3_to_center_body2'],
            )
            return

        try:
            if not hasattr(self, 'calibrate_body2_diameter_started'):
                self.calibration_in_progress = True
                self.calibrate_body2_diameter_started = True
                self.calibrate_body2_diameter_sensor3_buffer = []
                self.calibrate_body2_diameter_measurement_count = 0
                self.calibrate_body2_diameter_last_log_time = time.time()

                self.clear_serial_buffers()
                self.flush_sensor_queue()
                self.measurement_buffer['sensor3'].clear()

                reference_body2_diameter = self.read_reference_body2_diameter()
                print("üîß –ù–ê–ß–ê–õ–û –ö–ê–õ–ò–ë–†–û–í–ö–ò –î–ò–ê–ú–ï–¢–†–ê –ö–û–†–ü–£–°–ê 2 (CMD=108)")
                print(f" [CMD=108] –≠—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä: {reference_body2_diameter:.3f} –º–º")

                if reference_body2_diameter <= 0:
                    self.finalize_calibration_failure(
                        108,
                        " [CMD=108] –û–®–ò–ë–ö–ê: –≠—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –±–æ–ª—å—à–µ 0!",
                        cleanup_flags=['calibrate_body2_diameter_started'],
                        register_doublewords=[40040],
                        cache_attrs=['cached_distance_sensor3_to_center_body2'],
                    )
                    return

                self.write_cycle_flag(108)
                print(" [CMD=108] –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω—ã–π —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 3...")

            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor3_mm is None:
                time.sleep(0.002)
                return

            if self.is_valid_measurement(sensor3_mm):
                self.measurement_buffer['sensor3'].append(sensor3_mm)
                self.calibrate_body2_diameter_sensor3_buffer.append(sensor3_mm)
                self.calibrate_body2_diameter_measurement_count += 1

            current_time = time.time()
            if current_time - getattr(self, 'calibrate_body2_diameter_last_log_time', 0) >= 1.0:
                print(f" [CMD=108] –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö... –ò–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 3: {self.calibrate_body2_diameter_measurement_count}")
                self.calibrate_body2_diameter_last_log_time = current_time

        except Exception as e:
            print(f" [CMD=108] –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2: {e}")
            import traceback
            traceback.print_exc()
            self.finalize_calibration_failure(
                108,
                f" [CMD=108] –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_body2_diameter_started'],
                register_doublewords=[40040],
                cache_attrs=['cached_distance_sensor3_to_center_body2'],
            )

    def read_reference_body_diameter_separate(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ (—Ä–∞–∑–¥–µ–ª—å–Ω–æ) –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40034, 40035"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 34, 2)  # 40034-40035
                if values and len(values) >= 2:
                    high_word = int(values[0])
                    low_word = int(values[1])
                    return self.doubleword_to_float(low_word, high_word)
        except Exception as e:
            print(f" [CMD=107] –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞: {e}")
        return 0.0

    def read_reference_body2_diameter(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40036, 40037"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 36, 2)  # 40036-40037
                if values and len(values) >= 2:
                    high_word = int(values[0])
                    low_word = int(values[1])
                    return self.doubleword_to_float(low_word, high_word)
        except Exception as e:
            print(f" [CMD=108] –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞: {e}")
        return 0.0

    def write_distance_sensor3_to_center_body(self, distance: float):
        """–ó–∞–ø–∏—Å—å —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫ 3 - —Ü–µ–Ω—Ç—Ä –¥–ª—è —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ –≤ 40038-40039"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                low_word, high_word = self.float_to_doubleword(distance)
                self.modbus_server.slave_context.setValues(3, 38, [int(high_word)])
                self.modbus_server.slave_context.setValues(3, 39, [int(low_word)])
                print(f" [CMD=107] –ó–∞–ø–∏—Å–∞–Ω–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ 40038-40039: {distance:.3f} –º–º")
        except Exception as e:
            print(f" [CMD=107] –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è 40038-40039: {e}")

    def write_distance_sensor3_to_center_body2(self, distance: float):
        """–ó–∞–ø–∏—Å—å —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫ 3 - —Ü–µ–Ω—Ç—Ä –¥–ª—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 –≤ 40040-40041"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                low_word, high_word = self.float_to_doubleword(distance)
                self.modbus_server.slave_context.setValues(3, 40, [int(high_word)])
                self.modbus_server.slave_context.setValues(3, 41, [int(low_word)])
                print(f" [CMD=108] –ó–∞–ø–∏—Å–∞–Ω–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ 40040-40041: {distance:.3f} –º–º")
        except Exception as e:
            print(f" [CMD=108] –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è 40040-40041: {e}")
    
    def handle_configure_sensor3_range_state(self):
        """
        CMD=106: –ù–∞—Å—Ç—Ä–æ–π–∫–∞ –¥–∏–∞–ø–∞–∑–æ–Ω–æ–≤ –¥–ª—è –¥–∏—Å–∫—Ä–µ—Ç–Ω–æ–≥–æ —Å–∏–≥–Ω–∞–ª–∞ –¥–∞—Ç—á–∏–∫–∞ 3
        - –ß–∏—Ç–∞–µ–º –Ω–∞—á–∞–ª–æ –¥–∏–∞–ø–∞–∑–æ–Ω–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40404-40405 (–≤ –º–º)
        - –ß–∏—Ç–∞–µ–º –∫–æ–Ω–µ—Ü –¥–∏–∞–ø–∞–∑–æ–Ω–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40406-40407 (–≤ –º–º)
        - –í—ã—á–∏—Å–ª—è–µ–º –∑–Ω–∞—á–µ–Ω–∏—è –¥–ª—è –ø—Ä–æ—Ç–æ–∫–æ–ª–∞ RIFTEK –ø–æ —Ñ–æ—Ä–º—É–ª–µ
        - –ó–∞–ø–∏—Å—ã–≤–∞–µ–º –ø–∞—Ä–∞–º–µ—Ç—Ä—ã 0Ch-0Fh –≤ –¥–∞—Ç—á–∏–∫ 3
        - –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø–∞—Ä–∞–º–µ—Ç—Ä—ã –≤ FLASH –ø–∞–º—è—Ç—å –¥–∞—Ç—á–∏–∫–∞
        """
        if not self.sensors or not self.sensors.ser:
            print(" [CMD=106] –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            self.write_cycle_flag(-1)
            self.current_state = SystemState.ERROR
            return
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∑–∞–ø—É—Å–∫–µ
            if not hasattr(self, 'configure_sensor3_range_started'):
                self.configure_sensor3_range_started = True
                print(f"üîß –ù–ê–ß–ê–õ–û –ù–ê–°–¢–†–û–ô–ö–ò –î–ò–ê–ü–ê–ó–û–ù–û–í –î–ê–¢–ß–ò–ö–ê 3 (CMD=106)")
                
                # –ß–∏—Ç–∞–µ–º –Ω–∞—á–∞–ª–æ –¥–∏–∞–ø–∞–∑–æ–Ω–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40404-40405
                range_start_mm = self.read_range_start()
                print(f" [CMD=106] –ù–∞—á–∞–ª–æ –¥–∏–∞–ø–∞–∑–æ–Ω–∞: {range_start_mm:.3f} –º–º")
                
                # –ß–∏—Ç–∞–µ–º –∫–æ–Ω–µ—Ü –¥–∏–∞–ø–∞–∑–æ–Ω–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40406-40407
                range_end_mm = self.read_range_end()
                print(f" [CMD=106] –ö–æ–Ω–µ—Ü –¥–∏–∞–ø–∞–∑–æ–Ω–∞: {range_end_mm:.3f} –º–º")
                
                # –í—ã—á–∏—Å–ª—è–µ–º –∑–Ω–∞—á–µ–Ω–∏—è –¥–ª—è –ø—Ä–æ—Ç–æ–∫–æ–ª–∞ RIFTEK
                # –§–æ—Ä–º—É–ª–∞: riftek_value = int((16384/25) * (mm_value - 25))
                riftek_value_min = int((16384 / 25) * (range_start_mm - 20))
                riftek_value_max = int((16384 / 25) * (range_end_mm - 20))
                
                # –û–≥—Ä–∞–Ω–∏—á–∏–≤–∞–µ–º –∑–Ω–∞—á–µ–Ω–∏—è –¥–∏–∞–ø–∞–∑–æ–Ω–æ–º 0-16383
                riftek_value_min = max(0, min(16383, riftek_value_min))
                riftek_value_max = max(0, min(16383, riftek_value_max))
                
                print(f" [CMD=106] –ó–Ω–∞—á–µ–Ω–∏–µ RIFTEK (–Ω–∞—á–∞–ª–æ): {riftek_value_min}")
                print(f" [CMD=106] –ó–Ω–∞—á–µ–Ω–∏–µ RIFTEK (–∫–æ–Ω–µ—Ü): {riftek_value_max}")
                
                # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º –ø–∞—Ä–∞–º–µ—Ç—Ä—ã –≤ –¥–∞—Ç—á–∏–∫ 3 (–∞–¥—Ä–µ—Å 0x03)
                sensor_address = 3
                
                # –ü–∞—Ä–∞–º–µ—Ç—Ä 0Ch - –º–ª–∞–¥—à–∏–π –±–∞–π—Ç –Ω–∞—á–∞–ª–∞ –æ–∫–Ω–∞
                if not self.write_riftek_parameter(sensor_address, 0x0C, riftek_value_min & 0xFF):
                    raise Exception("–û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ 0Ch")
                
                # –ü–∞—Ä–∞–º–µ—Ç—Ä 0Dh - —Å—Ç–∞—Ä—à–∏–π –±–∞–π—Ç –Ω–∞—á–∞–ª–∞ –æ–∫–Ω–∞
                if not self.write_riftek_parameter(sensor_address, 0x0D, (riftek_value_min >> 8) & 0xFF):
                    raise Exception("–û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ 0Dh")
                
                # –ü–∞—Ä–∞–º–µ—Ç—Ä 0Eh - –º–ª–∞–¥—à–∏–π –±–∞–π—Ç –∫–æ–Ω—Ü–∞ –æ–∫–Ω–∞
                if not self.write_riftek_parameter(sensor_address, 0x0E, riftek_value_max & 0xFF):
                    raise Exception("–û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ 0Eh")
                
                # –ü–∞—Ä–∞–º–µ—Ç—Ä 0Fh - —Å—Ç–∞—Ä—à–∏–π –±–∞–π—Ç –∫–æ–Ω—Ü–∞ –æ–∫–Ω–∞
                if not self.write_riftek_parameter(sensor_address, 0x0F, (riftek_value_max >> 8) & 0xFF):
                    raise Exception("–û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ 0Fh")
                
                print(" [CMD=106] –ü–∞—Ä–∞–º–µ—Ç—Ä—ã –∑–∞–ø–∏—Å–∞–Ω—ã –≤ –¥–∞—Ç—á–∏–∫ 3")
                
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø–∞—Ä–∞–º–µ—Ç—Ä—ã –≤ FLASH –ø–∞–º—è—Ç—å
                if not self.save_riftek_parameters_to_flash(sensor_address):
                    raise Exception("–û—à–∏–±–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ –≤ FLASH")
                
                print(" [CMD=106] –ü–∞—Ä–∞–º–µ—Ç—Ä—ã —Å–æ—Ö—Ä–∞–Ω–µ–Ω—ã –≤ FLASH –ø–∞–º—è—Ç—å –¥–∞—Ç—á–∏–∫–∞ 3")
                
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å —É—Å–ø–µ—à–Ω–æ–≥–æ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è
                self.write_cycle_flag(106)
                print(" [CMD=106] –ù–∞—Å—Ç—Ä–æ–π–∫–∞ –¥–∏–∞–ø–∞–∑–æ–Ω–æ–≤ –∑–∞–≤–µ—Ä—à–µ–Ω–∞ —É—Å–ø–µ—à–Ω–æ")
                
                # –ó–∞–¥–µ—Ä–∂–∫–∞ –ø–µ—Ä–µ–¥ —Å–±—Ä–æ—Å–æ–º –∫–æ–º–∞–Ω–¥—ã
                time.sleep(1)
                
                # –ü–µ—Ä–µ—Ö–æ–¥–∏–º –≤ IDLE –∏ —Å–±—Ä–∞—Å—ã–≤–∞–µ–º –∫–æ–º–∞–Ω–¥—É
                self.current_state = SystemState.IDLE
                self.reset_command()
                if hasattr(self, 'configure_sensor3_range_started'):
                    delattr(self, 'configure_sensor3_range_started')
                
        except Exception as e:
            print(f" [CMD=106] –û—à–∏–±–∫–∞ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ –¥–∏–∞–ø–∞–∑–æ–Ω–æ–≤: {e}")
            import traceback
            traceback.print_exc()
            self.write_cycle_flag(-1)
            
            # –ó–∞–¥–µ—Ä–∂–∫–∞ –ø–µ—Ä–µ–¥ —Å–±—Ä–æ—Å–æ–º –∫–æ–º–∞–Ω–¥—ã –ø—Ä–∏ –æ—à–∏–±–∫–µ
            time.sleep(1)
            
            self.reset_command()
            self.current_state = SystemState.ERROR
            if hasattr(self, 'configure_sensor3_range_started'):
                delattr(self, 'configure_sensor3_range_started')
    
    def read_range_start(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ –Ω–∞—á–∞–ª–∞ –¥–∏–∞–ø–∞–∑–æ–Ω–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40404-40405"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # 40404-40405 -> –∏–Ω–¥–µ–∫—Å—ã 403-404
                # HMI: —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ –≤ 40404, –º–ª–∞–¥—à–µ–µ –≤ 40405
                values = self.modbus_server.slave_context.getValues(3, 404, 2)
                if values and len(values) >= 2:
                    high_word = int(values[0])  # 40404 - —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ
                    low_word = int(values[1])  # 40405 - –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ
                    return self.doubleword_to_float(low_word, high_word)
        except Exception as e:
            print(f" [CMD=106] –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è –Ω–∞—á–∞–ª–∞ –¥–∏–∞–ø–∞–∑–æ–Ω–∞: {e}")
        return 0.0
    
    def read_range_end(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ –∫–æ–Ω—Ü–∞ –¥–∏–∞–ø–∞–∑–æ–Ω–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40406-40407"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # 40406-40407 -> –∏–Ω–¥–µ–∫—Å—ã 405-406
                # HMI: —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ –≤ 40406, –º–ª–∞–¥—à–µ–µ –≤ 40407
                values = self.modbus_server.slave_context.getValues(3, 406, 2)
                if values and len(values) >= 2:
                    high_word = int(values[0])  # 40406 - —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ
                    low_word = int(values[1])  # 40407 - –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ
                    return self.doubleword_to_float(low_word, high_word)
        except Exception as e:
            print(f" [CMD=106] –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è –∫–æ–Ω—Ü–∞ –¥–∏–∞–ø–∞–∑–æ–Ω–∞: {e}")
        return 0.0
    
    def write_riftek_parameter(self, sensor_address: int, param_code: int, param_value: int) -> bool:
        """
        –ó–∞–ø–∏—Å—å –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ –≤ –¥–∞—Ç—á–∏–∫ –ø–æ –ø—Ä–æ—Ç–æ–∫–æ–ª—É RIFTEK
        
        Args:
            sensor_address: –ê–¥—Ä–µ—Å –¥–∞—Ç—á–∏–∫–∞ (1-4)
            param_code: –ö–æ–¥ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ (0x00-0xFF)
            param_value: –ó–Ω–∞—á–µ–Ω–∏–µ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ (0-255 –¥–ª—è –æ–¥–Ω–æ–±–∞–π—Ç–Ω—ã—Ö –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤)
            
        Returns:
            True –µ—Å–ª–∏ –∫–æ–º–∞–Ω–¥–∞ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω–∞ —É—Å–ø–µ—à–Ω–æ
        """
        try:
            if not self.sensors or not self.sensors.ser:
                return False
            
            # –§–æ—Ä–º–∞—Ç –∫–æ–º–∞–Ω–¥—ã –∑–∞–ø–∏—Å–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ (03h):
            # –ë–∞–π—Ç 0: 0|ADR (–∞–¥—Ä–µ—Å –¥–∞—Ç—á–∏–∫–∞, —Å—Ç–∞—Ä—à–∏–π –±–∏—Ç = 0)
            # –ë–∞–π—Ç 1: 1|000|COD (–∫–æ–¥ –∑–∞–ø—Ä–æ—Å–∞ 03h = 0x83)
            # –ë–∞–π—Ç 2: 1|SB|CNT|MSG[0] lo (–º–ª–∞–¥—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∫–æ–¥–∞ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞)
            # –ë–∞–π—Ç 3: 1|SB|CNT|MSG[0] hi (—Å—Ç–∞—Ä—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∫–æ–¥–∞ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞)
            # –ë–∞–π—Ç 4: 1|SB|CNT|MSG[1] lo (–º–ª–∞–¥—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∑–Ω–∞—á–µ–Ω–∏—è –ø–∞—Ä–∞–º–µ—Ç—Ä–∞)
            # –ë–∞–π—Ç 5: 1|SB|CNT|MSG[1] hi (—Å—Ç–∞—Ä—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∑–Ω–∞—á–µ–Ω–∏—è –ø–∞—Ä–∞–º–µ—Ç—Ä–∞)
            
            # –ö–æ–¥ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ –ø–µ—Ä–µ–¥–∞–µ—Ç—Å—è –ø–æ—Ç–µ—Ç—Ä–∞–¥–Ω–æ
            param_code_lo = param_code & 0x0F
            param_code_hi = (param_code >> 4) & 0x0F
            
            # –ó–Ω–∞—á–µ–Ω–∏–µ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ –ø–µ—Ä–µ–¥–∞–µ—Ç—Å—è –ø–æ—Ç–µ—Ç—Ä–∞–¥–Ω–æ
            param_value_lo = param_value & 0x0F
            param_value_hi = (param_value >> 4) & 0x0F
            
            # –§–æ—Ä–º–∏—Ä—É–µ–º –∫–æ–º–∞–Ω–¥—É
            command = bytes([
                sensor_address,  # –ë–∞–π—Ç 0: –∞–¥—Ä–µ—Å –¥–∞—Ç—á–∏–∫–∞
                0x83,  # –ë–∞–π—Ç 1: –∫–æ–¥ –∑–∞–ø—Ä–æ—Å–∞ 03h (0x83 = 1|000|0011)
                0x80 | param_code_lo,  # –ë–∞–π—Ç 2: –º–ª–∞–¥—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∫–æ–¥–∞ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞
                0x80 | param_code_hi,  # –ë–∞–π—Ç 3: —Å—Ç–∞—Ä—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∫–æ–¥–∞ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞
                0x80 | param_value_lo,  # –ë–∞–π—Ç 4: –º–ª–∞–¥—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∑–Ω–∞—á–µ–Ω–∏—è
                0x80 | param_value_hi,  # –ë–∞–π—Ç 5: —Å—Ç–∞—Ä—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∑–Ω–∞—á–µ–Ω–∏—è
            ])
            
            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∫–æ–º–∞–Ω–¥—É
            self.sensors.ser.write(command)
            time.sleep(0.01)  # –ù–µ–±–æ–ª—å—à–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ –∫–æ–º–∞–Ω–¥—ã –¥–∞—Ç—á–∏–∫–æ–º
            
            return True
            
        except Exception as e:
            print(f" [CMD=106] –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞ {param_code:02X}h: {e}")
            return False
    
    def save_riftek_parameters_to_flash(self, sensor_address: int) -> bool:
        """
        –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ –≤ FLASH –ø–∞–º—è—Ç—å –¥–∞—Ç—á–∏–∫–∞ (–∫–æ–º–∞–Ω–¥–∞ 04h —Å –∫–æ–Ω—Å—Ç–∞–Ω—Ç–æ–π 0xAA)
        
        Args:
            sensor_address: –ê–¥—Ä–µ—Å –¥–∞—Ç—á–∏–∫–∞ (1-4)
            
        Returns:
            True –µ—Å–ª–∏ –∫–æ–º–∞–Ω–¥–∞ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω–∞ —É—Å–ø–µ—à–Ω–æ
        """
        try:
            if not self.sensors or not self.sensors.ser:
                return False
            
            # –§–æ—Ä–º–∞—Ç –∫–æ–º–∞–Ω–¥—ã —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –≤ FLASH (04h):
            # –ë–∞–π—Ç 0: 0|ADR (–∞–¥—Ä–µ—Å –¥–∞—Ç—á–∏–∫–∞)
            # –ë–∞–π—Ç 1: 1|000|COD (–∫–æ–¥ –∑–∞–ø—Ä–æ—Å–∞ 04h = 0x84)
            # –ë–∞–π—Ç 2: 1|SB|CNT|MSG[0] lo (–º–ª–∞–¥—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∫–æ–Ω—Å—Ç–∞–Ω—Ç—ã 0xAA = 0xA)
            # –ë–∞–π—Ç 3: 1|SB|CNT|MSG[0] hi (—Å—Ç–∞—Ä—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ –∫–æ–Ω—Å—Ç–∞–Ω—Ç—ã 0xAA = 0xA)
            
            command = bytes([
                sensor_address,  # –ë–∞–π—Ç 0: –∞–¥—Ä–µ—Å –¥–∞—Ç—á–∏–∫–∞
                0x84,  # –ë–∞–π—Ç 1: –∫–æ–¥ –∑–∞–ø—Ä–æ—Å–∞ 04h (0x84 = 1|000|0100)
                0x8A,  # –ë–∞–π—Ç 2: –º–ª–∞–¥—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ 0xAA (0xA = 0x80 | 0x0A)
                0x8A,  # –ë–∞–π—Ç 3: —Å—Ç–∞—Ä—à–∞—è —Ç–µ—Ç—Ä–∞–¥–∞ 0xAA (0xA = 0x80 | 0x0A)
            ])
            
            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∫–æ–º–∞–Ω–¥—É
            self.sensors.ser.write(command)
            time.sleep(0.1)  # –ó–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –≤ FLASH (–º–æ–∂–µ—Ç –∑–∞–Ω—è—Ç—å –≤—Ä–µ–º—è)
            
            return True
            
        except Exception as e:
            print(f" [CMD=106] –û—à–∏–±–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –≤ FLASH: {e}")
            return False

    def read_recipe_flange_diameter(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ä–µ—Ü–µ–ø—Ç–∞ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40388, 40389"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 388, 2)  # 40388-40389
                if values and len(values) >= 2:
                    high_word = int(values[0])  # 40388 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40389 - –º–ª–∞–¥—à–∏–π
                    diameter = self.doubleword_to_float(low_word, high_word)
                    return diameter
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–µ—Ü–µ–ø—Ç–∞ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞: {e}")
        return 0.0
    
    def read_recipe_body_diameter(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ä–µ—Ü–µ–ø—Ç–∞ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40382, 40383"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 382, 2)  # 40382-40383
                if values and len(values) >= 2:
                    high_word = int(values[0])  # 40382 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40383 - –º–ª–∞–¥—à–∏–π
                    diameter = self.doubleword_to_float(low_word, high_word)
                    return diameter
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–µ—Ü–µ–ø—Ç–∞ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞: {e}")
        return 0.0

    def read_upper_wall_offset_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —Å–º–µ—â–µ–Ω–∏—è —Ç–æ–ª—â–∏–Ω—ã –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ (40500-40501)"""
        return self._read_offset_coeff(500)

    def read_lower_wall_offset_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —Å–º–µ—â–µ–Ω–∏—è —Ç–æ–ª—â–∏–Ω—ã –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ (40502-40503)"""
        return self._read_offset_coeff(502)

    def read_body_diameter_offset_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —Å–º–µ—â–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ (40504-40505)"""
        return self._read_offset_coeff(504)

    def read_flange_diameter_offset_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —Å–º–µ—â–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞ (40506-40507)"""
        return self._read_offset_coeff(506)

    def read_bottom_thickness_offset_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —Å–º–µ—â–µ–Ω–∏—è —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞ (40508-40509)"""
        return self._read_offset_coeff(508)
    
    def read_upper_wall_extrapolation_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ —Ç–æ–ª—â–∏–Ω—ã –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ (40511-40512)"""
        return self._read_offset_coeff(510)  # 40511 - 40000 = 511
    
    def read_bottom_wall_extrapolation_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ —Ç–æ–ª—â–∏–Ω—ã –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ (40513-40514)"""
        return self._read_offset_coeff(512)  # 40513 - 40000 = 513
    
    def read_body_diameter_extrapolation_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ (40515-40516)"""
        return self._read_offset_coeff(514)  # 40515 - 40000 = 515
    
    def read_flange_diameter_extrapolation_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞ (40517-40518)"""
        return self._read_offset_coeff(516)  # 40517 - 40000 = 517
    
    def read_bottom_thickness_extrapolation_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞ (40519-40520)"""
        return self._read_offset_coeff(518)  # 40519 - 40000 = 519

    def read_body2_diameter_extrapolation_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 (40521-40522)"""
        return self._read_offset_coeff(520)

    def read_body2_diameter_offset_coeff(self) -> float:
        """–ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —Å–º–µ—â–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 (40522-40523)"""
        return self._read_offset_coeff(522)

    def _read_offset_coeff(self, base_index: int) -> float:
        """–û–±—â–∏–π –º–µ—Ç–æ–¥ —á—Ç–µ–Ω–∏—è –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç–∞ —Å–º–µ—â–µ–Ω–∏—è"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, base_index, 2)
                if values and len(values) >= 2:
                    high_word = int(values[0])
                    low_word = int(values[1])
                    return self.doubleword_to_float(low_word, high_word)
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç–∞ —Å–º–µ—â–µ–Ω–∏—è (base_index={base_index}): {e}")
        return 0.0
    
    def write_distance_sensor3_to_center(self, distance: float):
        """–ó–∞–ø–∏—Å—å —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–æ–º 3 –∏ —Ü–µ–Ω—Ç—Ä–æ–º –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40032, 40033"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                low_word, high_word = self.float_to_doubleword(distance)
                # HMI —á–∏—Ç–∞–µ—Ç: —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ –∏–∑ 40032, –º–ª–∞–¥—à–µ–µ –∏–∑ 40033
                self.modbus_server.slave_context.setValues(3, 32, [int(high_word)])  # 40032 - —Å—Ç–∞—Ä—à–∏–π (–∏–Ω–¥–µ–∫—Å 31)
                self.modbus_server.slave_context.setValues(3, 33, [int(low_word)])   # 40033 - –º–ª–∞–¥—à–∏–π (–∏–Ω–¥–µ–∫—Å 32)
                # –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –ë–î
                # –û—Ç–∫–ª—é—á–µ–Ω–æ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –ë–î - —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ —É–∫–∞–∑–∞–Ω–Ω—ã–µ –≤–µ–ª–∏—á–∏–Ω—ã
                # if self.db_integration:
                #     self.db_integration.save_doubleword_register(40032, 'holding', distance, '–†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫ 3 - —Ü–µ–Ω—Ç—Ä')
                print(f" [CMD=105] –ó–∞–ø–∏—Å–∞–Ω–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫ 3 - —Ü–µ–Ω—Ç—Ä 40032-40033: {distance:.3f} –º–º (high: {int(high_word)}, low: {int(low_word)})")
        except Exception as e:
            print(f" [CMD=105] –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫ 3 - —Ü–µ–Ω—Ç—Ä: {e}")

    def read_register_40020_raw(self):
        """–ß—Ç–µ–Ω–∏–µ —Å—ã—Ä—ã—Ö –∏ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã—Ö –¥–∞–Ω–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40052-40053"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –°—ã—Ä—ã–µ –¥–∞–Ω–Ω—ã–µ
                raw_values = self.modbus_server.slave_context.getValues(3, 52, 2)
                # –î–ª—è 32-bit integer –ù–ï –ø–µ—Ä–µ–≤–æ—Ä–∞—á–∏–≤–∞–µ–º –∫–∞–∫ float
                if raw_values and len(raw_values) == 2:
                    # –ü—Ä–∞–≤–∏–ª—å–Ω–æ–µ –æ–±—ä–µ–¥–∏–Ω–µ–Ω–∏–µ: –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ (–Ω–∏–∑–∫–∏–µ –±–∏—Ç—ã) + —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ (–≤—ã—Å–æ–∫–∏–µ –±–∏—Ç—ã)
                    steps = (int(raw_values[0]) << 16) | int(raw_values[1])
                    return f"—Å—ã—Ä—ã–µ [52-53]: {raw_values} ‚Üí 32-bit int: {steps}"
        except Exception as e:
            return f"–û—à–∏–±–∫–∞: {e}"

    def read_register_40021_raw(self):
        """–ß—Ç–µ–Ω–∏–µ —Å—ã—Ä—ã—Ö –∏ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã—Ö –¥–∞–Ω–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–∞ 40054"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –°—ã—Ä—ã–µ –¥–∞–Ω–Ω—ã–µ - —á–∏—Ç–∞–µ–º –ø—Ä–∞–≤–∏–ª—å–Ω—ã–π —Ä–µ–≥–∏—Å—Ç—Ä
                raw_value = self.modbus_server.slave_context.getValues(3, 54, 1)
                if raw_value:
                    pulses = int(raw_value[0])
                    return f"—Å—ã—Ä—ã–µ [54]: {raw_value} ‚Üí 16-bit int: {pulses}"
        except Exception as e:
            return f"–û—à–∏–±–∫–∞: {e}"

    def read_reference_height_raw(self):
        """–ß—Ç–µ–Ω–∏–µ —Å—ã—Ä—ã—Ö –∏ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã—Ö –¥–∞–Ω–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40008-40009"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –°—ã—Ä—ã–µ –¥–∞–Ω–Ω—ã–µ
                raw_values = self.modbus_server.slave_context.getValues(3, 8, 2)
                # –û–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ
                height = self.read_reference_height()
                return f"—Å—ã—Ä—ã–µ [7-8]: {raw_values} ‚Üí float –º–º: {height:.3f}"
        except Exception as e:
            return f"–û—à–∏–±–∫–∞: {e}"

    def read_reference_height(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —ç—Ç–∞–ª–æ–Ω–Ω–æ–π –≤—ã—Å–æ—Ç—ã –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40008, 40009"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # HMI: —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ –≤ 40008, –º–ª–∞–¥—à–µ–µ –≤ 40009
                values = self.modbus_server.slave_context.getValues(3, 8, 2)
                if values and len(values) >= 2:
                    high_word = int(values[0])  # 40008 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40009 - –º–ª–∞–¥—à–∏–π
                    height = self.doubleword_to_float(low_word, high_word)
                    return height
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —ç—Ç–∞–ª–æ–Ω–Ω–æ–π –≤—ã—Å–æ—Ç—ã: {e}")
        return 0.0

    def write_distance_to_plane(self, distance: float):
        """–ó–∞–ø–∏—Å—å –¥–∏—Å—Ç–∞–Ω—Ü–∏–∏ –¥–æ –Ω–∞—á–∞–ª–∞ –ø–ª–æ—Å–∫–æ—Å—Ç–∏ –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40055, 40056"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                low_word, high_word = self.float_to_doubleword(distance)
                # HMI —á–∏—Ç–∞–µ—Ç: —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ –∏–∑ 40055, –º–ª–∞–¥—à–µ–µ –∏–∑ 40056
                self.modbus_server.slave_context.setValues(3, 55, [int(high_word)])  # 40055 - —Å—Ç–∞—Ä—à–∏–π
                self.modbus_server.slave_context.setValues(3, 56, [int(low_word)])   # 40056 - –º–ª–∞–¥—à–∏–π
                # –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –ë–î
                # –û—Ç–∫–ª—é—á–µ–Ω–æ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –ë–î - —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ —É–∫–∞–∑–∞–Ω–Ω—ã–µ –≤–µ–ª–∏—á–∏–Ω—ã
                # if self.db_integration:
                #     self.db_integration.save_doubleword_register(40055, 'holding', distance, '–î–∏—Å—Ç–∞–Ω—Ü–∏—è –¥–æ –ø–ª–æ—Å–∫–æ—Å—Ç–∏')
                print(f" –ó–∞–ø–∏—Å–∞–Ω–∞ –¥–∏—Å—Ç–∞–Ω—Ü–∏—è –¥–æ –ø–ª–æ—Å–∫–æ—Å—Ç–∏ 40055-40056: {distance:.3f} –º–º (high: {int(high_word)}, low: {int(low_word)})")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ –¥–∏—Å—Ç–∞–Ω—Ü–∏–∏ –¥–æ –ø–ª–æ—Å–∫–æ—Å—Ç–∏: {e}")
    
    def handle_idle_state(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Å–æ—Å—Ç–æ—è–Ω–∏—è –æ–∂–∏–¥–∞–Ω–∏—è"""
        try:
            # –†–∞–∑ –≤ —Å–µ–∫—É–Ω–¥—É –º–æ–Ω–∏—Ç–æ—Ä–∏–º –ø–æ–¥–∫–ª—é—á–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–æ–≤ –∏ –≤—ã–ø–æ–ª–Ω—è–µ–º –∞–≤—Ç–æ-–ø–µ—Ä–µ–ø–æ–¥–∫–ª—é—á–µ–Ω–∏–µ
            current_time = time.time()
            if current_time - self.idle_monitor_last_time >= 1.0:
                self.idle_monitor_last_time = current_time
                if not self.test_mode:
                    self.check_and_reconnect_sensors()
                    # –û–¥–Ω–æ—Å—Ç—Ä–æ—á–Ω—ã–π —Å—Ç–∞—Ç—É—Å –±–µ–∑ —Å–ø–∞–º–∞
                    connected = self._is_sensor_connection_alive()
                    print(f" [IDLE] –ú–æ–Ω–∏—Ç–æ—Ä–∏–Ω–≥ –ø–æ–¥–∫–ª—é—á–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤: {'OK' if connected else '–ù–ï–¢'}")
                # –£–±–µ–¥–∏–º—Å—è, —á—Ç–æ —Å—Ç–∞—Ç—É—Å –æ–∂–∏–¥–∞–Ω–∏—è –∫–æ–º–∞–Ω–¥—ã –æ—Å—Ç–∞—ë—Ç—Å—è 30009=0
                if self.modbus_server and self.modbus_server.slave_context:
                    self.modbus_server.slave_context.setValues(4, 8, [0])  # 30009 -> index 8
        except Exception as e:
            print(f" [IDLE] –û—à–∏–±–∫–∞ –º–æ–Ω–∏—Ç–æ—Ä–∏–Ω–≥–∞: {e}")
    
    def handle_calibrate_wall_state(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Å—Ç–µ–Ω–∫–∏ (CMD = 100)"""
        if not self.sensors:
            self.calibration_data['wall_distance_1_2'] = 0.0
            self.calibration_data['wall_distance_1_3'] = 0.0
            self.finalize_calibration_failure(
                100,
                " [CMD=100] –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!",
                register_doublewords=[40010, 40012],
                cache_attrs=['cached_distance_1_2', 'cached_distance_1_3'],
            )
            return

        try:
            if not hasattr(self, '_wall_calibration_initialized'):
                print("üîß –ù–ê–ß–ê–õ–û –ö–ê–õ–ò–ë–†–û–í–ö–ò –°–¢–ï–ù–ö–ò")
                self.calibration_in_progress = True
                self._wall_calibration_initialized = True
                self._wall_reference_thickness = self.read_reference_thickness()
                print(f" –¢–æ–ª—â–∏–Ω–∞ —ç—Ç–∞–ª–æ–Ω–∞: {self._wall_reference_thickness:.3f} –º–º")
                if self._wall_reference_thickness <= 0:
                    self.calibration_data['wall_distance_1_2'] = 0.0
                    self.calibration_data['wall_distance_1_3'] = 0.0
                    self.finalize_calibration_failure(
                        100,
                        " [CMD=100] –û–®–ò–ë–ö–ê: —Ç–æ–ª—â–∏–Ω–∞ —ç—Ç–∞–ª–æ–Ω–∞ –¥–æ–ª–∂–Ω–∞ –±—ã—Ç—å –±–æ–ª—å—à–µ 0!",
                        cleanup_flags=['_wall_calibration_initialized', '_wall_reference_thickness'],
                        register_doublewords=[40010, 40012],
                        cache_attrs=['cached_distance_1_2', 'cached_distance_1_3'],
                    )
                    return

                self.clear_serial_buffers()
                self.flush_sensor_queue()
                self.measurement_buffer['sensor1'].clear()
                self.measurement_buffer['sensor2'].clear()
                self.measurement_buffer['sensor3'].clear()
                self.wall_sample_count = 0
                self.wall_last_log_time = time.time()
                self.write_cycle_flag(100)

            sensor1_mm, sensor2_mm, sensor3_mm, _ = self.read_sensors_safe()
            if sensor1_mm is None or sensor2_mm is None or sensor3_mm is None:
                time.sleep(0.002)
                return

            if (self.is_valid_measurement(sensor1_mm) and
                    self.is_valid_measurement(sensor2_mm) and
                    self.is_valid_measurement(sensor3_mm)):
                self.measurement_buffer['sensor1'].append(sensor1_mm)
                self.measurement_buffer['sensor2'].append(sensor2_mm)
                self.measurement_buffer['sensor3'].append(sensor3_mm)
                self.wall_sample_count += 1

                current_time = time.time()
                if current_time - getattr(self, 'wall_last_log_time', 0) >= 1.0:
                    print(f" [CMD=100] –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö... –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.wall_sample_count}")
                    self.wall_last_log_time = current_time
        except Exception as e:
            self.calibration_data['wall_distance_1_2'] = 0.0
            self.calibration_data['wall_distance_1_3'] = 0.0
            self.finalize_calibration_failure(
                100,
                f" –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['_wall_calibration_initialized', '_wall_reference_thickness', 'wall_sample_count', 'wall_last_log_time'],
                register_doublewords=[40010, 40012],
                cache_attrs=['cached_distance_1_2', 'cached_distance_1_3'],
            )
    
    def handle_calibrate_bottom_state(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–Ω–∞ (CMD = 101)"""
        cleanup_flags = [
            'calibrate_bottom_started',
            'calibrate_bottom_start_time',
            'calibrate_bottom_sensor4_buffer',
            'calibrate_bottom_measurement_count',
            '_reference_bottom_thickness',
            'wall_last_log_time',
        ]

        if not self.sensors:
            self.calibration_data['bottom_distance_4'] = 0.0
            self.finalize_calibration_failure(
                101,
                " [CMD=101] –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!",
                cleanup_flags,
                register_doublewords=[40014],
                cache_attrs=['cached_distance_sensor4'],
            )
            return

        try:
            if not hasattr(self, 'calibrate_bottom_started'):
                self.calibration_in_progress = True
                self.calibrate_bottom_started = True
                self.calibrate_bottom_sensor4_buffer = []
                self.calibrate_bottom_measurement_count = 0
                self.calibrate_bottom_last_log_time = time.time()

                self.clear_serial_buffers()
                self.flush_sensor_queue()
                self.measurement_buffer['sensor4'].clear()

                reference_bottom_thickness = self.read_reference_bottom_thickness()
                print("üîß –ù–ê–ß–ê–õ–û –ö–ê–õ–ò–ë–†–û–í–ö–ò –î–ù–ê")
                print(f" –≠—Ç–∞–ª–æ–Ω–Ω–∞—è —Ç–æ–ª—â–∏–Ω–∞ –¥–Ω–∞: {reference_bottom_thickness:.3f} –º–º")
                if reference_bottom_thickness <= 0:
                    self.calibration_data['bottom_distance_4'] = 0.0
                    self.finalize_calibration_failure(
                        101,
                        " [CMD=101] –û–®–ò–ë–ö–ê: —ç—Ç–∞–ª–æ–Ω–Ω–∞—è —Ç–æ–ª—â–∏–Ω–∞ –¥–Ω–∞ –¥–æ–ª–∂–Ω–∞ –±—ã—Ç—å –±–æ–ª—å—à–µ 0!",
                        cleanup_flags,
                        register_doublewords=[40014],
                        cache_attrs=['cached_distance_sensor4'],
                    )
                    return
                self._reference_bottom_thickness = reference_bottom_thickness

                self.write_cycle_flag(101)

            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor4_mm is None:
                time.sleep(0.002)
                return

            if self.is_valid_measurement(sensor4_mm):
                self.measurement_buffer['sensor4'].append(sensor4_mm)
                self.calibrate_bottom_sensor4_buffer.append(sensor4_mm)
                self.calibrate_bottom_measurement_count += 1

                current_time = time.time()
                if current_time - getattr(self, 'calibrate_bottom_last_log_time', 0) >= 1.0:
                    print(f" [CMD=101] –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö... –ò–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 4: {self.calibrate_bottom_measurement_count}")
                    self.calibrate_bottom_last_log_time = current_time

        except Exception as e:
            self.calibration_data['bottom_distance_4'] = 0.0
            self.finalize_calibration_failure(
                101,
                f" –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–Ω–∞: {e}",
                cleanup_flags,
                register_doublewords=[40014],
                cache_attrs=['cached_distance_sensor4'],
            )
    
    def handle_calibrate_flange_state(self):
        """
        CMD=102: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞
        - –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40006-40007
        - –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω–æ —Å–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 1 –¥–æ –ø–æ–ª—É—á–µ–Ω–∏—è CMD=0
        - –†–∞—Å—á–µ—Ç –∏ –∑–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –≤—ã–ø–æ–ª–Ω—è–µ—Ç—Å—è –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ CMD -> 0
        """
        if not self.sensors:
            self.calibration_data['flange_distance_1_center'] = 0.0
            self.finalize_calibration_failure(
                102,
                " [CMD=102] –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!",
                cleanup_flags=['calibrate_flange_started'],
                register_doublewords=[40016],
                cache_attrs=['cached_distance_to_center'],
            )
            return
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∑–∞–ø—É—Å–∫–µ
            if not hasattr(self, 'calibrate_flange_started'):
                self.calibration_in_progress = True
                self.calibrate_flange_started = True
                self.calibrate_flange_sensor1_buffer = []
                self.calibrate_flange_measurement_count = 0
                self.calibrate_flange_last_log_time = time.time()
                
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π
                self.clear_serial_buffers()
                self.flush_sensor_queue()
                
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä –¥–∞—Ç—á–∏–∫–∞ 1
                self.measurement_buffer['sensor1'].clear()
                
                # –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40006-40007
                reference_diameter = self.read_reference_diameter()
                print(f"üîß –ù–ê–ß–ê–õ–û –ö–ê–õ–ò–ë–†–û–í–ö–ò –≠–¢–ê–õ–û–ù–ù–û–ì–û –î–ò–ê–ú–ï–¢–†–ê (CMD=102)")
                print(f" [CMD=102] –≠—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä: {reference_diameter:.3f} –º–º")
                
                if reference_diameter <= 0:
                    self.calibration_data['flange_distance_1_center'] = 0.0
                    self.finalize_calibration_failure(
                        102,
                        " [CMD=102] –û–®–ò–ë–ö–ê: –≠—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –±–æ–ª—å—à–µ 0!",
                        cleanup_flags=['calibrate_flange_started'],
                        register_doublewords=[40016],
                        cache_attrs=['cached_distance_to_center'],
                    )
                    return
                
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏
                self.write_cycle_flag(102)
                print(f" [CMD=102] –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω—ã–π —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 1...")
            
            # –ù–µ–ø—Ä–µ—Ä—ã–≤–Ω—ã–π —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö (–±–µ–∑ —Ç–∞–π–º–∞—É—Ç–∞)
            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor1_mm is None:
                time.sleep(0.002)
                return

            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –≤–∞–ª–∏–¥–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 1
            if self.is_valid_measurement(sensor1_mm):
                self.measurement_buffer['sensor1'].append(sensor1_mm)
                self.calibrate_flange_sensor1_buffer.append(sensor1_mm)
                self.calibrate_flange_measurement_count += 1

            # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø—Ä–æ–≥—Ä–µ—Å—Å —Ä–∞–∑ –≤ —Å–µ–∫—É–Ω–¥—É
            current_time = time.time()
            if current_time - getattr(self, 'calibrate_flange_last_log_time', 0) >= 1.0:
                print(f" [CMD=102] –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö... –ò–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1: {self.calibrate_flange_measurement_count}")
                self.calibrate_flange_last_log_time = current_time
                
        except Exception as e:
            print(f" [CMD=102] –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞: {e}")
            import traceback
            traceback.print_exc()
            self.calibration_data['flange_distance_1_center'] = 0.0
            self.finalize_calibration_failure(
                102,
                f" [CMD=102] –û—à–∏–±–∫–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_flange_started'],
                register_doublewords=[40016],
                cache_attrs=['cached_distance_to_center'],
            )
    
    def _finish_calibration_wall(self):
        """–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Å—Ç–µ–Ω–∫–∏ (CMD=100): —Ä–∞—Å—á–µ—Ç –∏ –∑–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤"""
        try:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∞ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞
            if not hasattr(self, '_wall_calibration_initialized'):
                print(" [CMD=100] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!")
                self.finalize_calibration_failure(
                    100,
                    " [CMD=100] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!",
                    cleanup_flags=['_wall_calibration_initialized', '_wall_reference_thickness'],
                    register_doublewords=[40010, 40012],
                    cache_attrs=['cached_distance_1_2', 'cached_distance_1_3'],
                )
                return
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞–ª–∏—á–∏–µ –¥–∞–Ω–Ω—ã—Ö
            if (len(self.measurement_buffer['sensor1']) == 0 or 
                len(self.measurement_buffer['sensor2']) == 0 or 
                len(self.measurement_buffer['sensor3']) == 0):
                print(f" [CMD=100] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!")
                print(f" [CMD=100] –î–∞–Ω–Ω—ã–µ: –î1={len(self.measurement_buffer['sensor1'])}, –î2={len(self.measurement_buffer['sensor2'])}, –î3={len(self.measurement_buffer['sensor3'])}")
                self.finalize_calibration_failure(
                    100,
                    " [CMD=100] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!",
                    cleanup_flags=['_wall_calibration_initialized', '_wall_reference_thickness'],
                    register_doublewords=[40010, 40012],
                    cache_attrs=['cached_distance_1_2', 'cached_distance_1_3'],
                )
                return
            
            # –£—Å—Ä–µ–¥–Ω—è–µ–º –∏–∑–º–µ—Ä–µ–Ω–∏—è
            try:
                avg_sensor1, avg_sensor2, avg_sensor3 = self.calculate_averages()
            except ValueError as e:
                print(f" [CMD=100] –û–®–ò–ë–ö–ê –ø—Ä–∏ —É—Å—Ä–µ–¥–Ω–µ–Ω–∏–∏: {e}")
                self.finalize_calibration_failure(
                    100,
                    f" [CMD=100] –û–®–ò–ë–ö–ê –ø—Ä–∏ —É—Å—Ä–µ–¥–Ω–µ–Ω–∏–∏: {e}",
                    cleanup_flags=['_wall_calibration_initialized', '_wall_reference_thickness'],
                    register_doublewords=[40010, 40012],
                    cache_attrs=['cached_distance_1_2', 'cached_distance_1_3'],
                )
                return
            
            print(f" [CMD=100] –°—Ä–µ–¥–Ω–∏–µ –∑–Ω–∞—á–µ–Ω–∏—è: –î1={avg_sensor1:.3f}–º–º, –î2={avg_sensor2:.3f}–º–º, –î3={avg_sensor3:.3f}–º–º")
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,2
            reference_thickness = getattr(self, '_wall_reference_thickness', 0.0)
            distance_1_2 = avg_sensor1 + avg_sensor2 + reference_thickness
            print(f" [CMD=100] –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,2: {distance_1_2:.3f} –º–º")
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,3
            distance_1_3 = avg_sensor1 - avg_sensor3
            print(f" [CMD=100] –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,3: {distance_1_3:.3f} –º–º")
            
            # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã
            self.write_calibration_result_1_2(distance_1_2)
            self.write_calibration_result_1_3(distance_1_3)
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º –≤ –ª–æ–∫–∞–ª—å–Ω—ã—Ö –¥–∞–Ω–Ω—ã—Ö
            self.calibration_data['wall_distance_1_2'] = distance_1_2
            self.calibration_data['wall_distance_1_3'] = distance_1_3
            
            # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã –ø–æ—Å–ª–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            
            # –û—á–∏—â–∞–µ–º —Ñ–ª–∞–≥–∏
            if hasattr(self, '_wall_calibration_initialized'):
                delattr(self, '_wall_calibration_initialized')
            if hasattr(self, '_wall_reference_thickness'):
                delattr(self, '_wall_reference_thickness')
            if hasattr(self, 'wall_sample_count'):
                delattr(self, 'wall_sample_count')
            if hasattr(self, 'wall_last_log_time'):
                delattr(self, 'wall_last_log_time')
            
            self.calibration_in_progress = False
            self.write_cycle_flag(0)
            print(" [CMD=100] –ö–ê–õ–ò–ë–†–û–í–ö–ê –°–¢–ï–ù–ö–ò –ó–ê–í–ï–†–®–ï–ù–ê –£–°–ü–ï–®–ù–û")
            
        except Exception as e:
            print(f" [CMD=100] –û–®–ò–ë–ö–ê –ø—Ä–∏ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–∏ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}")
            import traceback
            traceback.print_exc()
            self.calibration_data['wall_distance_1_2'] = 0.0
            self.calibration_data['wall_distance_1_3'] = 0.0
            self.finalize_calibration_failure(
                100,
                f" [CMD=100] –û—à–∏–±–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['_wall_calibration_initialized', '_wall_reference_thickness'],
                register_doublewords=[40010, 40012],
                cache_attrs=['cached_distance_1_2', 'cached_distance_1_3'],
            )
    
    def _finish_calibration_bottom(self):
        """–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–Ω–∞ (CMD=101): —Ä–∞—Å—á–µ—Ç –∏ –∑–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤"""
        try:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∞ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞
            if not hasattr(self, 'calibrate_bottom_started'):
                print(" [CMD=101] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!")
                self.finalize_calibration_failure(
                    101,
                    " [CMD=101] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!",
                    cleanup_flags=['calibrate_bottom_started', '_reference_bottom_thickness'],
                    register_doublewords=[40014],
                    cache_attrs=['cached_distance_sensor4'],
                )
                return
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞–ª–∏—á–∏–µ –¥–∞–Ω–Ω—ã—Ö
            if not hasattr(self, 'calibrate_bottom_sensor4_buffer') or len(self.calibrate_bottom_sensor4_buffer) == 0:
                print(f" [CMD=101] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!")
                self.finalize_calibration_failure(
                    101,
                    " [CMD=101] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!",
                    cleanup_flags=['calibrate_bottom_started', '_reference_bottom_thickness'],
                    register_doublewords=[40014],
                    cache_attrs=['cached_distance_sensor4'],
                )
                return
            
            # –£—Å—Ä–µ–¥–Ω—è–µ–º –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4
            avg_sensor4 = sum(self.calibrate_bottom_sensor4_buffer) / len(self.calibrate_bottom_sensor4_buffer)
            print(f" [CMD=101] –°—Ä–µ–¥–Ω–µ–µ –∑–Ω–∞—á–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 4: {avg_sensor4:.3f} –º–º (–∏–∑ {len(self.calibrate_bottom_sensor4_buffer)} –∏–∑–º–µ—Ä–µ–Ω–∏–π)")
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 4 –¥–æ –ø–æ–≤–µ—Ä—Ö–Ω–æ—Å—Ç–∏
            reference_bottom_thickness = getattr(self, '_reference_bottom_thickness', 0.0)
            distance_4_surface = avg_sensor4 + reference_bottom_thickness
            print(f" [CMD=101] –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –æ—Ç –¥–∞—Ç—á–∏–∫–∞ 4 –¥–æ –ø–æ–≤–µ—Ä—Ö–Ω–æ—Å—Ç–∏: {distance_4_surface:.3f} –º–º")
            
            # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç
            self.write_calibration_result_4_surface(distance_4_surface)
            self.calibration_data['bottom_distance_4'] = distance_4_surface
            self.cached_distance_sensor4 = distance_4_surface
            
            # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            
            # –û—á–∏—â–∞–µ–º —Ñ–ª–∞–≥–∏
            if hasattr(self, 'calibrate_bottom_started'):
                delattr(self, 'calibrate_bottom_started')
            if hasattr(self, '_reference_bottom_thickness'):
                delattr(self, '_reference_bottom_thickness')
            if hasattr(self, 'calibrate_bottom_sensor4_buffer'):
                delattr(self, 'calibrate_bottom_sensor4_buffer')
            if hasattr(self, 'calibrate_bottom_measurement_count'):
                delattr(self, 'calibrate_bottom_measurement_count')
            if hasattr(self, 'calibrate_bottom_last_log_time'):
                delattr(self, 'calibrate_bottom_last_log_time')
            
            self.calibration_in_progress = False
            self.write_cycle_flag(0)
            print(" [CMD=101] –ö–ê–õ–ò–ë–†–û–í–ö–ê –î–ù–ê –ó–ê–í–ï–†–®–ï–ù–ê –£–°–ü–ï–®–ù–û")
            
        except Exception as e:
            print(f" [CMD=101] –û–®–ò–ë–ö–ê –ø—Ä–∏ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–∏ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}")
            import traceback
            traceback.print_exc()
            self.calibration_data['bottom_distance_4'] = 0.0
            self.finalize_calibration_failure(
                101,
                f" [CMD=101] –û—à–∏–±–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_bottom_started', '_reference_bottom_thickness'],
                register_doublewords=[40014],
                cache_attrs=['cached_distance_sensor4'],
            )
    
    def _finish_calibration_flange(self):
        """–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Ñ–ª–∞–Ω—Ü–∞ (CMD=102): —Ä–∞—Å—á–µ—Ç –∏ –∑–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤"""
        try:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∞ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞
            if not hasattr(self, 'calibrate_flange_started'):
                print(" [CMD=102] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!")
                self.finalize_calibration_failure(
                    102,
                    " [CMD=102] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!",
                    cleanup_flags=['calibrate_flange_started'],
                    register_doublewords=[40016],
                    cache_attrs=['cached_distance_to_center'],
                )
                return
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞–ª–∏—á–∏–µ –¥–∞–Ω–Ω—ã—Ö
            if not hasattr(self, 'calibrate_flange_sensor1_buffer') or len(self.calibrate_flange_sensor1_buffer) == 0:
                print(f" [CMD=102] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!")
                self.finalize_calibration_failure(
                    102,
                    " [CMD=102] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!",
                    cleanup_flags=['calibrate_flange_started'],
                    register_doublewords=[40016],
                    cache_attrs=['cached_distance_to_center'],
                )
                return
            
            # –£—Å—Ä–µ–¥–Ω—è–µ–º –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 1
            avg_sensor1 = sum(self.calibrate_flange_sensor1_buffer) / len(self.calibrate_flange_sensor1_buffer)
            print(f" [CMD=102] –°—Ä–µ–¥–Ω–µ–µ –∑–Ω–∞—á–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 1: {avg_sensor1:.3f} –º–º (–∏–∑ {len(self.calibrate_flange_sensor1_buffer)} –∏–∑–º–µ—Ä–µ–Ω–∏–π)")
            
            # –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä
            reference_diameter = self.read_reference_diameter()
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–æ–º 1 –∏ —Ü–µ–Ω—Ç—Ä–æ–º –ø–µ—Ä–µ—Å–µ—á–µ–Ω–∏—è
            distance_1_center = (reference_diameter / 2) + avg_sensor1
            print(f" [CMD=102] –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–æ–º 1 –∏ —Ü–µ–Ω—Ç—Ä–æ–º: {distance_1_center:.3f} –º–º")
            print(f" [CMD=102] –§–æ—Ä–º—É–ª–∞: ({reference_diameter:.3f} / 2) + {avg_sensor1:.3f} = {distance_1_center:.3f}")
            
            # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç
            self.write_calibration_result_1_center(distance_1_center)
            self.cached_distance_to_center = distance_1_center
            self.calibration_data['flange_distance_1_center'] = distance_1_center
            
            # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            
            # –û—á–∏—â–∞–µ–º —Ñ–ª–∞–≥–∏
            if hasattr(self, 'calibrate_flange_started'):
                delattr(self, 'calibrate_flange_started')
            if hasattr(self, 'calibrate_flange_sensor1_buffer'):
                delattr(self, 'calibrate_flange_sensor1_buffer')
            if hasattr(self, 'calibrate_flange_measurement_count'):
                delattr(self, 'calibrate_flange_measurement_count')
            if hasattr(self, 'calibrate_flange_start_time'):
                delattr(self, 'calibrate_flange_start_time')
            if hasattr(self, 'calibrate_flange_measurement_duration'):
                delattr(self, 'calibrate_flange_measurement_duration')
            if hasattr(self, 'calibrate_flange_completed'):
                delattr(self, 'calibrate_flange_completed')
            
            self.calibration_in_progress = False
            self.write_cycle_flag(0)
            print(" [CMD=102] –ö–ê–õ–ò–ë–†–û–í–ö–ê –§–õ–ê–ù–¶–ê –ó–ê–í–ï–†–®–ï–ù–ê –£–°–ü–ï–®–ù–û")
            
        except Exception as e:
            print(f" [CMD=102] –û–®–ò–ë–ö–ê –ø—Ä–∏ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–∏ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}")
            import traceback
            traceback.print_exc()
            self.calibration_data['flange_distance_1_center'] = 0.0
            self.finalize_calibration_failure(
                102,
                f" [CMD=102] –û—à–∏–±–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_flange_started'],
                register_doublewords=[40016],
                cache_attrs=['cached_distance_to_center'],
            )
    
    def _finish_calibration_flange_diameter(self):
        """–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞ (CMD=105): —Ä–∞—Å—á–µ—Ç –∏ –∑–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤"""
        try:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∞ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞
            if not hasattr(self, 'calibrate_flange_diameter_started'):
                print(" [CMD=105] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!")
                self.finalize_calibration_failure(
                    105,
                    " [CMD=105] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!",
                    cleanup_flags=['calibrate_flange_diameter_started'],
                    register_doublewords=[40032],
                    cache_attrs=['cached_distance_sensor3_to_center'],
                )
                return
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞–ª–∏—á–∏–µ –¥–∞–Ω–Ω—ã—Ö
            if not hasattr(self, 'calibrate_flange_diameter_sensor3_buffer') or len(self.calibrate_flange_diameter_sensor3_buffer) == 0:
                print(f" [CMD=105] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!")
                self.finalize_calibration_failure(
                    105,
                    " [CMD=105] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!",
                    cleanup_flags=['calibrate_flange_diameter_started'],
                    register_doublewords=[40032],
                    cache_attrs=['cached_distance_sensor3_to_center'],
                )
                return
            
            # –£—Å—Ä–µ–¥–Ω—è–µ–º –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 3
            avg_sensor3 = sum(self.calibrate_flange_diameter_sensor3_buffer) / len(self.calibrate_flange_diameter_sensor3_buffer)
            print(f" [CMD=105] –°—Ä–µ–¥–Ω–µ–µ –∑–Ω–∞—á–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 3: {avg_sensor3:.3f} –º–º (–∏–∑ {len(self.calibrate_flange_diameter_sensor3_buffer)} –∏–∑–º–µ—Ä–µ–Ω–∏–π)")
            
            # –ß–∏—Ç–∞–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞
            reference_flange_diameter = self.read_reference_flange_diameter()
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–æ–º 3 –∏ —Ü–µ–Ω—Ç—Ä–æ–º –ø–µ—Ä–µ—Å–µ—á–µ–Ω–∏—è
            distance_sensor3_to_center = (reference_flange_diameter / 2) + avg_sensor3
            print(f" [CMD=105] –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–æ–º 3 –∏ —Ü–µ–Ω—Ç—Ä–æ–º: {distance_sensor3_to_center:.3f} –º–º")
            print(f" [CMD=105] –§–æ—Ä–º—É–ª–∞: ({reference_flange_diameter:.3f} / 2) + {avg_sensor3:.3f} = {distance_sensor3_to_center:.3f}")
            
            # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç
            self.write_distance_sensor3_to_center(distance_sensor3_to_center)
            self.cached_distance_sensor3_to_center = distance_sensor3_to_center
            
            # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            
            # –û—á–∏—â–∞–µ–º —Ñ–ª–∞–≥–∏
            if hasattr(self, 'calibrate_flange_diameter_started'):
                delattr(self, 'calibrate_flange_diameter_started')
            if hasattr(self, 'calibrate_flange_diameter_sensor3_buffer'):
                delattr(self, 'calibrate_flange_diameter_sensor3_buffer')
            if hasattr(self, 'calibrate_flange_diameter_measurement_count'):
                delattr(self, 'calibrate_flange_diameter_measurement_count')
            if hasattr(self, 'calibrate_flange_diameter_start_time'):
                delattr(self, 'calibrate_flange_diameter_start_time')
            if hasattr(self, 'calibrate_flange_diameter_measurement_duration'):
                delattr(self, 'calibrate_flange_diameter_measurement_duration')
            if hasattr(self, 'calibrate_flange_diameter_completed'):
                delattr(self, 'calibrate_flange_diameter_completed')
            
            self.calibration_in_progress = False
            self.write_cycle_flag(0)
            print(" [CMD=105] –ö–ê–õ–ò–ë–†–û–í–ö–ê –î–ò–ê–ú–ï–¢–†–ê –§–õ–ê–ù–¶–ê –ó–ê–í–ï–†–®–ï–ù–ê –£–°–ü–ï–®–ù–û")
            
        except Exception as e:
            print(f" [CMD=105] –û–®–ò–ë–ö–ê –ø—Ä–∏ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–∏ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}")
            import traceback
            traceback.print_exc()
            self.finalize_calibration_failure(
                105,
                f" [CMD=105] –û—à–∏–±–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_flange_diameter_started'],
                register_doublewords=[40032],
                cache_attrs=['cached_distance_sensor3_to_center'],
            )

    def _finish_calibration_body_diameter_separate(self):
        """–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ (CMD=107)"""
        try:
            if not hasattr(self, 'calibrate_body_diameter_separate_started'):
                self.finalize_calibration_failure(
                    107,
                    " [CMD=107] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!",
                    cleanup_flags=['calibrate_body_diameter_separate_started'],
                    register_doublewords=[40038],
                    cache_attrs=['cached_distance_sensor3_to_center_body'],
                )
                return

            if (not hasattr(self, 'calibrate_body_diameter_separate_sensor3_buffer') or
                len(self.calibrate_body_diameter_separate_sensor3_buffer) == 0):
                self.finalize_calibration_failure(
                    107,
                    " [CMD=107] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!",
                    cleanup_flags=['calibrate_body_diameter_separate_started'],
                    register_doublewords=[40038],
                    cache_attrs=['cached_distance_sensor3_to_center_body'],
                )
                return

            avg_sensor3 = sum(self.calibrate_body_diameter_separate_sensor3_buffer) / len(self.calibrate_body_diameter_separate_sensor3_buffer)
            reference_diameter = self.read_reference_body_diameter_separate()
            distance_sensor3_to_center = (reference_diameter / 2) + avg_sensor3

            print(f" [CMD=107] –°—Ä–µ–¥–Ω–µ–µ –¥–∞—Ç—á–∏–∫–∞ 3: {avg_sensor3:.3f} –º–º")
            print(f" [CMD=107] –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫ 3 - —Ü–µ–Ω—Ç—Ä: {distance_sensor3_to_center:.3f} –º–º")

            self.write_distance_sensor3_to_center_body(distance_sensor3_to_center)
            self.cached_distance_sensor3_to_center_body = distance_sensor3_to_center

            self.clear_measurement_buffers()
            self.clear_serial_buffers()

            for attr in [
                'calibrate_body_diameter_separate_started',
                'calibrate_body_diameter_separate_sensor3_buffer',
                'calibrate_body_diameter_separate_measurement_count',
                'calibrate_body_diameter_separate_last_log_time',
            ]:
                if hasattr(self, attr):
                    delattr(self, attr)

            self.calibration_in_progress = False
            self.write_cycle_flag(0)
            print(" [CMD=107] –ö–ê–õ–ò–ë–†–û–í–ö–ê –†–ê–ó–î–ï–õ–¨–ù–û–ì–û –î–ò–ê–ú–ï–¢–†–ê –ö–û–†–ü–£–°–ê –ó–ê–í–ï–†–®–ï–ù–ê –£–°–ü–ï–®–ù–û")

        except Exception as e:
            self.finalize_calibration_failure(
                107,
                f" [CMD=107] –û—à–∏–±–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_body_diameter_separate_started'],
                register_doublewords=[40038],
                cache_attrs=['cached_distance_sensor3_to_center_body'],
            )

    def _finish_calibration_body2_diameter(self):
        """–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 (CMD=108)"""
        try:
            if not hasattr(self, 'calibrate_body2_diameter_started'):
                self.finalize_calibration_failure(
                    108,
                    " [CMD=108] –û–®–ò–ë–ö–ê: –ö–∞–ª–∏–±—Ä–æ–≤–∫–∞ –Ω–µ –±—ã–ª–∞ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞!",
                    cleanup_flags=['calibrate_body2_diameter_started'],
                    register_doublewords=[40040],
                    cache_attrs=['cached_distance_sensor3_to_center_body2'],
                )
                return

            if (not hasattr(self, 'calibrate_body2_diameter_sensor3_buffer') or
                len(self.calibrate_body2_diameter_sensor3_buffer) == 0):
                self.finalize_calibration_failure(
                    108,
                    " [CMD=108] –û–®–ò–ë–ö–ê: –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞!",
                    cleanup_flags=['calibrate_body2_diameter_started'],
                    register_doublewords=[40040],
                    cache_attrs=['cached_distance_sensor3_to_center_body2'],
                )
                return

            avg_sensor3 = sum(self.calibrate_body2_diameter_sensor3_buffer) / len(self.calibrate_body2_diameter_sensor3_buffer)
            reference_diameter = self.read_reference_body2_diameter()
            distance_sensor3_to_center = (reference_diameter / 2) + avg_sensor3

            print(f" [CMD=108] –°—Ä–µ–¥–Ω–µ–µ –¥–∞—Ç—á–∏–∫–∞ 3: {avg_sensor3:.3f} –º–º")
            print(f" [CMD=108] –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫ 3 - —Ü–µ–Ω—Ç—Ä: {distance_sensor3_to_center:.3f} –º–º")

            self.write_distance_sensor3_to_center_body2(distance_sensor3_to_center)
            self.cached_distance_sensor3_to_center_body2 = distance_sensor3_to_center

            self.clear_measurement_buffers()
            self.clear_serial_buffers()

            for attr in [
                'calibrate_body2_diameter_started',
                'calibrate_body2_diameter_sensor3_buffer',
                'calibrate_body2_diameter_measurement_count',
                'calibrate_body2_diameter_last_log_time',
            ]:
                if hasattr(self, attr):
                    delattr(self, attr)

            self.calibration_in_progress = False
            self.write_cycle_flag(0)
            print(" [CMD=108] –ö–ê–õ–ò–ë–†–û–í–ö–ê –î–ò–ê–ú–ï–¢–†–ê –ö–û–†–ü–£–°–ê 2 –ó–ê–í–ï–†–®–ï–ù–ê –£–°–ü–ï–®–ù–û")

        except Exception as e:
            self.finalize_calibration_failure(
                108,
                f" [CMD=108] –û—à–∏–±–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏: {e}",
                cleanup_flags=['calibrate_body2_diameter_started'],
                register_doublewords=[40040],
                cache_attrs=['cached_distance_sensor3_to_center_body2'],
            )
    
    def read_reference_thickness(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ç–æ–ª—â–∏–Ω—ã —ç—Ç–∞–ª–æ–Ω–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40002, 40003"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ß–∏—Ç–∞–µ–º –¥–≤–∞ —Ä–µ–≥–∏—Å—Ç—Ä–∞ (40002, 40003) - HMI –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç –≤ –æ–±—Ä–∞—Ç–Ω–æ–º –ø–æ—Ä—è–¥–∫–µ
                values = self.modbus_server.slave_context.getValues(3, 2, 2)
                if values and len(values) >= 2:
                    # HMI –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç: —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –≤ 40002, –º–ª–∞–¥—à–∏–π –≤ 40003
                    high_word = int(values[0])  # 40002 - —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –æ—Ç HMI
                    low_word = int(values[1])   # 40003 - –º–ª–∞–¥—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –æ—Ç HMI
                    
                    # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º –∏–∑ –¥–≤—É—Ö 16-–±–∏—Ç–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ –≤ float
                    thickness_float = self.doubleword_to_float(low_word, high_word)
                    print(f" –ü—Ä–æ—á–∏—Ç–∞–Ω–∞ —Ç–æ–ª—â–∏–Ω–∞ —ç—Ç–∞–ª–æ–Ω–∞: {high_word}, {low_word} -> {thickness_float:.3f} –º–º")
                    return thickness_float
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —ç—Ç–∞–ª–æ–Ω–∞: {e}")
        return 0.0
    
    def measure_sensors_for_calibration(self):
        """–ò–∑–º–µ—Ä–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–æ–≤ –¥–ª—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –≤ —Ç–µ—á–µ–Ω–∏–µ 4 —Å–µ–∫—É–Ω–¥"""
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            return
        
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π
        self.clear_serial_buffers()
        self.flush_sensor_queue()
        
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π
        self.measurement_buffer['sensor1'].clear()
        self.measurement_buffer['sensor2'].clear()
        self.measurement_buffer['sensor3'].clear()
        self.measurement_buffer['sensor4'].clear()
            
        start_time = time.time()
        measurement_count = 0
        
        print(" –ù–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏–π...")
        
        while (time.time() - start_time) < self.measurement_duration:
            try:
                # –í—ã–ø–æ–ª–Ω—è–µ–º —á—Ç–µ–Ω–∏–µ –∏–∑–º–µ—Ä–µ–Ω–∏–π —Å –∑–∞—â–∏—Ç–æ–π –æ—Ç –∫–æ–Ω—Ñ–ª–∏–∫—Ç–æ–≤
                sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
                if sensor1_mm is None or sensor2_mm is None or sensor3_mm is None:
                    time.sleep(0.001)
                    continue
                
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –≤–∞–ª–∏–¥–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–ª—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Å—Ç–µ–Ω–∫–∏
                # –î–ª—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Å—Ç–µ–Ω–∫–∏ –Ω—É–∂–Ω—ã —Ç–æ–ª—å–∫–æ –¥–∞—Ç—á–∏–∫–∏ 1, 2 –∏ 3 (–¥–∞—Ç—á–∏–∫ 4 –Ω–µ –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è)
                # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è (None, 0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞)
                if (self.is_valid_measurement(sensor1_mm) and 
                    self.is_valid_measurement(sensor2_mm) and 
                    self.is_valid_measurement(sensor3_mm)):
                    # –î–∞—Ç—á–∏–∫ 4 –Ω–µ –æ–±—è–∑–∞—Ç–µ–ª–µ–Ω –¥–ª—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Å—Ç–µ–Ω–∫–∏, –Ω–æ —Å–æ—Ö—Ä–∞–Ω—è–µ–º –µ—Å–ª–∏ –≤–∞–ª–∏–¥–µ–Ω
                    self.measurement_buffer['sensor1'].append(sensor1_mm)
                    self.measurement_buffer['sensor2'].append(sensor2_mm)
                    self.measurement_buffer['sensor3'].append(sensor3_mm)
                    if self.is_valid_measurement(sensor4_mm):
                        self.measurement_buffer['sensor4'].append(sensor4_mm)
                    measurement_count += 1
                else:
                    # –û—Ç–ª–∞–¥–æ—á–Ω—ã–π –≤—ã–≤–æ–¥: –ø–æ–∫–∞–∑—ã–≤–∞–µ–º, –ø–æ—á–µ–º—É –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–µ –ø—Ä–æ—à–ª–∏ –≤–∞–ª–∏–¥–∞—Ü–∏—é
                    if measurement_count == 0 and int((time.time() - start_time)) % 2 == 0:  # –†–∞–∑ –≤ 2 —Å–µ–∫—É–Ω–¥—ã
                        invalid_reasons = []
                        if not self.is_valid_measurement(sensor1_mm):
                            invalid_reasons.append(f"–î1={sensor1_mm}")
                        if not self.is_valid_measurement(sensor2_mm):
                            invalid_reasons.append(f"–î2={sensor2_mm}")
                        if not self.is_valid_measurement(sensor3_mm):
                            invalid_reasons.append(f"–î3={sensor3_mm}")
                        if invalid_reasons:
                            print(f" ‚ö† –ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è: {', '.join(invalid_reasons)}")
                
                # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø—Ä–æ–≥—Ä–µ—Å—Å –∫–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
                elapsed = time.time() - start_time
                if int(elapsed) != int(elapsed - 0.1):  # –ö–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
                    print(f" –í—Ä–µ–º—è: {elapsed:.1f}—Å, –ò–∑–º–µ—Ä–µ–Ω–∏–π: {measurement_count}")
                    
            except Exception as e:
                print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è: {e}")
                # –£–±—Ä–∞–Ω sleep –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è
        
        print(f" –ò–∑–º–µ—Ä–µ–Ω–∏—è –∑–∞–≤–µ—Ä—à–µ–Ω—ã. –í—Å–µ–≥–æ: {measurement_count}")
    
    def calculate_averages(self) -> Tuple[float, float, float]:
        """–í—ã—á–∏—Å–ª–µ–Ω–∏–µ —Å—Ä–µ–¥–Ω–∏—Ö –∑–Ω–∞—á–µ–Ω–∏–π –¥–ª—è –¥–∞—Ç—á–∏–∫–æ–≤ 1, 2 –∏ 3"""
        if (len(self.measurement_buffer['sensor1']) == 0 or 
            len(self.measurement_buffer['sensor2']) == 0 or 
            len(self.measurement_buffer['sensor3']) == 0):
            raise ValueError("–ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è")
        
        avg_sensor1 = sum(self.measurement_buffer['sensor1']) / len(self.measurement_buffer['sensor1'])
        avg_sensor2 = sum(self.measurement_buffer['sensor2']) / len(self.measurement_buffer['sensor2'])
        avg_sensor3 = sum(self.measurement_buffer['sensor3']) / len(self.measurement_buffer['sensor3'])
        
        return avg_sensor1, avg_sensor2, avg_sensor3
    
    def write_calibration_result_1_2(self, distance: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è 1,2 –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40010, 40011"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º float –≤ –¥–≤–∞ 16-–±–∏—Ç–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–∞
                low_word, high_word = self.float_to_doubleword(distance)
                
                # HMI —á–∏—Ç–∞–µ—Ç: —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –∏–∑ 40010, –º–ª–∞–¥—à–∏–π –∏–∑ 40011
                # –ü–æ—ç—Ç–æ–º—É –∑–∞–ø–∏—Å—ã–≤–∞–µ–º –≤ –æ–±—Ä–∞—Ç–Ω–æ–º –ø–æ—Ä—è–¥–∫–µ
                self.modbus_server.slave_context.setValues(3, 10, [int(high_word)])  # 40010 - —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä
                self.modbus_server.slave_context.setValues(3, 11, [int(low_word)])   # 40011 - –º–ª–∞–¥—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä
                
                # –û—Ç–∫–ª—é—á–µ–Ω–æ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –ë–î - —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ —É–∫–∞–∑–∞–Ω–Ω—ã–µ –≤–µ–ª–∏—á–∏–Ω—ã
                # if self.db_integration:
                #     self.db_integration.save_doubleword_register(
                #         40010, 'holding', distance, '–†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,2'
                #     )
                
                print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç 1,2 –∑–∞–ø–∏—Å–∞–Ω –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40010, 40011: {distance:.3f} –º–º (high: {int(high_word)}, low: {int(low_word)})")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ 1,2: {e}")
    
    def write_calibration_result_1_3(self, distance: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è 1,3 –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40012, 40013"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º float –≤ –¥–≤–∞ 16-–±–∏—Ç–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–∞
                low_word, high_word = self.float_to_doubleword(distance)
                
                # HMI —á–∏—Ç–∞–µ—Ç: —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –∏–∑ 40012, –º–ª–∞–¥—à–∏–π –∏–∑ 40013
                # –ü–æ—ç—Ç–æ–º—É –∑–∞–ø–∏—Å—ã–≤–∞–µ–º –≤ –æ–±—Ä–∞—Ç–Ω–æ–º –ø–æ—Ä—è–¥–∫–µ
                self.modbus_server.slave_context.setValues(3, 12, [int(high_word)])  # 40012 - —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä
                self.modbus_server.slave_context.setValues(3, 13, [int(low_word)])   # 40013 - –º–ª–∞–¥—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä
                
                # –û—Ç–∫–ª—é—á–µ–Ω–æ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –ë–î - —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ —É–∫–∞–∑–∞–Ω–Ω—ã–µ –≤–µ–ª–∏—á–∏–Ω—ã
                # if self.db_integration:
                #     self.db_integration.save_doubleword_register(
                #         40012, 'holding', distance, '–†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,3'
                #     )
                
                print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç 1,3 –∑–∞–ø–∏—Å–∞–Ω –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40012, 40013: {distance:.3f} –º–º (high: {int(high_word)}, low: {int(low_word)})")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ 1,3: {e}")
    
    def read_reference_bottom_thickness(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —ç—Ç–∞–ª–æ–Ω–Ω–æ–π —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40004, 40005"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ß–∏—Ç–∞–µ–º –¥–≤–∞ —Ä–µ–≥–∏—Å—Ç—Ä–∞ (40004, 40005) - HMI –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç –≤ –æ–±—Ä–∞—Ç–Ω–æ–º –ø–æ—Ä—è–¥–∫–µ
                values = self.modbus_server.slave_context.getValues(3, 4, 2)
                if values and len(values) >= 2:
                    # HMI: —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –≤ 40004, –º–ª–∞–¥—à–∏–π –≤ 40005
                    high_word = int(values[0])  # 40004
                    low_word = int(values[1])   # 40005
                    
                    # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º –≤ float
                    thickness = self.doubleword_to_float(low_word, high_word)
                    return thickness
            return 0.0
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —ç—Ç–∞–ª–æ–Ω–Ω–æ–π —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞: {e}")
            return 0.0
    
    def measure_sensor4_for_calibration(self):
        """–ò–∑–º–µ—Ä–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 4 –¥–ª—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –≤ —Ç–µ—á–µ–Ω–∏–µ 4 —Å–µ–∫—É–Ω–¥"""
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            return
            
        start_time = time.time()
        measurement_count = 0
        
        print(" –ù–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 4...")
        
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π
        self.clear_serial_buffers()
        self.flush_sensor_queue()
        
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä –¥–∞—Ç—á–∏–∫–∞ 4
        self.measurement_buffer['sensor4'].clear()
        
        while (time.time() - start_time) < self.measurement_duration:
            try:
                # –í—ã–ø–æ–ª–Ω—è–µ–º —á—Ç–µ–Ω–∏–µ –∏–∑–º–µ—Ä–µ–Ω–∏–π —Å –∑–∞—â–∏—Ç–æ–π –æ—Ç –∫–æ–Ω—Ñ–ª–∏–∫—Ç–æ–≤
                sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
                if sensor4_mm is None:
                    time.sleep(0.001)
                    continue

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4
                # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è (None, 0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞)
                if self.is_valid_measurement(sensor4_mm):
                    self.measurement_buffer['sensor4'].append(sensor4_mm)
                    measurement_count += 1
                
                # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø—Ä–æ–≥—Ä–µ—Å—Å –∫–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
                elapsed = time.time() - start_time
                if int(elapsed) != int(elapsed - 0.1):  # –ö–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
                    print(f" –í—Ä–µ–º—è: {elapsed:.1f}—Å, –ò–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 4: {measurement_count}")
                    
            except Exception as e:
                print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4: {e}")
                # –£–±—Ä–∞–Ω sleep –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è
        
        print(f" –ò–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4 –∑–∞–≤–µ—Ä—à–µ–Ω—ã. –í—Å–µ–≥–æ: {measurement_count}")
    
    def calculate_sensor4_average(self) -> float:
        """–í—ã—á–∏—Å–ª–µ–Ω–∏–µ —Å—Ä–µ–¥–Ω–µ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è –¥–ª—è –¥–∞—Ç—á–∏–∫–∞ 4"""
        if len(self.measurement_buffer['sensor4']) == 0:
            raise ValueError("–ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4")
        
        avg_sensor4 = sum(self.measurement_buffer['sensor4']) / len(self.measurement_buffer['sensor4'])
        return round(avg_sensor4, 3)
    
    def write_calibration_result_4_surface(self, distance: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4 –¥–æ –ø–æ–≤–µ—Ä—Ö–Ω–æ—Å—Ç–∏ –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40014, 40015"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º float –≤ –¥–≤–∞ 16-–±–∏—Ç–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–∞
                low_word, high_word = self.float_to_doubleword(distance)
                
                # HMI —á–∏—Ç–∞–µ—Ç: —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –∏–∑ 40014, –º–ª–∞–¥—à–∏–π –∏–∑ 40015
                # –ü–æ—ç—Ç–æ–º—É –∑–∞–ø–∏—Å—ã–≤–∞–µ–º –≤ –æ–±—Ä–∞—Ç–Ω–æ–º –ø–æ—Ä—è–¥–∫–µ
                self.modbus_server.slave_context.setValues(3, 14, [int(high_word)])  # 40014 - —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä
                self.modbus_server.slave_context.setValues(3, 15, [int(low_word)])   # 40015 - –º–ª–∞–¥—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä
                
                # –û—Ç–∫–ª—é—á–µ–Ω–æ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –ë–î - —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ —É–∫–∞–∑–∞–Ω–Ω—ã–µ –≤–µ–ª–∏—á–∏–Ω—ã
                # if self.db_integration:
                #     self.db_integration.save_doubleword_register(
                #         40014, 'holding', distance, '–†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 4 –¥–æ –ø–æ–≤–µ—Ä—Ö–Ω–æ—Å—Ç–∏'
                #     )
                
                print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç –¥–∞—Ç—á–∏–∫–∞ 4 –∑–∞–ø–∏—Å–∞–Ω –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40014, 40015: {distance:.3f} –º–º (high: {int(high_word)}, low: {int(low_word)})")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –¥–∞—Ç—á–∏–∫–∞ 4: {e}")
    
    def read_reference_diameter(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40006, 40007"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ß–∏—Ç–∞–µ–º –¥–≤–∞ —Ä–µ–≥–∏—Å—Ç—Ä–∞ (40006, 40007) - HMI –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç –≤ –æ–±—Ä–∞—Ç–Ω–æ–º –ø–æ—Ä—è–¥–∫–µ
                values = self.modbus_server.slave_context.getValues(3, 6, 2)
                if values and len(values) >= 2:
                    # HMI: —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –≤ 40006, –º–ª–∞–¥—à–∏–π –≤ 40007
                    high_word = int(values[0])  # 40006
                    low_word = int(values[1])   # 40007
                    
                    # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º –≤ float
                    diameter = self.doubleword_to_float(low_word, high_word)
                    return diameter
            return 0.0
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —ç—Ç–∞–ª–æ–Ω–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞: {e}")
            return 0.0
    
    def measure_sensor1_for_calibration(self):
        """–ò–∑–º–µ—Ä–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 1 –¥–ª—è –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ –≤ —Ç–µ—á–µ–Ω–∏–µ 4 —Å–µ–∫—É–Ω–¥"""
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            return
            
        start_time = time.time()
        measurement_count = 0
        
        print(" –ù–∞—á–∞–ª–æ –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1...")
        
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π
        self.clear_serial_buffers()
        self.flush_sensor_queue()
        
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä –¥–∞—Ç—á–∏–∫–∞ 1
        self.measurement_buffer['sensor1'].clear()
        
        while (time.time() - start_time) < self.measurement_duration:
            try:
                # –í—ã–ø–æ–ª–Ω—è–µ–º —á—Ç–µ–Ω–∏–µ –∏–∑–º–µ—Ä–µ–Ω–∏–π —Å –∑–∞—â–∏—Ç–æ–π –æ—Ç –∫–æ–Ω—Ñ–ª–∏–∫—Ç–æ–≤
                sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
                if sensor1_mm is None:
                    time.sleep(0.001)
                    continue

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 1
                # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è (None, 0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞)
                if self.is_valid_measurement(sensor1_mm):
                    self.measurement_buffer['sensor1'].append(sensor1_mm)
                    measurement_count += 1
                
                # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø—Ä–æ–≥—Ä–µ—Å—Å –∫–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
                elapsed = time.time() - start_time
                if int(elapsed) != int(elapsed - 0.1):  # –ö–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
                    print(f" –í—Ä–µ–º—è: {elapsed:.1f}—Å, –ò–∑–º–µ—Ä–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–∞ 1: {measurement_count}")
                    
            except Exception as e:
                print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 1: {e}")
                # –£–±—Ä–∞–Ω sleep –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è
        
        print(f" –ò–∑–º–µ—Ä–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 1 –∑–∞–≤–µ—Ä—à–µ–Ω—ã. –í—Å–µ–≥–æ: {measurement_count}")
    
    def calculate_sensor1_average(self) -> float:
        """–í—ã—á–∏—Å–ª–µ–Ω–∏–µ —Å—Ä–µ–¥–Ω–µ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è –¥–ª—è –¥–∞—Ç—á–∏–∫–∞ 1"""
        if len(self.measurement_buffer['sensor1']) == 0:
            raise ValueError("–ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 1")
        
        avg_sensor1 = sum(self.measurement_buffer['sensor1']) / len(self.measurement_buffer['sensor1'])
        return round(avg_sensor1, 3)
    
    def write_calibration_result_1_center(self, distance: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –∫–∞–ª–∏–±—Ä–æ–≤–∫–∏ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 1 –¥–æ —Ü–µ–Ω—Ç—Ä–∞ –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40016, 40017"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º float –≤ –¥–≤–∞ 16-–±–∏—Ç–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–∞
                low_word, high_word = self.float_to_doubleword(distance)
                
                # HMI —á–∏—Ç–∞–µ—Ç: —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä –∏–∑ 40016, –º–ª–∞–¥—à–∏–π –∏–∑ 40017
                # –ü–æ—ç—Ç–æ–º—É –∑–∞–ø–∏—Å—ã–≤–∞–µ–º –≤ –æ–±—Ä–∞—Ç–Ω–æ–º –ø–æ—Ä—è–¥–∫–µ
                self.modbus_server.slave_context.setValues(3, 16, [int(high_word)])  # 40016 - —Å—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä
                self.modbus_server.slave_context.setValues(3, 17, [int(low_word)])   # 40017 - –º–ª–∞–¥—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä
                
                # –û—Ç–∫–ª—é—á–µ–Ω–æ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ –ë–î - —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–æ–ª—å–∫–æ —É–∫–∞–∑–∞–Ω–Ω—ã–µ –≤–µ–ª–∏—á–∏–Ω—ã
                # if self.db_integration:
                #     self.db_integration.save_doubleword_register(
                #         40016, 'holding', distance, '–†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 1 –¥–æ —Ü–µ–Ω—Ç—Ä–∞'
                #     )
                
                print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç –¥–∞—Ç—á–∏–∫–∞ 1 –∑–∞–ø–∏—Å–∞–Ω –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã 40016, 40017: {distance:.3f} –º–º (high: {int(high_word)}, low: {int(low_word)})")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –¥–∞—Ç—á–∏–∫–∞ 1: {e}")
    
    def float_to_doubleword(self, value: float) -> Tuple[int, int]:
        """–ö–æ–Ω–≤–µ—Ä—Ç–∞—Ü–∏—è float –≤ –¥–≤–∞ 16-–±–∏—Ç–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–∞"""
        packed = struct.pack('>f', value)  # Big-endian float
        high_word, low_word = struct.unpack('>HH', packed)
        return low_word, high_word  # –ú–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ –ø–µ—Ä–≤–æ–µ
    
    def doubleword_to_float(self, low_word: int, high_word: int) -> float:
        """–ö–æ–Ω–≤–µ—Ä—Ç–∞—Ü–∏—è –¥–≤—É—Ö 16-–±–∏—Ç–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ –≤ float"""
        packed = struct.pack('>HH', high_word, low_word)
        return struct.unpack('>f', packed)[0]
    
    def reset_command(self):
        """–°–±—Ä–æ—Å –∫–æ–º–∞–Ω–¥—ã –≤ —Ä–µ–≥–∏—Å—Ç—Ä 40001 –≤ 0"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # 40001 —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É–µ—Ç –∏–Ω–¥–µ–∫—Å—É 0 –≤ ModbusSequentialDataBlock
                self.modbus_server.slave_context.setValues(3, 0, [0])  # 40001 = 0
                print(" –ö–æ–º–∞–Ω–¥–∞ —Å–±—Ä–æ—à–µ–Ω–∞ –≤ 0")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —Å–±—Ä–æ—Å–∞ –∫–æ–º–∞–Ω–¥—ã: {e}")
    
    def handle_measure_wall_state(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ (CMD = 10) - –¢–û–õ–¨–ö–û –°–ë–û–† –î–ê–ù–ù–´–•"""
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            self.current_state = SystemState.ERROR
            return
        
        # –£–±—Ä–∞–Ω–∞ –ø—Ä–æ–≤–µ—Ä–∫–∞ get_current_command() –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è
        # –ö–æ–º–∞–Ω–¥–∞ –º–µ–Ω—è–µ—Ç—Å—è —á–µ—Ä–µ–∑ handle_command(), –∫–æ—Ç–æ—Ä—ã–π —É–∂–µ –º–µ–Ω—è–µ—Ç current_state
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è —Ç–∞–π–º–µ—Ä–∞ —á–∞—Å—Ç–æ—Ç—ã –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–∏
            if self.frequency_start_time is None:
                self.frequency_start_time = time.time()
                self.last_frequency_display = self.frequency_start_time
            
            # –°—Ç–∞—Ç—É—Å —É–∂–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –≤ manage_measurement_cycle_flag
            # –ü—Ä–æ—Å—Ç–æ –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
            
            # –í—ã–ø–æ–ª–Ω—è–µ–º QUAD –∏–∑–º–µ—Ä–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–æ–≤ 1 –∏ 2 (–±–µ–∑–æ–ø–∞—Å–Ω–æ–µ —á—Ç–µ–Ω–∏–µ —Å –±–ª–æ–∫–∏—Ä–æ–≤–∫–æ–π)
            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor1_mm is None or sensor2_mm is None:
                time.sleep(0.001)
                return
            
            # –£–≤–µ–ª–∏—á–∏–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫ –∏–∑–º–µ—Ä–µ–Ω–∏–π
            self.frequency_counter += 1
            
            # –í—ã–≤–æ–¥–∏–º —á–∞—Å—Ç–æ—Ç—É –∫–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
            current_time = time.time()
            if current_time - self.last_frequency_display >= 1.0:
                elapsed = current_time - self.frequency_start_time
                if elapsed > 0:
                    instant_freq = self.frequency_counter / elapsed
                    print(f" [CMD=10] –ß–∞—Å—Ç–æ—Ç–∞ –æ–ø—Ä–æ—Å–∞: {instant_freq:.1f} –ì—Ü | –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.frequency_counter}")
                self.last_frequency_display = current_time
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º —á—Ç–æ –ø–æ–ª—É—á–∏–ª–∏ –≤–∞–ª–∏–¥–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –æ—Ç –¥–∞—Ç—á–∏–∫–æ–≤ 1 –∏ 2
            # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è (None, 0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞)
            if (self.is_valid_measurement(sensor1_mm) and self.is_valid_measurement(sensor2_mm)):
                # –î–æ–±–∞–≤–ª—è–µ–º –≤ –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è
                self.temp_sensor1_buffer.append(sensor1_mm)
                self.temp_sensor2_buffer.append(sensor2_mm)
                
                # –ö–æ–≥–¥–∞ –Ω–∞–∫–æ–ø–∏–ª–æ—Å—å 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π - —É—Å—Ä–µ–¥–Ω—è–µ–º –∏ –∑–∞–ø–∏—Å—ã–≤–∞–µ–º
                if len(self.temp_sensor1_buffer) >= 10:
                    # –§–∏–ª—å—Ç—Ä—É–µ–º –∞–Ω–æ–º–∞–ª—å–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è –ø–µ—Ä–µ–¥ —É—Å—Ä–µ–¥–Ω–µ–Ω–∏–µ–º (–∏—Å–ø–æ–ª—å–∑—É–µ–º –º–µ–¥–∏–∞–Ω–Ω—ã–π —Ñ–∏–ª—å—Ç—Ä)
                    sorted_sensor1 = sorted(self.temp_sensor1_buffer)
                    sorted_sensor2 = sorted(self.temp_sensor2_buffer)
                    
                    # –í—ã—á–∏—Å–ª—è–µ–º –º–µ–¥–∏–∞–Ω—É
                    median_sensor1 = (sorted_sensor1[4] + sorted_sensor1[5]) / 2.0
                    median_sensor2 = (sorted_sensor2[4] + sorted_sensor2[5]) / 2.0
                    
                    # –§–∏–ª—å—Ç—Ä—É–µ–º –∑–Ω–∞—á–µ–Ω–∏—è, –∫–æ—Ç–æ—Ä—ã–µ –æ—Ç–∫–ª–æ–Ω—è—é—Ç—Å—è –æ—Ç –º–µ–¥–∏–∞–Ω—ã –±–æ–ª–µ–µ —á–µ–º –Ω–∞ 1.5–º–º
                    filtered_sensor1 = [v for v in self.temp_sensor1_buffer if abs(v - median_sensor1) <= 1.5]
                    filtered_sensor2 = [v for v in self.temp_sensor2_buffer if abs(v - median_sensor2) <= 1.5]
                    
                    # –ï—Å–ª–∏ –ø–æ—Å–ª–µ —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏–∏ –æ—Å—Ç–∞–ª–æ—Å—å –º–µ–Ω–µ–µ 5 –∑–Ω–∞—á–µ–Ω–∏–π - –∏—Å–ø–æ–ª—å–∑—É–µ–º –º–µ–¥–∏–∞–Ω—É
                    if len(filtered_sensor1) >= 5:
                        avg_sensor1 = sum(filtered_sensor1) / len(filtered_sensor1)
                    else:
                        avg_sensor1 = median_sensor1
                    
                    if len(filtered_sensor2) >= 5:
                        avg_sensor2 = sum(filtered_sensor2) / len(filtered_sensor2)
                    else:
                        avg_sensor2 = median_sensor2
                    
                    # –î–æ–±–∞–≤–ª—è–µ–º —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è –≤ –æ—Å–Ω–æ–≤–Ω—ã–µ –±—É—Ñ–µ—Ä—ã
                    self.sensor1_measurements.append(avg_sensor1)
                    self.sensor2_measurements.append(avg_sensor2)
                    
                    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∫–µ—à–∏—Ä–æ–≤–∞–Ω–Ω–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ (–≤–º–µ—Å—Ç–æ —á—Ç–µ–Ω–∏—è –∏–∑ Modbus)
                    distance_1_2 = self.cached_distance_1_2
                    wall_upper_offset = self.read_upper_wall_offset_coeff()
                    
                    if distance_1_2 is not None:
                        # –í—ã—á–∏—Å–ª—è–µ–º —Ç–æ–ª—â–∏–Ω—É —Å—Ç–µ–Ω–∫–∏ –ø–æ —Ñ–æ—Ä–º—É–ª–µ (–∏—Å–ø–æ–ª—å–∑—É–µ–º —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è)
                        wall_thickness = distance_1_2 - avg_sensor1 - avg_sensor2 + wall_upper_offset
                        self.wall_thickness_buffer.append(wall_thickness)
                        
                        # –í—ã–≤–æ–¥–∏–º —Ç–µ–∫—É—â–∏–µ –∑–Ω–∞—á–µ–Ω–∏—è –∫–∞–∂–¥—ã–µ 100 —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π (—É–º–µ–Ω—å—à–µ–Ω–∞ —á–∞—Å—Ç–æ—Ç–∞ –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è)
                        if len(self.wall_thickness_buffer) % 100 == 0:
                            print(f" –£—Å—Ä–µ–¥–Ω–µ–Ω–Ω–æ–µ –∏–∑–º–µ—Ä–µ–Ω–∏–µ #{len(self.wall_thickness_buffer)}: "
                                  f"–î–∞—Ç—á–∏–∫1={avg_sensor1:.3f}–º–º, –î–∞—Ç—á–∏–∫2={avg_sensor2:.3f}–º–º, "
                                  f"–¢–æ–ª—â–∏–Ω–∞={wall_thickness:.3f}–º–º")
                    else:
                        print(" –û—à–∏–±–∫–∞: –Ω–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ 1,2")
                    
                    # –û—á–∏—â–∞–µ–º –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —Å–ª–µ–¥—É—é—â–∏—Ö 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π
                    self.temp_sensor1_buffer = []
                    self.temp_sensor2_buffer = []
            else:
                print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è: –¥–∞—Ç—á–∏–∫1={sensor1_mm}, –¥–∞—Ç—á–∏–∫2={sensor2_mm}")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏: {e}")
            self.current_state = SystemState.ERROR
    
    def read_calibrated_distance_1_2(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,2 –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40010-40011"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ß–∏—Ç–∞–µ–º –¥–≤–∞ —Ä–µ–≥–∏—Å—Ç—Ä–∞ (HMI: —Å—Ç–∞—Ä—à–∏–π –≤ 40010, –º–ª–∞–¥—à–∏–π –≤ 40011)
                values = self.modbus_server.slave_context.getValues(3, 10, 2)  # 40010-40011
                if values and len(values) == 2:
                    high_word = int(values[0])  # 40010 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40011 - –º–ª–∞–¥—à–∏–π
                    distance = self.doubleword_to_float(low_word, high_word)
                    return distance
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è 1,2: {e}")
        return None
    
    def read_calibrated_distance_to_center(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–æ —Ü–µ–Ω—Ç—Ä–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40016-40017"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ß–∏—Ç–∞–µ–º –¥–≤–∞ —Ä–µ–≥–∏—Å—Ç—Ä–∞ (HMI: —Å—Ç–∞—Ä—à–∏–π –≤ 40016, –º–ª–∞–¥—à–∏–π –≤ 40017)
                values = self.modbus_server.slave_context.getValues(3, 16, 2)  # 40016-40017
                if values and len(values) == 2:
                    high_word = int(values[0])  # 40016 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40017 - –º–ª–∞–¥—à–∏–π
                    distance = self.doubleword_to_float(low_word, high_word)
                    return distance
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–æ —Ü–µ–Ω—Ç—Ä–∞: {e}")
        return None
    
    def read_calibrated_distance_1_3(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –º–µ–∂–¥—É –¥–∞—Ç—á–∏–∫–∞–º–∏ 1,3 –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40012-40013"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ß–∏—Ç–∞–µ–º –¥–≤–∞ —Ä–µ–≥–∏—Å—Ç—Ä–∞ (HMI: —Å—Ç–∞—Ä—à–∏–π –≤ 40012, –º–ª–∞–¥—à–∏–π –≤ 40013)
                values = self.modbus_server.slave_context.getValues(3, 12, 2)  # 40012-40013
                if values and len(values) == 2:
                    high_word = int(values[0])  # 40012 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40013 - –º–ª–∞–¥—à–∏–π
                    distance = self.doubleword_to_float(low_word, high_word)
                    return distance
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è 1,3: {e}")
        return None
    
    def read_calibrated_distance_sensor4(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4 –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40014-40015"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ß–∏—Ç–∞–µ–º –¥–≤–∞ —Ä–µ–≥–∏—Å—Ç—Ä–∞ (HMI: —Å—Ç–∞—Ä—à–∏–π –≤ 40014, –º–ª–∞–¥—à–∏–π –≤ 40015)
                values = self.modbus_server.slave_context.getValues(3, 14, 2)  # 40014-40015
                if values and len(values) == 2:
                    high_word = int(values[0])  # 40014 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40015 - –º–ª–∞–¥—à–∏–π
                    distance = self.doubleword_to_float(low_word, high_word)
                    return distance
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4: {e}")
        return None
    
    def read_calibrated_distance_sensor3_to_center(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ —Ü–µ–Ω—Ç—Ä–∞ –∏–∑ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40032, 40033"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 32, 2)  # 40032-40033 -> –∏–Ω–¥–µ–∫—Å—ã 32-33
                if values and len(values) >= 2:
                    high_word = int(values[0])  # 40032 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40033 - –º–ª–∞–¥—à–∏–π
                    distance = self.doubleword_to_float(low_word, high_word)
                    return distance
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ —Ü–µ–Ω—Ç—Ä–∞: {e}")
        return None

    def read_calibrated_distance_sensor3_to_center_body(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ —Ü–µ–Ω—Ç—Ä–∞ (—Ä–∞–∑–¥–µ–ª—å–Ω—ã–π –¥–∏–∞–º–µ—Ç—Ä –∫–æ—Ä–ø—É—Å–∞) –∏–∑ 40038-40039"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 38, 2)  # 40038-40039
                if values and len(values) >= 2:
                    high_word = int(values[0])
                    low_word = int(values[1])
                    return self.doubleword_to_float(low_word, high_word)
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ —Ü–µ–Ω—Ç—Ä–∞ (—Ä–∞–∑–¥–µ–ª—å–Ω—ã–π –∫–æ—Ä–ø—É—Å): {e}")
        return None

    def read_calibrated_distance_sensor3_to_center_body2(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ —Ü–µ–Ω—Ç—Ä–∞ (–∫–æ—Ä–ø—É—Å 2) –∏–∑ 40040-40041"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 40, 2)  # 40040-40041
                if values and len(values) >= 2:
                    high_word = int(values[0])
                    low_word = int(values[1])
                    return self.doubleword_to_float(low_word, high_word)
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ —Ü–µ–Ω—Ç—Ä–∞ (–∫–æ—Ä–ø—É—Å 2): {e}")
        return None
    
    def apply_extrapolation_to_buffer(self, buffer: list, extrapolation_coeff: float) -> list:
        """
        –ü—Ä–∏–º–µ–Ω–µ–Ω–∏–µ —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ –∫ –±—É—Ñ–µ—Ä—É –∏–∑–º–µ—Ä–µ–Ω–∏–π
        
        –§–æ—Ä–º—É–ª–∞: —ç–∫—Å—Ç—Ä–∞–ø–æ–ª–∏—Ä–æ–≤–∞–Ω–Ω–æ–µ_–∑–Ω–∞—á–µ–Ω–∏–µ = —Å—Ä–µ–¥–Ω–µ–µ + –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç * (–∏–∑–º–µ—Ä–µ–Ω–Ω–æ–µ - —Å—Ä–µ–¥–Ω–µ–µ)
        
        Args:
            buffer: –°–ø–∏—Å–æ–∫ –∏–∑–º–µ—Ä–µ–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π
            extrapolation_coeff: –ö–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏
            
        Returns:
            –°–ø–∏—Å–æ–∫ —ç–∫—Å—Ç—Ä–∞–ø–æ–ª–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π
        """
        if not buffer or len(buffer) == 0:
            return buffer
        
        # –í—ã—á–∏—Å–ª—è–µ–º —Å—Ä–µ–¥–Ω–µ–µ –∏–∑ –∏—Å—Ö–æ–¥–Ω–æ–≥–æ –±—É—Ñ–µ—Ä–∞
        avg_value = sum(buffer) / len(buffer)
        
        # –ü—Ä–∏–º–µ–Ω—è–µ–º —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏—é –∫ –∫–∞–∂–¥–æ–º—É –∑–Ω–∞—á–µ–Ω–∏—é
        extrapolated_buffer = []
        for value in buffer:
            extrapolated_value = avg_value + extrapolation_coeff * (value - avg_value)
            extrapolated_buffer.append(extrapolated_value)
        
        return extrapolated_buffer
    
    def process_wall_measurement_results(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Å—Ç–µ–Ω–∫–∏ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ 10‚Üí11"""
        try:
            if len(self.wall_thickness_buffer) == 0:
                print(" –û—à–∏–±–∫–∞: –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏")
                return
            
            # –ß–∏—Ç–∞–µ–º –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ –¥–ª—è –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
            extrapolation_coeff = self.read_upper_wall_extrapolation_coeff()
            
            # –ü—Ä–∏–º–µ–Ω—è–µ–º —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏—é –∫ –±—É—Ñ–µ—Ä—É
            if abs(extrapolation_coeff) > 0.0001:  # –ü—Ä–∏–º–µ–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –µ—Å–ª–∏ –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç –Ω–µ —Ä–∞–≤–µ–Ω –Ω—É–ª—é
                extrapolated_buffer = self.apply_extrapolation_to_buffer(self.wall_thickness_buffer, extrapolation_coeff)
                print(f" [–≠–ö–°–¢–†–ê–ü–û–õ–Ø–¶–ò–Ø] –ü—Ä–∏–º–µ–Ω–µ–Ω –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç {extrapolation_coeff:.6f} –∫ —Ç–æ–ª—â–∏–Ω–µ –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏")
            else:
                extrapolated_buffer = self.wall_thickness_buffer
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –∏–∑ —ç–∫—Å—Ç—Ä–∞–ø–æ–ª–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π
            max_thickness = max(extrapolated_buffer)
            min_thickness = min(extrapolated_buffer)
            avg_thickness = sum(extrapolated_buffer) / len(extrapolated_buffer)
            
            print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏:")
            print(f"   –ò–∑–º–µ—Ä–µ–Ω–∏–π: {len(self.wall_thickness_buffer)}")
            print(f"   –ú–∞–∫—Å–∏–º—É–º: {max_thickness:.3f}–º–º")
            print(f"   –°—Ä–µ–¥–Ω–µ–µ:  {avg_thickness:.3f}–º–º")
            print(f"   –ú–∏–Ω–∏–º—É–º:  {min_thickness:.3f}–º–º")
            
            # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã
            self.write_wall_measurement_results(max_thickness, avg_thickness, min_thickness)
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Å—Ç–µ–Ω–∫–∏: {e}")
    
    def write_wall_measurement_results(self, max_val: float, avg_val: float, min_val: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Å—Ç–µ–Ω–∫–∏ –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ú–∞–∫—Å–∏–º–∞–ª—å–Ω–∞—è —Ç–æ–ª—â–∏–Ω–∞ ‚Üí 30016-30017
                self.write_stream_result_to_input_registers(max_val, 30016)
                
                # –°—Ä–µ–¥–Ω—è—è —Ç–æ–ª—â–∏–Ω–∞ ‚Üí 30018-30019
                self.write_stream_result_to_input_registers(avg_val, 30018)
                
                # –ú–∏–Ω–∏–º–∞–ª—å–Ω–∞—è —Ç–æ–ª—â–∏–Ω–∞ ‚Üí 30020-30021
                self.write_stream_result_to_input_registers(min_val, 30020)
                
                print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –∑–∞–ø–∏—Å–∞–Ω—ã: –º–∞–∫—Å={max_val:.3f}, —Å—Ä–µ–¥={avg_val:.3f}, –º–∏–Ω={min_val:.3f}")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Å—Ç–µ–Ω–∫–∏: {e}")
    
    def process_flange_measurement_results(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ 11‚Üí12"""
        try:
            if (len(self.body_diameter_buffer) == 0 or len(self.flange_diameter_buffer) == 0 or
                len(self.bottom_thickness_buffer) == 0):
                print(" –û—à–∏–±–∫–∞: –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π —Ñ–ª–∞–Ω—Ü–∞ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏")
                return
            
            # –î–ò–ê–ì–ù–û–°–¢–ò–ö–ê: –í—ã–≤–æ–¥–∏–º –í–°–ï –±—É—Ñ–µ—Ä—ã —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π –¥–∞—Ç—á–∏–∫–æ–≤
            print(f"\n{'='*80}")
            print(f" –î–ò–ê–ì–ù–û–°–¢–ò–ö–ê: –í–°–ï –ë–£–§–ï–†–´ –£–°–†–ï–î–ù–ï–ù–ù–´–• –ó–ù–ê–ß–ï–ù–ò–ô –î–ê–¢–ß–ò–ö–û–í")
            print(f"{'='*80}")
            
            print(f"\n [–ë–£–§–ï–† –£–°–†–ï–î–ù–ï–ù–ù–´–• –ó–ù–ê–ß–ï–ù–ò–ô –î–ê–¢–ß–ò–ö–ê 1] –†–∞–∑–º–µ—Ä: {len(self.sensor1_flange_measurements)}")
            if len(self.sensor1_flange_measurements) > 0:
                print(f"   –í–°–ï –∑–Ω–∞—á–µ–Ω–∏—è: {[f'{x:.3f}' for x in self.sensor1_flange_measurements]}")
                # –°—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞ –ø–æ –¥–∞—Ç—á–∏–∫—É 1
                max_sensor1 = max(self.sensor1_flange_measurements)
                min_sensor1 = min(self.sensor1_flange_measurements)
                avg_sensor1 = sum(self.sensor1_flange_measurements) / len(self.sensor1_flange_measurements)
                print(f"   –°–¢–ê–¢–ò–°–¢–ò–ö–ê: –º–∞–∫—Å={max_sensor1:.3f}–º–º, —Å—Ä–µ–¥={avg_sensor1:.3f}–º–º, –º–∏–Ω={min_sensor1:.3f}–º–º")
            else:
                print(f"   –ë–£–§–ï–† –ü–£–°–¢!")
            
            print(f"\n [–ë–£–§–ï–† –£–°–†–ï–î–ù–ï–ù–ù–´–• –ó–ù–ê–ß–ï–ù–ò–ô –î–ê–¢–ß–ò–ö–ê 3] –†–∞–∑–º–µ—Ä: {len(self.sensor3_measurements)}")
            if len(self.sensor3_measurements) > 0:
                print(f"   –í–°–ï –∑–Ω–∞—á–µ–Ω–∏—è: {[f'{x:.3f}' for x in self.sensor3_measurements]}")
                # –°—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞ –ø–æ –¥–∞—Ç—á–∏–∫—É 3
                max_sensor3 = max(self.sensor3_measurements)
                min_sensor3 = min(self.sensor3_measurements)
                avg_sensor3 = sum(self.sensor3_measurements) / len(self.sensor3_measurements)
                print(f"   –°–¢–ê–¢–ò–°–¢–ò–ö–ê: –º–∞–∫—Å={max_sensor3:.3f}–º–º, —Å—Ä–µ–¥={avg_sensor3:.3f}–º–º, –º–∏–Ω={min_sensor3:.3f}–º–º")
            else:
                print(f"   –ë–£–§–ï–† –ü–£–°–¢!")
            
            print(f"\n [–ë–£–§–ï–† –£–°–†–ï–î–ù–ï–ù–ù–´–• –ó–ù–ê–ß–ï–ù–ò–ô –î–ê–¢–ß–ò–ö–ê 4] –†–∞–∑–º–µ—Ä: {len(self.sensor4_measurements)}")
            if len(self.sensor4_measurements) > 0:
                print(f"   –í–°–ï –∑–Ω–∞—á–µ–Ω–∏—è: {[f'{x:.3f}' for x in self.sensor4_measurements]}")
                # –°—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞ –ø–æ –¥–∞—Ç—á–∏–∫—É 4
                max_sensor4 = max(self.sensor4_measurements)
                min_sensor4 = min(self.sensor4_measurements)
                avg_sensor4 = sum(self.sensor4_measurements) / len(self.sensor4_measurements)
                print(f"   –°–¢–ê–¢–ò–°–¢–ò–ö–ê: –º–∞–∫—Å={max_sensor4:.3f}–º–º, —Å—Ä–µ–¥={avg_sensor4:.3f}–º–º, –º–∏–Ω={min_sensor4:.3f}–º–º")
            else:
                print(f"   –ë–£–§–ï–† –ü–£–°–¢!")
            
            # –î–ò–ê–ì–ù–û–°–¢–ò–ö–ê: –í—ã–≤–æ–¥–∏–º –í–°–ï –±—É—Ñ–µ—Ä—ã —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π
            print(f"\n{'='*80}")
            print(f" –î–ò–ê–ì–ù–û–°–¢–ò–ö–ê: –í–°–ï –ë–£–§–ï–†–´ –†–ê–°–°–ß–ò–¢–ê–ù–ù–´–• –ó–ù–ê–ß–ï–ù–ò–ô")
            print(f"{'='*80}")
            
            print(f"\n [–ë–£–§–ï–† –î–ò–ê–ú–ï–¢–† –ö–û–†–ü–£–°–ê] –†–∞–∑–º–µ—Ä: {len(self.body_diameter_buffer)}")
            if len(self.body_diameter_buffer) > 0:
                print(f"   –í–°–ï –∑–Ω–∞—á–µ–Ω–∏—è: {[f'{x:.3f}' for x in self.body_diameter_buffer]}")
            else:
                print(f"   –ë–£–§–ï–† –ü–£–°–¢!")
            
            print(f"\n [–ë–£–§–ï–† –î–ò–ê–ú–ï–¢–† –§–õ–ê–ù–¶–ê] –†–∞–∑–º–µ—Ä: {len(self.flange_diameter_buffer)}")
            if len(self.flange_diameter_buffer) > 0:
                print(f"   –í–°–ï –∑–Ω–∞—á–µ–Ω–∏—è: {[f'{x:.3f}' for x in self.flange_diameter_buffer]}")
            else:
                print(f"   –ë–£–§–ï–† –ü–£–°–¢!")
            
            print(f"\n [–ë–£–§–ï–† –¢–û–õ–©–ò–ù–ê –î–ù–ê] –†–∞–∑–º–µ—Ä: {len(self.bottom_thickness_buffer)}")
            if len(self.bottom_thickness_buffer) > 0:
                print(f"   –í–°–ï –∑–Ω–∞—á–µ–Ω–∏—è: {[f'{x:.3f}' for x in self.bottom_thickness_buffer]}")
            else:
                print(f"   –ë–£–§–ï–† –ü–£–°–¢!")
            
            print(f"\n{'='*80}\n")
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –¥–ª—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞
            # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è (0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, NaN, inf) –ø–µ—Ä–µ–¥ –≤—ã—á–∏—Å–ª–µ–Ω–∏–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏
            valid_body_radii = [r for r in self.body_diameter_buffer 
                               if r is not None and r > 0 and not (math.isnan(r) or math.isinf(r))]
            if len(valid_body_radii) == 0:
                print(" –û–®–ò–ë–ö–ê: –ù–µ—Ç –≤–∞–ª–∏–¥–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π —Ä–∞–¥–∏—É—Å–∞ –∫–æ—Ä–ø—É—Å–∞!")
                return
            
            # –ü—Ä–∏–º–µ–Ω—è–µ–º —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏—é –∫ —Ä–∞–¥–∏—É—Å–∞–º –∫–æ—Ä–ø—É—Å–∞
            body_extrapolation_coeff = self.read_body_diameter_extrapolation_coeff()
            if abs(body_extrapolation_coeff) > 0.0001:
                valid_body_radii = self.apply_extrapolation_to_buffer(valid_body_radii, body_extrapolation_coeff)
                print(f" [–≠–ö–°–¢–†–ê–ü–û–õ–Ø–¶–ò–Ø] –ü—Ä–∏–º–µ–Ω–µ–Ω –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç {body_extrapolation_coeff:.6f} –∫ —Ä–∞–¥–∏—É—Å—É –∫–æ—Ä–ø—É—Å–∞")
            
            # –í—ã—á–∏—Å–ª—è–µ–º –¥–∏–∞–º–µ—Ç—Ä—ã –∏–∑ –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫ (0¬∞ –∏ 180¬∞)
            # –§–æ—Ä–º—É–ª–∞: (distance_to_center - avg_sensor1_0deg) + (distance_to_center - avg_sensor1_180deg) + offset
            body_diameter_offset = self.read_body_diameter_offset_coeff()
            opposite_body_diameters = []
            
            if len(valid_body_radii) >= 2:
                # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –∏–∑–º–µ—Ä–µ–Ω–∏–π –∑–∞ –ø–æ–ª–Ω—ã–π –æ–±–æ—Ä–æ—Ç (360¬∞)
                # –î–ª—è –ª—é–±–æ–≥–æ –∫–æ–ª–∏—á–µ—Å—Ç–≤–∞ –æ–±–æ—Ä–æ—Ç–æ–≤ (360¬∞, 720¬∞, 1080¬∞ –∏ —Ç.–¥.)
                # –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω–∞—è —Ç–æ—á–∫–∞ –≤—Å–µ–≥–¥–∞ –Ω–∞ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–∏ N/2
                total_measurements = len(valid_body_radii)
                half_size = total_measurements // 2
                
                # –ë–µ—Ä–µ–º –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã–µ —Ç–æ—á–∫–∏ (—Å–º–µ—â–µ–Ω–Ω—ã–µ –Ω–∞ 180 –≥—Ä–∞–¥—É—Å–æ–≤)
                for i in range(half_size):
                    radius_0deg = valid_body_radii[i]          # –¢–æ—á–∫–∞ –Ω–∞ —Ç–µ–∫—É—â–µ–º —É–≥–ª–µ
                    radius_180deg = valid_body_radii[i + half_size]  # –¢–æ—á–∫–∞ –Ω–∞ 180¬∞ –æ—Ç —Ç–µ–∫—É—â–µ–π
                    # –§–æ—Ä–º—É–ª–∞: (distance_to_center - avg_sensor1_0deg) + (distance_to_center - avg_sensor1_180deg) + offset
                    body_diameter = radius_0deg + radius_180deg + body_diameter_offset
                    opposite_body_diameters.append(body_diameter)
                
                if len(opposite_body_diameters) > 0:
                    max_body_diameter = max(opposite_body_diameters)
                    min_body_diameter = min(opposite_body_diameters)
                    avg_body_diameter = sum(opposite_body_diameters) / len(opposite_body_diameters)
                    print(f" [DIAMETER] –í—ã—á–∏—Å–ª–µ–Ω–æ {len(opposite_body_diameters)} –¥–∏–∞–º–µ—Ç—Ä–æ–≤ –∫–æ—Ä–ø—É—Å–∞ –∏–∑ –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫")
                else:
                    # Fallback: –µ—Å–ª–∏ –Ω–µ —É–¥–∞–ª–æ—Å—å –≤—ã—á–∏—Å–ª–∏—Ç—å –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã–µ —Ç–æ—á–∫–∏
                    print(" [WARNING] –ù–µ —É–¥–∞–ª–æ—Å—å –≤—ã—á–∏—Å–ª–∏—Ç—å –¥–∏–∞–º–µ—Ç—Ä—ã –∫–æ—Ä–ø—É—Å–∞ –∏–∑ –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫, –∏—Å–ø–æ–ª—å–∑—É–µ–º —Å—Ç–∞—Ä—ã–π –º–µ—Ç–æ–¥")
                    max_body_diameter = max(valid_body_radii) * 2 + body_diameter_offset
                    min_body_diameter = min(valid_body_radii) * 2 + body_diameter_offset
                    avg_body_diameter = (sum(valid_body_radii) / len(valid_body_radii)) * 2 + body_diameter_offset
            else:
                # –ï—Å–ª–∏ –∏–∑–º–µ—Ä–µ–Ω–∏–π –Ω–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–ª—è –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫
                print(" [WARNING] –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–ª—è –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫ –∫–æ—Ä–ø—É—Å–∞, –∏—Å–ø–æ–ª—å–∑—É–µ–º —Å—Ç–∞—Ä—ã–π –º–µ—Ç–æ–¥")
                max_body_diameter = max(valid_body_radii) * 2 + body_diameter_offset
                min_body_diameter = min(valid_body_radii) * 2 + body_diameter_offset
                avg_body_diameter = (sum(valid_body_radii) / len(valid_body_radii)) * 2 + body_diameter_offset
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –¥–ª—è –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞
            # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è (0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, NaN, inf) –ø–µ—Ä–µ–¥ –≤—ã—á–∏—Å–ª–µ–Ω–∏–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏
            valid_flange_radii = [r for r in self.flange_diameter_buffer 
                                 if r is not None and r > 0 and not (math.isnan(r) or math.isinf(r))]
            if len(valid_flange_radii) == 0:
                print(" –û–®–ò–ë–ö–ê: –ù–µ—Ç –≤–∞–ª–∏–¥–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π —Ä–∞–¥–∏—É—Å–∞ —Ñ–ª–∞–Ω—Ü–∞!")
                return
            
            # –ü—Ä–∏–º–µ–Ω—è–µ–º —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏—é –∫ —Ä–∞–¥–∏—É—Å–∞–º —Ñ–ª–∞–Ω—Ü–∞
            flange_extrapolation_coeff = self.read_flange_diameter_extrapolation_coeff()
            if abs(flange_extrapolation_coeff) > 0.0001:
                valid_flange_radii = self.apply_extrapolation_to_buffer(valid_flange_radii, flange_extrapolation_coeff)
                print(f" [–≠–ö–°–¢–†–ê–ü–û–õ–Ø–¶–ò–Ø] –ü—Ä–∏–º–µ–Ω–µ–Ω –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç {flange_extrapolation_coeff:.6f} –∫ —Ä–∞–¥–∏—É—Å—É —Ñ–ª–∞–Ω—Ü–∞")
            
            # –í—ã—á–∏—Å–ª—è–µ–º –¥–∏–∞–º–µ—Ç—Ä—ã —Ñ–ª–∞–Ω—Ü–∞ –∏–∑ –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫ (0¬∞ –∏ 180¬∞)
            # –§–æ—Ä–º—É–ª–∞: (distance_to_center_flange - avg_sensor3_0deg) + (distance_to_center_flange - avg_sensor3_180deg) + offset
            flange_diameter_offset = self.read_flange_diameter_offset_coeff()
            opposite_flange_diameters = []
            
            if len(valid_flange_radii) >= 2:
                # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –∏–∑–º–µ—Ä–µ–Ω–∏–π –∑–∞ –ø–æ–ª–Ω—ã–π –æ–±–æ—Ä–æ—Ç (360¬∞)
                # –î–ª—è –ª—é–±–æ–≥–æ –∫–æ–ª–∏—á–µ—Å—Ç–≤–∞ –æ–±–æ—Ä–æ—Ç–æ–≤ (360¬∞, 720¬∞, 1080¬∞ –∏ —Ç.–¥.)
                # –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω–∞—è —Ç–æ—á–∫–∞ –≤—Å–µ–≥–¥–∞ –Ω–∞ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–∏ N/2
                total_measurements = len(valid_flange_radii)
                half_size = total_measurements // 2
                
                # –ë–µ—Ä–µ–º –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã–µ —Ç–æ—á–∫–∏ (—Å–º–µ—â–µ–Ω–Ω—ã–µ –Ω–∞ 180 –≥—Ä–∞–¥—É—Å–æ–≤)
                for i in range(half_size):
                    radius_0deg = valid_flange_radii[i]          # –¢–æ—á–∫–∞ –Ω–∞ —Ç–µ–∫—É—â–µ–º —É–≥–ª–µ
                    radius_180deg = valid_flange_radii[i + half_size]  # –¢–æ—á–∫–∞ –Ω–∞ 180¬∞ –æ—Ç —Ç–µ–∫—É—â–µ–π
                    # –§–æ—Ä–º—É–ª–∞: (distance_to_center_flange - avg_sensor3_0deg) + (distance_to_center_flange - avg_sensor3_180deg) + offset
                    flange_diameter = radius_0deg + radius_180deg + flange_diameter_offset
                    opposite_flange_diameters.append(flange_diameter)
                
                if len(opposite_flange_diameters) > 0:
                    max_flange_diameter = max(opposite_flange_diameters)
                    min_flange_diameter = min(opposite_flange_diameters)
                    avg_flange_diameter = sum(opposite_flange_diameters) / len(opposite_flange_diameters)
                    print(f" [DIAMETER] –í—ã—á–∏—Å–ª–µ–Ω–æ {len(opposite_flange_diameters)} –¥–∏–∞–º–µ—Ç—Ä–æ–≤ —Ñ–ª–∞–Ω—Ü–∞ –∏–∑ –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫")
                else:
                    # Fallback: –µ—Å–ª–∏ –Ω–µ —É–¥–∞–ª–æ—Å—å –≤—ã—á–∏—Å–ª–∏—Ç—å –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã–µ —Ç–æ—á–∫–∏
                    print(" [WARNING] –ù–µ —É–¥–∞–ª–æ—Å—å –≤—ã—á–∏—Å–ª–∏—Ç—å –¥–∏–∞–º–µ—Ç—Ä—ã —Ñ–ª–∞–Ω—Ü–∞ –∏–∑ –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫, –∏—Å–ø–æ–ª—å–∑—É–µ–º —Å—Ç–∞—Ä—ã–π –º–µ—Ç–æ–¥")
                    max_flange_diameter = max(valid_flange_radii) * 2 + flange_diameter_offset
                    min_flange_diameter = min(valid_flange_radii) * 2 + flange_diameter_offset
                    avg_flange_diameter = (sum(valid_flange_radii) / len(valid_flange_radii)) * 2 + flange_diameter_offset
            else:
                # –ï—Å–ª–∏ –∏–∑–º–µ—Ä–µ–Ω–∏–π –Ω–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–ª—è –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫
                print(" [WARNING] –ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –∏–∑–º–µ—Ä–µ–Ω–∏–π –¥–ª—è –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫ —Ñ–ª–∞–Ω—Ü–∞, –∏—Å–ø–æ–ª—å–∑—É–µ–º —Å—Ç–∞—Ä—ã–π –º–µ—Ç–æ–¥")
                max_flange_diameter = max(valid_flange_radii) * 2 + flange_diameter_offset
                min_flange_diameter = min(valid_flange_radii) * 2 + flange_diameter_offset
                avg_flange_diameter = (sum(valid_flange_radii) / len(valid_flange_radii)) * 2 + flange_diameter_offset
            
            # –¢–æ–ª—â–∏–Ω–∞ —Ñ–ª–∞–Ω—Ü–∞ —Ç–µ–ø–µ—Ä—å –ø–µ—Ä–µ–¥–∞—ë—Ç—Å—è —Å –ü–ö, –Ω–µ —Ä–∞—Å—Å—á–∏—Ç—ã–≤–∞–µ—Ç—Å—è –∑–¥–µ—Å—å
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –¥–ª—è —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞
            # –ü—Ä–∏–º–µ–Ω—è–µ–º —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏—é –∫ —Ç–æ–ª—â–∏–Ω–µ –¥–Ω–∞
            bottom_extrapolation_coeff = self.read_bottom_thickness_extrapolation_coeff()
            if abs(bottom_extrapolation_coeff) > 0.0001:
                extrapolated_bottom_thickness = self.apply_extrapolation_to_buffer(self.bottom_thickness_buffer, bottom_extrapolation_coeff)
                print(f" [–≠–ö–°–¢–†–ê–ü–û–õ–Ø–¶–ò–Ø] –ü—Ä–∏–º–µ–Ω–µ–Ω –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç {bottom_extrapolation_coeff:.6f} –∫ —Ç–æ–ª—â–∏–Ω–µ –¥–Ω–∞")
            else:
                extrapolated_bottom_thickness = self.bottom_thickness_buffer
            
            max_bottom_thickness = max(extrapolated_bottom_thickness)
            min_bottom_thickness = min(extrapolated_bottom_thickness)
            avg_bottom_thickness = sum(extrapolated_bottom_thickness) / len(extrapolated_bottom_thickness)
            
            print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞:")
            print(f"   –ò–∑–º–µ—Ä–µ–Ω–∏–π: {len(self.body_diameter_buffer)}")
            print(f"   –î–∏–∞–º–µ—Ç—Ä –∫–æ—Ä–ø—É—Å–∞: –º–∞–∫—Å={max_body_diameter:.3f}–º–º, —Å—Ä–µ–¥={avg_body_diameter:.3f}–º–º, –º–∏–Ω={min_body_diameter:.3f}–º–º")
            print(f"   –î–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞: –º–∞–∫—Å={max_flange_diameter:.3f}–º–º, —Å—Ä–µ–¥={avg_flange_diameter:.3f}–º–º, –º–∏–Ω={min_flange_diameter:.3f}–º–º")
            print(f"   –¢–æ–ª—â–∏–Ω–∞ –¥–Ω–∞: –º–∞–∫—Å={max_bottom_thickness:.3f}–º–º, —Å—Ä–µ–¥={avg_bottom_thickness:.3f}–º–º, –º–∏–Ω={min_bottom_thickness:.3f}–º–º")
            
            # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã
            self.write_flange_measurement_results(
                max_body_diameter, avg_body_diameter, min_body_diameter,
                max_flange_diameter, avg_flange_diameter, min_flange_diameter,
                max_bottom_thickness, avg_bottom_thickness, min_bottom_thickness
            )
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞: {e}")
    
    def write_flange_measurement_results(self, 
                                       max_body_diameter: float, avg_body_diameter: float, min_body_diameter: float,
                                       max_flange_diameter: float, avg_flange_diameter: float, min_flange_diameter: float,
                                       max_bottom_thickness: float, avg_bottom_thickness: float, min_bottom_thickness: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞ –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –î–∏–∞–º–µ—Ç—Ä –∫–æ—Ä–ø—É—Å–∞ ‚Üí 30046-30051
                self.write_stream_result_to_input_registers(max_body_diameter, 30046)   # –ú–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–µ
                self.write_stream_result_to_input_registers(avg_body_diameter, 30048)   # –°—Ä–µ–¥–Ω–µ–µ
                self.write_stream_result_to_input_registers(min_body_diameter, 30050)   # –ú–∏–Ω–∏–º–∞–ª—å–Ω–æ–µ
                
                # –î–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞ ‚Üí 30052-30057
                self.write_stream_result_to_input_registers(max_flange_diameter, 30054) # –ú–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–µ
                self.write_stream_result_to_input_registers(avg_flange_diameter, 30052) # –°—Ä–µ–¥–Ω–µ–µ
                self.write_stream_result_to_input_registers(min_flange_diameter, 30056) # –ú–∏–Ω–∏–º–∞–ª—å–Ω–æ–µ
                
                # –¢–æ–ª—â–∏–Ω–∞ —Ñ–ª–∞–Ω—Ü–∞ —Ç–µ–ø–µ—Ä—å –ø–µ—Ä–µ–¥–∞—ë—Ç—Å—è —Å –ü–ö, –Ω–µ –∑–∞–ø–∏—Å—ã–≤–∞–µ—Ç—Å—è –∑–¥–µ—Å—å
                
                # –¢–æ–ª—â–∏–Ω–∞ –¥–Ω–∞ ‚Üí 30028-30033
                self.write_stream_result_to_input_registers(max_bottom_thickness, 30028) # –ú–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–µ
                self.write_stream_result_to_input_registers(avg_bottom_thickness, 30030) # –°—Ä–µ–¥–Ω–µ–µ
                self.write_stream_result_to_input_registers(min_bottom_thickness, 30032) # –ú–∏–Ω–∏–º–∞–ª—å–Ω–æ–µ
                
                print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã —Ñ–ª–∞–Ω—Ü–∞ –∑–∞–ø–∏—Å–∞–Ω—ã –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞: {e}")

    def calculate_diameter_stats_from_radii(self, radii_buffer: list, extrapolation_coeff: float, offset_coeff: float, label: str):
        """–†–∞—Å—á—ë—Ç max/avg/min –¥–∏–∞–º–µ—Ç—Ä–∞ –ø–æ –±—É—Ñ–µ—Ä—É —Ä–∞–¥–∏—É—Å–æ–≤"""
        valid_radii = [r for r in radii_buffer if r is not None and r > 0 and not (math.isnan(r) or math.isinf(r))]
        if len(valid_radii) == 0:
            print(f" –û–®–ò–ë–ö–ê: –ù–µ—Ç –≤–∞–ª–∏–¥–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π —Ä–∞–¥–∏—É—Å–∞ ({label})!")
            return None

        if abs(extrapolation_coeff) > 0.0001:
            valid_radii = self.apply_extrapolation_to_buffer(valid_radii, extrapolation_coeff)
            print(f" [–≠–ö–°–¢–†–ê–ü–û–õ–Ø–¶–ò–Ø] {label}: –ø—Ä–∏–º–µ–Ω–µ–Ω –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç {extrapolation_coeff:.6f}")

        if len(valid_radii) >= 2:
            total_measurements = len(valid_radii)
            half_size = total_measurements // 2
            opposite_diameters = []
            for i in range(half_size):
                diameter_val = valid_radii[i] + valid_radii[i + half_size] + offset_coeff
                opposite_diameters.append(diameter_val)

            if len(opposite_diameters) > 0:
                max_val = max(opposite_diameters)
                min_val = min(opposite_diameters)
                avg_val = sum(opposite_diameters) / len(opposite_diameters)
            else:
                max_val = max(valid_radii) * 2 + offset_coeff
                min_val = min(valid_radii) * 2 + offset_coeff
                avg_val = (sum(valid_radii) / len(valid_radii)) * 2 + offset_coeff
        else:
            max_val = max(valid_radii) * 2 + offset_coeff
            min_val = min(valid_radii) * 2 + offset_coeff
            avg_val = (sum(valid_radii) / len(valid_radii)) * 2 + offset_coeff

        return max_val, avg_val, min_val

    def process_flange_only_measurement_results(self):
        """–ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞ –∏ —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞ (CMD=21)"""
        if len(self.flange_diameter_buffer) == 0:
            print(" –û—à–∏–±–∫–∞: –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞")
            return
        if len(self.bottom_thickness_buffer) == 0:
            print(" –û—à–∏–±–∫–∞: –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞")
            return

        flange_stats = self.calculate_diameter_stats_from_radii(
            self.flange_diameter_buffer,
            self.read_flange_diameter_extrapolation_coeff(),
            self.read_flange_diameter_offset_coeff(),
            "—Ñ–ª–∞–Ω–µ—Ü (—Ä–∞–∑–¥–µ–ª—å–Ω–æ)"
        )
        if not flange_stats:
            return
        max_flange, avg_flange, min_flange = flange_stats

        bottom_extrapolation_coeff = self.read_bottom_thickness_extrapolation_coeff()
        if abs(bottom_extrapolation_coeff) > 0.0001:
            extrapolated_bottom = self.apply_extrapolation_to_buffer(self.bottom_thickness_buffer, bottom_extrapolation_coeff)
            print(f" [–≠–ö–°–¢–†–ê–ü–û–õ–Ø–¶–ò–Ø] –¢–æ–ª—â–∏–Ω–∞ –¥–Ω–∞ (—Ä–∞–∑–¥–µ–ª—å–Ω–æ): –ø—Ä–∏–º–µ–Ω–µ–Ω –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç {bottom_extrapolation_coeff:.6f}")
        else:
            extrapolated_bottom = self.bottom_thickness_buffer

        max_bottom = max(extrapolated_bottom)
        avg_bottom = sum(extrapolated_bottom) / len(extrapolated_bottom)
        min_bottom = min(extrapolated_bottom)

        self.write_flange_only_measurement_results(
            max_flange, avg_flange, min_flange,
            max_bottom, avg_bottom, min_bottom
        )
        print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ —Ñ–ª–∞–Ω—Ü–∞: –º–∞–∫—Å={max_flange:.3f}, —Å—Ä–µ–¥={avg_flange:.3f}, –º–∏–Ω={min_flange:.3f}")
        print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞ (—Ä–∞–∑–¥–µ–ª—å–Ω–æ): –º–∞–∫—Å={max_bottom:.3f}, —Å—Ä–µ–¥={avg_bottom:.3f}, –º–∏–Ω={min_bottom:.3f}")

    def process_body_only_measurement_results(self):
        """–ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ (CMD=31)"""
        if len(self.body_only_diameter_buffer) == 0:
            print(" –û—à–∏–±–∫–∞: –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞")
            return

        stats = self.calculate_diameter_stats_from_radii(
            self.body_only_diameter_buffer,
            self.read_body_diameter_extrapolation_coeff(),
            self.read_body_diameter_offset_coeff(),
            "–∫–æ—Ä–ø—É—Å (—Ä–∞–∑–¥–µ–ª—å–Ω–æ)"
        )
        if not stats:
            return
        max_val, avg_val, min_val = stats
        self.write_body_only_measurement_results(max_val, avg_val, min_val)
        print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞: –º–∞–∫—Å={max_val:.3f}, —Å—Ä–µ–¥={avg_val:.3f}, –º–∏–Ω={min_val:.3f}")

    def process_body2_measurement_results(self):
        """–ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 (CMD=41)"""
        if len(self.body2_diameter_buffer) == 0:
            print(" –û—à–∏–±–∫–∞: –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2")
            return

        stats = self.calculate_diameter_stats_from_radii(
            self.body2_diameter_buffer,
            self.read_body2_diameter_extrapolation_coeff(),
            self.read_body2_diameter_offset_coeff(),
            "–∫–æ—Ä–ø—É—Å 2"
        )
        if not stats:
            return
        max_val, avg_val, min_val = stats
        self.write_body2_measurement_results(max_val, avg_val, min_val)
        print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2: –º–∞–∫—Å={max_val:.3f}, —Å—Ä–µ–¥={avg_val:.3f}, –º–∏–Ω={min_val:.3f}")

    def write_flange_only_measurement_results(
        self,
        max_flange: float, avg_flange: float, min_flange: float,
        max_bottom: float, avg_bottom: float, min_bottom: float
    ):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞ –∏ —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞"""
        # –î–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞ ‚Üí 30052-30057
        self.write_stream_result_to_input_registers(max_flange, 30054)
        self.write_stream_result_to_input_registers(avg_flange, 30052)
        self.write_stream_result_to_input_registers(min_flange, 30056)
        # –¢–æ–ª—â–∏–Ω–∞ –¥–Ω–∞ ‚Üí 30028-30033
        self.write_stream_result_to_input_registers(max_bottom, 30028)
        self.write_stream_result_to_input_registers(avg_bottom, 30030)
        self.write_stream_result_to_input_registers(min_bottom, 30032)

    def write_body_only_measurement_results(self, max_val: float, avg_val: float, min_val: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ –≤ 30046-30051"""
        self.write_stream_result_to_input_registers(max_val, 30046)
        self.write_stream_result_to_input_registers(avg_val, 30048)
        self.write_stream_result_to_input_registers(min_val, 30050)

    def write_body2_measurement_results(self, max_val: float, avg_val: float, min_val: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 –≤ 30059-30064"""
        self.write_stream_result_to_input_registers(max_val, 30059)
        self.write_stream_result_to_input_registers(avg_val, 30061)
        self.write_stream_result_to_input_registers(min_val, 30063)
    
    def process_bottom_wall_measurement_results(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ –ø—Ä–∏ –ø–µ—Ä–µ—Ö–æ–¥–µ 12‚Üí0"""
        try:
            if len(self.bottom_wall_thickness_buffer) == 0:
                print(" –û—à–∏–±–∫–∞: –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏")
                return
            
            # –ß–∏—Ç–∞–µ–º –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏–∏ –¥–ª—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
            extrapolation_coeff = self.read_bottom_wall_extrapolation_coeff()
            
            # –ü—Ä–∏–º–µ–Ω—è–µ–º —ç–∫—Å—Ç—Ä–∞–ø–æ–ª—è—Ü–∏—é –∫ –±—É—Ñ–µ—Ä—É
            if abs(extrapolation_coeff) > 0.0001:  # –ü—Ä–∏–º–µ–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –µ—Å–ª–∏ –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç –Ω–µ —Ä–∞–≤–µ–Ω –Ω—É–ª—é
                extrapolated_buffer = self.apply_extrapolation_to_buffer(self.bottom_wall_thickness_buffer, extrapolation_coeff)
                print(f" [–≠–ö–°–¢–†–ê–ü–û–õ–Ø–¶–ò–Ø] –ü—Ä–∏–º–µ–Ω–µ–Ω –∫–æ—ç—Ñ—Ñ–∏—Ü–∏–µ–Ω—Ç {extrapolation_coeff:.6f} –∫ —Ç–æ–ª—â–∏–Ω–µ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏")
            else:
                extrapolated_buffer = self.bottom_wall_thickness_buffer
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –∏–∑ —ç–∫—Å—Ç—Ä–∞–ø–æ–ª–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π
            max_bottom_wall_thickness = max(extrapolated_buffer)
            min_bottom_wall_thickness = min(extrapolated_buffer)
            avg_bottom_wall_thickness = sum(extrapolated_buffer) / len(extrapolated_buffer)
            
            print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏:")
            print(f"   –ò–∑–º–µ—Ä–µ–Ω–∏–π: {len(self.bottom_wall_thickness_buffer)}")
            print(f"   –ú–∞–∫—Å–∏–º—É–º: {max_bottom_wall_thickness:.3f}–º–º")
            print(f"   –°—Ä–µ–¥–Ω–µ–µ:  {avg_bottom_wall_thickness:.3f}–º–º")
            print(f"   –ú–∏–Ω–∏–º—É–º:  {min_bottom_wall_thickness:.3f}–º–º")
            
            # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã
            self.write_bottom_wall_measurement_results(max_bottom_wall_thickness, avg_bottom_wall_thickness, min_bottom_wall_thickness)
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏: {e}")
    
    def write_bottom_wall_measurement_results(self, max_val: float, avg_val: float, min_val: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ú–∞–∫—Å–∏–º–∞–ª—å–Ω–∞—è —Ç–æ–ª—â–∏–Ω–∞ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ ‚Üí 30022-30023
                self.write_stream_result_to_input_registers(max_val, 30022)
                
                # –°—Ä–µ–¥–Ω—è—è —Ç–æ–ª—â–∏–Ω–∞ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ ‚Üí 30024-30025
                self.write_stream_result_to_input_registers(avg_val, 30024)
                
                # –ú–∏–Ω–∏–º–∞–ª—å–Ω–∞—è —Ç–æ–ª—â–∏–Ω–∞ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ ‚Üí 30026-30027
                self.write_stream_result_to_input_registers(min_val, 30026)
                
                print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ –∑–∞–ø–∏—Å–∞–Ω—ã: –º–∞–∫—Å={max_val:.3f}, —Å—Ä–µ–¥={avg_val:.3f}, –º–∏–Ω={min_val:.3f}")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏: {e}")
    
    def write_doubleword_to_input_registers(self, value: float, base_address: int):
        """–ó–∞–ø–∏—Å—å DoubleWord float –≤ Input —Ä–µ–≥–∏—Å—Ç—Ä—ã"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º –≤ DoubleWord
                low_word, high_word = self.float_to_doubleword(value)
                
                # –í—ã—á–∏—Å–ª—è–µ–º –∏–Ω–¥–µ–∫—Å—ã (–º–ª–∞–¥—à–∏–π –≤ base_address, —Å—Ç–∞—Ä—à–∏–π –≤ base_address+1)
                reg_index_low = base_address - 30000       # –ú–ª–∞–¥—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä (base_address)
                reg_index_high = base_address - 30000 + 1  # –°—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä (base_address+1)
                
                # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º (–º–ª–∞–¥—à–∏–π –≤ –ø–µ—Ä–≤—ã–π —Ä–µ–≥–∏—Å—Ç—Ä, —Å—Ç–∞—Ä—à–∏–π –≤–æ –≤—Ç–æ—Ä–æ–π)
                self.modbus_server.slave_context.setValues(4, reg_index_low, [int(low_word)])
                self.modbus_server.slave_context.setValues(4, reg_index_high, [int(high_word)])
                
                print(f" –ó–∞–ø–∏—Å–∞–Ω–æ –≤ {base_address}-{base_address+1}: {value:.3f}–º–º (low={low_word}, high={high_word})")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ DoubleWord –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã {base_address}-{base_address+1}: {e}")
    
    def handle_measure_flange_state(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞ (CMD = 12) - –¢–û–õ–¨–ö–û –°–ë–û–† –î–ê–ù–ù–´–•"""
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            self.current_state = SystemState.ERROR
            return
        
        # –£–±—Ä–∞–Ω–∞ –ø—Ä–æ–≤–µ—Ä–∫–∞ get_current_command() –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è
        # –ö–æ–º–∞–Ω–¥–∞ –º–µ–Ω—è–µ—Ç—Å—è —á–µ—Ä–µ–∑ handle_command(), –∫–æ—Ç–æ—Ä—ã–π —É–∂–µ –º–µ–Ω—è–µ—Ç current_state
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è —Ç–∞–π–º–µ—Ä–∞ —á–∞—Å—Ç–æ—Ç—ã –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–∏ —Ñ–ª–∞–Ω—Ü–∞
            if self.frequency_start_time is None:
                self.frequency_start_time = time.time()
                self.last_frequency_display = self.frequency_start_time
                self.frequency_counter = 0
            
            # –°—Ç–∞—Ç—É—Å —É–∂–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –≤ manage_measurement_cycle_flag
            # –ü—Ä–æ—Å—Ç–æ –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
            
            # –í—ã–ø–æ–ª–Ω—è–µ–º QUAD –∏–∑–º–µ—Ä–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–æ–≤ 1, 3 –∏ 4 (–±–µ–∑–æ–ø–∞—Å–Ω–æ–µ —á—Ç–µ–Ω–∏–µ —Å –±–ª–æ–∫–∏—Ä–æ–≤–∫–æ–π)
            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor1_mm is None or sensor3_mm is None or sensor4_mm is None:
                time.sleep(0.001)
                return
            
            # –£–≤–µ–ª–∏—á–∏–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫ –∏–∑–º–µ—Ä–µ–Ω–∏–π
            self.frequency_counter += 1
            
            # –í—ã–≤–æ–¥–∏–º —á–∞—Å—Ç–æ—Ç—É –∫–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
            current_time = time.time()
            if current_time - self.last_frequency_display >= 1.0:
                elapsed = current_time - self.frequency_start_time
                if elapsed > 0:
                    instant_freq = self.frequency_counter / elapsed
                    print(f" [CMD=12] –ß–∞—Å—Ç–æ—Ç–∞ –æ–ø—Ä–æ—Å–∞: {instant_freq:.1f} –ì—Ü | –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.frequency_counter}")
                self.last_frequency_display = current_time
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º —á—Ç–æ –ø–æ–ª—É—á–∏–ª–∏ –≤–∞–ª–∏–¥–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –æ—Ç –¥–∞—Ç—á–∏–∫–æ–≤ 1, 3 –∏ 4
            # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è (None, 0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞)
            if (self.is_valid_measurement(sensor1_mm) and 
                self.is_valid_measurement(sensor3_mm) and 
                self.is_valid_measurement(sensor4_mm)):
                # –î–æ–±–∞–≤–ª—è–µ–º –≤ –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è
                self.temp_sensor1_flange_buffer.append(sensor1_mm)
                self.temp_sensor3_buffer.append(sensor3_mm)
                self.temp_sensor4_buffer.append(sensor4_mm)
                
                # –ö–æ–≥–¥–∞ –Ω–∞–∫–æ–ø–∏–ª–æ—Å—å 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π - —É—Å—Ä–µ–¥–Ω—è–µ–º –∏ –∑–∞–ø–∏—Å—ã–≤–∞–µ–º
                if len(self.temp_sensor1_flange_buffer) >= 10:
                    # –§–∏–ª—å—Ç—Ä—É–µ–º –∞–Ω–æ–º–∞–ª—å–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è –ø–µ—Ä–µ–¥ —É—Å—Ä–µ–¥–Ω–µ–Ω–∏–µ–º (–∏—Å–ø–æ–ª—å–∑—É–µ–º –º–µ–¥–∏–∞–Ω–Ω—ã–π —Ñ–∏–ª—å—Ç—Ä)
                    # –°–æ—Ä—Ç–∏—Ä—É–µ–º –∑–Ω–∞—á–µ–Ω–∏—è –∏ –±–µ—Ä–µ–º –º–µ–¥–∏–∞–Ω—É –¥–ª—è –∫–∞–∂–¥–æ–≥–æ –¥–∞—Ç—á–∏–∫–∞
                    sorted_sensor1 = sorted(self.temp_sensor1_flange_buffer)
                    sorted_sensor3 = sorted(self.temp_sensor3_buffer)
                    sorted_sensor4 = sorted(self.temp_sensor4_buffer)
                    
                    # –í—ã—á–∏—Å–ª—è–µ–º –º–µ–¥–∏–∞–Ω—É (—Å—Ä–µ–¥–Ω–µ–µ –∏–∑ –¥–≤—É—Ö —Ü–µ–Ω—Ç—Ä–∞–ª—å–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π –¥–ª—è —á–µ—Ç–Ω–æ–≥–æ –∫–æ–ª–∏—á–µ—Å—Ç–≤–∞)
                    median_sensor1 = (sorted_sensor1[4] + sorted_sensor1[5]) / 2.0
                    median_sensor3 = (sorted_sensor3[4] + sorted_sensor3[5]) / 2.0
                    median_sensor4 = (sorted_sensor4[4] + sorted_sensor4[5]) / 2.0
                    
                    # –§–∏–ª—å—Ç—Ä—É–µ–º –∑–Ω–∞—á–µ–Ω–∏—è, –∫–æ—Ç–æ—Ä—ã–µ –æ—Ç–∫–ª–æ–Ω—è—é—Ç—Å—è –æ—Ç –º–µ–¥–∏–∞–Ω—ã –±–æ–ª–µ–µ —á–µ–º –Ω–∞ 1.5–º–º
                    # –≠—Ç–æ –ø–æ–º–æ–∂–µ—Ç –æ—Ç–±—Ä–æ—Å–∏—Ç—å –∞–Ω–æ–º–∞–ª—å–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è
                    filtered_sensor1 = [v for v in self.temp_sensor1_flange_buffer if abs(v - median_sensor1) <= 1.5]
                    filtered_sensor3 = [v for v in self.temp_sensor3_buffer if abs(v - median_sensor3) <= 1.5]
                    filtered_sensor4 = [v for v in self.temp_sensor4_buffer if abs(v - median_sensor4) <= 1.5]
                    
                    # –ï—Å–ª–∏ –ø–æ—Å–ª–µ —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏–∏ –æ—Å—Ç–∞–ª–æ—Å—å –º–µ–Ω–µ–µ 5 –∑–Ω–∞—á–µ–Ω–∏–π - –∏—Å–ø–æ–ª—å–∑—É–µ–º –º–µ–¥–∏–∞–Ω—É
                    # –ò–Ω–∞—á–µ –∏—Å–ø–æ–ª—å–∑—É–µ–º —Å—Ä–µ–¥–Ω–µ–µ –æ—Ç—Ñ–∏–ª—å—Ç—Ä–æ–≤–∞–Ω–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π
                    if len(filtered_sensor1) >= 5:
                        avg_sensor1 = sum(filtered_sensor1) / len(filtered_sensor1)
                    else:
                        avg_sensor1 = median_sensor1
                    
                    if len(filtered_sensor3) >= 5:
                        avg_sensor3 = sum(filtered_sensor3) / len(filtered_sensor3)
                    else:
                        avg_sensor3 = median_sensor3
                    
                    if len(filtered_sensor4) >= 5:
                        avg_sensor4 = sum(filtered_sensor4) / len(filtered_sensor4)
                    else:
                        avg_sensor4 = median_sensor4
                    
                    # –î–æ–±–∞–≤–ª—è–µ–º —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è –≤ –æ—Å–Ω–æ–≤–Ω—ã–µ –±—É—Ñ–µ—Ä—ã
                    self.sensor1_flange_measurements.append(avg_sensor1)
                    self.sensor3_measurements.append(avg_sensor3)
                    self.sensor4_measurements.append(avg_sensor4)
                    
                    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∫–µ—à–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è
                    distance_to_center = self.cached_distance_to_center
                    distance_to_center_flange = self.cached_distance_sensor3_to_center  # –†–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ —Ü–µ–Ω—Ç—Ä–∞ (–∏–∑ –∫–æ–º–∞–Ω–¥—ã 105)
                    distance_1_3 = self.cached_distance_1_3
                    distance_sensor4 = self.cached_distance_sensor4
                    recipe_diametr_body = self.read_recipe_body_diameter()
                    recipe_diametr_flange = self.read_recipe_flange_diameter()
                    body_diameter_offset = self.read_body_diameter_offset_coeff()
                    flange_diameter_offset = self.read_flange_diameter_offset_coeff()
                    bottom_thickness_offset = self.read_bottom_thickness_offset_coeff()
                    
                    if (distance_to_center is not None and distance_to_center_flange is not None and 
                        distance_1_3 is not None and distance_sensor4 is not None):
                        
                        # 1) –†–∞–¥–∏—É—Å –∫–æ—Ä–ø—É—Å–∞ (–î–∞—Ç—á–∏–∫ 1) - —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ä–∞–¥–∏—É—Å –¥–ª—è –ø–æ—Å–ª–µ–¥—É—é—â–µ–≥–æ —Ä–∞—Å—á–µ—Ç–∞ –∏–∑ –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫
                        # –§–æ—Ä–º—É–ª–∞: (—Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–æ —Ü–µ–Ω—Ç—Ä–∞ - –ø–æ–∫–∞–∑–∞–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 1)
                        body_radius = distance_to_center - avg_sensor1
                        self.body_diameter_buffer.append(body_radius)  # –í—Ä–µ–º–µ–Ω–Ω–æ –∏—Å–ø–æ–ª—å–∑—É–µ–º —Ç–æ—Ç –∂–µ –±—É—Ñ–µ—Ä –¥–ª—è —Ä–∞–¥–∏—É—Å–æ–≤
                        
                        # 2) –†–∞–¥–∏—É—Å —Ñ–ª–∞–Ω—Ü–∞ (–î–∞—Ç—á–∏–∫ 3) - —Å–æ—Ö—Ä–∞–Ω—è–µ–º —Ä–∞–¥–∏—É—Å –¥–ª—è –ø–æ—Å–ª–µ–¥—É—é—â–µ–≥–æ —Ä–∞—Å—á–µ—Ç–∞ –∏–∑ –ø—Ä–æ—Ç–∏–≤–æ–ø–æ–ª–æ–∂–Ω—ã—Ö —Ç–æ—á–µ–∫
                        # –§–æ—Ä–º—É–ª–∞: (—Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 3 –¥–æ —Ü–µ–Ω—Ç—Ä–∞ - –ø–æ–∫–∞–∑–∞–Ω–∏–µ –¥–∞—Ç—á–∏–∫–∞ 3)
                        flange_radius = distance_to_center_flange - avg_sensor3
                        self.flange_diameter_buffer.append(flange_radius)  # –í—Ä–µ–º–µ–Ω–Ω–æ –∏—Å–ø–æ–ª—å–∑—É–µ–º —Ç–æ—Ç –∂–µ –±—É—Ñ–µ—Ä –¥–ª—è —Ä–∞–¥–∏—É—Å–æ–≤
                        
                        # 3) –¢–æ–ª—â–∏–Ω–∞ —Ñ–ª–∞–Ω—Ü–∞ - —Ç–µ–ø–µ—Ä—å –ø–µ—Ä–µ–¥–∞—ë—Ç—Å—è —Å –ü–ö, –Ω–µ —Ä–∞—Å—Å—á–∏—Ç—ã–≤–∞–µ—Ç—Å—è –∑–¥–µ—Å—å
                        
                        # 4) –¢–æ–ª—â–∏–Ω–∞ –¥–Ω–∞ (–î–∞—Ç—á–∏–∫ 4)
                        bottom_thickness = distance_sensor4 - avg_sensor4 + bottom_thickness_offset
                        self.bottom_thickness_buffer.append(bottom_thickness)
                        
                        # –í—ã–≤–æ–¥–∏–º –ø—Ä–æ–≥—Ä–µ—Å—Å –∫–∞–∂–¥—ã–µ 100 —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π
                        if len(self.body_diameter_buffer) % 100 == 0:
                            print(f" [CMD=12] –°–æ–±—Ä–∞–Ω–æ: {len(self.body_diameter_buffer)} –∏–∑–º–µ—Ä–µ–Ω–∏–π")
                            print(f"   –†–∞–¥–∏—É—Å –∫–æ—Ä–ø—É—Å–∞={body_radius:.3f}–º–º, –†–∞–¥–∏—É—Å —Ñ–ª–∞–Ω—Ü–∞={flange_radius:.3f}–º–º")
                            print(f"   –¢–æ–ª—â–∏–Ω–∞ –¥–Ω–∞={bottom_thickness:.3f}–º–º")
                    else:
                        print(" –û—à–∏–±–∫–∞: –Ω–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è")
                    
                    # –û—á–∏—â–∞–µ–º –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —Å–ª–µ–¥—É—é—â–∏—Ö 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π
                    self.temp_sensor1_flange_buffer = []
                    self.temp_sensor3_buffer = []
                    self.temp_sensor4_buffer = []
            else:
                print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è: –¥–∞—Ç—á–∏–∫1={sensor1_mm}, –¥–∞—Ç—á–∏–∫3={sensor3_mm}, –¥–∞—Ç—á–∏–∫4={sensor4_mm}")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞: {e}")
            self.current_state = SystemState.ERROR
    
    def handle_measure_bottom_state(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ (CMD = 14) - –¢–û–õ–¨–ö–û –°–ë–û–† –î–ê–ù–ù–´–•"""
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            self.current_state = SystemState.ERROR
            return
        
        # –£–±—Ä–∞–Ω–∞ –ø—Ä–æ–≤–µ—Ä–∫–∞ get_current_command() –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è
        # –ö–æ–º–∞–Ω–¥–∞ –º–µ–Ω—è–µ—Ç—Å—è —á–µ—Ä–µ–∑ handle_command(), –∫–æ—Ç–æ—Ä—ã–π —É–∂–µ –º–µ–Ω—è–µ—Ç current_state
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è —Ç–∞–π–º–µ—Ä–∞ —á–∞—Å—Ç–æ—Ç—ã –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–∏ –¥–Ω–∞
            # –í–∞–∂–Ω–æ: –≤—Å–µ–≥–¥–∞ —Å–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫–∏ –ø—Ä–∏ –≤—Ö–æ–¥–µ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
            # —á—Ç–æ–±—ã –Ω–µ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å –¥–∞–Ω–Ω—ã–µ –æ—Ç –ø—Ä–µ–¥—ã–¥—É—â–µ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
            if not hasattr(self, '_bottom_frequency_initialized') or self.frequency_start_time is None:
                self.frequency_start_time = time.time()
                self.last_frequency_display = self.frequency_start_time
                self.frequency_counter = 0
                self._bottom_frequency_initialized = True
            
            # –°—Ç–∞—Ç—É—Å —É–∂–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –≤ manage_measurement_cycle_flag
            # –ü—Ä–æ—Å—Ç–æ –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
            
            # –í—ã–ø–æ–ª–Ω—è–µ–º QUAD –∏–∑–º–µ—Ä–µ–Ω–∏–µ –¥–∞—Ç—á–∏–∫–æ–≤ 1 –∏ 2 (–±–µ–∑–æ–ø–∞—Å–Ω–æ–µ —á—Ç–µ–Ω–∏–µ —Å –±–ª–æ–∫–∏—Ä–æ–≤–∫–æ–π)
            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor1_mm is None or sensor2_mm is None:
                time.sleep(0.001)
                return
            
            # –£–≤–µ–ª–∏—á–∏–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫ –∏–∑–º–µ—Ä–µ–Ω–∏–π
            self.frequency_counter += 1
            
            # –í—ã–≤–æ–¥–∏–º —á–∞—Å—Ç–æ—Ç—É –∫–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
            current_time = time.time()
            if current_time - self.last_frequency_display >= 1.0:
                elapsed = current_time - self.frequency_start_time
                if elapsed > 0:
                    instant_freq = self.frequency_counter / elapsed
                    print(f" [CMD=14] –ß–∞—Å—Ç–æ—Ç–∞ –æ–ø—Ä–æ—Å–∞: {instant_freq:.1f} –ì—Ü | –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.frequency_counter}")
                self.last_frequency_display = current_time
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º —á—Ç–æ –ø–æ–ª—É—á–∏–ª–∏ –≤–∞–ª–∏–¥–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –æ—Ç –¥–∞—Ç—á–∏–∫–æ–≤ 1 –∏ 2
            # –§–∏–ª—å—Ç—Ä—É–µ–º –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è (None, 0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞)
            if (self.is_valid_measurement(sensor1_mm) and self.is_valid_measurement(sensor2_mm)):
                # –î–æ–±–∞–≤–ª—è–µ–º –≤ –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è
                self.temp_sensor1_bottom_buffer.append(sensor1_mm)
                self.temp_sensor2_bottom_buffer.append(sensor2_mm)
                
                # –ö–æ–≥–¥–∞ –Ω–∞–∫–æ–ø–∏–ª–æ—Å—å 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π - —É—Å—Ä–µ–¥–Ω—è–µ–º –∏ –∑–∞–ø–∏—Å—ã–≤–∞–µ–º
                if len(self.temp_sensor1_bottom_buffer) >= 10:
                    # –§–∏–ª—å—Ç—Ä—É–µ–º –∞–Ω–æ–º–∞–ª—å–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è –ø–µ—Ä–µ–¥ —É—Å—Ä–µ–¥–Ω–µ–Ω–∏–µ–º (–∏—Å–ø–æ–ª—å–∑—É–µ–º –º–µ–¥–∏–∞–Ω–Ω—ã–π —Ñ–∏–ª—å—Ç—Ä)
                    sorted_sensor1 = sorted(self.temp_sensor1_bottom_buffer)
                    sorted_sensor2 = sorted(self.temp_sensor2_bottom_buffer)
                    
                    # –í—ã—á–∏—Å–ª—è–µ–º –º–µ–¥–∏–∞–Ω—É
                    median_sensor1 = (sorted_sensor1[4] + sorted_sensor1[5]) / 2.0
                    median_sensor2 = (sorted_sensor2[4] + sorted_sensor2[5]) / 2.0
                    
                    # –§–∏–ª—å—Ç—Ä—É–µ–º –∑–Ω–∞—á–µ–Ω–∏—è, –∫–æ—Ç–æ—Ä—ã–µ –æ—Ç–∫–ª–æ–Ω—è—é—Ç—Å—è –æ—Ç –º–µ–¥–∏–∞–Ω—ã –±–æ–ª–µ–µ —á–µ–º –Ω–∞ 1.5–º–º
                    filtered_sensor1 = [v for v in self.temp_sensor1_bottom_buffer if abs(v - median_sensor1) <= 1.5]
                    filtered_sensor2 = [v for v in self.temp_sensor2_bottom_buffer if abs(v - median_sensor2) <= 1.5]
                    
                    # –ï—Å–ª–∏ –ø–æ—Å–ª–µ —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏–∏ –æ—Å—Ç–∞–ª–æ—Å—å –º–µ–Ω–µ–µ 5 –∑–Ω–∞—á–µ–Ω–∏–π - –∏—Å–ø–æ–ª—å–∑—É–µ–º –º–µ–¥–∏–∞–Ω—É
                    if len(filtered_sensor1) >= 5:
                        avg_sensor1 = sum(filtered_sensor1) / len(filtered_sensor1)
                    else:
                        avg_sensor1 = median_sensor1
                    
                    if len(filtered_sensor2) >= 5:
                        avg_sensor2 = sum(filtered_sensor2) / len(filtered_sensor2)
                    else:
                        avg_sensor2 = median_sensor2
                    
                    # –î–æ–±–∞–≤–ª—è–µ–º —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è –≤ –æ—Å–Ω–æ–≤–Ω—ã–µ –±—É—Ñ–µ—Ä—ã
                    self.sensor1_bottom_measurements.append(avg_sensor1)
                    self.sensor2_bottom_measurements.append(avg_sensor2)
                    
                    lower_wall_offset = self.read_lower_wall_offset_coeff()
                    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∫–µ—à–∏—Ä–æ–≤–∞–Ω–Ω–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ (–≤–º–µ—Å—Ç–æ —á—Ç–µ–Ω–∏—è –∏–∑ Modbus)
                    distance_1_2 = self.cached_distance_1_2
                    
                    if distance_1_2 is not None:
                        # –í—ã—á–∏—Å–ª—è–µ–º —Ç–æ–ª—â–∏–Ω—É –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏ –ø–æ —Ñ–æ—Ä–º—É–ª–µ (–∏—Å–ø–æ–ª—å–∑—É–µ–º —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è)
                        bottom_wall_thickness = distance_1_2 - avg_sensor1 - avg_sensor2 + lower_wall_offset
                        self.bottom_wall_thickness_buffer.append(bottom_wall_thickness)
                        
                        # –í—ã–≤–æ–¥–∏–º —Ç–µ–∫—É—â–∏–µ –∑–Ω–∞—á–µ–Ω–∏—è –∫–∞–∂–¥—ã–µ 100 —É—Å—Ä–µ–¥–Ω–µ–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏–π (—É–º–µ–Ω—å—à–µ–Ω–∞ —á–∞—Å—Ç–æ—Ç–∞ –¥–ª—è —É—Å–∫–æ—Ä–µ–Ω–∏—è)
                        if len(self.bottom_wall_thickness_buffer) % 100 == 0:
                            print(f" –£—Å—Ä–µ–¥–Ω–µ–Ω–Ω–æ–µ –∏–∑–º–µ—Ä–µ–Ω–∏–µ #{len(self.bottom_wall_thickness_buffer)}: "
                                  f"–î–∞—Ç—á–∏–∫1={avg_sensor1:.3f}–º–º, –î–∞—Ç—á–∏–∫2={avg_sensor2:.3f}–º–º, "
                                  f"–¢–æ–ª—â–∏–Ω–∞ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏={bottom_wall_thickness:.3f}–º–º")
                    else:
                        print(" –û—à–∏–±–∫–∞: –Ω–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ 1,2")
                    
                    # –û—á–∏—â–∞–µ–º –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —Å–ª–µ–¥—É—é—â–∏—Ö 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π
                    self.temp_sensor1_bottom_buffer = []
                    self.temp_sensor2_bottom_buffer = []
            else:
                print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è: –¥–∞—Ç—á–∏–∫1={sensor1_mm}, –¥–∞—Ç—á–∏–∫2={sensor2_mm}")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏: {e}")
            self.current_state = SystemState.ERROR
    
    # ===== –ù–û–í–´–ï –ú–ï–¢–û–î–´ –î–õ–Ø –†–ê–ó–î–ï–õ–Å–ù–ù–û–ì–û –¶–ò–ö–õ–ê –ò–ó–ú–ï–†–ï–ù–ò–Ø =====
    
    def handle_measure_height_process_state(self):
        """
        CMD=9: –ò–∑–º–µ—Ä–µ–Ω–∏–µ –≤—ã—Å–æ—Ç—ã - –ø–æ–∏—Å–∫ –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏—è –∏ —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö
        """
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            self.current_state = SystemState.ERROR
            return
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è —Ç–∞–π–º–µ—Ä–∞ —á–∞—Å—Ç–æ—Ç—ã –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–∏
            if self.frequency_start_time is None:
                self.frequency_start_time = time.time()
                self.last_frequency_display = self.frequency_start_time
                self.frequency_counter = 0
                self.obstacle_detected = False
                self.obstacle_filter_count = 0
                self.height_measurements = []
                print(" [CMD=9] –ù–∞—á–∞–ª–æ –ø–æ–∏—Å–∫–∞ –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏—è...")
            
            # –ß–∏—Ç–∞–µ–º —Ç–æ–ª—å–∫–æ –¥–∞—Ç—á–∏–∫ 1 –¥–ª—è –ø–æ–∏—Å–∫–∞ –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏—è
            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if sensor1_mm is None:
                time.sleep(0.001)
                return
            
            # –£–≤–µ–ª–∏—á–∏–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫ –∏–∑–º–µ—Ä–µ–Ω–∏–π
            self.frequency_counter += 1
            
            # –í—ã–≤–æ–¥–∏–º —á–∞—Å—Ç–æ—Ç—É –∫–∞–∂–¥—É—é —Å–µ–∫—É–Ω–¥—É
            current_time = time.time()
            if current_time - self.last_frequency_display >= 1.0:
                elapsed = current_time - self.frequency_start_time
                if elapsed > 0:
                    instant_freq = self.frequency_counter / elapsed
                    status = "–ü–æ–∏—Å–∫ –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏—è" if not self.obstacle_detected else "–°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –≤—ã—Å–æ—Ç—ã"
                    print(f" [CMD=9] {status}: {instant_freq:.1f} –ì—Ü | –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.frequency_counter}")
                self.last_frequency_display = current_time
            
            # –ü–†–ê–í–ò–õ–¨–ù–ê–Ø –õ–û–ì–ò–ö–ê: 
            # –ü–æ–∫–∞ –¥–∞—Ç—á–∏–∫ = 0 ‚Üí –∏—â–µ–º –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏–µ
            # –ö–∞–∫ —Ç–æ–ª—å–∫–æ –¥–∞—Ç—á–∏–∫ != 0 (5 —Ä–∞–∑ –ø–æ–¥—Ä—è–¥) ‚Üí –Ω–∞—à–ª–∏ –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏–µ, —Å—á–∏—Ç–∞–µ–º –≤—ã—Å–æ—Ç—É
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –≤–∞–ª–∏–¥–Ω–æ—Å—Ç—å –∏–∑–º–µ—Ä–µ–Ω–∏—è (—Ñ–∏–ª—å—Ç—Ä—É–µ–º None, 0, –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–µ, –≤–Ω–µ –¥–∏–∞–ø–∞–∑–æ–Ω–∞)
            if self.is_valid_measurement(sensor1_mm):
                # –î–∞—Ç—á–∏–∫ –ø–æ–∫–∞–∑—ã–≤–∞–µ—Ç –≤–∞–ª–∏–¥–Ω–æ–µ –Ω–µ–Ω—É–ª–µ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ - –µ—Å—Ç—å –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏–µ!
                self.obstacle_filter_count += 1
                
                if not self.obstacle_detected and self.obstacle_filter_count >= 5:
                    # 5 –Ω–µ–Ω—É–ª–µ–≤—ã—Ö –ø–æ–∫–∞–∑–∞–Ω–∏–π –ø–æ–¥—Ä—è–¥ - –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏–µ –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–æ!
                    self.obstacle_detected = True
                    self.write_cycle_flag(90)  # –°—Ç–∞—Ç—É—Å: –ø—Ä–µ–ø—è—Ç—Å—Ç–≤–∏–µ –Ω–∞–π–¥–µ–Ω–æ
                    print(f" [CMD=9] –ü—Ä–µ–ø—è—Ç—Å—Ç–≤–∏–µ –Ω–∞–π–¥–µ–Ω–æ! –î–∞—Ç—á–∏–∫ 1 = {sensor1_mm:.3f}–º–º (5 –ø–æ–∫–∞–∑–∞–Ω–∏–π –ø–æ–¥—Ä—è–¥)")
                
                if self.obstacle_detected:
                    # –°–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞ –≤—ã—Å–æ—Ç—ã
                    self.collect_height_data()
            else:
                # –î–∞—Ç—á–∏–∫ = 0 –∏–ª–∏ None ‚Üí —Å–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫, –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º –ø–æ–∏—Å–∫
                self.obstacle_filter_count = 0
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤—ã—Å–æ—Ç—ã: {e}")
            self.current_state = SystemState.ERROR
    
    def collect_height_data(self):
        """–°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞ –≤—ã—Å–æ—Ç—ã"""
        try:
            # –ß–∏—Ç–∞–µ–º —Ä–µ–≥–∏—Å—Ç—Ä—ã –∫–∞–∫ –≤ –∫–æ–º–∞–Ω–¥–∞—Ö 103 –∏ 104
            steps = self.read_register_40020()  # –ö–æ–ª–∏—á–µ—Å—Ç–≤–æ —à–∞–≥–æ–≤ (40052-40053)
            pulses_per_mm = self.read_register_40021()  # –ò–º–ø—É–ª—å—Å–æ–≤ –Ω–∞ 1 –º–º (40054)
            distance_to_plane = self.read_register_40022_40023()  # –î–∏—Å—Ç–∞–Ω—Ü–∏—è –¥–æ –ø–ª–æ—Å–∫–æ—Å—Ç–∏ (40055-40056)
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º —á—Ç–æ –≤—Å–µ –∑–Ω–∞—á–µ–Ω–∏—è –≤–∞–ª–∏–¥–Ω—ã (–Ω–µ None –∏ –Ω–µ 0 –¥–ª—è pulses_per_mm)
            if steps is not None and pulses_per_mm is not None and pulses_per_mm > 0 and distance_to_plane is not None:
                # –†–∞—Å—Å—á–∏—Ç—ã–≤–∞–µ–º –≤—ã—Å–æ—Ç—É –ø–æ —Ñ–æ—Ä–º—É–ª–µ: height = distance_to_plane - (steps/pulses)
                height = distance_to_plane - (steps / pulses_per_mm)
                self.height_measurements.append(height)
                
                # –í—ã–≤–æ–¥–∏–º –ø—Ä–æ–≥—Ä–µ—Å—Å –∫–∞–∂–¥—ã–µ 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π —Å –æ—Ç–ª–∞–¥–æ—á–Ω–æ–π –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–µ–π
                if len(self.height_measurements) % 10 == 0:
                    # –ß–∏—Ç–∞–µ–º —Å—ã—Ä—ã–µ –¥–∞–Ω–Ω—ã–µ —à–∞–≥–æ–≤ –¥–ª—è –æ—Ç–ª–∞–¥–∫–∏
                    try:
                        raw_steps_values = self.modbus_server.slave_context.getValues(3, 52, 2)
                        raw_steps_str = f"—Å—ã—Ä—ã–µ [52-53]: {raw_steps_values}" if raw_steps_values else "None"
                    except:
                        raw_steps_str = "–û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è"
                    
                    print(f" [CMD=9] –°–æ–±—Ä–∞–Ω–æ: {len(self.height_measurements)} | "
                          f"–®–∞–≥–∏: {raw_steps_str} ‚Üí {steps}, –ò–º–ø/–º–º={pulses_per_mm}, –î–∏—Å—Ç={distance_to_plane:.3f}–º–º ‚Üí –í—ã—Å–æ—Ç–∞={height:.3f}–º–º")
                
                # –ï—Å–ª–∏ —Å–æ–±—Ä–∞–ª–∏ –¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö (–Ω–∞–ø—Ä–∏–º–µ—Ä, 50 –∏–∑–º–µ—Ä–µ–Ω–∏–π) –ò —Ñ–ª–∞–≥ –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω
                if len(self.height_measurements) >= 50 and not self.height_calculated:
                    self.calculate_and_save_height()
                    
        except Exception as e:
            # –ù–µ –≤—ã–≤–æ–¥–∏–º –æ—à–∏–±–∫—É –µ—Å–ª–∏ –ø—Ä–æ—Å—Ç–æ –Ω—É–ª–µ–≤—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è
            if "division by zero" not in str(e):
                print(f" [CMD=9] –û—à–∏–±–∫–∞ —Å–±–æ—Ä–∞ –¥–∞–Ω–Ω—ã—Ö –≤—ã—Å–æ—Ç—ã: {e}")
    
    def calculate_and_save_height(self):
        """–†–∞—Å—á–µ—Ç –∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤—ã—Å–æ—Ç—ã"""
        try:
            if len(self.height_measurements) == 0:
                print(" –û—à–∏–±–∫–∞: –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Ä–∞—Å—á–µ—Ç–∞ –≤—ã—Å–æ—Ç—ã")
                return
            
            # –í—ã—á–∏—Å–ª—è–µ–º —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É
            max_height = max(self.height_measurements)
            min_height = min(self.height_measurements)
            avg_height = sum(self.height_measurements) / len(self.height_measurements)
            
            # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã
            self.write_height_measurement_results(max_height, avg_height, min_height)
            
            # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è
            self.write_cycle_flag(91)  # –°—Ç–∞—Ç—É—Å: –≤—ã—Å–æ—Ç–∞ —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–∞
            self.height_calculated = True
            
            # –í—ã–≤–æ–¥–∏–º —Ñ–æ—Ä–º—É–ª—É —Ä–∞—Å—á–µ—Ç–∞ –≤—ã—Å–æ—Ç—ã
            print(f" [CMD=9] –§–û–†–ú–£–õ–ê: height = distance_to_plane - (steps / pulses_per_mm)")
            print(f" [CMD=9] –ì–¥–µ: distance_to_plane={self.read_register_40022_40023():.3f}–º–º (40055-40056)")
            print(f" [CMD=9]       steps={self.read_register_40020()} (40052-40053)")
            print(f" [CMD=9]       pulses_per_mm={self.read_register_40021()} (40054)")
            print(f" [CMD=9] –í—ã—Å–æ—Ç–∞ —Ä–∞—Å—Å—á–∏—Ç–∞–Ω–∞: –º–∞–∫—Å={max_height:.3f}–º–º, "
                  f"—Å—Ä–µ–¥={avg_height:.3f}–º–º, –º–∏–Ω={min_height:.3f}–º–º")
            print(f" [CMD=9] –ì–æ—Ç–æ–≤ –∫ —Å–ª–µ–¥—É—é—â–µ–π –∫–æ–º–∞–Ω–¥–µ (CMD=10)")
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —Ä–∞—Å—á–µ—Ç–∞ –≤—ã—Å–æ—Ç—ã: {e}")
    
    def read_register_40020(self) -> int:
        """–ß—Ç–µ–Ω–∏–µ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40052-40053 (–∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —à–∞–≥–æ–≤ - DoubleWord integer)"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 52, 2)  # 40052-40053
                if values and len(values) == 2:
                    # –ü—Ä–∞–≤–∏–ª—å–Ω–æ–µ –æ–±—ä–µ–¥–∏–Ω–µ–Ω–∏–µ 32-bit integer: —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ + –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ
                    steps = (int(values[0]) << 16) | int(values[1])
                    return steps
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40052-40053: {e}")
        return None
    
    def read_register_40021(self) -> int:
        """–ß—Ç–µ–Ω–∏–µ —Ä–µ–≥–∏—Å—Ç—Ä–∞ 40054 (–∏–º–ø—É–ª—å—Å–æ–≤ –Ω–∞ 1 –º–º)"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                value = self.modbus_server.slave_context.getValues(3, 54, 1)[0]  # 40054
                return int(value)
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–µ–≥–∏—Å—Ç—Ä–∞ 40054: {e}")
        return None
    
    def read_register_40022_40023(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40055-40056 (–¥–∏—Å—Ç–∞–Ω—Ü–∏—è –¥–æ –ø–ª–æ—Å–∫–æ—Å—Ç–∏)"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 55, 2)  # 40055-40056
                if values and len(values) == 2:
                    high_word = int(values[0])  # 40055 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40056 - –º–ª–∞–¥—à–∏–π
                    distance = self.doubleword_to_float(low_word, high_word)
                    return distance
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40055-40056: {e}")
        return None
    
    def read_measured_height(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40057-40058 (–∏–∑–º–µ—Ä–µ–Ω–Ω–∞—è –≤—ã—Å–æ—Ç–∞ –∑–∞–≥–æ—Ç–æ–≤–∫–∏)"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 57, 2)  # 40057-40058
                if values and len(values) == 2:
                    high_word = int(values[0])  # 40057 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40058 - –º–ª–∞–¥—à–∏–π
                    height = self.doubleword_to_float(low_word, high_word)
                    return height
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40057-40058: {e}")
        return None
    
    def read_measured_flange_thickness(self) -> float:
        """–ß—Ç–µ–Ω–∏–µ —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40059-40060 (–∏–∑–º–µ—Ä–µ–Ω–Ω–∞—è —Ç–æ–ª—â–∏–Ω–∞ —Ñ–ª–∞–Ω—Ü–∞)"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                values = self.modbus_server.slave_context.getValues(3, 59, 2)  # 40059-40060
                if values and len(values) == 2:
                    high_word = int(values[0])  # 40059 - —Å—Ç–∞—Ä—à–∏–π
                    low_word = int(values[1])   # 40060 - –º–ª–∞–¥—à–∏–π
                    thickness = self.doubleword_to_float(low_word, high_word)
                    return thickness
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤ 40059-40060: {e}")
        return None
    
    def write_height_measurement_results(self, max_val: float, avg_val: float, min_val: float):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤—ã—Å–æ—Ç—ã –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ú–∞–∫—Å–∏–º–∞–ª—å–Ω–∞—è –≤—ã—Å–æ—Ç–∞ ‚Üí 30040-30041
                self.write_stream_result_to_input_registers(max_val, 30040)
                
                # –°—Ä–µ–¥–Ω—è—è –≤—ã—Å–æ—Ç–∞ ‚Üí 30042-30043
                self.write_stream_result_to_input_registers(avg_val, 30042)
                
                # –ú–∏–Ω–∏–º–∞–ª—å–Ω–∞—è –≤—ã—Å–æ—Ç–∞ ‚Üí 30044-30045
                self.write_stream_result_to_input_registers(min_val, 30044)
                
                print(f" –†–µ–∑—É–ª—å—Ç–∞—Ç—ã –≤—ã—Å–æ—Ç—ã –∑–∞–ø–∏—Å–∞–Ω—ã: –º–∞–∫—Å={max_val:.3f}, —Å—Ä–µ–¥={avg_val:.3f}, –º–∏–Ω={min_val:.3f}")
                
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤—ã—Å–æ—Ç—ã: {e}")
    
    def handle_measure_wall_process_state(self):
        """
        CMD=10: –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        –ü—Ä–æ—Å—Ç–æ —Å–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ, –Ω–µ –¥–µ–ª–∞–µ–º –ø–æ–¥—Å—á—ë—Ç
        """
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π (–æ–¥–∏–Ω —Ä–∞–∑ –ø—Ä–∏ –≤—Ö–æ–¥–µ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–µ)
        if not hasattr(self, '_wall_measurement_started'):
            self._wall_measurement_started = True
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            print(f" [CMD=10] –ë—É—Ñ–µ—Ä—ã –æ—á–∏—â–µ–Ω—ã –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π")
        
        # –ó–∞–≥—Ä—É–∂–∞–µ–º –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –≤ –∫–µ—à (–æ–¥–∏–Ω —Ä–∞–∑ –ø—Ä–∏ –≤—Ö–æ–¥–µ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–µ)
        if self.cached_distance_1_2 is None:
            self.cached_distance_1_2 = self.read_calibrated_distance_1_2()
            print(f" –ó–∞–≥—Ä—É–∂–µ–Ω–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ 1-2: {self.cached_distance_1_2:.3f}–º–º")
        
        # –ü–µ—Ä–µ–Ω–∞–ø—Ä–∞–≤–ª—è–µ–º –Ω–∞ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π –º–µ—Ç–æ–¥ (–æ–Ω —É–∂–µ —Ç–æ–ª—å–∫–æ —Å–æ–±–∏—Ä–∞–µ—Ç –¥–∞–Ω–Ω—ã–µ)
        self.handle_measure_wall_state()
    
    def handle_calculate_wall_state(self):
        """
        CMD=11: –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        """
        try:
            # –í—ã–ø–æ–ª–Ω—è–µ–º –ø–æ–¥—Å—á—ë—Ç —Ç–æ–ª—å–∫–æ –æ–¥–∏–Ω —Ä–∞–∑
            if not self.wall_calculated:
                print(" [CMD=11] –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –≤–µ—Ä—Ö–Ω–µ–π —Å—Ç–µ–Ω–∫–∏...")
                
                # –í—ã–∑—ã–≤–∞–µ–º —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π –º–µ—Ç–æ–¥ –ø–æ–¥—Å—á—ë—Ç–∞
                self.process_wall_measurement_results()
                
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å "–≥–æ—Ç–æ–≤–æ –∫ —Å–ª–µ–¥—É—é—â–µ–π –∫–æ–º–∞–Ω–¥–µ"
                self.write_cycle_flag(110)
                print(" [STATUS=110] –ü–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –≥–æ—Ç–æ–≤ –∫ CMD=12")
                
                # –û—Ç–º–µ—á–∞–µ–º —á—Ç–æ —Ä–∞—Å—á—ë—Ç –≤—ã–ø–æ–ª–Ω–µ–Ω
                self.wall_calculated = True
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –ø–æ–¥—Å—á—ë—Ç–∞ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Å—Ç–µ–Ω–∫–∏: {e}")
            self.current_state = SystemState.ERROR
    
    def handle_measure_flange_process_state(self):
        """
        CMD=12: –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞
        –ü—Ä–æ—Å—Ç–æ —Å–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ, –Ω–µ –¥–µ–ª–∞–µ–º –ø–æ–¥—Å—á—ë—Ç
        """
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π (–æ–¥–∏–Ω —Ä–∞–∑ –ø—Ä–∏ –≤—Ö–æ–¥–µ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–µ)
        if not hasattr(self, '_flange_measurement_started'):
            self._flange_measurement_started = True
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            print(f" [CMD=12] –ë—É—Ñ–µ—Ä—ã –æ—á–∏—â–µ–Ω—ã –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π")
        
        # –ó–∞–≥—Ä—É–∂–∞–µ–º –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω—ã–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –≤ –∫–µ—à (–æ–¥–∏–Ω —Ä–∞–∑ –ø—Ä–∏ –≤—Ö–æ–¥–µ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–µ)
        if self.cached_distance_to_center is None:
            self.cached_distance_to_center = self.read_calibrated_distance_to_center()
            self.cached_distance_1_3 = self.read_calibrated_distance_1_3()
            self.cached_distance_sensor4 = self.read_calibrated_distance_sensor4()
            self.cached_distance_sensor3_to_center = self.read_calibrated_distance_sensor3_to_center()
            print(f" –ó–∞–≥—Ä—É–∂–µ–Ω—ã —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è: —Ü–µ–Ω—Ç—Ä={self.cached_distance_to_center:.3f}–º–º, "
                  f"1-3={self.cached_distance_1_3:.3f}–º–º, sensor4={self.cached_distance_sensor4:.3f}–º–º, "
                  f"sensor3_to_center={self.cached_distance_sensor3_to_center:.3f}–º–º")
        
        # –ü–µ—Ä–µ–Ω–∞–ø—Ä–∞–≤–ª—è–µ–º –Ω–∞ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π –º–µ—Ç–æ–¥
        self.handle_measure_flange_state()
    
    def handle_calculate_flange_state(self):
        """
        CMD=13: –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ñ–ª–∞–Ω—Ü–∞
        """
        try:
            # –í—ã–ø–æ–ª–Ω—è–µ–º –ø–æ–¥—Å—á—ë—Ç —Ç–æ–ª—å–∫–æ –æ–¥–∏–Ω —Ä–∞–∑
            if not self.flange_calculated:
                print(" [CMD=13] –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ñ–ª–∞–Ω—Ü–∞...")
                
                # –í—ã–∑—ã–≤–∞–µ–º —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π –º–µ—Ç–æ–¥ –ø–æ–¥—Å—á—ë—Ç–∞
                self.process_flange_measurement_results()
                
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å "–≥–æ—Ç–æ–≤–æ –∫ —Å–ª–µ–¥—É—é—â–µ–π –∫–æ–º–∞–Ω–¥–µ"
                self.write_cycle_flag(112)
                print(" [STATUS=112] –ü–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –≥–æ—Ç–æ–≤ –∫ CMD=14")
                
                # –û—Ç–º–µ—á–∞–µ–º —á—Ç–æ —Ä–∞—Å—á—ë—Ç –≤—ã–ø–æ–ª–Ω–µ–Ω
                self.flange_calculated = True
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –ø–æ–¥—Å—á—ë—Ç–∞ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ —Ñ–ª–∞–Ω—Ü–∞: {e}")
            self.current_state = SystemState.ERROR

    def collect_sensor3_radius_measurement(self, temp_buffer: list, radii_buffer: list, distance_to_center: float, cmd_label: str):
        """–°–±–æ—Ä –∏ —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏—è —Ä–∞–¥–∏—É—Å–∞ –ø–æ –¥–∞—Ç—á–∏–∫—É 3 —Å —É—Å—Ä–µ–¥–Ω–µ–Ω–∏–µ–º –ø–æ 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π"""
        sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
        if sensor3_mm is None:
            time.sleep(0.001)
            return

        self.frequency_counter += 1
        current_time = time.time()
        if self.frequency_start_time is not None and current_time - self.last_frequency_display >= 1.0:
            elapsed = current_time - self.frequency_start_time
            if elapsed > 0:
                instant_freq = self.frequency_counter / elapsed
                print(f" [{cmd_label}] –ß–∞—Å—Ç–æ—Ç–∞ –æ–ø—Ä–æ—Å–∞: {instant_freq:.1f} –ì—Ü | –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.frequency_counter}")
            self.last_frequency_display = current_time

        if not self.is_valid_measurement(sensor3_mm):
            return

        temp_buffer.append(sensor3_mm)
        if len(temp_buffer) < 10:
            return

        sorted_sensor3 = sorted(temp_buffer)
        median_sensor3 = (sorted_sensor3[4] + sorted_sensor3[5]) / 2.0
        filtered_sensor3 = [v for v in temp_buffer if abs(v - median_sensor3) <= 1.5]
        if len(filtered_sensor3) >= 5:
            avg_sensor3 = sum(filtered_sensor3) / len(filtered_sensor3)
        else:
            avg_sensor3 = median_sensor3

        radius = distance_to_center - avg_sensor3
        radii_buffer.append(radius)
        temp_buffer.clear()

    def collect_flange_and_bottom_measurement(self):
        """–°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –¥–ª—è CMD=20: –¥–∏–∞–º–µ—Ç—Ä —Ñ–ª–∞–Ω—Ü–∞ (–¥–∞—Ç—á–∏–∫ 3) –∏ —Ç–æ–ª—â–∏–Ω–∞ –¥–Ω–∞ (–¥–∞—Ç—á–∏–∫ 4)"""
        sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
        if sensor3_mm is None or sensor4_mm is None:
            time.sleep(0.001)
            return

        self.frequency_counter += 1
        current_time = time.time()
        if self.frequency_start_time is not None and current_time - self.last_frequency_display >= 1.0:
            elapsed = current_time - self.frequency_start_time
            if elapsed > 0:
                instant_freq = self.frequency_counter / elapsed
                print(f" [CMD=20] –ß–∞—Å—Ç–æ—Ç–∞ –æ–ø—Ä–æ—Å–∞: {instant_freq:.1f} –ì—Ü | –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.frequency_counter}")
            self.last_frequency_display = current_time

        if not (self.is_valid_measurement(sensor3_mm) and self.is_valid_measurement(sensor4_mm)):
            return

        self.temp_sensor3_flange_only_buffer.append(sensor3_mm)
        self.temp_sensor4_buffer.append(sensor4_mm)
        if len(self.temp_sensor3_flange_only_buffer) < 10:
            return

        sorted_sensor3 = sorted(self.temp_sensor3_flange_only_buffer)
        sorted_sensor4 = sorted(self.temp_sensor4_buffer)
        median_sensor3 = (sorted_sensor3[4] + sorted_sensor3[5]) / 2.0
        median_sensor4 = (sorted_sensor4[4] + sorted_sensor4[5]) / 2.0

        filtered_sensor3 = [v for v in self.temp_sensor3_flange_only_buffer if abs(v - median_sensor3) <= 1.5]
        filtered_sensor4 = [v for v in self.temp_sensor4_buffer if abs(v - median_sensor4) <= 1.5]

        if len(filtered_sensor3) >= 5:
            avg_sensor3 = sum(filtered_sensor3) / len(filtered_sensor3)
        else:
            avg_sensor3 = median_sensor3

        if len(filtered_sensor4) >= 5:
            avg_sensor4 = sum(filtered_sensor4) / len(filtered_sensor4)
        else:
            avg_sensor4 = median_sensor4

        flange_radius = self.cached_distance_sensor3_to_center - avg_sensor3
        bottom_thickness = self.cached_distance_sensor4 - avg_sensor4 + self.read_bottom_thickness_offset_coeff()
        self.flange_diameter_buffer.append(flange_radius)
        self.bottom_thickness_buffer.append(bottom_thickness)

        self.temp_sensor3_flange_only_buffer.clear()
        self.temp_sensor4_buffer.clear()

    def handle_measure_flange_only_process_state(self):
        """CMD=20: –†–∞–∑–¥–µ–ª—å–Ω—ã–π —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞ –∏ —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞"""
        if not hasattr(self, '_flange_only_measurement_started'):
            self._flange_only_measurement_started = True
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            self.frequency_counter = 0
            self.frequency_start_time = time.time()
            self.last_frequency_display = self.frequency_start_time
            self.cached_distance_sensor3_to_center = self.read_calibrated_distance_sensor3_to_center()
            self.cached_distance_sensor4 = self.read_calibrated_distance_sensor4()
            print(" [CMD=20] –ë—É—Ñ–µ—Ä—ã –æ—á–∏—â–µ–Ω—ã, —Å—Ç–∞—Ä—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è —Ñ–ª–∞–Ω—Ü–∞ –∏ —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞")

        if self.cached_distance_sensor3_to_center is None:
            print(" [CMD=20] –û—à–∏–±–∫–∞: –Ω–µ—Ç –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫3-—Ü–µ–Ω—Ç—Ä –¥–ª—è —Ñ–ª–∞–Ω—Ü–∞")
            return
        if self.cached_distance_sensor4 is None:
            print(" [CMD=20] –û—à–∏–±–∫–∞: –Ω–µ—Ç –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫–∞ 4 –¥–ª—è —Ç–æ–ª—â–∏–Ω—ã –¥–Ω–∞")
            return

        self.collect_flange_and_bottom_measurement()

    def handle_calculate_flange_only_state(self):
        """CMD=21: –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞"""
        try:
            if not self.flange_only_calculated:
                print(" [CMD=21] –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞...")
                self.process_flange_only_measurement_results()
                self.write_cycle_flag(212)
                self.flange_only_calculated = True
                print(" [STATUS=212] –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ —Ñ–ª–∞–Ω—Ü–∞ –∑–∞–≤–µ—Ä—à—ë–Ω")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –ø–æ–¥—Å—á—ë—Ç–∞ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ —Ñ–ª–∞–Ω—Ü–∞: {e}")
            self.current_state = SystemState.ERROR

    def handle_measure_body_only_process_state(self):
        """CMD=30: –†–∞–∑–¥–µ–ª—å–Ω—ã–π —Å–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞"""
        if not hasattr(self, '_body_only_measurement_started'):
            self._body_only_measurement_started = True
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            self.frequency_counter = 0
            self.frequency_start_time = time.time()
            self.last_frequency_display = self.frequency_start_time
            self.cached_distance_sensor3_to_center_body = self.read_calibrated_distance_sensor3_to_center_body()
            print(" [CMD=30] –ë—É—Ñ–µ—Ä—ã –æ—á–∏—â–µ–Ω—ã, —Å—Ç–∞—Ä—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞")

        if self.cached_distance_sensor3_to_center_body is None:
            print(" [CMD=30] –û—à–∏–±–∫–∞: –Ω–µ—Ç –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫3-—Ü–µ–Ω—Ç—Ä (—Ä–∞–∑–¥–µ–ª—å–Ω—ã–π –∫–æ—Ä–ø—É—Å)")
            return

        self.collect_sensor3_radius_measurement(
            self.temp_sensor3_body_only_buffer,
            self.body_only_diameter_buffer,
            self.cached_distance_sensor3_to_center_body,
            "CMD=30",
        )

    def handle_calculate_body_only_state(self):
        """CMD=31: –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞"""
        try:
            if not self.body_only_calculated:
                print(" [CMD=31] –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞...")
                self.process_body_only_measurement_results()
                self.write_cycle_flag(312)
                self.body_only_calculated = True
                print(" [STATUS=312] –ü–æ–¥—Å—á—ë—Ç —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ –∑–∞–≤–µ—Ä—à—ë–Ω")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –ø–æ–¥—Å—á—ë—Ç–∞ —Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞: {e}")
            self.current_state = SystemState.ERROR

    def handle_measure_body2_process_state(self):
        """CMD=40: –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2"""
        if not hasattr(self, '_body2_measurement_started'):
            self._body2_measurement_started = True
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            self.frequency_counter = 0
            self.frequency_start_time = time.time()
            self.last_frequency_display = self.frequency_start_time
            self.cached_distance_sensor3_to_center_body2 = self.read_calibrated_distance_sensor3_to_center_body2()
            self.body2_quality_required = True
            print(" [CMD=40] –ë—É—Ñ–µ—Ä—ã –æ—á–∏—â–µ–Ω—ã, —Å—Ç–∞—Ä—Ç –∏–∑–º–µ—Ä–µ–Ω–∏—è –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2")

        if self.cached_distance_sensor3_to_center_body2 is None:
            print(" [CMD=40] –û—à–∏–±–∫–∞: –Ω–µ—Ç –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏—è –¥–∞—Ç—á–∏–∫3-—Ü–µ–Ω—Ç—Ä (–∫–æ—Ä–ø—É—Å 2)")
            return

        self.collect_sensor3_radius_measurement(
            self.temp_sensor3_body2_buffer,
            self.body2_diameter_buffer,
            self.cached_distance_sensor3_to_center_body2,
            "CMD=40",
        )

    def handle_calculate_body2_state(self):
        """CMD=41: –ü–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2"""
        try:
            if not self.body2_calculated:
                print(" [CMD=41] –ü–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2...")
                self.process_body2_measurement_results()
                self.write_cycle_flag(412)
                self.body2_calculated = True
                print(" [STATUS=412] –ü–æ–¥—Å—á—ë—Ç –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2 –∑–∞–≤–µ—Ä—à—ë–Ω")
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –ø–æ–¥—Å—á—ë—Ç–∞ –¥–∏–∞–º–µ—Ç—Ä–∞ –∫–æ—Ä–ø—É—Å–∞ 2: {e}")
            self.current_state = SystemState.ERROR
    
    def handle_measure_bottom_process_state(self):
        """
        CMD=14: –°–±–æ—Ä –¥–∞–Ω–Ω—ã—Ö –∏–∑–º–µ—Ä–µ–Ω–∏—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        –ü—Ä–æ—Å—Ç–æ —Å–æ–±–∏—Ä–∞–µ–º –¥–∞–Ω–Ω—ã–µ, –Ω–µ –¥–µ–ª–∞–µ–º –ø–æ–¥—Å—á—ë—Ç
        """
        # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π (–æ–¥–∏–Ω —Ä–∞–∑ –ø—Ä–∏ –≤—Ö–æ–¥–µ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–µ)
        if not hasattr(self, '_bottom_measurement_started'):
            self._bottom_measurement_started = True
            self.clear_measurement_buffers()
            self.clear_serial_buffers()
            # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—á–µ—Ç—á–∏–∫–∏ —á–∞—Å—Ç–æ—Ç—ã –¥–ª—è –Ω–æ–≤–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è (–≤–∞–∂–Ω–æ!)
            self.frequency_counter = 0
            self.frequency_start_time = None
            self.last_frequency_display = 0
            # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ñ–ª–∞–≥ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏ —á–∞—Å—Ç–æ—Ç—ã –¥–ª—è –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
            if hasattr(self, '_bottom_frequency_initialized'):
                delattr(self, '_bottom_frequency_initialized')
            print(f" [CMD=14] –ë—É—Ñ–µ—Ä—ã –æ—á–∏—â–µ–Ω—ã –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º –∏–∑–º–µ—Ä–µ–Ω–∏–π, —Å—á–µ—Ç—á–∏–∫–∏ —á–∞—Å—Ç–æ—Ç—ã —Å–±—Ä–æ—à–µ–Ω—ã")
        
        # –ó–∞–≥—Ä—É–∂–∞–µ–º –∫–∞–ª–∏–±—Ä–æ–≤–∞–Ω–Ω–æ–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ –≤ –∫–µ—à (–æ–¥–∏–Ω —Ä–∞–∑ –ø—Ä–∏ –≤—Ö–æ–¥–µ –≤ —Å–æ—Å—Ç–æ—è–Ω–∏–µ)
        # –ò—Å–ø–æ–ª—å–∑—É–µ–º —Ç–æ—Ç –∂–µ –∫–µ—à —á—Ç–æ –∏ –¥–ª—è CMD=10, —Ç.–∫. —ç—Ç–æ —Ç–æ –∂–µ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ 1-2
        if self.cached_distance_1_2 is None:
            self.cached_distance_1_2 = self.read_calibrated_distance_1_2()
            print(f" –ó–∞–≥—Ä—É–∂–µ–Ω–æ —Ä–∞—Å—Å—Ç–æ—è–Ω–∏–µ 1-2: {self.cached_distance_1_2:.3f}–º–º")
        
        # –ü–µ—Ä–µ–Ω–∞–ø—Ä–∞–≤–ª—è–µ–º –Ω–∞ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π –º–µ—Ç–æ–¥
        self.handle_measure_bottom_state()
    
    def handle_calculate_bottom_state(self):
        """
        CMD=15: –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏
        """
        try:
            # –í—ã–ø–æ–ª–Ω—è–µ–º –ø–æ–¥—Å—á—ë—Ç —Ç–æ–ª—å–∫–æ –æ–¥–∏–Ω —Ä–∞–∑
            if not self.bottom_calculated:
                print(" [CMD=15] –ü–æ–¥—Å—á—ë—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏...")
                
                # –í—ã–∑—ã–≤–∞–µ–º —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π –º–µ—Ç–æ–¥ –ø–æ–¥—Å—á—ë—Ç–∞
                self.process_bottom_wall_measurement_results()
                
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å "–≥–æ—Ç–æ–≤–æ –∫ —Å–ª–µ–¥—É—é—â–µ–π –∫–æ–º–∞–Ω–¥–µ"
                self.write_cycle_flag(114)
                print(" [STATUS=114] –ü–æ–¥—Å—á—ë—Ç –∑–∞–≤–µ—Ä—à—ë–Ω, –≥–æ—Ç–æ–≤ –∫ CMD=16")
                
                # –û—Ç–º–µ—á–∞–µ–º —á—Ç–æ —Ä–∞—Å—á—ë—Ç –≤—ã–ø–æ–ª–Ω–µ–Ω
                self.bottom_calculated = True
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –ø–æ–¥—Å—á—ë—Ç–∞ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –Ω–∏–∂–Ω–µ–π —Å—Ç–µ–Ω–∫–∏: {e}")
            self.current_state = SystemState.ERROR
    
    def handle_quality_evaluation_state(self):
        """
        CMD=16: –û—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞ –∏–∑–¥–µ–ª–∏—è
        """
        try:
            # –í—ã–ø–æ–ª–Ω—è–µ–º –æ—Ü–µ–Ω–∫—É —Ç–æ–ª—å–∫–æ –æ–¥–∏–Ω —Ä–∞–∑
            if not self.quality_evaluated:
                print(" [CMD=16] –û—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞ –∏–∑–¥–µ–ª–∏—è...")
                
                # –ß–∏—Ç–∞–µ–º –∏–∑–º–µ—Ä–µ–Ω–Ω—É—é –≤—ã—Å–æ—Ç—É –∑–∞–≥–æ—Ç–æ–≤–∫–∏
                measured_height = self.read_measured_height()
                if measured_height is not None:
                    print(f" [CMD=16] –ò–∑–º–µ—Ä–µ–Ω–Ω–∞—è –≤—ã—Å–æ—Ç–∞ –∑–∞–≥–æ—Ç–æ–≤–∫–∏ (40057-40058): {measured_height:.3f}–º–º")
                else:
                    print(f" [CMD=16] –û–®–ò–ë–ö–ê: –ù–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å –∏–∑–º–µ—Ä–µ–Ω–Ω—É—é –≤—ã—Å–æ—Ç—É –∑–∞–≥–æ—Ç–æ–≤–∫–∏!")
                
                # –û—Ü–µ–Ω–∫–∞ –∫–∞—á–µ—Å—Ç–≤–∞ –∏–∑–¥–µ–ª–∏—è
                quality_result = self.evaluate_product_quality()
                
                # –ò–∑–≤–ª–µ–∫–∞–µ–º –∏—Ç–æ–≥–æ–≤—ã–π —Ä–µ–∑—É–ª—å—Ç–∞—Ç
                result = quality_result.get('result', 'BAD') if isinstance(quality_result, dict) else quality_result
                
                # –û–±–Ω–æ–≤–ª–µ–Ω–∏–µ —Å—á—ë—Ç—á–∏–∫–æ–≤ –∏–∑–¥–µ–ª–∏–π
                self.update_product_counters(result)
                
                # –ò–Ω–∫—Ä–µ–º–µ–Ω—Ç–∞—Ü–∏—è —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤
                if isinstance(quality_result, dict):
                    self.increment_parameter_statistics(quality_result)
                
                # –£–≤–µ–ª–∏—á–∏–≤–∞–µ–º –Ω–æ–º–µ—Ä –∏–∑–¥–µ–ª–∏—è
                self.increment_product_number()
                
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å "–≥–æ—Ç–æ–≤–æ –∫ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—é"
                self.write_cycle_flag(116)
                print(f" [STATUS=116] –û—Ü–µ–Ω–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∞ ({result}), –≥–æ—Ç–æ–≤ –∫ CMD=0")
                
                # –û—Ç–º–µ—á–∞–µ–º —á—Ç–æ –æ—Ü–µ–Ω–∫–∞ –≤—ã–ø–æ–ª–Ω–µ–Ω–∞
                self.quality_evaluated = True
                
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã –ø–æ—Å–ª–µ —Ä–∞—Å—á–µ—Ç–æ–≤ –∫–æ–º–∞–Ω–¥—ã 16
                self.clear_measurement_buffers()
                self.clear_serial_buffers()
                print(f" [CMD=16] –ë—É—Ñ–µ—Ä—ã –æ—á–∏—â–µ–Ω—ã –ø–æ—Å–ª–µ –æ—Ü–µ–Ω–∫–∏ –∫–∞—á–µ—Å—Ç–≤–∞")
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ –æ—Ü–µ–Ω–∫–∏ –∫–∞—á–µ—Å—Ç–≤–∞: {e}")
            self.current_state = SystemState.ERROR
    
    # ===== –ö–û–ù–ï–¶ –ù–û–í–´–• –ú–ï–¢–û–î–û–í =====
    
    def handle_stream_quad_state(self):
        """–ü–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º QUAD (CMD=200) - –≤—Å–µ 4 –¥–∞—Ç—á–∏–∫–∞ –æ–¥–Ω–æ–≤—Ä–µ–º–µ–Ω–Ω–æ"""
        if not self.sensors:
            print(" –û—à–∏–±–∫–∞: –¥–∞—Ç—á–∏–∫–∏ –Ω–µ –ø–æ–¥–∫–ª—é—á–µ–Ω—ã!")
            self.current_state = SystemState.ERROR
            return
        
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å–º–µ–Ω—É –∫–æ–º–∞–Ω–¥—ã
        current_cmd = self.get_current_command()
        if current_cmd != 200:
            print(f" –ö–æ–º–∞–Ω–¥–∞ –∏–∑–º–µ–Ω–∏–ª–∞—Å—å —Å 200 –Ω–∞ {current_cmd}. –í—ã—Ö–æ–¥–∏–º –∏–∑ QUAD –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ —Ä–µ–∂–∏–º–∞")
            # –û—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º QUAD —Ä–µ–∂–∏–º
            self.stream_active_quad = False
            self.handle_command(current_cmd)
            return
        
        try:
            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –ø—Ä–∏ –ø–µ—Ä–≤–æ–º –∑–∞–ø—É—Å–∫–µ
            if not self.stream_active_quad:
                # –û—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º QUAD —Ä–µ–∂–∏–º–∞
                self.clear_serial_buffers()
                self.flush_sensor_queue()
                self.stream_active_quad = True
                self.stream_measurement_count = 0
                self.stream_start_time = time.time()
                self.stream_temp_sensor1_buffer = []
                self.stream_temp_sensor2_buffer = []
                self.stream_temp_sensor3_buffer = []
                self.stream_temp_sensor4_buffer = []
                print(" –ó–∞–ø—É—â–µ–Ω QUAD –ø–æ—Ç–æ–∫–æ–≤—ã–π —Ä–µ–∂–∏–º (–≤—Å–µ 4 –¥–∞—Ç—á–∏–∫–∞)")
            
            # –ó–∞–±–∏—Ä–∞–µ–º –æ—á–µ—Ä–µ–¥–Ω–æ–µ –∏–∑–º–µ—Ä–µ–Ω–∏–µ –∏–∑ –ø–æ—Ç–æ–∫–∞ —á—Ç–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤
            sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm = self.read_sensors_safe()
            if None in (sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm):
                # –î–∞–Ω–Ω—ã–µ –µ—â—ë –Ω–µ –≥–æ—Ç–æ–≤—ã ‚Äì –¥–µ–ª–∞–µ–º –∫–æ—Ä–æ—Ç–∫—É—é –ø–∞—É–∑—É –∏ –ø–æ–≤—Ç–æ—Ä—è–µ–º —Ü–∏–∫–ª
                time.sleep(0.001)
                return
            
            self.stream_measurement_count += 1
            
            # –ü–µ—Ä–∏–æ–¥–∏—á–µ—Å–∫–∞—è –æ—á–∏—Å—Ç–∫–∞ –±—É—Ñ–µ—Ä–æ–≤ —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ (–∫–∞–∂–¥—ã–µ 10 —Å–µ–∫—É–Ω–¥) –¥–ª—è –ø—Ä–µ–¥–æ—Ç–≤—Ä–∞—â–µ–Ω–∏—è –ø–µ—Ä–µ–ø–æ–ª–Ω–µ–Ω–∏—è
            # –ù–∞ —Å–ª–∞–±—ã—Ö –ø—Ä–æ—Ü–µ—Å—Å–æ—Ä–∞—Ö –±—É—Ñ–µ—Ä—ã –º–æ–≥—É—Ç –ø–µ—Ä–µ–ø–æ–ª–Ω—è—Ç—å—Å—è –±—ã—Å—Ç—Ä–µ–µ
            if self.stream_measurement_count % 1000 == 0:  # ~10 —Å–µ–∫—É–Ω–¥ –ø—Ä–∏ 109 –ì—Ü
                self.clear_serial_buffers()
                if self.stream_measurement_count % 3000 == 0:  # –°–æ–æ–±—â–µ–Ω–∏–µ —Ç–æ–ª—å–∫–æ –∫–∞–∂–¥—ã–µ 30 —Å–µ–∫—É–Ω–¥
                    print(f" [CMD=200] –ü–µ—Ä–∏–æ–¥–∏—á–µ—Å–∫–∞—è –æ—á–∏—Å—Ç–∫–∞ –±—É—Ñ–µ—Ä–æ–≤ —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞ (–∏–∑–º–µ—Ä–µ–Ω–∏–µ #{self.stream_measurement_count})")
            
            # –ü—Ä–æ–≤–µ—Ä–∫–∞ –≤–∞–ª–∏–¥–Ω–æ—Å—Ç–∏ –¥–∞–Ω–Ω—ã—Ö: —Ñ–∏–ª—å—Ç—Ä—É–µ–º —è–≤–Ω–æ –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è
            # –î–∏–∞–ø–∞–∑–æ–Ω: 0-50 –º–º (–±–∞–∑–æ–≤–æ–µ 20 + –¥–∏–∞–ø–∞–∑–æ–Ω 25)
            max_valid_value = 50.0
            min_valid_value = 0.0
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –≤–∞–ª–∏–¥–Ω–æ—Å—Ç—å –∫–∞–∂–¥–æ–≥–æ –∑–Ω–∞—á–µ–Ω–∏—è
            values_valid = True
            invalid_sensors = []
            
            if sensor1_mm is not None and (sensor1_mm < min_valid_value or sensor1_mm > max_valid_value):
                values_valid = False
                invalid_sensors.append(1)
            if sensor2_mm is not None and (sensor2_mm < min_valid_value or sensor2_mm > max_valid_value):
                values_valid = False
                invalid_sensors.append(2)
            if sensor3_mm is not None and (sensor3_mm < min_valid_value or sensor3_mm > max_valid_value):
                values_valid = False
                invalid_sensors.append(3)
            if sensor4_mm is not None and (sensor4_mm < min_valid_value or sensor4_mm > max_valid_value):
                values_valid = False
                invalid_sensors.append(4)
            
            # –í—ã–≤–æ–¥–∏–º –æ—à–∏–±–∫–∏ –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π –Ω–µ —á–∞—â–µ —Ä–∞–∑–∞ –≤ 5 —Å–µ–∫—É–Ω–¥
            if invalid_sensors and self.stream_measurement_count % 500 == 0:
                print(f" [CMD=200] ‚ö† –ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ {invalid_sensors}: "
                      f"–î1={sensor1_mm:.3f}–º–º –î2={sensor2_mm:.3f}–º–º –î3={sensor3_mm:.3f}–º–º –î4={sensor4_mm:.3f}–º–º (–¥–æ–ª–∂–Ω–æ –±—ã—Ç—å 0-50)")
                # –ü—Ä–∏ –æ–±–Ω–∞—Ä—É–∂–µ–Ω–∏–∏ –Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã—Ö –∑–Ω–∞—á–µ–Ω–∏–π –æ—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã
                self.clear_serial_buffers()
            
            # –ï—Å–ª–∏ –ø–æ–ª—É—á–∏–ª–∏ –≤–∞–ª–∏–¥–Ω—ã–µ –∏–∑–º–µ—Ä–µ–Ω–∏—è –æ—Ç –≤—Å–µ—Ö –¥–∞—Ç—á–∏–∫–æ–≤ –ò –∑–Ω–∞—á–µ–Ω–∏—è –≤ –¥–æ–ø—É—Å—Ç–∏–º–æ–º –¥–∏–∞–ø–∞–∑–æ–Ω–µ
            if (all(v is not None for v in [sensor1_mm, sensor2_mm, sensor3_mm, sensor4_mm]) and values_valid):
                # –î–æ–±–∞–≤–ª—è–µ–º –≤ –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —É—Å—Ä–µ–¥–Ω–µ–Ω–∏—è
                self.stream_temp_sensor1_buffer.append(sensor1_mm)
                self.stream_temp_sensor2_buffer.append(sensor2_mm)
                self.stream_temp_sensor3_buffer.append(sensor3_mm)
                self.stream_temp_sensor4_buffer.append(sensor4_mm)
                
                # –ö–æ–≥–¥–∞ –Ω–∞–∫–æ–ø–∏–ª–æ—Å—å 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π - —É—Å—Ä–µ–¥–Ω—è–µ–º –∏ –∑–∞–ø–∏—Å—ã–≤–∞–µ–º –≤ —Ä–µ–≥–∏—Å—Ç—Ä—ã
                if len(self.stream_temp_sensor1_buffer) >= 10:
                    # –í—ã—á–∏—Å–ª—è–µ–º —Å—Ä–µ–¥–Ω–∏–µ –∑–Ω–∞—á–µ–Ω–∏—è –¥–ª—è –∫–∞–∂–¥–æ–≥–æ –¥–∞—Ç—á–∏–∫–∞
                    avg_sensor1 = sum(self.stream_temp_sensor1_buffer) / len(self.stream_temp_sensor1_buffer)
                    avg_sensor2 = sum(self.stream_temp_sensor2_buffer) / len(self.stream_temp_sensor2_buffer)
                    avg_sensor3 = sum(self.stream_temp_sensor3_buffer) / len(self.stream_temp_sensor3_buffer)
                    avg_sensor4 = sum(self.stream_temp_sensor4_buffer) / len(self.stream_temp_sensor4_buffer)
                    
                    # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º –≤—Å–µ 4 —Ä–µ–≥–∏—Å—Ç—Ä–∞ –æ–¥–Ω–æ–≤—Ä–µ–º–µ–Ω–Ω–æ
                    # –û–ø—Ç–∏–º–∏–∑–∞—Ü–∏—è: –≥—Ä—É–ø–ø–∏—Ä—É–µ–º –∑–∞–ø–∏—Å–∏ –¥–ª—è —Å–Ω–∏–∂–µ–Ω–∏—è –Ω–∞–≥—Ä—É–∑–∫–∏ –Ω–∞ Modbus
                    try:
                        # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º –≤—Å–µ —Ä–µ–≥–∏—Å—Ç—Ä—ã –ø–æ—Å–ª–µ–¥–æ–≤–∞—Ç–µ–ª—å–Ω–æ (–±—ã—Å—Ç—Ä–µ–µ —á–µ–º –ø–∞—Ä–∞–ª–ª–µ–ª—å–Ω–æ)
                        self.write_stream_result_to_input_registers(avg_sensor1, 30001)  # –î–∞—Ç—á–∏–∫ 1
                        self.write_stream_result_to_input_registers(avg_sensor2, 30003)  # –î–∞—Ç—á–∏–∫ 2
                        self.write_stream_result_to_input_registers(avg_sensor3, 30005)  # –î–∞—Ç—á–∏–∫ 3
                        self.write_stream_result_to_input_registers(avg_sensor4, 30007)  # –î–∞—Ç—á–∏–∫ 4
                    except Exception as e:
                        print(f" –û–®–ò–ë–ö–ê –ó–ê–ü–ò–°–ò –í –†–ï–ì–ò–°–¢–†–´: {e}")
                        # –ü—Ä–∏ –æ—à–∏–±–∫–µ –∑–∞–ø–∏—Å–∏ –æ—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã —Å–µ—Ä–∏–π–Ω–æ–≥–æ –ø–æ—Ä—Ç–∞
                        self.clear_serial_buffers()
                    
                    # –í—ã–≤–æ–¥–∏–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç —Ä–∞–∑ –≤ —Å–µ–∫—É–Ω–¥—É
                    current_time = time.time()
                    if not hasattr(self, '_last_stream_quad_print'):
                        self._last_stream_quad_print = current_time
                    
                    if current_time - self._last_stream_quad_print >= 1.0:
                        elapsed = current_time - self.stream_start_time
                        frequency = self.stream_measurement_count / elapsed if elapsed > 0 else 0
                        
                        print(f" [CMD=200] QUAD: {elapsed:5.1f}—Å | –ò–∑–º–µ—Ä–µ–Ω–∏–π: {self.stream_measurement_count:6d} | "
                              f"–ß–∞—Å—Ç–æ—Ç–∞: {frequency:7.1f} –ì—Ü | "
                              f"–î1={avg_sensor1:.3f}–º–º –î2={avg_sensor2:.3f}–º–º –î3={avg_sensor3:.3f}–º–º –î4={avg_sensor4:.3f}–º–º")
                        
                        self._last_stream_quad_print = current_time
                    
                    # –û—á–∏—â–∞–µ–º –≤—Ä–µ–º–µ–Ω–Ω—ã–µ –±—É—Ñ–µ—Ä—ã –¥–ª—è —Å–ª–µ–¥—É—é—â–∏—Ö 10 –∏–∑–º–µ—Ä–µ–Ω–∏–π
                    self.stream_temp_sensor1_buffer = []
                    self.stream_temp_sensor2_buffer = []
                    self.stream_temp_sensor3_buffer = []
                    self.stream_temp_sensor4_buffer = []
            else:
                # –û—à–∏–±–∫–∞ –ø–æ–ª—É—á–µ–Ω–∏—è –¥–∞–Ω–Ω—ã—Ö (None –∑–Ω–∞—á–µ–Ω–∏—è)
                # –ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è —É–∂–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã –≤—ã—à–µ —Å –æ—á–∏—Å—Ç–∫–æ–π –±—É—Ñ–µ—Ä–æ–≤
                if not values_valid:
                    # –ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è —É–∂–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã –≤—ã—à–µ - –ø—Ä–æ—Å—Ç–æ –ø—Ä–æ–ø—É—Å–∫–∞–µ–º
                    pass
                elif self.stream_measurement_count % 500 == 0:  # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –æ—à–∏–±–∫—É –Ω–µ —á–∞—â–µ —Ä–∞–∑–∞ –≤ 5 —Å–µ–∫—É–Ω–¥
                    print(f" [CMD=200] –û—à–∏–±–∫–∞ –∏–∑–º–µ—Ä–µ–Ω–∏—è (None): –î1={sensor1_mm}, –î2={sensor2_mm}, –î3={sensor3_mm}, –î4={sensor4_mm}")
                    # –ü—Ä–∏ None –∑–Ω–∞—á–µ–Ω–∏—è—Ö —Ç–∞–∫–∂–µ –æ—á–∏—â–∞–µ–º –±—É—Ñ–µ—Ä—ã
                    self.clear_serial_buffers()
            
        except Exception as e:
            print(f" –û—à–∏–±–∫–∞ QUAD –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ —Ä–µ–∂–∏–º–∞: {e}")
            # –û—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º QUAD —Ä–µ–∂–∏–º –ø—Ä–∏ –æ—à–∏–±–∫–µ
            self.stream_active_quad = False
            self.current_state = SystemState.ERROR
    def write_stream_result_to_input_registers(self, value: float, base_address: int):
        """–ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –ø–æ—Ç–æ–∫–æ–≤–æ–≥–æ –∏–∑–º–µ—Ä–µ–Ω–∏—è –≤ Input —Ä–µ–≥–∏—Å—Ç—Ä—ã"""
        try:
            if self.modbus_server and self.modbus_server.slave_context:
                # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º float –≤ –¥–≤–∞ 16-–±–∏—Ç–Ω—ã—Ö —Ä–µ–≥–∏—Å—Ç—Ä–∞
                low_word, high_word = self.float_to_doubleword(value)
                
                # –í—ã—á–∏—Å–ª—è–µ–º –∏–Ω–¥–µ–∫—Å—ã —Ä–µ–≥–∏—Å—Ç—Ä–æ–≤
                # –ü–æ –æ–ø–∏—Å–∞–Ω–∏—é: base_address - 1 —Å–æ–¥–µ—Ä–∂–∏—Ç –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ, base_address —Å–æ–¥–µ—Ä–∂–∏—Ç —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ
                # –ù–∞–ø—Ä–∏–º–µ—Ä, –¥–ª—è 30052: –∏–Ω–¥–µ–∫—Å 51 = –º–ª–∞–¥—à–µ–µ —Å–ª–æ–≤–æ, –∏–Ω–¥–µ–∫—Å 52 = —Å—Ç–∞—Ä—à–µ–µ —Å–ª–æ–≤–æ
                reg_index_low = base_address - 30000    # –ú–ª–∞–¥—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä (base_address - 1)
                reg_index_high = base_address - 30000 + 1      # –°—Ç–∞—Ä—à–∏–π —Ä–µ–≥–∏—Å—Ç—Ä (base_address)
                
                # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º –≤ Input —Ä–µ–≥–∏—Å—Ç—Ä—ã (—Ñ—É–Ω–∫—Ü–∏—è 4)
                self.modbus_server.slave_context.setValues(4, reg_index_low, [int(high_word)])    # –ú–ª–∞–¥—à–∏–π
                self.modbus_server.slave_context.setValues(4, reg_index_high, [int(low_word)])  # –°—Ç–∞—Ä—à
                
        except Exception as e:
            print(f" –û–®–ò–ë–ö–ê –ó–ê–ü–ò–°–ò –í INPUT –†–ï–ì–ò–°–¢–†–´ {base_address}-{base_address+1}: {e}")
            import traceback
            traceback.print_exc()
    
    def handle_error_state(self):
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ —Å–æ—Å—Ç–æ—è–Ω–∏—è –æ—à–∏–±–∫–∏"""
        print(" –°–æ—Å—Ç–æ—è–Ω–∏–µ –æ—à–∏–±–∫–∏. –ü—Ä–æ–≤–µ—Ä—å—Ç–µ —Å–∏—Å—Ç–µ–º—É.")


def check_single_instance():
    """–ü—Ä–æ–≤–µ—Ä–∫–∞, —á—Ç–æ –∑–∞–ø—É—â–µ–Ω —Ç–æ–ª—å–∫–æ –æ–¥–∏–Ω —ç–∫–∑–µ–º–ø–ª—è—Ä –ø—Ä–æ–≥—Ä–∞–º–º—ã"""
    current_pid = os.getpid()
    script_name = os.path.basename(__file__)
    
    # –ú–µ—Ç–æ–¥ 1: –ü—Ä–æ–≤–µ—Ä–∫–∞ —á–µ—Ä–µ–∑ –ø–æ–∏—Å–∫ –¥—Ä—É–≥–∏—Ö –ø—Ä–æ—Ü–µ—Å—Å–æ–≤ —Å —Ç–µ–º –∂–µ —Å–∫—Ä–∏–ø—Ç–æ–º
    if HAS_PSUTIL:
        try:
            found_processes = []
            for proc in psutil.process_iter(['pid', 'name', 'cmdline']):
                try:
                    if proc.info['pid'] == current_pid:
                        continue
                    
                    cmdline = proc.info.get('cmdline', [])
                    if cmdline and any(script_name in str(arg) for arg in cmdline):
                        # –ù–∞–π–¥–µ–Ω –¥—Ä—É–≥–æ–π –ø—Ä–æ—Ü–µ—Å—Å —Å —Ç–µ–º –∂–µ —Å–∫—Ä–∏–ø—Ç–æ–º
                        other_pid = proc.info['pid']
                        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ —ç—Ç–æ –¥–µ–π—Å—Ç–≤–∏—Ç–µ–ª—å–Ω–æ –Ω–∞—à —Å–∫—Ä–∏–ø—Ç
                        if any('laser_geometry_system.py' in str(arg) for arg in cmdline):
                            found_processes.append(other_pid)
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    continue
            
            # –ï—Å–ª–∏ –Ω–∞–π–¥–µ–Ω—ã –¥—Ä—É–≥–∏–µ –ø—Ä–æ—Ü–µ—Å—Å—ã, –∂–¥–µ–º –Ω–µ–º–Ω–æ–≥–æ –∏ –ø—Ä–æ–≤–µ—Ä—è–µ–º —Å–Ω–æ–≤–∞
            # (–≤–æ–∑–º–æ–∂–Ω–æ, —ç—Ç–æ —Å—Ç–∞—Ä—ã–µ –ø—Ä–æ—Ü–µ—Å—Å—ã, –∫–æ—Ç–æ—Ä—ã–µ –∑–∞–≤–µ—Ä—à–∞—é—Ç—Å—è)
            if found_processes:
                print(f"[–ü–†–û–í–ï–†–ö–ê] –ù–∞–π–¥–µ–Ω—ã –¥—Ä—É–≥–∏–µ –ø—Ä–æ—Ü–µ—Å—Å—ã: {found_processes}, –æ–∂–∏–¥–∞–Ω–∏–µ 2 —Å–µ–∫—É–Ω–¥—ã...")
                time.sleep(2)
                
                # –ü–æ–≤—Ç–æ—Ä–Ω–∞—è –ø—Ä–æ–≤–µ—Ä–∫–∞
                still_running = []
                for pid in found_processes:
                    if os.path.exists(f'/proc/{pid}'):
                        try:
                            proc = psutil.Process(pid)
                            cmdline = proc.cmdline()
                            if any('laser_geometry_system.py' in str(arg) for arg in cmdline):
                                still_running.append(pid)
                        except (psutil.NoSuchProcess, psutil.AccessDenied):
                            pass  # –ü—Ä–æ—Ü–µ—Å—Å –∑–∞–≤–µ—Ä—à–∏–ª—Å—è
                
                if still_running:
                    return False, f"–ü—Ä–æ–≥—Ä–∞–º–º–∞ —É–∂–µ –∑–∞–ø—É—â–µ–Ω–∞ (PID: {still_running})"
        except Exception as e:
            pass  # –ï—Å–ª–∏ –Ω–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ–≤–µ—Ä–∏—Ç—å —á–µ—Ä–µ–∑ psutil, –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º
    
    # –ú–µ—Ç–æ–¥ 2: –ü—Ä–æ–≤–µ—Ä–∫–∞ —á–µ—Ä–µ–∑ socket (–ø–æ—Ä—Ç 502)
    try:
        test_sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        test_sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        test_sock.settimeout(0.1)
        result = test_sock.connect_ex(('192.168.1.50', 502))
        test_sock.close()
        if result == 0:
            # –ü–æ—Ä—Ç –∑–∞–Ω—è—Ç - –≤–æ–∑–º–æ–∂–Ω–æ, –¥—Ä—É–≥–æ–π –ø—Ä–æ—Ü–µ—Å—Å —É–∂–µ –∑–∞–ø—É—â–µ–Ω
            return False, "–ü–æ—Ä—Ç 502 —É–∂–µ –∑–∞–Ω—è—Ç (–≤–æ–∑–º–æ–∂–Ω–æ, –ø—Ä–æ–≥—Ä–∞–º–º–∞ —É–∂–µ –∑–∞–ø—É—â–µ–Ω–∞)"
    except Exception:
        pass  # –ï—Å–ª–∏ –Ω–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ–≤–µ—Ä–∏—Ç—å –ø–æ—Ä—Ç, –ø—Ä–æ–¥–æ–ª–∂–∞–µ–º
    
    # –ú–µ—Ç–æ–¥ 3: Lock —Ñ–∞–π–ª —Å fcntl (—Å–∞–º—ã–π –Ω–∞–¥–µ–∂–Ω—ã–π –¥–ª—è Linux)
    lock_file_path = '/tmp/laser_geometry_system.lock'
    lock_file = None
    
    try:
        # –ü—ã—Ç–∞–µ–º—Å—è –æ—Ç–∫—Ä—ã—Ç—å lock —Ñ–∞–π–ª
        lock_file = open(lock_file_path, 'w')
        
        # –ü—ã—Ç–∞–µ–º—Å—è –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞—Ç—å —Ñ–∞–π–ª (–Ω–µ–±–ª–æ–∫–∏—Ä—É—é—â–∏–π —Ä–µ–∂–∏–º)
        fcntl.flock(lock_file.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
        
        # –ó–∞–ø–∏—Å—ã–≤–∞–µ–º PID —Ç–µ–∫—É—â–µ–≥–æ –ø—Ä–æ—Ü–µ—Å—Å–∞
        lock_file.write(str(current_pid))
        lock_file.flush()
        
        # –§—É–Ω–∫—Ü–∏—è –¥–ª—è –æ—á–∏—Å—Ç–∫–∏ lock —Ñ–∞–π–ª–∞ –ø—Ä–∏ –≤—ã—Ö–æ–¥–µ
        def cleanup_lock():
            try:
                if lock_file:
                    fcntl.flock(lock_file.fileno(), fcntl.LOCK_UN)
                    lock_file.close()
                if os.path.exists(lock_file_path):
                    os.remove(lock_file_path)
            except:
                pass
        
        # –†–µ–≥–∏—Å—Ç—Ä–∏—Ä—É–µ–º –æ—á–∏—Å—Ç–∫—É –ø—Ä–∏ –≤—ã—Ö–æ–¥–µ
        import atexit
        atexit.register(cleanup_lock)
        
        return True, None
        
    except (IOError, OSError) as e:
        # –§–∞–π–ª —É–∂–µ –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω - –¥—Ä—É–≥–æ–π –ø—Ä–æ—Ü–µ—Å—Å –∑–∞–ø—É—â–µ–Ω
        if lock_file:
            lock_file.close()
        
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –¥–µ–π—Å—Ç–≤–∏—Ç–µ–ª—å–Ω–æ –ª–∏ –ø—Ä–æ—Ü–µ—Å—Å –µ—â–µ —Ä–∞–±–æ—Ç–∞–µ—Ç
        try:
            if os.path.exists(lock_file_path):
                with open(lock_file_path, 'r') as f:
                    pid_str = f.read().strip()
                    if pid_str:
                        old_pid = int(pid_str)
                        
                        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —Å—É—â–µ—Å—Ç–≤—É–µ—Ç –ª–∏ –ø—Ä–æ—Ü–µ—Å—Å —Å —ç—Ç–∏–º PID
                        if os.path.exists(f'/proc/{old_pid}'):
                            return False, f"–ü—Ä–æ–≥—Ä–∞–º–º–∞ —É–∂–µ –∑–∞–ø—É—â–µ–Ω–∞ (PID: {old_pid})"
                        else:
                            # –ü—Ä–æ—Ü–µ—Å—Å –Ω–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç, —É–¥–∞–ª—è–µ–º —Å—Ç–∞—Ä—ã–π lock —Ñ–∞–π–ª
                            try:
                                os.remove(lock_file_path)
                                # –ü—ã—Ç–∞–µ–º—Å—è –∑–∞–ø—É—Å—Ç–∏—Ç—å—Å—è —Å–Ω–æ–≤–∞
                                return check_single_instance()
                            except:
                                return False, "–ù–µ —É–¥–∞–ª–æ—Å—å —É–¥–∞–ª–∏—Ç—å —Å—Ç–∞—Ä—ã–π lock —Ñ–∞–π–ª"
                    else:
                        # Lock —Ñ–∞–π–ª –ø—É—Å—Ç–æ–π, —É–¥–∞–ª—è–µ–º –µ–≥–æ
                        try:
                            os.remove(lock_file_path)
                            return check_single_instance()
                        except:
                            return False, "–ù–µ —É–¥–∞–ª–æ—Å—å —É–¥–∞–ª–∏—Ç—å –ø—É—Å—Ç–æ–π lock —Ñ–∞–π–ª"
            else:
                # Lock —Ñ–∞–π–ª –Ω–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç, –Ω–æ fcntl –Ω–µ —Å–º–æ–≥ –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞—Ç—å
                return False, "–ù–µ —É–¥–∞–ª–æ—Å—å –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞—Ç—å —Ñ–∞–π–ª (–≤–æ–∑–º–æ–∂–Ω–æ, –¥—Ä—É–≥–æ–π –ø—Ä–æ—Ü–µ—Å—Å –∏—Å–ø–æ–ª—å–∑—É–µ—Ç –µ–≥–æ)"
        except (ValueError, OSError) as e:
            # –û—à–∏–±–∫–∞ –ø—Ä–∏ —á—Ç–µ–Ω–∏–∏/–ø—Ä–æ–≤–µ—Ä–∫–µ lock —Ñ–∞–π–ª–∞
            try:
                if os.path.exists(lock_file_path):
                    os.remove(lock_file_path)
                    return check_single_instance()
            except:
                pass
            return False, f"–û—à–∏–±–∫–∞ –ø—Ä–æ–≤–µ—Ä–∫–∏ lock —Ñ–∞–π–ª–∞: {e}"
    
    except Exception as e:
        if lock_file:
            lock_file.close()
        return False, f"–û—à–∏–±–∫–∞ –ø—Ä–æ–≤–µ—Ä–∫–∏ –µ–¥–∏–Ω—Å—Ç–≤–µ–Ω–Ω–æ–≥–æ —ç–∫–∑–µ–º–ø–ª—è—Ä–∞: {e}"


def main():
    """–ì–ª–∞–≤–Ω–∞—è —Ñ—É–Ω–∫—Ü–∏—è"""
    # –í–ê–ñ–ù–û: –ü—Ä–æ–≤–µ—Ä–∫–∞ –Ω–∞ –µ–¥–∏–Ω—Å—Ç–≤–µ–Ω–Ω—ã–π —ç–∫–∑–µ–º–ø–ª—è—Ä –¥–æ–ª–∂–Ω–∞ –±—ã—Ç—å –ü–ï–†–í–û–ô –æ–ø–µ—Ä–∞—Ü–∏–µ–π
    # –¥–æ –ª—é–±—ã—Ö –¥—Ä—É–≥–∏—Ö –¥–µ–π—Å—Ç–≤–∏–π, —á—Ç–æ–±—ã –ø—Ä–µ–¥–æ—Ç–≤—Ä–∞—Ç–∏—Ç—å race condition
    can_run, error_msg = check_single_instance()
    if not can_run:
        print(f"–û–®–ò–ë–ö–ê: {error_msg}")
        print("–ü—Ä–æ–≥—Ä–∞–º–º–∞ —É–∂–µ –∑–∞–ø—É—â–µ–Ω–∞. –ó–∞–≤–µ—Ä—à–∞–µ–º...")
        sys.exit(1)
    
    # –ù–µ–±–æ–ª—å—à–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞, —á—Ç–æ–±—ã –¥–∞—Ç—å –≤—Ä–µ–º—è –ø–µ—Ä–≤–æ–º—É –ø—Ä–æ—Ü–µ—Å—Å—É –ø–æ–ª–Ω–æ—Å—Ç—å—é –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞—Ç—å—Å—è
    # –≠—Ç–æ –ø–æ–º–æ–≥–∞–µ—Ç –ø—Ä–µ–¥–æ—Ç–≤—Ä–∞—Ç–∏—Ç—å —Å–∏—Ç—É–∞—Ü–∏—é, –∫–æ–≥–¥–∞ –¥–≤–∞ –ø—Ä–æ—Ü–µ—Å—Å–∞ –ø—Ä–æ—Ö–æ–¥—è—Ç –ø—Ä–æ–≤–µ—Ä–∫—É –ø–æ—á—Ç–∏ –æ–¥–Ω–æ–≤—Ä–µ–º–µ–Ω–Ω–æ
    time.sleep(0.5)
    
    # –ü–æ–≤—Ç–æ—Ä–Ω–∞—è –ø—Ä–æ–≤–µ—Ä–∫–∞ –ø–æ—Å–ª–µ –∑–∞–¥–µ—Ä–∂–∫–∏ (–Ω–∞ —Å–ª—É—á–∞–π, –µ—Å–ª–∏ –¥—Ä—É–≥–æ–π –ø—Ä–æ—Ü–µ—Å—Å —É—Å–ø–µ–ª –∑–∞–ø—É—Å—Ç–∏—Ç—å—Å—è)
    can_run, error_msg = check_single_instance()
    if not can_run:
        print(f"–û–®–ò–ë–ö–ê: {error_msg}")
        print("–û–±–Ω–∞—Ä—É–∂–µ–Ω –¥—Ä—É–≥–æ–π —ç–∫–∑–µ–º–ø–ª—è—Ä –ø–æ—Å–ª–µ –∑–∞–¥–µ—Ä–∂–∫–∏. –ó–∞–≤–µ—Ä—à–∞–µ–º...")
        sys.exit(1)
    
    print("–°–ò–°–¢–ï–ú–ê –õ–ê–ó–ï–†–ù–û–ô –ì–ï–û–ú–ï–¢–†–ò–ò")
    print("–ò–Ω—Ç–µ–≥—Ä–∞—Ü–∏—è –¥–∞—Ç—á–∏–∫–æ–≤ –†–§602 + Modbus + –ê–≤—Ç–æ–º–∞—Ç —Å–æ—Å—Ç–æ—è–Ω–∏–π")
    print("=" * 60)
    
    # –ù–∞—Å—Ç—Ä–æ–π–∫–∏ —Å–∏—Å—Ç–µ–º—ã
    PORT = '/dev/ttyUSB0'  # –ò–∑–º–µ–Ω–∏—Ç–µ –ø—Ä–∏ –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ—Å—Ç–∏ –Ω–∞ –¥—Ä—É–≥–æ–π ttyUSB/ttyACM –ø–æ—Ä—Ç
    BAUDRATE = 921600
    MODBUS_PORT = 502
    TEST_MODE = False  # –†–µ–∂–∏–º —Å —Ä–µ–∞–ª—å–Ω—ã–º–∏ –¥–∞—Ç—á–∏–∫–∞–º–∏
    
    # –°–æ–∑–¥–∞–Ω–∏–µ –∏ –∑–∞–ø—É—Å–∫ —Å–∏—Å—Ç–µ–º—ã
    system = LaserGeometrySystem(PORT, BAUDRATE, MODBUS_PORT, test_mode=TEST_MODE)
    
    try:
        system.start_system()
    except KeyboardInterrupt:
        print("\n –û—Å—Ç–∞–Ω–æ–≤–∫–∞ –ø–æ –∑–∞–ø—Ä–æ—Å—É –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è")
    except Exception as e:
        print(f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞: {e}")
    finally:
        try:
            system.stop_system()
        except Exception:
            pass
        # –§–∏–Ω–∞–ª—å–Ω–∞—è –æ—á–∏—Å—Ç–∫–∞ –æ–ø—Ç–∏–º–∏–∑–∞—Ü–∏–π
        cleanup_laser_system_optimizations()


if __name__ == "__main__":
    main()
 